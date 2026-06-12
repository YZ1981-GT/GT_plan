"""MetadataCodec 单元测试

验证 embed/extract 在 xlsx 和 docx 格式下的正确性。
"""

from __future__ import annotations

import tempfile
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from app.schemas.wp_export_schemas import MetadataBundle
from app.services.wp_export.metadata_codec import MetadataCodec


@pytest.fixture
def codec() -> MetadataCodec:
    return MetadataCodec()


@pytest.fixture
def sample_metadata() -> MetadataBundle:
    return MetadataBundle(
        wp_code="D1",
        project_id=uuid4(),
        file_version=5,
        export_timestamp=datetime(2025, 6, 1, 12, 0, 0, tzinfo=timezone.utc),
        preparer="张三",
        reviewer="李四",
        review_status="approved",
    )


@pytest.fixture
def minimal_metadata() -> MetadataBundle:
    return MetadataBundle(
        wp_code="E2-1",
        project_id=uuid4(),
        file_version=1,
        export_timestamp=datetime(2025, 1, 15, 8, 30, 0, tzinfo=timezone.utc),
        preparer=None,
        reviewer=None,
        review_status=None,
    )


class TestEmbedExtractXlsx:
    """xlsx embed/extract 测试"""

    def test_round_trip_full_metadata(
        self, codec: MetadataCodec, sample_metadata: MetadataBundle
    ) -> None:
        """完整元数据 embed→save→load→extract 应返回相同值"""
        wb = Workbook()
        codec.embed_xlsx(wb, sample_metadata)

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            path = Path(f.name)
        wb.save(str(path))

        try:
            wb2 = load_workbook(str(path))
            extracted = codec.extract_xlsx(wb2)

            assert extracted is not None
            assert extracted.wp_code == sample_metadata.wp_code
            assert extracted.project_id == sample_metadata.project_id
            assert extracted.file_version == sample_metadata.file_version
            assert extracted.export_timestamp == sample_metadata.export_timestamp
            assert extracted.preparer == sample_metadata.preparer
            assert extracted.reviewer == sample_metadata.reviewer
            assert extracted.review_status == sample_metadata.review_status
        finally:
            path.unlink(missing_ok=True)

    def test_round_trip_optional_none(
        self, codec: MetadataCodec, minimal_metadata: MetadataBundle
    ) -> None:
        """可选字段为 None 时 embed→extract 应保持 None"""
        wb = Workbook()
        codec.embed_xlsx(wb, minimal_metadata)

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            path = Path(f.name)
        wb.save(str(path))

        try:
            wb2 = load_workbook(str(path))
            extracted = codec.extract_xlsx(wb2)

            assert extracted is not None
            assert extracted.wp_code == minimal_metadata.wp_code
            assert extracted.preparer is None
            assert extracted.reviewer is None
            assert extracted.review_status is None
        finally:
            path.unlink(missing_ok=True)

    def test_extract_no_metadata_returns_none(self, codec: MetadataCodec) -> None:
        """无元数据的 xlsx 提取应返回 None"""
        wb = Workbook()
        result = codec.extract_xlsx(wb)
        assert result is None

    def test_embed_idempotent(
        self, codec: MetadataCodec, sample_metadata: MetadataBundle
    ) -> None:
        """重复 embed 同一 workbook 不应产生重复属性"""
        wb = Workbook()
        codec.embed_xlsx(wb, sample_metadata)
        codec.embed_xlsx(wb, sample_metadata)

        prop_names = wb.custom_doc_props.names
        # 每个字段只应出现一次
        for field in MetadataCodec.XLSX_PROPS:
            assert prop_names.count(field) == 1


class TestEmbedExtractDocx:
    """docx embed/extract 测试"""

    def test_round_trip_full_metadata(
        self, codec: MetadataCodec, sample_metadata: MetadataBundle
    ) -> None:
        """完整元数据 embed→save→load→extract 应返回相同值"""
        doc = Document()
        codec.embed_docx(doc, sample_metadata)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        doc.save(str(path))

        try:
            doc2 = Document(str(path))
            extracted = codec.extract_docx(doc2)

            assert extracted is not None
            assert extracted.wp_code == sample_metadata.wp_code
            assert extracted.project_id == sample_metadata.project_id
            assert extracted.file_version == sample_metadata.file_version
            assert extracted.export_timestamp == sample_metadata.export_timestamp
            assert extracted.preparer == sample_metadata.preparer
            assert extracted.reviewer == sample_metadata.reviewer
            assert extracted.review_status == sample_metadata.review_status
        finally:
            path.unlink(missing_ok=True)

    def test_round_trip_optional_none(
        self, codec: MetadataCodec, minimal_metadata: MetadataBundle
    ) -> None:
        """可选字段为 None 时 embed→extract 应保持 None"""
        doc = Document()
        codec.embed_docx(doc, minimal_metadata)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = Path(f.name)
        doc.save(str(path))

        try:
            doc2 = Document(str(path))
            extracted = codec.extract_docx(doc2)

            assert extracted is not None
            assert extracted.preparer is None
            assert extracted.reviewer is None
            assert extracted.review_status is None
        finally:
            path.unlink(missing_ok=True)

    def test_extract_empty_comments_returns_none(
        self, codec: MetadataCodec
    ) -> None:
        """空 comments 的 docx 提取应返回 None"""
        doc = Document()
        result = codec.extract_docx(doc)
        assert result is None

    def test_extract_invalid_json_returns_none(
        self, codec: MetadataCodec
    ) -> None:
        """非 JSON 内容的 comments 提取应返回 None"""
        doc = Document()
        doc.core_properties.comments = "这不是 JSON 内容"
        result = codec.extract_docx(doc)
        assert result is None

    def test_extract_missing_required_fields_returns_none(
        self, codec: MetadataCodec
    ) -> None:
        """JSON 但缺少必要字段时应返回 None"""
        import json

        doc = Document()
        doc.core_properties.comments = json.dumps({"some_field": "value"})
        result = codec.extract_docx(doc)
        assert result is None

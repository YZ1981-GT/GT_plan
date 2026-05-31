"""Tests for knowledge document content_text extraction (Task 9: content_text 填充保障).

Validates:
- PDF/docx upload uses MinerU OCR (recognize_for_ocr) for full text extraction
- Graceful degradation to PyPDF2/python-docx when MinerU unavailable
- content_text is populated after upload (保障 spec B 向量索引有内容)

**Validates: Requirements 6.3**
"""

import io
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.routers.knowledge_folders import _extract_text_with_ocr


# ═══════════════════════════════════════════════════════════
# Unit tests for _extract_text_with_ocr
# ═══════════════════════════════════════════════════════════


@pytest.mark.asyncio
async def test_extract_text_mineru_success():
    """MinerU 可用且返回文本时，使用 MinerU 结果"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = True
    mock_mineru.recognize_for_ocr.return_value = {
        "text": "MinerU extracted full text content",
        "engine": "mineru",
        "regions": [],
    }

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.pdf", b"fake pdf content", "test.pdf"
        )

    assert result == "MinerU extracted full text content"
    mock_mineru.is_available.assert_awaited_once()
    mock_mineru.recognize_for_ocr.assert_awaited_once_with("/tmp/test.pdf")


@pytest.mark.asyncio
async def test_extract_text_mineru_unavailable_fallback_pdf():
    """MinerU 不可用时，PDF 降级到 PyPDF2"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = False

    # Create a minimal valid PDF content
    import PyPDF2

    pdf_buffer = io.BytesIO()
    writer = PyPDF2.PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(pdf_buffer)
    pdf_content = pdf_buffer.getvalue()

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.pdf", pdf_content, "test.pdf"
        )

    # Blank page has no text, so result should be None
    assert result is None


@pytest.mark.asyncio
async def test_extract_text_mineru_unavailable_fallback_docx():
    """MinerU 不可用时，docx 降级到 python-docx"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = False

    # Create a minimal valid docx
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Hello from docx fallback")
    doc.add_paragraph("Second paragraph")
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_content = docx_buffer.getvalue()

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.docx", docx_content, "test.docx"
        )

    assert result is not None
    assert "Hello from docx fallback" in result
    assert "Second paragraph" in result


@pytest.mark.asyncio
async def test_extract_text_mineru_exception_fallback():
    """MinerU 抛异常时，降级到 fallback"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.side_effect = RuntimeError("MinerU crashed")

    # Create a minimal valid docx
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Fallback content after crash")
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_content = docx_buffer.getvalue()

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.docx", docx_content, "test.docx"
        )

    assert result is not None
    assert "Fallback content after crash" in result


@pytest.mark.asyncio
async def test_extract_text_mineru_empty_text_fallback():
    """MinerU 返回空文本时，降级到 fallback"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = True
    mock_mineru.recognize_for_ocr.return_value = {
        "text": "",  # empty
        "engine": "mineru",
        "regions": [],
    }

    # Create a minimal valid docx
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Content from fallback after empty MinerU")
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_content = docx_buffer.getvalue()

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.docx", docx_content, "test.docx"
        )

    assert result is not None
    assert "Content from fallback after empty MinerU" in result


@pytest.mark.asyncio
async def test_extract_text_truncates_at_50000():
    """content_text 截断到 50000 字符"""
    long_text = "A" * 60000
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = True
    mock_mineru.recognize_for_ocr.return_value = {
        "text": long_text,
        "engine": "mineru",
        "regions": [],
    }

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.pdf", b"fake", "test.pdf"
        )

    assert result is not None
    assert len(result) == 50000


@pytest.mark.asyncio
async def test_extract_text_docx_with_mineru():
    """docx 文件也走 MinerU OCR"""
    mock_mineru = AsyncMock()
    mock_mineru.is_available.return_value = True
    mock_mineru.recognize_for_ocr.return_value = {
        "text": "MinerU docx full text",
        "engine": "mineru",
        "regions": [],
    }

    with patch(
        "app.services.mineru_service.MinerUService",
        return_value=mock_mineru,
    ):
        result = await _extract_text_with_ocr(
            "/tmp/test.docx", b"fake docx", "test.docx"
        )

    assert result == "MinerU docx full text"
    mock_mineru.recognize_for_ocr.assert_awaited_once_with("/tmp/test.docx")

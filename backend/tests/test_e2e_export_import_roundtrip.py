"""E2E Round-Trip Integration Tests

覆盖全部底稿类型: 表格/文字/程序表/审定表
使用真实 openpyxl/python-docx 读写 + SHA-256 验证

Requirements: 10.1, 10.2, 10.3, 10.4
"""

from __future__ import annotations

import hashlib
import tempfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

from app.services.wp_export.export_engine import WpExportEngine
from app.services.wp_export.metadata_codec import MetadataCodec
from app.services.wp_export.serialization import (
    compute_snapshot_hash,
    deserialize_cell_value,
    serialize_cell_value,
)
from app.schemas.wp_export_schemas import MetadataBundle


# ─── Fixtures ─────────────────────────────────────────────────────────────────


@pytest.fixture
def metadata_codec() -> MetadataCodec:
    return MetadataCodec()


@pytest.fixture
def sample_metadata() -> MetadataBundle:
    return MetadataBundle(
        wp_code="D1",
        project_id=uuid4(),
        file_version=3,
        export_timestamp=datetime(2026, 6, 10, 8, 30, 0, tzinfo=timezone.utc),
        preparer="张三",
        reviewer="李四",
        review_status="approved",
    )


@pytest.fixture
def sample_table_data() -> list[list]:
    """Sample table workpaper data (mixed types)."""
    return [
        ["A001", "银行存款", 1500000.1234, "2026-01-15", "已核实"],
        ["A002", "应收账款", 2300000.5678, "2026-02-20", "待确认"],
        ["A003", "存货", 890000.0, "2026-03-10", "已核实"],
    ]


@pytest.fixture
def sample_procedures() -> list[dict]:
    """Sample program sheet procedures."""
    return [
        {
            "procedure_code": "D1A-001",
            "description": "获取银行对账单并核对余额",
            "execution_status": "已完成",
            "execution_conclusion": "余额一致，无差异",
            "executor": "王五",
        },
        {
            "procedure_code": "D1A-002",
            "description": "检查期末大额存款的真实性",
            "execution_status": "进行中",
            "execution_conclusion": "",
            "executor": "赵六",
        },
        {
            "procedure_code": "D1A-003",
            "description": "抽查银行流水异常交易",
            "execution_status": "未开始",
            "execution_conclusion": "",
            "executor": "",
        },
    ]


@pytest.fixture
def sample_accounts() -> list[dict]:
    """Sample audit sheet accounts."""
    return [
        {
            "account_code": "1001",
            "account_name": "库存现金",
            "unadjusted_amount": 50000.0,
            "adjustment_amount": -2000.0,
            "audited_amount": 48000.0,
        },
        {
            "account_code": "1002",
            "account_name": "银行存款",
            "unadjusted_amount": 1500000.0,
            "adjustment_amount": 0.0,
            "audited_amount": 1500000.0,
        },
        {
            "account_code": "1122",
            "account_name": "应收账款",
            "unadjusted_amount": 2300000.0,
            "adjustment_amount": -150000.0,
            "audited_amount": 2150000.0,
        },
    ]


# ─── Test 1: Table Workpaper Round-Trip ───────────────────────────────────────


class TestTableWorkpaperRoundTrip:
    """表格底稿 Round-Trip: xlsx 创建 → 序列化 → 哈希 → 反序列化 → 验证一致

    Requirements: 10.1, 10.2, 10.3
    """

    def test_xlsx_data_roundtrip(self, sample_table_data: list[list]) -> None:
        """Create xlsx with sample data, save to tempfile, reload, verify identical."""
        # ─── Step 1: Create xlsx with serialized data ────────────────────
        wb = Workbook()
        ws = wb.active
        ws.title = "数据表"

        headers = ["科目编码", "科目名称", "金额", "日期", "状态"]
        col_types = ["text", "text", "number", "date", "text"]

        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=1, column=col_idx, value=header)

        for row_idx, row_data in enumerate(sample_table_data, start=2):
            for col_idx, (value, col_type) in enumerate(
                zip(row_data, col_types), start=1
            ):
                serialized = serialize_cell_value(value, col_type)
                ws.cell(row=row_idx, column=col_idx, value=serialized)

        # ─── Step 2: Save to tempfile ────────────────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = Path(f.name)
            wb.save(f.name)
        wb.close()

        # ─── Step 3: Reload and deserialize ──────────────────────────────
        try:
            wb2 = load_workbook(str(tmp_path))
            ws2 = wb2.active

            for row_idx, row_data in enumerate(sample_table_data, start=2):
                for col_idx, (original, col_type) in enumerate(
                    zip(row_data, col_types), start=1
                ):
                    cell_value = ws2.cell(row=row_idx, column=col_idx).value
                    deserialized = deserialize_cell_value(cell_value, col_type)
                    expected = serialize_cell_value(original, col_type)

                    assert deserialized == expected, (
                        f"Row {row_idx} Col {col_idx}: "
                        f"expected {expected!r}, got {deserialized!r}"
                    )
            wb2.close()
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_xlsx_snapshot_hash_after_roundtrip(
        self, sample_table_data: list[list]
    ) -> None:
        """Compute hash before save, reload and recompute, verify identical."""
        col_types = ["text", "text", "number", "date", "text"]

        # Build workbook_data dict for hash
        serialized_rows = []
        for row_data in sample_table_data:
            serialized_row = [
                serialize_cell_value(v, t) for v, t in zip(row_data, col_types)
            ]
            serialized_rows.append(serialized_row)

        workbook_data = {"数据表": serialized_rows}
        hash_before = compute_snapshot_hash(workbook_data)

        # Save xlsx and reload
        wb = Workbook()
        ws = wb.active
        ws.title = "数据表"
        for row_idx, row in enumerate(serialized_rows, start=1):
            for col_idx, val in enumerate(row, start=1):
                ws.cell(row=row_idx, column=col_idx, value=val)

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = Path(f.name)
            wb.save(f.name)
        wb.close()

        try:
            wb2 = load_workbook(str(tmp_path))
            ws2 = wb2.active

            reloaded_rows = []
            for row_idx in range(1, len(serialized_rows) + 1):
                row = []
                for col_idx, col_type in enumerate(col_types, start=1):
                    cell_val = ws2.cell(row=row_idx, column=col_idx).value
                    row.append(deserialize_cell_value(cell_val, col_type))
                reloaded_rows.append(row)

            workbook_data_after = {"数据表": reloaded_rows}
            hash_after = compute_snapshot_hash(workbook_data_after)
            wb2.close()

            assert hash_before == hash_after, (
                f"Hash mismatch: before={hash_before}, after={hash_after}"
            )
        finally:
            tmp_path.unlink(missing_ok=True)


# ─── Test 2: Text Workpaper Round-Trip (docx) ─────────────────────────────────


class TestTextWorkpaperRoundTrip:
    """文字底稿 Round-Trip: docx 创建 → 嵌入元数据 → 提取元数据 → 验证一致

    Requirements: 10.1, 10.3
    """

    def test_docx_metadata_roundtrip(
        self,
        metadata_codec: MetadataCodec,
        sample_metadata: MetadataBundle,
    ) -> None:
        """Create docx, embed metadata, save, reload, extract, verify identical."""
        # ─── Step 1: Create docx and embed metadata ──────────────────────
        doc = Document()
        doc.add_heading("D1 银行存款", level=1)
        doc.add_paragraph("本底稿用于记录银行存款审计程序执行情况。")
        doc.add_paragraph("检查期末银行对账单，核对余额一致性。")

        metadata_codec.embed_docx(doc, sample_metadata)

        # ─── Step 2: Save to tempfile ────────────────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = Path(f.name)
            doc.save(f.name)

        # ─── Step 3: Reload and extract metadata ─────────────────────────
        try:
            doc2 = Document(str(tmp_path))
            extracted = metadata_codec.extract_docx(doc2)

            assert extracted is not None, "Metadata extraction returned None"
            assert extracted.wp_code == sample_metadata.wp_code
            assert extracted.project_id == sample_metadata.project_id
            assert extracted.file_version == sample_metadata.file_version
            assert extracted.export_timestamp == sample_metadata.export_timestamp
            assert extracted.preparer == sample_metadata.preparer
            assert extracted.reviewer == sample_metadata.reviewer
            assert extracted.review_status == sample_metadata.review_status
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_docx_content_preserved_after_metadata_embed(
        self,
        metadata_codec: MetadataCodec,
        sample_metadata: MetadataBundle,
    ) -> None:
        """Embedding metadata does not corrupt document content."""
        paragraphs_text = [
            "本底稿用于记录银行存款审计程序执行情况。",
            "检查期末银行对账单，核对余额一致性。",
            "已获取全部 12 个月的银行对账单。",
        ]

        doc = Document()
        doc.add_heading("测试底稿", level=1)
        for text in paragraphs_text:
            doc.add_paragraph(text)

        metadata_codec.embed_docx(doc, sample_metadata)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = Path(f.name)
            doc.save(f.name)

        try:
            doc2 = Document(str(tmp_path))
            # First paragraph is the heading
            actual_paragraphs = [
                p.text for p in doc2.paragraphs if p.text and p.text != "测试底稿"
            ]
            assert actual_paragraphs == paragraphs_text
        finally:
            tmp_path.unlink(missing_ok=True)


# ─── Test 3: Program Sheet Round-Trip ─────────────────────────────────────────


class TestProgramSheetRoundTrip:
    """程序表 Round-Trip: 生成 → 保存 → 加载 → 验证 5 列和数据完整性

    Requirements: 10.1, 10.3
    """

    def test_program_sheet_columns_and_data(
        self, sample_procedures: list[dict]
    ) -> None:
        """Build program sheet bytes, save to file, reload, verify structure."""
        # ─── Step 1: Build program sheet ─────────────────────────────────
        buf = WpExportEngine.build_program_sheet_bytes(sample_procedures)

        # ─── Step 2: Save to tempfile and reload ─────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = Path(f.name)
            f.write(buf.getvalue())

        try:
            wb = load_workbook(str(tmp_path))
            ws = wb.active

            # ─── Step 3: Verify 5 columns in header ──────────────────────
            assert ws.title == "程序表"
            expected_headers = [
                "[只读] 程序编号",
                "[只读] 程序描述",
                "[可编辑] 执行状态",
                "[可编辑] 执行结论",
                "[可编辑] 执行人",
            ]
            actual_headers = [ws.cell(row=1, column=i).value for i in range(1, 6)]
            assert actual_headers == expected_headers

            # ─── Step 4: Verify data rows ────────────────────────────────
            fields = [
                "procedure_code",
                "description",
                "execution_status",
                "execution_conclusion",
                "executor",
            ]
            for row_idx, proc in enumerate(sample_procedures, start=2):
                for col_idx, field in enumerate(fields, start=1):
                    actual = ws.cell(row=row_idx, column=col_idx).value
                    expected = proc.get(field, "")
                    # openpyxl normalizes empty strings to None on save/load
                    actual_norm = actual if actual is not None else ""
                    assert actual_norm == expected, (
                        f"Row {row_idx} Col {col_idx} ({field}): "
                        f"expected {expected!r}, got {actual_norm!r}"
                    )

            # ─── Step 5: Verify row count ────────────────────────────────
            data_row_count = ws.max_row - 1  # minus header
            assert data_row_count == len(sample_procedures)

            wb.close()
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_program_sheet_protection_markers(
        self, sample_procedures: list[dict]
    ) -> None:
        """Verify readonly columns have protection set."""
        buf = WpExportEngine.build_program_sheet_bytes(sample_procedures)

        wb = load_workbook(BytesIO(buf.getvalue()))
        ws = wb.active

        # Sheet should be protected
        assert ws.protection.sheet is True

        # Col 1 & 2 (readonly) should be locked
        for row_idx in range(2, len(sample_procedures) + 2):
            assert ws.cell(row=row_idx, column=1).protection.locked is True
            assert ws.cell(row=row_idx, column=2).protection.locked is True
            # Col 3, 4, 5 (editable) should be unlocked
            assert ws.cell(row=row_idx, column=3).protection.locked is False
            assert ws.cell(row=row_idx, column=4).protection.locked is False
            assert ws.cell(row=row_idx, column=5).protection.locked is False

        wb.close()


# ─── Test 4: Audit Sheet Round-Trip ───────────────────────────────────────────


class TestAuditSheetRoundTrip:
    """审定表 Round-Trip: 生成 → 保存 → 加载 → 验证 5 列 + 汇总行 + 数据完整性

    Requirements: 10.1, 10.3
    """

    def test_audit_sheet_columns_and_summary(
        self, sample_accounts: list[dict]
    ) -> None:
        """Build audit sheet bytes, save, reload, verify 5 cols + summary row."""
        # ─── Step 1: Build audit sheet ───────────────────────────────────
        buf = WpExportEngine.build_audit_sheet_bytes(sample_accounts)

        # ─── Step 2: Save to tempfile and reload ─────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = Path(f.name)
            f.write(buf.getvalue())

        try:
            wb = load_workbook(str(tmp_path))
            ws = wb.active

            # ─── Step 3: Verify sheet title ──────────────────────────────
            assert ws.title == "审定表"

            # ─── Step 4: Verify 5 columns in header ──────────────────────
            expected_headers = ["科目编码", "科目名称", "未审数", "调整数", "审定数"]
            actual_headers = [ws.cell(row=1, column=i).value for i in range(1, 6)]
            assert actual_headers == expected_headers

            # ─── Step 5: Verify data rows ────────────────────────────────
            fields = [
                "account_code",
                "account_name",
                "unadjusted_amount",
                "adjustment_amount",
                "audited_amount",
            ]
            for row_idx, account in enumerate(sample_accounts, start=2):
                for col_idx, field in enumerate(fields, start=1):
                    actual = ws.cell(row=row_idx, column=col_idx).value
                    expected = account.get(field, "")
                    assert actual == expected, (
                        f"Row {row_idx} Col {col_idx} ({field}): "
                        f"expected {expected!r}, got {actual!r}"
                    )

            # ─── Step 6: Verify summary row ──────────────────────────────
            summary_row = len(sample_accounts) + 2
            assert ws.cell(row=summary_row, column=1).value == "合计"

            expected_total_unadj = sum(
                a["unadjusted_amount"] for a in sample_accounts
            )
            expected_total_adj = sum(
                a["adjustment_amount"] for a in sample_accounts
            )
            expected_total_aud = sum(
                a["audited_amount"] for a in sample_accounts
            )

            assert ws.cell(row=summary_row, column=3).value == expected_total_unadj
            assert ws.cell(row=summary_row, column=4).value == expected_total_adj
            assert ws.cell(row=summary_row, column=5).value == expected_total_aud

            wb.close()
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_audit_sheet_data_integrity_after_save_reload(
        self, sample_accounts: list[dict]
    ) -> None:
        """Verify data is byte-identical after save → reload cycle."""
        buf1 = WpExportEngine.build_audit_sheet_bytes(sample_accounts)
        content1 = buf1.getvalue()

        # Reload the bytes back
        wb = load_workbook(BytesIO(content1))
        ws = wb.active

        # All numeric values should be preserved exactly
        for row_idx, account in enumerate(sample_accounts, start=2):
            unadj = ws.cell(row=row_idx, column=3).value
            adj = ws.cell(row=row_idx, column=4).value
            aud = ws.cell(row=row_idx, column=5).value

            assert unadj == account["unadjusted_amount"]
            assert adj == account["adjustment_amount"]
            assert aud == account["audited_amount"]

        wb.close()


# ─── Test 5: SHA-256 Consistency ──────────────────────────────────────────────


class TestSHA256Consistency:
    """SHA-256 一致性: 同内容两次导出产生相同哈希

    Requirements: 10.2
    """

    def test_same_content_produces_identical_hash(self) -> None:
        """Export same content twice, both produce identical SHA-256 hash."""
        workbook_data = {
            "Sheet1": [
                [1234.5678, "测试", "2026-01-01"],
                [9999.0, "数据", "2026-12-31"],
            ],
            "Sheet2": [
                ["文字内容", None, 0.0],
            ],
        }

        hash1 = compute_snapshot_hash(workbook_data)
        hash2 = compute_snapshot_hash(workbook_data)

        assert hash1 == hash2
        assert len(hash1) == 64  # SHA-256 produces 64 hex chars

    def test_different_content_produces_different_hash(self) -> None:
        """Different content must produce different SHA-256 hash."""
        data_a = {"Sheet1": [[1.0, "hello"]]}
        data_b = {"Sheet1": [[1.0, "world"]]}

        hash_a = compute_snapshot_hash(data_a)
        hash_b = compute_snapshot_hash(data_b)

        assert hash_a != hash_b

    def test_sheet_order_does_not_affect_hash(self) -> None:
        """Hash is deterministic regardless of dict insertion order."""
        data_forward = {
            "Alpha": [[1, 2]],
            "Beta": [[3, 4]],
        }
        data_reverse = {
            "Beta": [[3, 4]],
            "Alpha": [[1, 2]],
        }

        assert compute_snapshot_hash(data_forward) == compute_snapshot_hash(
            data_reverse
        )

    def test_program_sheet_hash_consistency(
        self, sample_procedures: list[dict]
    ) -> None:
        """Same program data exported twice produces identical file SHA-256."""
        buf1 = WpExportEngine.build_program_sheet_bytes(sample_procedures)
        buf2 = WpExportEngine.build_program_sheet_bytes(sample_procedures)

        hash1 = hashlib.sha256(buf1.getvalue()).hexdigest()
        hash2 = hashlib.sha256(buf2.getvalue()).hexdigest()

        assert hash1 == hash2

    def test_audit_sheet_hash_consistency(
        self, sample_accounts: list[dict]
    ) -> None:
        """Same audit data exported twice produces identical file SHA-256."""
        buf1 = WpExportEngine.build_audit_sheet_bytes(sample_accounts)
        buf2 = WpExportEngine.build_audit_sheet_bytes(sample_accounts)

        hash1 = hashlib.sha256(buf1.getvalue()).hexdigest()
        hash2 = hashlib.sha256(buf2.getvalue()).hexdigest()

        assert hash1 == hash2


# ─── Test 6: Metadata Round-Trip Across Save/Load ─────────────────────────────


class TestMetadataRoundTripSaveLoad:
    """元数据 Round-Trip: 嵌入 → 保存文件 → 加载文件 → 提取 → 验证一致

    Requirements: 10.1, 10.4
    """

    def test_xlsx_metadata_survives_file_save_load(
        self,
        metadata_codec: MetadataCodec,
        sample_metadata: MetadataBundle,
    ) -> None:
        """Embed metadata in xlsx, save to disk, reload from disk, extract, verify."""
        # ─── Step 1: Create workbook and embed metadata ──────────────────
        wb = Workbook()
        ws = wb.active
        ws.title = "测试"
        ws.cell(row=1, column=1, value="测试数据")

        metadata_codec.embed_xlsx(wb, sample_metadata)

        # ─── Step 2: Save to actual file ─────────────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            tmp_path = Path(f.name)
            wb.save(f.name)
        wb.close()

        # ─── Step 3: Reload from disk and extract ────────────────────────
        try:
            wb2 = load_workbook(str(tmp_path))
            extracted = metadata_codec.extract_xlsx(wb2)
            wb2.close()

            assert extracted is not None, "Metadata extraction returned None"
            assert extracted.wp_code == sample_metadata.wp_code
            assert extracted.project_id == sample_metadata.project_id
            assert extracted.file_version == sample_metadata.file_version
            assert extracted.export_timestamp == sample_metadata.export_timestamp
            assert extracted.preparer == sample_metadata.preparer
            assert extracted.reviewer == sample_metadata.reviewer
            assert extracted.review_status == sample_metadata.review_status
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_docx_metadata_survives_file_save_load(
        self,
        metadata_codec: MetadataCodec,
        sample_metadata: MetadataBundle,
    ) -> None:
        """Embed metadata in docx, save to disk, reload from disk, extract, verify."""
        # ─── Step 1: Create document and embed metadata ──────────────────
        doc = Document()
        doc.add_heading("审计底稿", level=1)
        doc.add_paragraph("第一段内容")

        metadata_codec.embed_docx(doc, sample_metadata)

        # ─── Step 2: Save to actual file ─────────────────────────────────
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = Path(f.name)
            doc.save(f.name)

        # ─── Step 3: Reload from disk and extract ────────────────────────
        try:
            doc2 = Document(str(tmp_path))
            extracted = metadata_codec.extract_docx(doc2)

            assert extracted is not None, "Metadata extraction returned None"
            assert extracted.wp_code == sample_metadata.wp_code
            assert extracted.project_id == sample_metadata.project_id
            assert extracted.file_version == sample_metadata.file_version
            assert extracted.export_timestamp == sample_metadata.export_timestamp
            assert extracted.preparer == sample_metadata.preparer
            assert extracted.reviewer == sample_metadata.reviewer
            assert extracted.review_status == sample_metadata.review_status
        finally:
            tmp_path.unlink(missing_ok=True)

    def test_xlsx_metadata_idempotent_embed(
        self,
        metadata_codec: MetadataCodec,
        sample_metadata: MetadataBundle,
    ) -> None:
        """Embedding metadata twice produces same result (idempotent)."""
        wb = Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="data")

        # Embed twice
        metadata_codec.embed_xlsx(wb, sample_metadata)
        metadata_codec.embed_xlsx(wb, sample_metadata)

        # Save and reload
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()

        wb2 = load_workbook(buf)
        extracted = metadata_codec.extract_xlsx(wb2)
        wb2.close()

        assert extracted is not None
        assert extracted.wp_code == sample_metadata.wp_code
        assert extracted.file_version == sample_metadata.file_version

    def test_metadata_with_optional_fields_none(
        self,
        metadata_codec: MetadataCodec,
    ) -> None:
        """Metadata with None optional fields round-trips correctly."""
        metadata = MetadataBundle(
            wp_code="E2",
            project_id=uuid4(),
            file_version=1,
            export_timestamp=datetime(2026, 3, 15, 10, 0, 0, tzinfo=timezone.utc),
            preparer=None,
            reviewer=None,
            review_status=None,
        )

        # Test xlsx path
        wb = Workbook()
        metadata_codec.embed_xlsx(wb, metadata)
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()

        wb2 = load_workbook(buf)
        extracted = metadata_codec.extract_xlsx(wb2)
        wb2.close()

        assert extracted is not None
        assert extracted.wp_code == "E2"
        assert extracted.file_version == 1
        assert extracted.preparer is None
        assert extracted.reviewer is None
        assert extracted.review_status is None

        # Test docx path
        doc = Document()
        metadata_codec.embed_docx(doc, metadata)
        buf2 = BytesIO()
        doc.save(buf2)
        buf2.seek(0)

        doc2 = Document(buf2)
        extracted2 = metadata_codec.extract_docx(doc2)

        assert extracted2 is not None
        assert extracted2.wp_code == "E2"
        assert extracted2.file_version == 1
        assert extracted2.preparer is None
        assert extracted2.reviewer is None
        assert extracted2.review_status is None

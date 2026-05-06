"""EQCR 备忘录 Word 生成单元测试

测试纯函数 ``build_memo_docx_bytes``（不依赖 DB、不依赖 LibreOffice）。
"""

from __future__ import annotations

import io

import pytest

from app.services.eqcr_memo_service import (
    MEMO_SECTIONS,
    build_memo_docx_bytes,
    eqcr_memo_pdf_generator,
)


def test_build_memo_docx_returns_valid_bytes():
    """build_memo_docx_bytes 返回非空的 docx 字节流。"""
    sections = {
        "项目概况": "项目名称：Test Project\n客户：Test Client",
        "重要性判断": "EQCR 认可",
        "EQCR 总评与结论": "总体结论：通过",
    }
    data = build_memo_docx_bytes(
        project_name="Test Project",
        client_name="Test Client",
        sections=sections,
    )
    assert isinstance(data, bytes)
    assert len(data) > 0
    # Word docx 本质是 zip，前四字节是 'PK'
    assert data[:2] == b"PK"


def test_build_memo_docx_parseable_by_python_docx():
    """生成的 docx 字节可以被 python-docx 反向解析。"""
    import docx

    sections = {
        name: f"{name} 的内容" for name in MEMO_SECTIONS
    }
    data = build_memo_docx_bytes(
        project_name="A",
        client_name="B",
        sections=sections,
    )
    doc = docx.Document(io.BytesIO(data))

    # 应该包含大标题 "独立复核备忘录"
    headings = [p.text for p in doc.paragraphs]
    assert any("独立复核备忘录" in t for t in headings)

    # 应该包含所有 10 个章节标题
    for section_name in MEMO_SECTIONS:
        assert any(section_name in t for t in headings), (
            f"Section '{section_name}' not found in document"
        )


def test_build_memo_docx_handles_missing_sections():
    """章节字典缺字段时回填"（未填写）"。"""
    import docx

    sections = {"项目概况": "仅填一个"}
    data = build_memo_docx_bytes(
        project_name="X", client_name="Y", sections=sections,
    )
    doc = docx.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "（未填写）" in text


def test_eqcr_memo_pdf_generator_no_files_returns_none():
    """wizard_state 无 memo 文件时返回 None。"""
    import uuid

    result = eqcr_memo_pdf_generator(uuid.uuid4(), {})
    assert result is None

    result2 = eqcr_memo_pdf_generator(
        uuid.uuid4(),
        {"eqcr_memo": {"files": {"pdf": None}}},
    )
    assert result2 is None


def test_eqcr_memo_pdf_generator_nonexistent_path_returns_none():
    """PDF 文件路径不存在时返回 None，不抛异常。"""
    import uuid

    result = eqcr_memo_pdf_generator(
        uuid.uuid4(),
        {"eqcr_memo": {"files": {"pdf": "/nonexistent/path/foo.pdf"}}},
    )
    assert result is None

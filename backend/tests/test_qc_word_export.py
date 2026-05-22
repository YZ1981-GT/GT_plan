"""Tests for qc_report_export router — Phase 7 F5

Validates: Requirements F5.2, F5.6, F5.7
"""

import io
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.routers.qc_report_export import (
    _build_qc_report,
    _check_qc_admin,
    router,
)


class TestQCReportExportRouter:
    """Test QC report Word export endpoint."""

    def test_router_prefix(self):
        """Router has correct prefix."""
        assert "/qc-report" in router.prefix

    def test_router_tags(self):
        """Router has correct tags."""
        assert "qc-report" in router.tags

    def test_permission_check_admin(self):
        """Admin user passes permission check."""
        user = MagicMock()
        user.role.value = "admin"
        # Should not raise
        _check_qc_admin(user)

    def test_permission_check_qc(self):
        """QC user passes permission check."""
        user = MagicMock()
        user.role.value = "qc"
        _check_qc_admin(user)

    def test_permission_check_auditor_rejected(self):
        """Non-QC/admin user is rejected."""
        from fastapi import HTTPException

        user = MagicMock()
        user.role.value = "auditor"
        with pytest.raises(HTTPException) as exc_info:
            _check_qc_admin(user)
        assert exc_info.value.status_code == 403

    def test_permission_check_manager_rejected(self):
        """Manager user is rejected."""
        from fastapi import HTTPException

        user = MagicMock()
        user.role.value = "manager"
        with pytest.raises(HTTPException) as exc_info:
            _check_qc_admin(user)
        assert exc_info.value.status_code == 403

    def test_build_empty_report(self):
        """Empty data produces valid docx with headers but no data rows."""
        buffer = _build_qc_report("测试项目", [], [], [])
        assert isinstance(buffer, io.BytesIO)
        # Verify it's a valid docx (starts with PK zip header)
        content = buffer.read()
        assert content[:2] == b"PK"
        assert len(content) > 100

    def test_build_report_with_data(self):
        """Report with data produces valid docx."""
        risk_data = [["D", "5", "3", "2", "10"]]
        opinion_data = [["1", "blocker", "D", "测试问题", "open", "2026-01-01"]]
        rect_data = [["1", "测试问题", "closed", "张三", "2026-01-15"]]

        buffer = _build_qc_report("项目A", risk_data, opinion_data, rect_data)
        content = buffer.read()
        assert content[:2] == b"PK"
        assert len(content) > 200

    def test_build_report_three_chapters(self):
        """Report contains three chapters (verify via python-docx parsing)."""
        from docx import Document

        buffer = _build_qc_report("测试", [], [], [])
        buffer.seek(0)
        doc = Document(buffer)

        # Find headings
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("风险汇总" in h for h in headings)
        assert any("意见清单" in h for h in headings)
        assert any("整改状态" in h for h in headings)

    def test_build_report_has_tables(self):
        """Report contains 3 tables (one per chapter)."""
        from docx import Document

        buffer = _build_qc_report("测试", [["D", "3", "2", "1", "6"]], [], [])
        buffer.seek(0)
        doc = Document(buffer)

        assert len(doc.tables) == 3

    def test_router_has_export_endpoint(self):
        """Router has GET /export endpoint."""
        paths = [r.path for r in router.routes]
        assert any("/export" in p for p in paths)

"""Sprint 6 Tests: Word Export + Export Package + Filename Sanitization

Tests cover:
  - Word page setup (致同 margins)
  - TOC presence
  - Table formatting (borders, header bold, amount right-align)
  - Empty sections generate "本期无此项业务"
  - ZIP contents (xlsx + docx + manifest.json)
  - Manifest structure
  - Filename sanitization
  - Export package consistency gate integration
  - HTML preview
  - Heading numbering
"""
from __future__ import annotations

import json
import zipfile
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest
import pytest_asyncio
from sqlalchemy import Column, Integer, String, create_engine, text
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import Session, sessionmaker

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


class FakeDisclosureNote:
    """Fake DisclosureNote for testing without DB."""

    def __init__(self, note_section: str, section_title: str, text_content: str = "",
                 table_data: dict | None = None, year: int = 2025, project_id=None):
        self.id = uuid4()
        self.note_section = note_section
        self.section_title = section_title
        self.text_content = text_content
        self.table_data = table_data
        self.year = year
        self.project_id = project_id or uuid4()


class FakeProject:
    """Fake Project for testing."""

    def __init__(self, name="测试公司", client_name="测试客户有限公司",
                 template_type="soe", report_scope="standalone", wizard_state=None):
        self.id = uuid4()
        self.name = name
        self.client_name = client_name
        self.template_type = template_type
        self.report_scope = report_scope
        self.wizard_state = wizard_state or {}


def _make_table_data(headers, rows_data):
    """Helper to create table_data dict."""
    rows = []
    for label, values in rows_data:
        rows.append({"label": label, "values": values})
    return {"headers": headers, "rows": rows}


# ---------------------------------------------------------------------------
# Test: Word page setup (致同 margins)
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_word_page_margins():
    """Verify 致同 standard page margins: top=3.2cm, bottom=2.54cm, left=3cm, right=3.18cm."""
    from docx.shared import Cm

    from app.services.note_word_exporter import NoteWordExporter

    # Mock DB session
    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = []
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    # Parse the generated document
    from docx import Document
    doc = Document(output)
    section = doc.sections[0]

    # Allow small EMU rounding tolerance (±200 EMU ≈ 0.0002cm)
    assert abs(section.top_margin - Cm(3.2)) < 200
    assert abs(section.bottom_margin - Cm(2.54)) < 200
    assert abs(section.left_margin - Cm(3)) < 200
    assert abs(section.right_margin - Cm(3.18)) < 200


@pytest.mark.asyncio
async def test_word_header_footer_margins():
    """Verify header=1.3cm, footer=1.3cm."""
    from docx.shared import Cm

    from app.services.note_word_exporter import NoteWordExporter

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = []
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    from docx import Document
    doc = Document(output)
    section = doc.sections[0]

    # Allow small EMU rounding tolerance
    assert abs(section.header_distance - Cm(1.3)) < 200
    assert abs(section.footer_distance - Cm(1.3)) < 200


# ---------------------------------------------------------------------------
# Test: TOC presence
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_word_toc_present():
    """Verify TOC field code is present in the document."""
    from app.services.note_word_exporter import NoteWordExporter

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = []
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(output)

    # Search for TOC field code in document XML
    body = doc.element.body
    xml_str = body.xml
    assert "TOC" in xml_str, "TOC field code should be present in document"


# ---------------------------------------------------------------------------
# Test: Table formatting
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_word_table_header_bold():
    """Verify table header row cells are bold."""
    from app.services.note_word_exporter import NoteWordExporter

    notes = [
        FakeDisclosureNote(
            note_section="1",
            section_title="货币资金",
            table_data=_make_table_data(
                ["项目", "期末余额", "期初余额"],
                [("库存现金", [1000.00, 800.00]), ("银行存款", [50000.00, 45000.00])],
            ),
        )
    ]

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = notes
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    from docx import Document
    doc = Document(output)

    # Find the first table
    tables = doc.tables
    assert len(tables) >= 1, "Should have at least one table"

    # Check header row has bold runs (check XML for bold attribute)
    header_row = tables[0].rows[0]
    found_bold = False
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    found_bold = True
                    break
    assert found_bold, "Header row should contain bold runs"


# ---------------------------------------------------------------------------
# Test: Empty sections
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_word_empty_section_placeholder():
    """Verify empty sections generate '本期无此项业务'."""
    from app.services.note_word_exporter import NoteWordExporter

    notes = [
        FakeDisclosureNote(
            note_section="1",
            section_title="应收票据",
            text_content="",
            table_data={"headers": ["项目", "金额"], "rows": [{"label": "合计", "values": [0]}]},
        )
    ]

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = notes
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    from docx import Document
    doc = Document(output)

    # Search for placeholder text
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "本期无此项业务" in full_text


# ---------------------------------------------------------------------------
# Test: skip_empty parameter
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_word_skip_empty_sections():
    """Verify skip_empty=True removes empty sections entirely."""
    from app.services.note_word_exporter import NoteWordExporter

    notes = [
        FakeDisclosureNote(
            note_section="1",
            section_title="货币资金",
            text_content="有内容的章节",
        ),
        FakeDisclosureNote(
            note_section="2",
            section_title="应收票据",
            text_content="",
            table_data={"headers": ["项目"], "rows": []},
        ),
    ]

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = notes
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025, skip_empty=True)

    from docx import Document
    doc = Document(output)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "货币资金" in full_text
    assert "应收票据" not in full_text


# ---------------------------------------------------------------------------
# Test: ZIP contents
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_zip_contains_required_files():
    """Verify ZIP contains xlsx + docx + manifest.json."""
    from app.services.export_package_service import ExportPackageService

    mock_db = AsyncMock(spec=AsyncSession)

    # Mock project query
    project = FakeProject()
    mock_project_result = MagicMock()
    mock_project_result.scalar_one_or_none.return_value = project

    # Mock notes query (empty) - will be called by NoteWordExporter
    mock_notes_result = MagicMock()
    mock_notes_result.scalars.return_value.all.return_value = []

    # execute will be called multiple times: project lookup + notes lookup
    mock_db.execute = AsyncMock(side_effect=[
        mock_project_result,  # Project query in ExportPackageService
        mock_notes_result,    # Notes query in NoteWordExporter
    ])

    # Patch the excel exporter's export method
    with patch("app.services.report_excel_exporter.ReportExcelExporter.export",
               new_callable=AsyncMock, return_value=BytesIO(b"fake xlsx content")):
        service = ExportPackageService(mock_db)
        output = await service.export_package(project_id=project.id, year=2025)

    # Verify ZIP contents
    with zipfile.ZipFile(output) as zf:
        names = zf.namelist()
        assert any(n.endswith(".xlsx") for n in names), "ZIP should contain .xlsx file"
        assert any(n.endswith(".docx") for n in names), "ZIP should contain .docx file"
        assert "manifest.json" in names, "ZIP should contain manifest.json"


# ---------------------------------------------------------------------------
# Test: Manifest structure
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_zip_manifest_structure():
    """Verify manifest.json has correct structure."""
    from app.services.export_package_service import ExportPackageService

    mock_db = AsyncMock(spec=AsyncSession)

    project = FakeProject(name="测试项目", client_name="测试客户")
    mock_project_result = MagicMock()
    mock_project_result.scalar_one_or_none.return_value = project

    mock_notes_result = MagicMock()
    mock_notes_result.scalars.return_value.all.return_value = []

    mock_db.execute = AsyncMock(side_effect=[
        mock_project_result,
        mock_notes_result,
    ])

    with patch("app.services.report_excel_exporter.ReportExcelExporter.export",
               new_callable=AsyncMock, return_value=BytesIO(b"fake xlsx")):
        service = ExportPackageService(mock_db)
        output = await service.export_package(project_id=project.id, year=2025)

    with zipfile.ZipFile(output) as zf:
        manifest_bytes = zf.read("manifest.json")
        manifest = json.loads(manifest_bytes)

    assert manifest["version"] == "1.0"
    assert manifest["year"] == 2025
    assert "files" in manifest
    assert "consistency_check" in manifest
    assert manifest["consistency_check"]["overall"] == "pass"
    assert len(manifest["files"]) >= 2  # At least xlsx + docx


# ---------------------------------------------------------------------------
# Test: Filename sanitization
# ---------------------------------------------------------------------------


def test_filename_sanitization():
    """Verify special characters are replaced with underscores."""
    from app.services.export_package_service import sanitize_filename

    assert sanitize_filename("公司/名称") == "公司_名称"
    assert sanitize_filename('文件"名') == "文件_名"
    assert sanitize_filename("a:b*c?d") == "a_b_c_d"
    assert sanitize_filename("正常文件名.xlsx") == "正常文件名.xlsx"
    assert sanitize_filename("a<b>c|d") == "a_b_c_d"


def test_filename_convention_reports():
    """Verify report filename follows convention: {公司简称}_{年度}年度财务报表.xlsx"""
    from app.services.export_package_service import get_company_short_name, sanitize_filename

    project = FakeProject(client_name="四川物流有限公司")
    short = get_company_short_name(project)
    filename = sanitize_filename(f"{short}_{2025}年度财务报表.xlsx")
    assert "2025年度财务报表.xlsx" in filename
    assert "/" not in filename
    assert "\\" not in filename


def test_filename_convention_notes():
    """Verify notes filename follows convention: {公司简称}_{年度}年度财务报表附注.docx"""
    from app.routers.note_export import sanitize_export_filename

    filename = sanitize_export_filename("测试公司_2025年度财务报表附注.docx")
    assert filename == "测试公司_2025年度财务报表附注.docx"

    filename2 = sanitize_export_filename("公司/名:称_2025年度财务报表附注.docx")
    assert "/" not in filename2
    assert ":" not in filename2


def test_company_short_name_from_wizard_state():
    """Verify company short name reads from wizard_state first."""
    from app.services.export_package_service import get_company_short_name

    project = FakeProject(
        client_name="四川物流有限公司",
        wizard_state={"company_short_name": "川物流"},
    )
    assert get_company_short_name(project) == "川物流"


def test_company_short_name_fallback():
    """Verify company short name falls back to client_name."""
    from app.services.export_package_service import get_company_short_name

    project = FakeProject(client_name="短名", wizard_state={})
    assert get_company_short_name(project) == "短名"


# ---------------------------------------------------------------------------
# Test: HTML preview
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_html_preview():
    """Verify HTML preview generates valid HTML."""
    from app.services.note_word_exporter import NoteWordExporter

    notes = [
        FakeDisclosureNote(
            note_section="1",
            section_title="货币资金",
            text_content="现金及银行存款",
        ),
    ]

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = notes
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    html = await exporter.preview_html(project_id=uuid4(), year=2025)

    assert "note-preview" in html
    assert "货币资金" in html
    assert "现金及银行存款" in html
    assert "一、" in html  # Level 1 heading prefix


# ---------------------------------------------------------------------------
# Test: Heading numbering
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_heading_numbering():
    """Verify heading numbering follows 致同 convention."""
    from app.services.note_word_exporter import NoteWordExporter

    notes = [
        FakeDisclosureNote(note_section="1", section_title="公司基本情况"),
        FakeDisclosureNote(note_section="2", section_title="财务报表编制基础"),
        FakeDisclosureNote(note_section="2.1", section_title="编制基础"),
        FakeDisclosureNote(note_section="2.1.1", section_title="具体说明"),
    ]

    mock_db = AsyncMock(spec=AsyncSession)
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = notes
    mock_db.execute = AsyncMock(return_value=mock_result)

    exporter = NoteWordExporter(mock_db)
    output = await exporter.export(project_id=uuid4(), year=2025)

    from docx import Document
    doc = Document(output)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "一、公司基本情况" in full_text
    assert "二、财务报表编制基础" in full_text
    assert "（一）编制基础" in full_text
    assert "1. 具体说明" in full_text


# ---------------------------------------------------------------------------
# Test: Amount formatting
# ---------------------------------------------------------------------------


def test_format_amount():
    """Verify amount formatting with thousands separator."""
    from app.services.note_word_exporter import _format_amount

    assert _format_amount(1234567.89) == "1,234,567.89"
    assert _format_amount(0) == "-"
    assert _format_amount(None) == "-"
    assert _format_amount("") == "-"
    assert _format_amount(100.5) == "100.50"

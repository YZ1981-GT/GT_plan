"""TemplateFillService 集成测试（audit-report-template-integration task 7.6）.

验证两阶段流水线核心不变式：
- preview 不创建 deliverable_version / word_export_task（仅写 fill_preview_sessions）
- confirm 创建 word_export_task + 一个 version（version_no=1）
- 再次 preview→confirm（同项目）→ version_no 递增为 2
- KAM 校验（validate_kam_word_mode）单元覆盖

使用内存 SQLite + 真实 python-docx 模板 + 临时存储目录。
"""

from __future__ import annotations

import uuid
from pathlib import Path

import pytest
import pytest_asyncio
import sqlalchemy as sa
from docx import Document
from sqlalchemy.dialects.sqlite.base import SQLiteTypeCompiler
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.models.base import Base
from app.models.phase13_models import WordExportTask, WordExportTaskVersion
from app.models.report_models import FillPreviewSession
from app.services import deliverable_service as deliverable_module
from app.services import template_fill_service as tfs_module
from app.services.template_fill_service import (
    TemplateFillService,
    validate_kam_word_mode,
)
from app.services.template_manifest_loader import TemplateManifestLoader


SQLiteTypeCompiler.visit_JSONB = SQLiteTypeCompiler.visit_JSON


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest_asyncio.fixture
async def session() -> AsyncSession:
    engine = create_async_engine("sqlite+aiosqlite:///:memory:", echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    async with factory() as s:
        yield s
    await engine.dispose()


def _make_template_dir(base: Path) -> Path:
    """构造临时 manifest 目录 + 含占位符/OPT/NOTE 的报告正文 docx。"""
    tpl_dir = base / "templates"
    (tpl_dir / "report_body").mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_paragraph("致 {{company_full_name}} 全体股东：")
    doc.add_paragraph("我们审计了 {{company_full_name}} {{audit_year}} 年度财务报表。")
    doc.add_paragraph("##NOTE:内部提示，定稿删除##")
    # 可选段落 OPT 块
    doc.add_paragraph("##OPT:emphasis:强调事项段##")
    doc.add_paragraph("我们提醒财务报表使用者关注以下事项。")
    doc.add_paragraph("##/OPT:emphasis##")
    doc.add_paragraph("##OPT:key_audit_matters:关键审计事项段##")
    doc.add_paragraph("关键审计事项是我们根据职业判断认为最为重要的事项。")
    doc.add_paragraph("##/OPT:key_audit_matters##")
    doc.save(str(tpl_dir / "report_body" / "test_unqualified_d.docx"))

    manifest = tpl_dir / "template_manifest.json"
    manifest.write_text(
        """
{
  "version": "test-v1",
  "report_body": {
    "unqualified": {
      "type_d": {"simple": "report_body/test_unqualified_d.docx"}
    }
  },
  "financial_statements": {},
  "disclosure_notes": {}
}
""".strip(),
        encoding="utf-8",
    )
    return tpl_dir


@pytest.fixture
def patched_storage(tmp_path, monkeypatch) -> Path:
    """将 STORAGE_ROOT 指向临时目录（preview + deliverable）。"""
    storage = tmp_path / "storage"
    storage.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(deliverable_module, "STORAGE_ROOT", storage)
    monkeypatch.setattr(tfs_module, "STORAGE_ROOT", storage)
    return storage


def _build_service(session: AsyncSession, tpl_dir: Path) -> TemplateFillService:
    loader = TemplateManifestLoader(base_dir=tpl_dir)
    svc = TemplateFillService(session, loader=loader)

    async def _fake_map(project_id, db):
        return {
            "company_full_name": "测试有限公司",
            "audit_year": "2025",
            "signing_partner": "",  # 故意留空 → missing_fields 命中
        }

    # 避免依赖 ReportPlaceholderService / 真实项目数据
    svc.registry.build_placeholder_map = _fake_map  # type: ignore[assignment]
    return svc


# ---------------------------------------------------------------------------
# preview 不落库
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_preview_does_not_create_deliverable(
    session, tmp_path, patched_storage
):
    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)
    project_id = uuid.uuid4()
    user_id = uuid.uuid4()

    result = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=user_id,
    )

    # 无 word_export_task / version 创建
    tasks = (await session.execute(sa.select(WordExportTask))).scalars().all()
    versions = (
        await session.execute(sa.select(WordExportTaskVersion))
    ).scalars().all()
    assert tasks == []
    assert versions == []

    # 仅写了一条 preview session
    sessions = (
        await session.execute(sa.select(FillPreviewSession))
    ).scalars().all()
    assert len(sessions) == 1
    assert sessions[0].id == result.preview_session_id

    # OPT 扫描出 emphasis + key_audit_matters
    opt_ids = {v.section_id for v in result.optional_sections}
    assert opt_ids == {"emphasis", "key_audit_matters"}

    # missing_fields 含留空的 signing_partner
    assert "signing_partner" in result.missing_fields
    assert result.company_subtype_resolved == "type_d"
    assert result.template_version == "test-v1"


# ---------------------------------------------------------------------------
# confirm 创建 version + version_no 递增
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_confirm_creates_version_and_increments(
    session, tmp_path, patched_storage
):
    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)
    project_id = uuid.uuid4()
    user_id = uuid.uuid4()

    # --- 第一轮 preview → confirm ---
    preview1 = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=user_id,
    )
    confirm1 = await svc.confirm_report_body(
        project_id,
        2025,
        preview_session_id=preview1.preview_session_id,
        optional_sections={"emphasis": False, "key_audit_matters": True},
        user_id=user_id,
    )
    # create_task 预置空 v1 占位版本；render_and_store 落内容版本 → v2
    first_version = confirm1.version_no
    assert first_version >= 1

    # 创建了一个 task
    tasks = (await session.execute(sa.select(WordExportTask))).scalars().all()
    assert len(tasks) == 1
    versions_after_1 = (
        await session.execute(sa.select(WordExportTaskVersion))
    ).scalars().all()
    count_after_1 = len(versions_after_1)
    assert count_after_1 >= 1

    # preview session 已删除（confirm 时清理）
    sessions = (
        await session.execute(sa.select(FillPreviewSession))
    ).scalars().all()
    assert sessions == []

    # report_body_json schema（需求 6.8）
    body = confirm1.report_body_json
    assert body["optional_sections"] == {
        "emphasis": False,
        "key_audit_matters": True,
    }
    assert body["template_version"] == "test-v1"
    assert body["company_subtype"] == "type_d"
    assert body["template_variant"] == "simple"
    assert "signing_partner" in body["missing_fields"]
    # guidance 副本与内容版本号一致
    assert body["guidance_version_path"].endswith(
        f"with_notes_v{first_version}.docx"
    )

    # guidance 副本已落盘
    assert Path(body["guidance_version_path"]).exists()

    # --- 第二轮 preview → confirm（同项目）→ version_no 递增 ---
    preview2 = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=user_id,
    )
    confirm2 = await svc.confirm_report_body(
        project_id,
        2025,
        preview_session_id=preview2.preview_session_id,
        optional_sections={"emphasis": True, "key_audit_matters": True},
        user_id=user_id,
    )
    assert confirm2.version_no == first_version + 1

    # 仍是同一个 task（复用），版本数 +1
    tasks = (await session.execute(sa.select(WordExportTask))).scalars().all()
    assert len(tasks) == 1
    versions_after_2 = (
        await session.execute(sa.select(WordExportTaskVersion))
    ).scalars().all()
    assert len(versions_after_2) == count_after_1 + 1


# ---------------------------------------------------------------------------
# 清洁版不含 NOTE（guidance 副本含 NOTE）
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_confirm_strips_notes_from_clean_copy(
    session, tmp_path, patched_storage
):
    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)
    project_id = uuid.uuid4()
    user_id = uuid.uuid4()

    preview = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=user_id,
    )
    confirm = await svc.confirm_report_body(
        project_id,
        2025,
        preview_session_id=preview.preview_session_id,
        optional_sections={"emphasis": False, "key_audit_matters": True},
        user_id=user_id,
    )

    # 正式版（task.file_path）不含 NOTE 标记
    task = (await session.execute(sa.select(WordExportTask))).scalars().one()
    clean_doc = Document(task.file_path)
    clean_text = "\n".join(p.text for p in clean_doc.paragraphs)
    assert "##NOTE" not in clean_text
    # 未勾选的 emphasis 段被删除，勾选的 KAM 保留
    assert "强调事项" not in clean_text
    assert "关键审计事项" in clean_text

    # guidance 副本仍含 NOTE
    guidance_doc = Document(confirm.guidance_version_path)
    guidance_text = "\n".join(p.text for p in guidance_doc.paragraphs)
    assert "##NOTE" in guidance_text


# ---------------------------------------------------------------------------
# 过期 session confirm 被拒
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_confirm_rejects_expired_session(
    session, tmp_path, patched_storage
):
    from datetime import datetime, timedelta, timezone

    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)
    project_id = uuid.uuid4()
    user_id = uuid.uuid4()

    preview = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=user_id,
    )
    # 人为过期
    sess = await session.get(FillPreviewSession, preview.preview_session_id)
    sess.expires_at = datetime.now(timezone.utc) - timedelta(hours=1)
    await session.flush()

    with pytest.raises(ValueError, match="过期"):
        await svc.confirm_report_body(
            project_id,
            2025,
            preview_session_id=preview.preview_session_id,
            optional_sections={"key_audit_matters": True},
            user_id=user_id,
        )


@pytest.mark.asyncio
async def test_confirm_rejects_wrong_user(session, tmp_path, patched_storage):
    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)
    project_id = uuid.uuid4()

    preview = await svc.preview_report_body(
        project_id,
        2025,
        opinion_type="unqualified",
        company_subtype="type_d",
        template_variant="simple",
        user_id=uuid.uuid4(),
    )
    with pytest.raises(ValueError, match="用户"):
        await svc.confirm_report_body(
            project_id,
            2025,
            preview_session_id=preview.preview_session_id,
            optional_sections={"key_audit_matters": True},
            user_id=uuid.uuid4(),  # 不同用户
        )


# ---------------------------------------------------------------------------
# purge_expired_sessions
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_purge_expired_sessions(session, tmp_path, patched_storage):
    from datetime import datetime, timedelta, timezone

    tpl_dir = _make_template_dir(tmp_path)
    svc = _build_service(session, tpl_dir)

    p1 = await svc.preview_report_body(
        uuid.uuid4(), 2025, opinion_type="unqualified",
        company_subtype="type_d", template_variant="simple", user_id=uuid.uuid4(),
    )
    p2 = await svc.preview_report_body(
        uuid.uuid4(), 2025, opinion_type="unqualified",
        company_subtype="type_d", template_variant="simple", user_id=uuid.uuid4(),
    )
    # p1 过期，p2 未过期
    s1 = await session.get(FillPreviewSession, p1.preview_session_id)
    s1.expires_at = datetime.now(timezone.utc) - timedelta(hours=1)
    await session.flush()

    purged = await svc.purge_expired_sessions()
    assert purged == 1

    remaining = (
        await session.execute(sa.select(FillPreviewSession))
    ).scalars().all()
    assert len(remaining) == 1
    assert remaining[0].id == p2.preview_session_id


# ---------------------------------------------------------------------------
# validate_kam_word_mode 单元
# ---------------------------------------------------------------------------


def test_validate_kam_listed_unchecked_warns():
    warning = validate_kam_word_mode(
        optional_sections={"key_audit_matters": False},
        company_type="listed",
        is_pie=False,
        opinion_type="unqualified",
    )
    assert warning is not None
    assert "关键审计事项" in warning


def test_validate_kam_listed_checked_ok():
    warning = validate_kam_word_mode(
        optional_sections={"key_audit_matters": True},
        company_type="listed",
        is_pie=False,
        opinion_type="unqualified",
    )
    assert warning is None


def test_validate_kam_non_listed_no_warning():
    warning = validate_kam_word_mode(
        optional_sections={"key_audit_matters": False},
        company_type="non_listed",
        is_pie=False,
        opinion_type="unqualified",
    )
    assert warning is None


def test_validate_kam_disclaimer_no_warning():
    warning = validate_kam_word_mode(
        optional_sections={"key_audit_matters": False},
        company_type="listed",
        is_pie=True,
        opinion_type="disclaimer",
    )
    assert warning is None

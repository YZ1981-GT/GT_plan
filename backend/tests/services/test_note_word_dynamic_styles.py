"""Tests for Sprint C.4 — Word 导出动态样式增强.

Covers:
- C.4.1: GTNoteDynamicRow / GTNoteDynamicCol 样式
- C.4.3: 空表替换 + 空章节跳过
- C.4.5: 抵销双列构造
- C.4.9: Jinja ref() 预渲染
"""
from __future__ import annotations

import pytest
from docx import Document

from app.services.note_word_dynamic_styles import (
    DYNAMIC_COL_FILL,
    DYNAMIC_ROW_FILL,
    ELIMINATION_FILL,
    apply_consol_elimination_style,
    apply_dynamic_col_style,
    apply_dynamic_row_style,
    build_consol_dual_column_table_data,
    get_table_render_mode,
    is_empty_table,
    pre_render_jinja_refs,
    replace_empty_table_with_paragraph,
    should_skip_empty_section,
)


# ---------------------------------------------------------------------------
# C.4.1: Dynamic Row/Col Styles
# ---------------------------------------------------------------------------


class TestDynamicStyles:
    """C.4.1 — GTNoteDynamicRow / GTNoteDynamicCol."""

    def test_apply_dynamic_row_style(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[0]
        apply_dynamic_row_style(cell)
        # Verify shading XML is set
        tc_pr = cell._tc.tcPr
        assert tc_pr is not None
        from docx.oxml.ns import qn
        shd = tc_pr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == DYNAMIC_ROW_FILL

    def test_apply_dynamic_col_style(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[1]
        apply_dynamic_col_style(cell)
        from docx.oxml.ns import qn
        shd = cell._tc.tcPr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == DYNAMIC_COL_FILL

    def test_apply_elimination_style(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[1]
        apply_consol_elimination_style(cell, eliminated=True)
        from docx.oxml.ns import qn
        shd = cell._tc.tcPr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == ELIMINATION_FILL

    def test_elimination_style_skipped_when_false(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.rows[1].cells[1]
        apply_consol_elimination_style(cell, eliminated=False)
        from docx.oxml.ns import qn
        # No shading should be added
        if cell._tc.tcPr is not None:
            shd = cell._tc.tcPr.find(qn("w:shd"))
            assert shd is None or shd.get(qn("w:fill")) != ELIMINATION_FILL


# ---------------------------------------------------------------------------
# C.4.3: Empty Section / Table
# ---------------------------------------------------------------------------


class TestEmptyHandling:
    """C.4.3 — 空表替换 + 空章节跳过."""

    def test_skip_deleted_section(self):
        assert should_skip_empty_section({"is_deleted": True}) is True

    def test_skip_not_applicable_section(self):
        assert should_skip_empty_section({"status": "not_applicable"}) is True

    def test_normal_section_not_skipped(self):
        # design §7.1 ④ 生效后：仅 status=active 但无任何内容的章节视为空，应跳过。
        # 一个「正常」章节须有内容才不跳过。
        assert should_skip_empty_section(
            {"status": "active", "text_content": "本期内容"}
        ) is False

    def test_skip_is_empty_flag(self):
        """design §7.1 ③: is_empty=True 整节删除."""
        assert should_skip_empty_section({"is_empty": True}) is True

    def test_skip_empty_text_and_empty_tables(self):
        """design §7.1 ④: text 空 + 所有表全空 → 跳过."""
        note = {
            "text_content": "",
            "table_data": {"rows": [{"cells": [0, 0, "-"]}]},
        }
        assert should_skip_empty_section(note) is True

    def test_skip_empty_text_and_empty_multi_tables(self):
        """design §7.1 ④: 多表 _tables 全空 → 跳过."""
        note = {
            "text_content": None,
            "table_data": {
                "_tables": [
                    {"rows": [{"cells": [0]}]},
                    {"rows": []},
                ]
            },
        }
        assert should_skip_empty_section(note) is True

    def test_no_skip_with_text_content(self):
        """design §7.1 ④: 有文本内容则不跳过（即使表为空）."""
        note = {
            "text_content": "本期货币资金构成如下",
            "table_data": {"rows": [{"cells": [0]}]},
        }
        assert should_skip_empty_section(note) is False

    def test_no_skip_with_nonempty_table(self):
        """design §7.1 ④: 表有数据则不跳过（即使无文本）."""
        note = {
            "text_content": "",
            "table_data": {"rows": [{"cells": [0, 1234.5]}]},
        }
        assert should_skip_empty_section(note) is False

    def test_empty_table_all_zeros(self):
        td = {"rows": [{"cells": [0, 0, 0]}, {"cells": ["-", None, ""]}]}
        assert is_empty_table(td) is True

    def test_non_empty_table(self):
        td = {"rows": [{"cells": [0, 100, 0]}]}
        assert is_empty_table(td) is False

    def test_empty_table_no_rows(self):
        assert is_empty_table({"rows": []}) is True

    def test_render_mode_no_business(self):
        td = {"_render_as": "no_business_paragraph", "rows": [{"cells": [100]}]}
        assert get_table_render_mode(td) == "no_business_paragraph"

    def test_render_mode_empty(self):
        td = {"rows": [{"cells": [0, 0]}]}
        assert get_table_render_mode(td) == "no_business_paragraph"

    def test_render_mode_normal(self):
        td = {"rows": [{"cells": [100]}]}
        assert get_table_render_mode(td) == "normal"

    def test_replace_empty_table_paragraph(self):
        doc = Document()
        p = replace_empty_table_with_paragraph(doc, "本期无此项业务")
        assert p is not None
        assert "本期无此项业务" in p.text


# ---------------------------------------------------------------------------
# C.4.5: Consol Dual Column
# ---------------------------------------------------------------------------


class TestConsolDualColumn:
    """C.4.5 — 抵销前后双列."""

    def test_simple_dual_column(self):
        base = {
            "headers": ["项目", "金额"],
            "rows": [{"cells": ["银行存款", 100.0]}],
        }
        result = build_consol_dual_column_table_data(base, {0: {1: 20.0}})
        # Headers: 项目 + 金额（抵销前）+ 金额（抵销后）
        assert len(result["headers"]) == 3
        assert "抵销前" in result["headers"][1]
        assert "抵销后" in result["headers"][2]
        # Row cells: 银行存款 / 100 / 80
        assert result["rows"][0]["cells"] == ["银行存款", 100.0, 80.0]
        assert result["_has_elimination_columns"] is True

    def test_no_elimination_data(self):
        base = {
            "headers": ["项目", "金额"],
            "rows": [{"cells": ["现金", 50.0]}],
        }
        result = build_consol_dual_column_table_data(base, None)
        # Without elimination data, post = pre
        assert result["rows"][0]["cells"] == ["现金", 50.0, 50.0]

    def test_columns_meta_propagation(self):
        base = {
            "headers": ["项目", "金额"],
            "rows": [{"cells": ["A", 100.0]}],
            "_columns_meta": [
                {"id": "col_label", "label": "项目"},
                {"id": "col_amount", "label": "金额"},
            ],
        }
        result = build_consol_dual_column_table_data(base, {})
        meta = result["_columns_meta"]
        assert len(meta) == 3
        assert meta[1]["id"] == "col_amount_pre"
        assert meta[2]["id"] == "col_amount_post"
        assert meta[2]["is_elimination"] is True

    def test_dict_value_extraction(self):
        base = {
            "headers": ["项目", "金额"],
            "rows": [{"cells": ["A", {"value": 200.0, "mode": "manual"}]}],
        }
        result = build_consol_dual_column_table_data(base, {0: {1: 50.0}})
        assert result["rows"][0]["cells"][1] == 200.0  # pre
        assert result["rows"][0]["cells"][2] == 150.0  # post


# ---------------------------------------------------------------------------
# C.4.9: Jinja ref() Pre-rendering
# ---------------------------------------------------------------------------


class TestJinjaRefRendering:
    """C.4.9 — Word 内部引用 ref() 渲染."""

    def test_simple_ref(self):
        text = '详见 {{ ref("section_revenue") }}'
        result = pre_render_jinja_refs(text, {"section_revenue": "八、（一）2."})
        assert result == "详见 八、（一）2."

    def test_multiple_refs(self):
        text = '见 {{ ref("s1") }} 和 {{ ref("s2") }}'
        result = pre_render_jinja_refs(text, {"s1": "一、", "s2": "二、"})
        assert result == "见 一、 和 二、"

    def test_unknown_ref(self):
        text = '{{ ref("unknown") }}'
        result = pre_render_jinja_refs(text, {})
        assert "[未知章节: unknown]" in result

    def test_no_refs(self):
        text = "普通文本无 ref"
        assert pre_render_jinja_refs(text, {}) == text

    def test_double_quotes_and_single_quotes(self):
        text = '{{ ref("a") }} vs {{ref(\'b\')}}'
        result = pre_render_jinja_refs(text, {"a": "A.", "b": "B."})
        assert "A." in result
        assert "B." in result

    def test_empty_text(self):
        assert pre_render_jinja_refs("", {"a": "A"}) == ""

    def test_empty_numbers(self):
        text = '{{ ref("a") }}'
        # No mapping returns the original text (empty rendered_numbers short-circuits)
        result = pre_render_jinja_refs(text, {})
        assert "未知章节" in result

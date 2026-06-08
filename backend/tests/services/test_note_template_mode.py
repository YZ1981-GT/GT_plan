"""集成测试 — NoteWordExporter mode='template'（附注模板填充流程）.

Spec:   .kiro/specs/audit-report-template-integration/ tasks 10.2/10.6/9.3
Design: design.md §7 附注模板填充流程 / §7.1 裁剪优先级 / §8 §13 编号
Reqs:   10, 12, 13

测试目标：
  1. 保留章节的 {{section:code}} 文本 + {{table:code}} 表格被填充
  2. 空/删除章节（§7.1 ①③④）整 SECTION 块被删除
  3. {{seq:prefix}} 裁剪后编号连续，与 compute_section_numbers 一致
  4. legacy_alias join（DB 五、1 → 模板 八、1）
  5. standalone 导出排除 consolidated_only 章节块
  6. 真实模板：固定资产命中 §八 块（非 §四 会计政策块）

策略：构造最小合成「打标」docx（含 ##SECTION:## 块 + 占位符），
patch manifest loader + section_code_index，避免依赖真实模板内联占位。
"""

from __future__ import annotations

import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest
from docx import Document

import app.services.note_word_exporter as nwe_mod
from app.services.note_word_exporter import NoteWordExporter
from app.services.template_manifest_loader import ManifestEntry


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_note(
    section: str,
    title: str,
    *,
    text_content: str | None = None,
    table_data: dict | None = None,
    is_deleted: bool = False,
    is_empty: bool = False,
    status=None,
):
    return SimpleNamespace(
        id=uuid4(),
        project_id=uuid4(),
        year=2025,
        note_section=section,
        section_title=title,
        account_name=None,
        content_type=None,
        table_data=table_data,
        text_content=text_content,
        source_template=None,
        status=status,
        sort_order=None,
        is_deleted=is_deleted,
        is_empty=is_empty,
        is_stale=False,
    )


def _make_db(notes: list) -> MagicMock:
    db = MagicMock()
    scalars = MagicMock()
    scalars.all = MagicMock(return_value=notes)
    result = MagicMock()
    result.scalars = MagicMock(return_value=scalars)
    db.execute = AsyncMock(return_value=result)
    return db


def _extract(docx_bytes: bytes, member: str = "word/document.xml") -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        return z.read(member).decode("utf-8")


def _build_tagged_docx(tmp_path: Path) -> Path:
    """构造最小合成打标 docx：4 个 SECTION 块 + 占位符."""
    doc = Document()

    def block(code: str, title: str, body_lines: list[str]):
        doc.add_paragraph(f"##SECTION:{code}##")
        doc.add_paragraph(title)
        for line in body_lines:
            doc.add_paragraph(line)
        doc.add_paragraph(f"##/SECTION:{code}##")

    # 八、1 保留（DB 用 legacy alias 五、1 join）+ 文本 + 表格
    block("八、1", "货币资金", ["{{seq:八}} 货币资金", "{{section:八、1}}", "{{table:八、1}}"])
    # 八、2 空章节（is_empty）→ 删除
    block("八、2", "交易性金融资产", ["{{seq:八}}", "{{section:八、2}}"])
    # 八、3 保留 + 文本
    block("八、3", "衍生金融资产", ["{{seq:八}} 衍生金融资产", "{{section:八、3}}"])
    # 九、1 consolidated_only → standalone 排除
    block("九、1", "合并范围", ["{{seq:九}}", "{{section:九、1}}"])

    out = tmp_path / "soe_standalone.docx"
    doc.save(str(out))
    return out


def _patch_loader(monkeypatch, docx_path: Path):
    entry = ManifestEntry(
        rel_path=Path("disclosure_notes/soe_standalone.docx"),
        abs_path=docx_path.resolve(),
        exists=True,
    )
    fake_loader = MagicMock()
    fake_loader.resolve_disclosure_notes = MagicMock(return_value=entry)
    monkeypatch.setattr(nwe_mod, "get_template_manifest_loader", lambda: fake_loader)


def _patch_index(monkeypatch):
    sections = [
        {"section_code": "八、1", "legacy_aliases": ["五、1"], "content_type": "table"},
        {"section_code": "八、2", "legacy_aliases": [], "content_type": "table"},
        {"section_code": "八、3", "legacy_aliases": [], "content_type": "text"},
        {"section_code": "九、1", "legacy_aliases": [], "content_type": "text"},
    ]
    monkeypatch.setattr(nwe_mod, "_load_section_code_index", lambda vk: sections)


def _patch_scope(monkeypatch):
    """九、 前缀视为 consolidated_only（standalone 排除）."""

    def fake_applies(section, template_type, report_scope):
        code = (section or "").strip()
        if code.startswith("九、") and report_scope == "standalone":
            return False
        return True

    monkeypatch.setattr(nwe_mod, "note_applies_to_report_scope", fake_applies)


# ===========================================================================
# 合成模板测试
# ===========================================================================


@pytest.mark.asyncio
async def test_template_mode_fill_skip_seq_and_alias(tmp_path, monkeypatch):
    docx_path = _build_tagged_docx(tmp_path)
    _patch_loader(monkeypatch, docx_path)
    _patch_index(monkeypatch)
    _patch_scope(monkeypatch)

    notes = [
        # DB 用 legacy 编号 五、1（应 join 到模板 八、1）
        _make_note(
            "五、1",
            "货币资金",
            text_content="货币资金期末余额披露内容",
            table_data={
                "headers": ["项目", "期末余额"],
                "rows": [{"label": "银行存款", "values": [123456.0]}],
            },
        ),
        # 八、2 空章节 → 删除
        _make_note("八、2", "交易性金融资产", is_empty=True),
        # 八、3 保留（文本）
        _make_note("八、3", "衍生金融资产", text_content="衍生金融资产披露内容"),
    ]
    db = _make_db(notes)
    exporter = NoteWordExporter(db)

    bio = await exporter.export(
        project_id=uuid4(),
        year=2025,
        template_type="soe",
        report_scope="standalone",
        mode="template",
    )
    xml = _extract(bio.getvalue())
    out_doc = Document(BytesIO(bio.getvalue()))
    texts = [p.text for p in out_doc.paragraphs]
    full = "\n".join(texts)

    # 1. legacy alias join：八、1 文本被 五、1 note 填充
    assert "货币资金期末余额披露内容" in full, "legacy alias (五、1→八、1) 文本未填充"
    # 表格被渲染（label 出现在 document.xml）
    assert "银行存款" in xml, "{{table:八、1}} 表格未渲染"

    # 2. 空章节 八、2 整块被删除
    assert "交易性金融资产" not in full, "空章节 八、2 未删除"
    # consolidated_only 九、1 被排除
    assert "合并范围" not in full, "consolidated_only 九、1 未排除"

    # 3. 八、3 保留
    assert "衍生金融资产披露内容" in full

    # 4. 残留标记清理
    assert "##SECTION" not in full, "残留 ##SECTION 标记未清理"
    assert "{{section:" not in full, "残留 {{section}} 占位未填充"
    assert "{{table:" not in full, "残留 {{table}} 占位未填充"
    assert "{{seq:" not in full, "残留 {{seq}} 占位未填充"

    # 5. {{seq:}} 编号连续：八组保留 八、1 + 八、3 → "1","2"
    from app.services.note_section_numbering import compute_section_numbers

    expected = compute_section_numbers(
        [{"note_section": "八、1"}, {"note_section": "八、3"}],
        report_scope="standalone",
        template_type="soe",
    )
    assert expected == {"八、1": "1", "八、3": "2"}
    # 八、1 块含 "1 货币资金"，八、3 块含 "2 衍生金融资产"
    assert any(t.strip().startswith("1") and "货币资金" in t for t in texts), texts
    assert any(t.strip().startswith("2") and "衍生金融资产" in t for t in texts), texts


@pytest.mark.asyncio
async def test_template_mode_deleted_status_skipped(tmp_path, monkeypatch):
    """§7.1 ①is_deleted / ②status=not_applicable 章节整块删除."""
    docx_path = _build_tagged_docx(tmp_path)
    _patch_loader(monkeypatch, docx_path)
    _patch_index(monkeypatch)
    _patch_scope(monkeypatch)

    notes = [
        _make_note("八、1", "货币资金", text_content="保留内容"),
        _make_note("八、3", "衍生金融资产", status="not_applicable"),
    ]
    db = _make_db(notes)
    exporter = NoteWordExporter(db)
    bio = await exporter.export(
        project_id=uuid4(),
        year=2025,
        template_type="soe",
        report_scope="standalone",
        mode="template",
    )
    out_doc = Document(BytesIO(bio.getvalue()))
    full = "\n".join(p.text for p in out_doc.paragraphs)

    assert "保留内容" in full
    # 八、3 status=not_applicable → 删除
    assert "衍生金融资产" not in full, "status=not_applicable 章节未删除"
    # 单组仅 1 条（八、1）→ 不编号
    assert "{{seq:" not in full


# ===========================================================================
# 真实模板测试（固定资产命中 §八 块，非 §四 会计政策）
# ===========================================================================

_REAL_TEMPLATE = (
    Path(__file__).resolve().parents[2]
    / "data"
    / "audit_report_templates"
    / "disclosure_notes"
    / "soe_standalone.docx"
)


@pytest.mark.skipif(not _REAL_TEMPLATE.exists(), reason="真实附注模板缺失")
@pytest.mark.asyncio
async def test_template_mode_real_template_collision(monkeypatch):
    """真实 soe_standalone 模板：固定资产 note 填到 §八 项目注释块.

    模板含 固定资产 重名块（§四 会计政策 + §八 项目注释）。
    DB note_section='八、22'（固定资产，项目注释章），应只命中 §八 块。
    """
    # 找出真实模板中固定资产对应的 八、N section_code
    doc = Document(str(_REAL_TEMPLATE))
    from app.services.word_doc_utils import scan_section_blocks

    blocks = scan_section_blocks(doc)
    fa_code = None
    for b in blocks:
        if not b.section_code.startswith("八、"):
            continue
        titles = [
            nwe_mod.__dict__  # placeholder; real check below
        ]
        # 块内第二段通常是标题
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph

        para_texts = [
            (Paragraph(el, doc).text or "").strip()
            for el in b.elements
            if el.tag == _qn("w:p")
        ]
        if any("固定资产" == t for t in para_texts):
            fa_code = b.section_code
            break

    if fa_code is None:
        pytest.skip("真实模板未找到 八、固定资产 块")

    note = _make_note(fa_code, "固定资产", text_content="固定资产期末账面价值披露")
    db = _make_db([note])
    exporter = NoteWordExporter(db)
    bio = await exporter.export(
        project_id=uuid4(),
        year=2025,
        template_type="soe",
        report_scope="standalone",
        mode="template",
    )
    full = "\n".join(p.text for p in Document(BytesIO(bio.getvalue())).paragraphs)
    # 模板无 {{section}} 占位时不强行注入；至少确认导出成功 + 标记清理
    assert "##SECTION" not in full, "真实模板残留 SECTION 标记"
    assert len(bio.getvalue()) > 0

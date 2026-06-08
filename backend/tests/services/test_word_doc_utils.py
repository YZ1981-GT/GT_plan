"""word_doc_utils 单元测试."""

from pathlib import Path

from docx import Document

from app.services.word_doc_utils import (
    apply_optional_sections,
    merge_runs_for_replace,
    replace_placeholders_in_doc,
    scan_optional_sections,
    strip_guidance_notes,
)


def _doc_with_paras(*lines: str) -> Document:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def test_merge_runs_for_replace_splits():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{comp")
    p.add_run("any_name}}")
    merge_runs_for_replace(p)
    assert p.text == "{{company_name}}"


def test_replace_placeholders_in_doc():
    doc = _doc_with_paras("{{company_full_name}} 审计报告")
    n = replace_placeholders_in_doc(doc, {"company_full_name": "测试公司"})
    assert "测试公司" in doc.paragraphs[0].text
    assert "{{" not in doc.paragraphs[0].text
    assert n >= 1


def test_scan_optional_sections():
    doc = _doc_with_paras(
        "##OPT:emphasis:强调事项段##",
        "强调事项正文",
        "##/OPT:emphasis##",
    )
    secs = scan_optional_sections(doc)
    assert len(secs) == 1
    assert secs[0].section_id == "emphasis"
    assert secs[0].description == "强调事项段"


def test_apply_optional_sections_removes_unselected():
    doc = _doc_with_paras(
        "保留段",
        "##OPT:emphasis:强调##",
        "强调正文",
        "##/OPT:emphasis##",
        "结尾",
    )
    removed = apply_optional_sections(doc, {"emphasis": False})
    texts = [p.text for p in doc.paragraphs]
    assert "强调正文" not in texts
    assert removed >= 3


def test_apply_optional_sections_keeps_selected():
    doc = _doc_with_paras(
        "##OPT:kam:KAM##",
        "KAM 正文",
        "##/OPT:kam##",
    )
    apply_optional_sections(doc, {"kam": True})
    texts = [p.text for p in doc.paragraphs]
    assert "KAM 正文" in texts
    assert not any("##OPT" in t for t in texts)


def test_strip_guidance_notes_block():
    doc = _doc_with_paras(
        "正文",
        "##NOTE:编制提示##",
        "请填写签字合伙人",
        "##/NOTE:编制提示##",
        "结尾",
    )
    removed = strip_guidance_notes(doc)
    texts = [p.text for p in doc.paragraphs]
    assert "请填写签字合伙人" not in texts
    assert "正文" in texts
    assert removed >= 1


def test_strip_guidance_preserves_legal_footnote():
    doc = _doc_with_paras("附注 [注1] 披露说明")
    strip_guidance_notes(doc)
    assert "[注1]" in doc.paragraphs[0].text

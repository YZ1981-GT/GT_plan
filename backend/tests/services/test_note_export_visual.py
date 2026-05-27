"""视觉回归测试 — 致同 Word 排版规范 11 项断言（Sprint 2 Task 2.6）.

Spec:    .kiro/specs/disclosure-note-full-revamp/ Sprint 2 Task 2.2 + 2.6
Design:  D7 致同 Word 排版规范单一真源（GTNote* 命名空间）

策略：
- 不连真 PG。用 ``MagicMock`` 替换 ``db.execute`` 返回伪造 ``DisclosureNote`` 列表
- 直接 ``await NoteWordExporter.export()`` 拿 BytesIO
- 用 ``zipfile`` 解 ``word/document.xml`` 和 ``word/styles.xml`` 做 XML 字符串断言

11 项断言（CI 卡点）：
    1. 字体名（中文 仿宋_GB2312 / 数字 Arial Narrow）
    2. 字号小四 12pt 统一（sz=24 半磅）
    3. 章节标题左缩进 -2 字符（GTNoteHeading1: leftChars=-200）
    4. 正文首行不缩进（GTNoteBody: firstLine=0）
    5. 段前 0、段后 0.9 行（spacing after=216 twip）
    6. 表格三线（顶/底 sz=8 / 表头下 sz=4）
    7. 表格内行高（trHeight 397 twip = 0.7cm exact）
    8. 空值/零值留白（``fmt_amount_gt(0) == ""`` 且 ``fmt_amount_gt(None) == ""``）
    9. 标题行不重复（无 ``w:tblHeader`` 元素）
    10. 页边距上 3.2 / 下 2.54 / 左 3 / 右 3.18 cm
    11. docx 样式 grep 含 ``GTNote*`` 前缀 ≥ 7 个（6 段落 + 1 字符 + 1 表格 + 2 sidecar = 10）
"""

from __future__ import annotations

import zipfile
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest

from app.services.note_word_exporter import (
    NoteWordExporter,
    TEMPLATE_PATH,
    apply_gt_dual_font,
    apply_gt_row_height,
    apply_gt_three_line,
    fill_multi_header,
    fmt_amount_gt,
)


# ---------------------------------------------------------------------------
# Fixtures：伪造 db + DisclosureNote 列表
# ---------------------------------------------------------------------------


def _make_fake_note(
    section: str,
    title: str,
    table_data: dict | None = None,
    text_content: str | None = None,
):
    """构造最小 SimpleNamespace 模拟 DisclosureNote ORM 对象."""
    return SimpleNamespace(
        id=uuid4(),
        project_id=uuid4(),
        year=2025,
        note_section=section,
        section_title=title,
        account_name=None,
        content_type=None,
        table_data=table_data,
        text_content=text_content,
        source_template=None,
        status=None,
        sort_order=None,
        is_deleted=False,
        is_stale=False,
    )


def _make_fake_db(notes: list) -> MagicMock:
    """模拟 db.execute(query) → result.scalars().all() = notes."""
    db = MagicMock()
    scalars = MagicMock()
    scalars.all = MagicMock(return_value=notes)
    result = MagicMock()
    result.scalars = MagicMock(return_value=scalars)
    db.execute = AsyncMock(return_value=result)
    return db


@pytest.fixture
def fake_notes():
    """构造一个最小附注集合（1 章节 1 表 + 含 0 / None 值用于测试 fmt_amount_gt）."""
    table = {
        "headers": ["项目", "期末数", "期初数"],
        "rows": [
            {"label": "应收账款", "values": [1234567.89, 0]},
            {"label": "其他", "values": [None, 500.0]},
            {"label": "合计", "values": [1234567.89, 500.0], "is_total": True},
        ],
    }
    return [
        _make_fake_note(
            "1",
            "货币资金",
            table_data=table,
            text_content="本期末货币资金期末余额为 1,234,567.89 元。",
        ),
    ]


@pytest.fixture
async def exported_docx_bytes(fake_notes):
    """运行真实 NoteWordExporter.export → 返回 BytesIO 内容."""
    db = _make_fake_db(fake_notes)
    exporter = NoteWordExporter(db)
    bio = await exporter.export(project_id=uuid4(), year=2025)
    return bio.getvalue()


def _extract_xml(docx_bytes: bytes, member: str) -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        return z.read(member).decode("utf-8")


# ===========================================================================
# 0. 模板 docx 必须存在（Sprint 2 Task 2.1 产出）
# ===========================================================================


def test_template_docx_exists():
    """note_export_template.docx 必须由 scripts/build_note_export_template.py 生成."""
    assert TEMPLATE_PATH.exists(), (
        f"模板 docx 不存在: {TEMPLATE_PATH}。"
        "请运行: python scripts/build_note_export_template.py --apply"
    )


# ===========================================================================
# 1. 字体名（中文 仿宋_GB2312 / 数字 Arial Narrow）
# ===========================================================================


def test_assert_1_fonts_present_in_styles(exported_docx_bytes):
    """断言 1：styles.xml 中含 仿宋_GB2312 + Arial Narrow."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    assert "仿宋_GB2312" in styles, "styles.xml 缺少 仿宋_GB2312"
    assert "Arial Narrow" in styles, "styles.xml 缺少 Arial Narrow"


def test_assert_1_dual_font_applied_in_document(exported_docx_bytes):
    """断言 1b：document.xml 数字 cell 同时含 ascii=Arial Narrow + eastAsia=仿宋_GB2312."""
    doc = _extract_xml(exported_docx_bytes, "word/document.xml")
    assert 'w:ascii="Arial Narrow"' in doc
    assert 'w:eastAsia="仿宋_GB2312"' in doc


# ===========================================================================
# 2. 字号小四 12pt 统一（sz=24 半磅）
# ===========================================================================


def test_assert_2_font_size_xiaosi_12pt(exported_docx_bytes):
    """断言 2：所有 GTNote* 段落样式均 sz=24（小四 12pt）."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    # 抽样断言：GTNoteBody 块内必含 sz="24"
    body_block_start = styles.find('w:styleId="GTNoteBody"')
    assert body_block_start > -1
    body_block_end = styles.find("</w:style>", body_block_start)
    body_block = styles[body_block_start:body_block_end]
    assert 'w:sz w:val="24"' in body_block or 'w:val="24"' in body_block


# ===========================================================================
# 3. 章节标题左缩进 -2 字符（GTNoteHeading1: leftChars=-200）
# ===========================================================================


def test_assert_3_heading1_left_indent_minus_2_char(exported_docx_bytes):
    """断言 3：GTNoteHeading1 含 ind leftChars=-200."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    h1_start = styles.find('w:styleId="GTNoteHeading1"')
    h1_end = styles.find("</w:style>", h1_start)
    h1_block = styles[h1_start:h1_end]
    assert 'w:leftChars="-200"' in h1_block, "GTNoteHeading1 缺少左缩进 -2 字符"


# ===========================================================================
# 4. 正文首行不缩进（GTNoteBody: firstLine=0）
# ===========================================================================


def test_assert_4_body_no_first_line_indent(exported_docx_bytes):
    """断言 4：GTNoteBody 含 firstLine=0（不缩进）."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    body_start = styles.find('w:styleId="GTNoteBody"')
    body_end = styles.find("</w:style>", body_start)
    body_block = styles[body_start:body_end]
    assert 'w:firstLine="0"' in body_block, "GTNoteBody 首行缩进未关闭"


# ===========================================================================
# 5. 段前 0、段后 0.9 行（spacing after=216 twip）
# ===========================================================================


def test_assert_5_spacing_before_zero_after_216(exported_docx_bytes):
    """断言 5：GTNoteBody / GTNoteHeading1 段前 0 + 段后 ≥ 216 twip（0.9 行 ≈ 216）."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    for style_id in ("GTNoteHeading1", "GTNoteBody"):
        s = styles.find(f'w:styleId="{style_id}"')
        e = styles.find("</w:style>", s)
        block = styles[s:e]
        assert 'w:before="0"' in block, f"{style_id} 段前不为 0"
        # 段后 0.9 行 = 216 twip（12pt × 0.9 × 20）
        assert 'w:after="216"' in block, f"{style_id} 段后不为 216 twip"


# ===========================================================================
# 6. 表格三线（顶/底 sz=8 / 表头下 sz=4）
# ===========================================================================


def test_assert_6_table_three_line_borders(exported_docx_bytes):
    """断言 6：document.xml 至少一个表格 tblBorders 含 top/bottom sz=8 + 表头 cell tcBorders.bottom sz=4."""
    doc = _extract_xml(exported_docx_bytes, "word/document.xml")
    # 顶/底 1pt
    assert 'w:val="single" w:sz="8"' in doc, "三线表顶/底 1pt 边框缺失"
    # 表头行 cell 下边框 1/2pt
    assert 'w:val="single" w:sz="4"' in doc, "三线表表头 cell 下边框 1/2pt 缺失"


# ===========================================================================
# 7. 表格内行高（trHeight 397 twip = 0.7cm exact）
# ===========================================================================


def test_assert_7_row_height_397_twip_exact(exported_docx_bytes):
    """断言 7：document.xml 表格行含 trHeight val=397 hRule=exact."""
    doc = _extract_xml(exported_docx_bytes, "word/document.xml")
    assert 'w:val="397"' in doc and 'w:hRule="exact"' in doc, (
        "表格行高未固定为 397 twip exact"
    )


# ===========================================================================
# 8. 空值/零值留白（fmt_amount_gt 行为）
# ===========================================================================


def test_assert_8_fmt_amount_gt_blank_for_zero_and_none():
    """断言 8：fmt_amount_gt(0) / fmt_amount_gt(None) / fmt_amount_gt("") 均返回 ""."""
    assert fmt_amount_gt(0) == ""
    assert fmt_amount_gt(0.0) == ""
    assert fmt_amount_gt(None) == ""
    assert fmt_amount_gt("") == ""
    # 非零数值正常格式化
    assert fmt_amount_gt(1234.5) == "1,234.50"
    assert fmt_amount_gt(-100) == "-100.00"
    # 非数值 fallback 到原值字符串
    assert fmt_amount_gt("abc") == "abc"


# ===========================================================================
# 9. 标题行不重复（无 w:tblHeader 元素）
# ===========================================================================


def test_assert_9_no_tbl_header_repeat(exported_docx_bytes):
    """断言 9：document.xml 不含 ``<w:tblHeader/>``（标题行不重复）."""
    doc = _extract_xml(exported_docx_bytes, "word/document.xml")
    assert "<w:tblHeader" not in doc, "出现 w:tblHeader 元素，标题行会重复（违反致同规范）"


# ===========================================================================
# 10. 页边距上 3.2 / 下 2.54 / 左 3 / 右 3.18 cm
# ===========================================================================


def test_assert_10_page_margins_twip(exported_docx_bytes):
    """断言 10：document.xml 第一节 pgMar 含致同标准 twip 值."""
    doc = _extract_xml(exported_docx_bytes, "word/document.xml")
    assert 'w:top="1814"' in doc, "页边距 top 不为 1814 twip (3.2cm)"
    assert 'w:bottom="1440"' in doc, "页边距 bottom 不为 1440 twip (2.54cm)"
    assert 'w:left="1701"' in doc, "页边距 left 不为 1701 twip (3cm)"
    assert 'w:right="1803"' in doc, "页边距 right 不为 1803 twip (3.18cm)"


# ===========================================================================
# 11. docx 样式 grep 含 GTNote* 前缀 ≥ 7 个
# ===========================================================================


def test_assert_11_gt_note_styles_count_at_least_7(exported_docx_bytes):
    """断言 11：styles.xml 含 ≥ 7 个 GTNote* 命名空间样式（6 段落 + 1 字符；表格可选）."""
    styles = _extract_xml(exported_docx_bytes, "word/styles.xml")
    expected = [
        "GTNoteHeading1",
        "GTNoteHeading2",
        "GTNoteHeading3",
        "GTNoteBody",
        "GTNoteAfterTable",
        "GTNoteUnit",
        "GTNoteNumberRun",
    ]
    found = [name for name in expected if f'w:styleId="{name}"' in styles]
    assert len(found) >= 7, (
        f"GTNote* 样式数 {len(found)} < 7（必需 6 段落 + 1 字符），实际命中: {found}"
    )
    # 表格样式可选但本 spec 模板带，断言一下
    assert 'w:styleId="GTNoteThreeLine"' in styles


# ===========================================================================
# 附加：fill_multi_header rowspan/colspan grid 二阶段填充正确性单测
# ===========================================================================


def test_fill_multi_header_basic_grid():
    """fill_multi_header: 简单二级表头无合并，每个 cell 文本写入正确."""
    from docx import Document as _Doc

    doc = _Doc()
    table = doc.add_table(rows=2, cols=4)
    header_rows = [
        [
            {"text": "项目", "rowspan": 2},
            {"text": "本期增加", "colspan": 2},
            {"text": "本期减少", "rowspan": 2},
        ],
        [
            {"text": "购置"},
            {"text": "在建转入"},
        ],
    ]
    fill_multi_header(table, header_rows, total_cols=4)
    # 第一行：项目（合并 2 行）/ 本期增加（合并 2 列）/ 本期减少（合并 2 行）
    assert table.rows[0].cells[0].text == "项目"
    assert "本期增加" in table.rows[0].cells[1].text
    assert table.rows[0].cells[3].text == "本期减少"
    # 第二行：购置 / 在建转入
    assert table.rows[1].cells[1].text == "购置"
    assert table.rows[1].cells[2].text == "在建转入"


def test_apply_gt_row_height_no_tbl_header():
    """apply_gt_row_height 不应注入 w:tblHeader 元素（断言 9 防御）."""
    from docx import Document as _Doc
    from lxml import etree

    doc = _Doc()
    table = doc.add_table(rows=1, cols=3)
    apply_gt_row_height(table.rows[0], cm=0.7)
    tr_xml = etree.tostring(table.rows[0]._tr, encoding="unicode")
    assert 'w:tblHeader' not in tr_xml
    assert 'w:val="397"' in tr_xml
    assert 'w:hRule="exact"' in tr_xml


def test_apply_gt_three_line_borders_only_top_bottom():
    """apply_gt_three_line: 顶/底 sz=8 + 表头 cell 下 sz=4，不含 left/right/insideV."""
    from docx import Document as _Doc
    from lxml import etree

    doc = _Doc()
    table = doc.add_table(rows=2, cols=3)
    apply_gt_three_line(table)
    tbl_xml = etree.tostring(table._tbl, encoding="unicode")
    assert 'w:val="single" w:sz="8"' in tbl_xml  # 顶/底 1pt
    assert 'w:val="single" w:sz="4"' in tbl_xml  # 表头 0.5pt
    # left / right / insideV 应为 nil
    assert 'w:val="nil"' in tbl_xml or 'w:val="none"' in tbl_xml


def test_apply_gt_dual_font_runs_both_ascii_and_east_asia():
    """apply_gt_dual_font: rPr 同时注入 ascii + eastAsia."""
    from docx import Document as _Doc
    from lxml import etree

    doc = _Doc()
    p = doc.add_paragraph()
    run = p.add_run("1,234.56")
    apply_gt_dual_font(run)
    r_xml = etree.tostring(run._r, encoding="unicode")
    assert 'w:ascii="Arial Narrow"' in r_xml
    assert 'w:eastAsia="仿宋_GB2312"' in r_xml

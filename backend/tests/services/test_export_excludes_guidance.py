"""Property 3/4: 判空与导出不含 guidance — note-guidance-text-separation Task 3.2/3.3."""

from __future__ import annotations

from types import SimpleNamespace

import pytest

from app.services.note_word_exporter import NoteWordExporter

BODY_SAMPLE = "本公司采用成本模式计量。"
GUIDANCE_SAMPLE = "（注：应披露公允价值确认依据。）"


def _make_note(**kwargs):
    defaults = {
        "is_deleted": False,
        "status": "draft",
        "is_empty": False,
        "text_content": "",
        "table_data": {},
        "guidance_text": "（注：应披露公允价值确认依据。）",
    }
    defaults.update(kwargs)
    return SimpleNamespace(**defaults)


class _FakeDb:
    pass


def test_note_to_skip_dict_omits_guidance():
    exporter = NoteWordExporter(_FakeDb())  # type: ignore[arg-type]
    note = _make_note(text_content=BODY_SAMPLE)
    skip_dict = exporter._note_to_skip_dict(note)  # noqa: SLF001
    assert "guidance_text" not in skip_dict
    assert skip_dict["text_content"] == BODY_SAMPLE


def test_should_skip_section_with_guidance_only():
    exporter = NoteWordExporter(_FakeDb())  # type: ignore[arg-type]
    note = _make_note(text_content="", table_data={})
    assert exporter._should_skip_section(note) is True  # noqa: SLF001


@pytest.mark.asyncio
async def test_property_4_html_export_excludes_guidance():
    """Feature: note-guidance-text-separation, Property 4: 导出输出不含 guidance (html)"""
    exporter = NoteWordExporter(_FakeDb())  # type: ignore[arg-type]
    note = _make_note(
        text_content=BODY_SAMPLE,
        table_data=None,
    )
    note.note_section = "八、1"
    note.section_title = "货币资金"

    async def _load_notes(*_args, **_kwargs):
        return [note]

    exporter._load_notes = _load_notes  # type: ignore[method-assign]
    html = await exporter.preview_html(__import__("uuid").uuid4(), 2025)
    assert GUIDANCE_SAMPLE not in html
    assert BODY_SAMPLE in html


def test_property_4_fill_section_block_uses_text_content_only():
    """Feature: note-guidance-text-separation, Property 4: template 路径不含 guidance"""
    from docx import Document

    exporter = NoteWordExporter(_FakeDb())  # type: ignore[arg-type]
    note = _make_note(text_content=BODY_SAMPLE)
    note.note_section = "八、1"

    class _Block:
        section_code = "八、1"
        elements = []

    doc = Document()
    p = doc.add_paragraph("{{section:八、1}}")
    block = _Block()
    block.elements = [p._element]
    exporter._fill_section_block(doc, block, note)  # noqa: SLF001
    assert GUIDANCE_SAMPLE not in doc.paragraphs[0].text
    assert BODY_SAMPLE in doc.paragraphs[0].text

"""Tests for Sprint C.4.7 — 多公司基线对比 PDF 工具."""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from app.services.note_baseline_diff_pdf import (
    BaselineComparison,
    compute_baseline_diff_data,
    generate_baseline_comparison_report,
)


def _make_baseline(name: str, sections: list[dict]) -> dict:
    return {"id": name, "name": name, "sections": sections}


class TestComputeBaselineDiff:
    """Test diff computation across baselines."""

    def test_empty_baselines(self):
        result = compute_baseline_diff_data([])
        assert result == []

    def test_single_baseline(self):
        b = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {"headers": ["项目"], "rows": []}},
        ])
        result = compute_baseline_diff_data([b])
        assert len(result) == 1
        assert result[0].baselines_present == ["v1"]
        assert result[0].baselines_absent == []

    def test_section_presence(self):
        b1 = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {}},
            {"section_id": "s2", "section_title": "应收账款", "table_data": {}},
        ])
        b2 = _make_baseline("v2", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {}},
            # no s2
            {"section_id": "s3", "section_title": "存货", "table_data": {}},
        ])
        result = compute_baseline_diff_data([b1, b2])
        # 3 unique sections
        assert len(result) == 3
        # s1 present in both
        s1_comp = next(c for c in result if c.section_id == "s1")
        assert "v1" in s1_comp.baselines_present
        assert "v2" in s1_comp.baselines_present
        # s2 only in v1
        s2_comp = next(c for c in result if c.section_id == "s2")
        assert s2_comp.baselines_present == ["v1"]
        assert s2_comp.baselines_absent == ["v2"]
        # s3 only in v2
        s3_comp = next(c for c in result if c.section_id == "s3")
        assert s3_comp.baselines_present == ["v2"]
        assert s3_comp.baselines_absent == ["v1"]

    def test_header_diff_detected(self):
        b1 = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "x", "table_data": {"headers": ["A", "B"], "rows": []}},
        ])
        b2 = _make_baseline("v2", [
            {"section_id": "s1", "section_title": "x", "table_data": {"headers": ["A", "B", "C"], "rows": []}},
        ])
        result = compute_baseline_diff_data([b1, b2])
        s1 = result[0]
        assert any(d["type"] == "headers_differ" for d in s1.field_diffs)

    def test_row_count_diff_detected(self):
        b1 = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "x",
             "table_data": {"headers": ["A"], "rows": [{"cells": [1]}]}},
        ])
        b2 = _make_baseline("v2", [
            {"section_id": "s1", "section_title": "x",
             "table_data": {"headers": ["A"], "rows": [{"cells": [1]}, {"cells": [2]}]}},
        ])
        result = compute_baseline_diff_data([b1, b2])
        s1 = result[0]
        assert any(d["type"] == "row_count_differs" for d in s1.field_diffs)

    def test_no_diff_when_identical(self):
        sections = [
            {"section_id": "s1", "section_title": "x",
             "table_data": {"headers": ["A"], "rows": [{"cells": [1]}]}},
        ]
        b1 = _make_baseline("v1", sections)
        b2 = _make_baseline("v2", sections)
        result = compute_baseline_diff_data([b1, b2])
        assert result[0].field_diffs == []


class TestReportGeneration:
    """Test Word report generation."""

    def test_generate_empty_report(self):
        bytes_data = generate_baseline_comparison_report([])
        assert len(bytes_data) > 0
        # Verify it's a valid Word document
        doc = Document(BytesIO(bytes_data))
        assert doc is not None

    def test_generate_single_baseline_report(self):
        b = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {}},
        ])
        bytes_data = generate_baseline_comparison_report([b])
        doc = Document(BytesIO(bytes_data))
        # Verify title
        title_para = doc.paragraphs[0]
        assert "对比" in title_para.text

    def test_generate_multi_baseline_report(self):
        b1 = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {"headers": ["A"], "rows": []}},
            {"section_id": "s2", "section_title": "应收账款", "table_data": {}},
        ])
        b2 = _make_baseline("v2", [
            {"section_id": "s1", "section_title": "货币资金", "table_data": {"headers": ["A", "B"], "rows": []}},
        ])
        bytes_data = generate_baseline_comparison_report([b1, b2], title="测试报告")
        doc = Document(BytesIO(bytes_data))
        # Should have presence matrix table
        assert len(doc.tables) >= 1
        # First table should have header row + 2 data rows (s1, s2)
        presence_table = doc.tables[0]
        assert len(presence_table.rows) == 3  # header + 2

    def test_report_includes_diff_details(self):
        b1 = _make_baseline("v1", [
            {"section_id": "s1", "section_title": "test",
             "table_data": {"headers": ["A"], "rows": []}},
        ])
        b2 = _make_baseline("v2", [
            {"section_id": "s1", "section_title": "test",
             "table_data": {"headers": ["A", "B"], "rows": []}},
        ])
        bytes_data = generate_baseline_comparison_report([b1, b2])
        doc = Document(BytesIO(bytes_data))
        # Check report mentions "字段级差异"
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "字段级差异" in all_text


class TestBaselineComparisonClass:
    """Test BaselineComparison data class."""

    def test_init_with_minimal_args(self):
        comp = BaselineComparison(
            section_id="s1",
            section_title="x",
            baselines_present=["v1"],
            baselines_absent=[],
        )
        assert comp.section_id == "s1"
        assert comp.field_diffs == []

    def test_init_with_field_diffs(self):
        comp = BaselineComparison(
            section_id="s1",
            section_title="x",
            baselines_present=["v1", "v2"],
            baselines_absent=[],
            field_diffs=[{"type": "x", "description": "y"}],
        )
        assert len(comp.field_diffs) == 1

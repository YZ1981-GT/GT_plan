#!/usr/bin/env python3
"""POC：整理报告正文模板 1.1 简版 — 占位符 + OPT + 删【】说明.

Task 0.0.1

Usage:
    python backend/scripts/prepare_report_body_poc.py --dry-run
    python backend/scripts/prepare_report_body_poc.py --write
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

_BACKEND = Path(__file__).resolve().parent.parent
REPORT_BODY_DIR = _BACKEND / "data" / "audit_report_templates" / "report_body"
POC_FILE = (
    "1.1 模板A-无保留意见审计报告模板（上市公司、三板创新层及公开发债）-简版.docx"
)

# 长模式优先
TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    ("ABC股份有限公司全体股东：", "{{company_full_name}}全体股东："),
    ("ABC股份有限公司", "{{company_full_name}}"),
    ("ABC公司", "{{company_short_name}}"),
    ("致同审字（2026）第110ASXXXX号", "{{report_number}}"),
    ("2025年12月31日", "{{audit_period_end}}"),
    ("2025年度", "{{audit_year}}年度"),
    ("2025年", "{{audit_year}}年"),
    ("二〇二五年度", "{{audit_year}}年度"),
    ("致同会计师事务所（特殊普通合伙）", "{{firm_name}}"),
]

GUIDANCE_LINE_RE = re.compile(r"^【")
INLINE_GUIDANCE_RE = re.compile(r"【[^】]*】")

OPT_BLOCKS: list[tuple[str, str, str]] = [
    # (start_heading_prefix, section_id, description)
    ("三、关键审计事项", "key_audit_matters", "关键审计事项"),
    ("四、其他信息", "other_information", "其他信息"),
]


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _apply_replacements(text: str) -> str:
    out = text
    for old, new in TEXT_REPLACEMENTS:
        out = out.replace(old, new)
    return out


def _strip_inline_guidance(text: str) -> str:
    return INLINE_GUIDANCE_RE.sub("", text).strip()


def _insert_before(ref: Paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    ref._element.addprevious(new_p)
    para = Paragraph(new_p, ref._parent)
    para.add_run(text)


def _insert_after(ref: Paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    ref._element.addnext(new_p)
    para = Paragraph(new_p, ref._parent)
    para.add_run(text)


def _delete_paragraph(para: Paragraph) -> None:
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _find_heading_index(paras: list[Paragraph], prefix: str) -> int | None:
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t.startswith(prefix):
            return i
    return None


def prepare_doc(doc: Document, *, write: bool) -> dict[str, int]:
    stats = {"replaced": 0, "deleted": 0, "stripped_inline": 0, "opt_blocks": 0}

    # 1) 全文替换 + 标记删【】行 / 剥行内【】
    to_delete: list[Paragraph] = []
    for para in list(doc.paragraphs):
        raw = para.text or ""
        if not raw.strip():
            continue
        if GUIDANCE_LINE_RE.match(raw.strip()):
            to_delete.append(para)
            stats["deleted"] += 1
            continue
        new_text = _apply_replacements(raw)
        if INLINE_GUIDANCE_RE.search(new_text):
            cleaned = _strip_inline_guidance(new_text)
            if not cleaned:
                to_delete.append(para)
                stats["deleted"] += 1
                continue
            new_text = cleaned
            stats["stripped_inline"] += 1
        if new_text != raw:
            stats["replaced"] += 1
            if write:
                _set_paragraph_text(para, new_text)

    if write:
        for para in to_delete:
            _delete_paragraph(para)

    # 2) OPT 块（从后往前插入避免索引漂移）
    paras = doc.paragraphs
    for start_prefix, sec_id, desc in reversed(OPT_BLOCKS):
        start_idx = _find_heading_index(list(paras), start_prefix)
        if start_idx is None:
            continue
        # 找下一「X、」大节作为结束
        end_idx = len(paras)
        for j in range(start_idx + 1, len(paras)):
            t = (paras[j].text or "").strip()
            if re.match(r"^[一二三四五六七八九十]+、", t) and not t.startswith(start_prefix):
                end_idx = j
                break
        open_m = f"##OPT:{sec_id}:{desc}##"
        close_m = f"##/OPT:{sec_id}##"
        # 已打标则跳过
        if start_idx > 0 and open_m in (paras[start_idx - 1].text or ""):
            continue
        stats["opt_blocks"] += 1
        if write:
            _insert_before(paras[start_idx], open_m)
            paras = doc.paragraphs
            end_para = paras[min(end_idx, len(paras) - 1)]
            _insert_after(end_para, close_m)

    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare report body POC template")
    parser.add_argument("--write", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--file", default=POC_FILE)
    args = parser.parse_args()
    if not args.write:
        args.dry_run = True

    path = REPORT_BODY_DIR / args.file
    if not path.is_file():
        print(f"Missing: {path}", file=sys.stderr)
        raise SystemExit(1)

    doc = Document(path)
    stats = prepare_doc(doc, write=args.write)
    print(f"file: {path.name}")
    print(f"  replaced={stats['replaced']} deleted={stats['deleted']} "
          f"stripped_inline={stats['stripped_inline']} opt_blocks={stats['opt_blocks']}")
    if args.write:
        doc.save(path)
        print(f"Wrote {path}")
    else:
        print("(dry-run; pass --write to apply)")


if __name__ == "__main__":
    main()

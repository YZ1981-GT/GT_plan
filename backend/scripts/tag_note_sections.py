#!/usr/bin/env python3
"""章节作用域（chapter-scoped）附注 Word 模板 ##SECTION: 自动打标.

解决可行性分析发现的标题重名歧义：账户标题（固定资产/长期股权投资/存货…）
在同一份 docx 中出现两次——一次作 §四/§三「会计政策」标题，一次作 §八/§五
「项目注释」标题。纯标题匹配无法区分（63.6–69% 无歧义）。

本脚本用 seed `section_number` 已编码的章节前缀（`八、22`=固定资产在项目注释章
`八`；`四、固定资产`=同名标题在会计政策章 `四`）将每个 seed 章节**仅在其正确的
Heading-1 章节范围内**匹配，消除跨章节歧义。

机制（沿用 tag_note_section_poc.py）：
- `_find_section_range`：按 heading rank 边界确定 [start, end)
- `_insert_paragraph_before/after` + 从后往前插入避免索引漂移
- `_already_tagged`：幂等，可重复运行

Usage:
    python scripts/tag_note_sections.py --variant soe_standalone --dry-run
    python scripts/tag_note_sections.py --variant soe_standalone --write
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

_BACKEND = Path(__file__).resolve().parent.parent
DATA = _BACKEND / "data"
NOTES_DIR = DATA / "audit_report_templates" / "disclosure_notes"
BACKUP_DIR = NOTES_DIR / "_backup"

VARIANT_SEED: dict[str, str] = {
    "soe_standalone": "note_template_soe.json",
    "soe_consolidated": "note_template_soe.json",
    "listed_standalone": "note_template_listed.json",
    "listed_consolidated": "note_template_listed.json",
}

CN_NUM_CHARS = set("一二三四五六七八九十百零")
# CN numeral → ordinal（用于章节单调性约束）
_CN_ORDINAL = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7,
    "八": 8, "九": 9, "十": 10, "十一": 11, "十二": 12, "十三": 13,
    "十四": 14, "十五": 15, "十六": 16, "十七": 17, "十八": 18,
    "十九": 19, "二十": 20,
}

FUZZY_THRESHOLD = 0.85       # 章节内标题模糊匹配阈值
CHAPTER_MATCH_THRESHOLD = 0.55  # docx Heading-1 → seed 章节标题匹配阈值


# ---------------------------------------------------------------------------
# heading rank helpers
# ---------------------------------------------------------------------------
def _heading_rank(style_name: str) -> int | None:
    """返回标题级别 (1=最高)，非标题段落返回 None."""
    name = style_name or ""
    m = re.match(r"Heading\s*(\d+)", name, re.I)
    if m:
        return int(m.group(1))
    m2 = re.match(r"标题\s*(\d+)", name)
    if m2:
        return int(m2.group(1))
    # 上市模板自定义样式「标题3的样式」= rank 3；国企「附注三级」= rank 3
    if name in ("标题3的样式", "附注三级"):
        return 3
    return None


def _norm_title(text: str) -> str:
    """归一标题：去空白、去【…】提示、去尾部（…）说明、注释≡附注."""
    t = (text or "").strip()
    # 去除【…】提示括号（含内容）
    t = re.sub(r"【[^】]*】", "", t)
    # 去除尾部 （…）/(...) 说明（可能多段，循环剥离）
    while True:
        nt = re.sub(r"[（(][^）)]*[）)]\s*$", "", t)
        if nt == t:
            break
        t = nt
    t = re.sub(r"\s+", "", t)
    # 章节标题口径统一：注释 ↔ 附注
    t = t.replace("注释", "附注")
    return t


def _ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()


# ---------------------------------------------------------------------------
# seed loading
# ---------------------------------------------------------------------------
def _variant_scope(variant_key: str) -> str:
    return "consolidated" if variant_key.endswith("_consolidated") else "standalone"


def _section_applies(scope_field: str, variant_scope: str) -> bool:
    """consolidated_only 仅适用于 consolidated 变体."""
    sf = (scope_field or "both").strip()
    if sf in ("both", ""):
        return True
    if sf == "consolidated_only":
        return variant_scope == "consolidated"
    if sf == "standalone_only":
        return variant_scope == "standalone"
    return True


def load_seed(seed_file: str) -> dict[str, Any]:
    return json.loads((DATA / seed_file).read_text(encoding="utf-8"))


def build_chapter_map(seed: dict[str, Any]) -> dict[str, str]:
    """bare-numeral 章节条目 → {cn_numeral: chapter_title}."""
    out: dict[str, str] = {}
    for s in seed.get("sections", []):
        num = (s.get("section_number") or "").strip()
        if num and all(ch in CN_NUM_CHARS for ch in num):
            out[num] = s.get("section_title", "")
    return out


def leaf_sections(seed: dict[str, Any], variant_scope: str) -> list[dict[str, Any]]:
    """返回需打标的 leaf 章节（section_number 含「、」），按 scope 过滤."""
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for s in seed.get("sections", []):
        num = (s.get("section_number") or "").strip()
        if "、" not in num:
            continue
        if num in seen:  # 去重（seed 偶有重复 section_number）
            continue
        if not _section_applies(s.get("scope", "both"), variant_scope):
            continue
        seen.add(num)
        out.append(s)
    return out


def _is_numbered_leaf(section_number: str) -> bool:
    """形如「八、22」的项目注释账户节（章节前缀 + 阿拉伯数字）."""
    return bool(re.match(r"^[一二三四五六七八九十]+、\d+$", section_number.strip()))


# ---------------------------------------------------------------------------
# chapter anchoring：docx Heading-1 → seed 章节
# ---------------------------------------------------------------------------
@dataclass
class ChapterAnchor:
    cn_numeral: str
    chapter_title: str
    para_start: int   # Heading-1 段落索引
    para_end: int     # 下一个被识别章节的 Heading-1 索引（exclusive）
    docx_text: str
    method: str


def anchor_chapters(doc: Document, chapter_map: dict[str, str]) -> tuple[list[ChapterAnchor], list[str]]:
    """遍历 docx Heading-1，匹配到 seed 章节标题，强制单调递增.

    返回 (anchors, warnings)。anchors 按文档顺序。
    """
    warnings: list[str] = []
    # 收集所有 Heading-1 段落索引
    h1_indices: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        if _heading_rank(style) == 1:
            h1_indices.append(i)

    chapters_norm = {cn: _norm_title(title) for cn, title in chapter_map.items()}

    anchors: list[ChapterAnchor] = []
    last_ordinal = 0
    for idx in h1_indices:
        text = (doc.paragraphs[idx].text or "").strip()
        norm = _norm_title(text)
        if not norm:
            continue
        # 分层匹配：exact > loose(互含) > fuzzy。
        # 含包关系（loose）必须优先于 fuzzy，避免「财务报表项目附注」(单体省略「合并」)
        # 被 fuzzy 误配到「母公司财务报表主要项目注释」(十六)，而非正确的「合并…」(五)。
        eligible = [
            (cn, cnorm) for cn, cnorm in chapters_norm.items()
            if _CN_ORDINAL.get(cn, 999) > last_ordinal  # 单调约束
        ]
        best_cn: str | None = None
        best_score = 0.0
        best_method = ""
        # tier 1: exact
        for cn, cnorm in eligible:
            if norm == cnorm:
                best_cn, best_score, best_method = cn, 1.0, "exact"
                break
        # tier 2: loose containment（取重叠占比最高者）
        if best_cn is None:
            for cn, cnorm in eligible:
                if cnorm and (norm in cnorm or cnorm in norm):
                    shorter = min(len(norm), len(cnorm))
                    longer = max(len(norm), len(cnorm))
                    score = (shorter / longer) * 0.95 if longer else 0.0
                    if score > best_score:
                        best_cn, best_score, best_method = cn, score, "loose"
        # tier 3: fuzzy
        if best_cn is None:
            for cn, cnorm in eligible:
                score = _ratio(norm, cnorm)
                if score > best_score:
                    best_cn, best_score, best_method = cn, score, "fuzzy"
        if best_cn is not None and best_score >= CHAPTER_MATCH_THRESHOLD:
            anchors.append(
                ChapterAnchor(
                    cn_numeral=best_cn,
                    chapter_title=chapter_map[best_cn],
                    para_start=idx,
                    para_end=-1,  # 稍后填充
                    docx_text=text,
                    method=f"{best_method}({best_score:.2f})",
                )
            )
            last_ordinal = _CN_ORDINAL.get(best_cn, last_ordinal)
        else:
            warnings.append(
                f"Heading-1 [{idx}] {text!r} 未匹配到任何 seed 章节 (best={best_score:.2f})"
            )

    # 填充 para_end：下一 anchor 的 para_start，最后一个到文末
    n_paras = len(doc.paragraphs)
    for k, a in enumerate(anchors):
        a.para_end = anchors[k + 1].para_start if k + 1 < len(anchors) else n_paras

    return anchors, warnings


# ---------------------------------------------------------------------------
# within-chapter title matching
# ---------------------------------------------------------------------------
@dataclass
class MatchResult:
    section_code: str
    section_title: str
    chapter: str
    method: str
    start: int
    end: int
    rank: int
    candidate_indices: list[int] = field(default_factory=list)


def _candidate_headings(doc: Document, lo: int, hi: int) -> list[tuple[int, int, str]]:
    """返回 [lo, hi) 内的标题段落 (idx, rank, normalized_text)."""
    out: list[tuple[int, int, str]] = []
    for i in range(lo, min(hi, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        style = p.style.name if p.style else ""
        rank = _heading_rank(style)
        if rank is not None and rank >= 2:  # 章节标题为 H1，账户/子节为 H2+
            out.append((i, rank, _norm_title(p.text or "")))
    return out


def _match_within_chapter(
    doc: Document,
    seed_title: str,
    anchor: ChapterAnchor,
) -> tuple[int | None, str, int, list[int]]:
    """在章节范围内按 exact→loose→fuzzy 匹配标题，返回 (start_idx, method, rank, ambiguous_candidates).

    若同一最优 tier + 最小 rank 仍有多个候选 → 歧义，start_idx=None 且返回候选列表。
    """
    target = _norm_title(seed_title)
    cands = _candidate_headings(doc, anchor.para_start + 1, anchor.para_end)

    def _pick(matched: list[tuple[int, int, str]], method: str):
        if not matched:
            return None
        min_rank = min(r for _, r, _ in matched)
        top = [(i, r) for i, r, _ in matched if r == min_rank]
        if len(top) == 1:
            return (top[0][0], method, min_rank, [])
        # 同 rank 多候选 → 歧义
        return (None, method, min_rank, [i for i, _ in top])

    # exact
    exact = [(i, r, t) for i, r, t in cands if t == target]
    res = _pick(exact, "exact")
    if res is not None:
        return res
    # loose: 互为包含（长度比 >= 0.6 防止过短误匹配）
    loose = []
    for i, r, t in cands:
        if not t or not target:
            continue
        if (target in t or t in target):
            shorter, longer = min(len(t), len(target)), max(len(t), len(target))
            if longer and shorter / longer >= 0.6:
                loose.append((i, r, t))
    res = _pick(loose, "loose")
    if res is not None:
        return res
    # fuzzy
    fuzzy = [(i, r, t) for i, r, t in cands if _ratio(t, target) >= FUZZY_THRESHOLD]
    res = _pick(fuzzy, "fuzzy")
    if res is not None:
        return res
    return (None, "unmatched", -1, [])


def _find_section_end(doc: Document, start: int, rank: int, chapter_end: int) -> int:
    """从 start+1 起，遇到 rank<=当前 的标题即为 end；不超过章节边界."""
    for i in range(start + 1, min(chapter_end, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        hr = _heading_rank(p.style.name if p.style else "")
        if hr is not None and hr <= rank:
            return i
    return min(chapter_end, len(doc.paragraphs))


# ---------------------------------------------------------------------------
# tagging core
# ---------------------------------------------------------------------------
def _insert_paragraph_before(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addprevious(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


def _insert_paragraph_after(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


def _already_tagged_codes(doc: Document) -> set[str]:
    codes: set[str] = set()
    open_re = re.compile(r"^##SECTION:([^#]+)##\s*$")
    for p in doc.paragraphs:
        m = open_re.match((p.text or "").strip())
        if m:
            codes.add(m.group(1).strip())
    return codes


@dataclass
class VariantPlan:
    variant: str
    matched: list[MatchResult]
    ambiguous: list[dict[str, Any]]
    unmatched: list[dict[str, Any]]
    skipped_tagged: list[str]
    chapter_anchors: list[ChapterAnchor]
    chapter_warnings: list[str]
    total_leaf: int
    numbered_total: int
    numbered_matched: int
    numbered_ambiguous: int


def plan_variant(variant_key: str) -> VariantPlan:
    seed_file = VARIANT_SEED[variant_key]
    seed = load_seed(seed_file)
    variant_scope = _variant_scope(variant_key)
    chapter_map = build_chapter_map(seed)

    docx_path = NOTES_DIR / f"{variant_key}.docx"
    doc = Document(docx_path)

    anchors, ch_warnings = anchor_chapters(doc, chapter_map)
    anchor_by_cn = {a.cn_numeral: a for a in anchors}
    already = _already_tagged_codes(doc)

    leaves = leaf_sections(seed, variant_scope)
    matched: list[MatchResult] = []
    ambiguous: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    skipped: list[str] = []

    numbered_total = numbered_matched = numbered_ambiguous = 0

    for s in leaves:
        code = s["section_number"].strip()
        title = s.get("section_title", "")
        prefix = code.split("、", 1)[0]
        is_numbered = _is_numbered_leaf(code)
        if is_numbered:
            numbered_total += 1

        if code in already:
            skipped.append(code)
            if is_numbered:
                numbered_matched += 1  # 已打标视为已解决
            continue

        anchor = anchor_by_cn.get(prefix)
        if anchor is None:
            unmatched.append({"code": code, "title": title,
                              "reason": f"章节 {prefix} 在本变体 docx 中无 Heading-1（可能为单体版省略章 / 未识别）"})
            continue

        start, method, rank, amb_cands = _match_within_chapter(doc, title, anchor)
        if start is None and amb_cands:
            ambiguous.append({"code": code, "title": title, "chapter": prefix,
                              "method": method, "candidate_indices": amb_cands})
            if is_numbered:
                numbered_ambiguous += 1
            continue
        if start is None:
            unmatched.append({"code": code, "title": title,
                              "reason": f"章节 {prefix} 范围 [{anchor.para_start},{anchor.para_end}) 内无标题匹配 {title!r}"})
            continue

        end = _find_section_end(doc, start, rank, anchor.para_end)
        matched.append(MatchResult(code, title, prefix, method, start, end, rank))
        if is_numbered:
            numbered_matched += 1

    return VariantPlan(
        variant=variant_key,
        matched=matched,
        ambiguous=ambiguous,
        unmatched=unmatched,
        skipped_tagged=skipped,
        chapter_anchors=anchors,
        chapter_warnings=ch_warnings,
        total_leaf=len(leaves),
        numbered_total=numbered_total,
        numbered_matched=numbered_matched,
        numbered_ambiguous=numbered_ambiguous,
    )


def print_plan(plan: VariantPlan, *, verbose: bool = True) -> None:
    print(f"\n{'='*70}")
    print(f"变体: {plan.variant}  (scope={_variant_scope(plan.variant)})")
    print(f"{'='*70}")
    if not plan.chapter_anchors:
        print("  🛑 章节锚定失败：未能将任何 Heading-1 映射到 seed 章节！")
    else:
        print(f"章节锚定 ({len(plan.chapter_anchors)} 章):")
        for a in plan.chapter_anchors:
            print(f"  {a.cn_numeral:>3} {a.chapter_title!r} ← docx[{a.para_start}] {a.docx_text!r} "
                  f"[{a.para_start},{a.para_end}) {a.method}")
    if plan.chapter_warnings:
        for w in plan.chapter_warnings:
            print(f"  ⚠ {w}")

    if verbose:
        print(f"\n打标计划 (matched={len(plan.matched)}):")
        for m in sorted(plan.matched, key=lambda x: x.start):
            print(f"  PLAN {m.section_code:<14} paras [{m.start},{m.end}) "
                  f"章{m.chapter} rank{m.rank} {m.method} title={m.section_title!r}")

    if plan.ambiguous:
        print(f"\n歧义 (ambiguous={len(plan.ambiguous)}) — 不打标:")
        for a in plan.ambiguous:
            print(f"  AMBIG {a['code']:<14} 章{a['chapter']} {a['method']} "
                  f"candidates={a['candidate_indices']} title={a['title']!r}")

    if plan.unmatched:
        print(f"\n未匹配 (unmatched={len(plan.unmatched)}):")
        for u in plan.unmatched:
            print(f"  MISS  {u['code']:<14} {u['title']!r} :: {u['reason']}")

    if plan.skipped_tagged:
        print(f"\n已打标跳过 (idempotent skip={len(plan.skipped_tagged)}): {plan.skipped_tagged}")

    # 汇总
    leaf_n = plan.total_leaf
    unamb = len(plan.matched) + len(plan.skipped_tagged)
    print(f"\n--- 汇总 ({plan.variant}) ---")
    print(f"  seed leaf 章节总数        : {leaf_n}")
    print(f"  无歧义已匹配(含已打标)    : {unamb}  ({unamb/leaf_n*100:.1f}%)" if leaf_n else "  n/a")
    print(f"  歧义残留                  : {len(plan.ambiguous)}")
    print(f"  未匹配                    : {len(plan.unmatched)}")
    nt = plan.numbered_total
    print(f"  --- GATE 指标 (项目注释数字节 形如「{plan.chapter_anchors and ''}八、22」) ---")
    if nt:
        rate = plan.numbered_matched / nt * 100
        print(f"  数字节总数                : {nt}")
        print(f"  数字节无歧义匹配          : {plan.numbered_matched}  ({rate:.1f}%)")
        print(f"  数字节章内歧义            : {plan.numbered_ambiguous}")
    else:
        print("  数字节总数: 0")


def gate_pass(plan: VariantPlan) -> tuple[bool, str]:
    """GATE：数字节无歧义匹配率 ≥95% 且 章内零歧义."""
    if not plan.chapter_anchors:
        return False, "章节锚定失败"
    if plan.numbered_total == 0:
        return False, "无数字节可评估"
    rate = plan.numbered_matched / plan.numbered_total * 100
    if plan.numbered_ambiguous > 0:
        return False, f"存在 {plan.numbered_ambiguous} 个章内歧义数字节"
    if rate < 95.0:
        return False, f"数字节匹配率 {rate:.1f}% < 95%"
    return True, f"数字节匹配率 {rate:.1f}%，章内零歧义"


def apply_tags(variant_key: str, plan: VariantPlan) -> int:
    """写入 SECTION 标记（前置已备份）。返回打标节数."""
    docx_path = NOTES_DIR / f"{variant_key}.docx"
    doc = Document(docx_path)

    # 从后往前（按 start 降序）插入避免索引漂移
    for m in sorted(plan.matched, key=lambda x: x.start, reverse=True):
        end_idx = m.end if m.end > m.start else m.start + 1
        end_para = doc.paragraphs[min(end_idx, len(doc.paragraphs)) - 1]
        start_para = doc.paragraphs[m.start]
        _insert_paragraph_after(end_para, f"##/SECTION:{m.section_code}##")
        _insert_paragraph_before(start_para, f"##SECTION:{m.section_code}##")

    doc.save(docx_path)
    return len(plan.matched)


def count_section_markers(variant_key: str) -> int:
    docx_path = NOTES_DIR / f"{variant_key}.docx"
    if not docx_path.is_file():
        return 0
    doc = Document(docx_path)
    open_re = re.compile(r"^##SECTION:([^#]+)##\s*$")
    return sum(1 for p in doc.paragraphs if open_re.match((p.text or "").strip()))


def backup_docx(variant_key: str) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    src = NOTES_DIR / f"{variant_key}.docx"
    dst = BACKUP_DIR / f"{variant_key}.docx"
    shutil.copy2(src, dst)
    return dst


def main() -> None:
    parser = argparse.ArgumentParser(description="Chapter-scoped ##SECTION: tagger for disclosure notes")
    parser.add_argument("--variant", required=True, choices=list(VARIANT_SEED.keys()))
    parser.add_argument("--write", action="store_true", help="应用打标（先备份）")
    parser.add_argument("--dry-run", action="store_true", help="仅打印计划（默认）")
    parser.add_argument("--quiet", action="store_true", help="精简输出（仅汇总）")
    args = parser.parse_args()
    if not args.write:
        args.dry_run = True

    docx_path = NOTES_DIR / f"{args.variant}.docx"
    if not docx_path.is_file():
        print(f"Missing docx: {docx_path}")
        raise SystemExit(1)

    plan = plan_variant(args.variant)
    print_plan(plan, verbose=not args.quiet)

    ok, reason = gate_pass(plan)
    print(f"\n  GATE: {'✅ PASS' if ok else '🛑 BLOCK'} — {reason}")

    if args.dry_run:
        print("\n(dry-run；通过 --write 应用)")
        return

    # write path
    if not plan.chapter_anchors:
        print("\n🛑 章节锚定失败，拒绝写入（不回退朴素标题匹配）。")
        raise SystemExit(2)
    if not ok:
        print(f"\n🛑 GATE 未通过（{reason}），拒绝写入。请人工处理。")
        raise SystemExit(3)

    before = count_section_markers(args.variant)
    bkp = backup_docx(args.variant)
    print(f"\n已备份 → {bkp}")
    n = apply_tags(args.variant, plan)
    after = count_section_markers(args.variant)
    print(f"已写入 {docx_path}：打标 {n} 节；##SECTION: 计数 {before} → {after}")


if __name__ == "__main__":
    main()

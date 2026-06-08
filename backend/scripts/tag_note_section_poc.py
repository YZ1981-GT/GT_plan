#!/usr/bin/env python3
"""POC：在附注 Word 模板中插入 ##SECTION:code## 块标记.

Task 0.0.2a — soe_standalone 必做节：一、1 / 二、1 / 八、1

Usage:
    python backend/scripts/tag_note_section_poc.py --variant soe_standalone --dry-run
    python backend/scripts/tag_note_section_poc.py --variant soe_standalone --write
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

_BACKEND = Path(__file__).resolve().parent.parent
DATA = _BACKEND / "data"
NOTES_DIR = DATA / "audit_report_templates" / "disclosure_notes"

# POC 章节：section_code → 在 Word 中定位用的 section_title
POC_TARGETS: dict[str, list[dict[str, str]]] = {
    "soe_standalone": [
        # Word 标题为「公司（企业）基本情况」，JSON section_title 为「公司基本情况」
        {"section_code": "一、1", "section_title": "公司（企业）基本情况", "heading_level": "1"},
        {"section_code": "二、1", "section_title": "财务报表编制基础", "heading_level": "1"},
        {"section_code": "八、1", "section_title": "货币资金", "heading_level": "3"},
    ],
    "soe_consolidated": [
        {"section_code": "一、1", "section_title": "公司（企业）基本情况", "heading_level": "1"},
        {"section_code": "七、本期纳入合并报表", "section_title": "本期纳入合并报表范围的子公司基本情况", "heading_level": "2"},
        {"section_code": "八、1", "section_title": "货币资金", "heading_level": "3"},
    ],
    "listed_standalone": [
        # 上市版 Word 标题无「（企业）」；货币资金为 Heading 2（国企为 Heading 3）
        {"section_code": "一、1", "section_title": "公司基本情况", "heading_level": "1"},
        {"section_code": "五、1", "section_title": "货币资金", "heading_level": "2"},
    ],
    "listed_consolidated": [
        {"section_code": "一、1", "section_title": "公司基本情况", "heading_level": "1"},
        {"section_code": "五、1", "section_title": "货币资金", "heading_level": "2"},
        # JSON section_title=母公司…；Word 实际为「公司财务报表主要项目注释」
        {"section_code": "十六", "section_title": "公司财务报表主要项目注释", "heading_level": "1"},
    ],
}

HEADING_STYLE = {
    "1": "Heading 1",
    "2": "Heading 2",
    "3": "Heading 3",
    "4": "Heading 4",
}


def _heading_rank(style_name: str) -> int | None:
    m = re.match(r"Heading\s*(\d+)", style_name or "", re.I)
    return int(m.group(1)) if m else None


def _insert_paragraph_before(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addprevious(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


def _insert_paragraph_after(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


def _find_section_range(
    doc: Document,
    *,
    section_title: str,
    heading_level: str,
) -> tuple[int, int] | None:
    """返回 [start_idx, end_idx) 段落索引（end 为下一同级/更高级标题）."""
    expected_style = HEADING_STYLE.get(heading_level, "Heading 1")
    rank = int(heading_level)
    start: int | None = None
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        style = para.style.name if para.style else ""
        if start is None:
            if section_title in text and style == expected_style:
                start = i
            continue
        # 已找到起点：遇到同级或更高级标题则结束
        hr = _heading_rank(style)
        if hr is not None and hr <= rank:
            return start, i
    if start is not None:
        return start, len(doc.paragraphs)
    return None


def _already_tagged(doc: Document, section_code: str) -> bool:
    open_m = f"##SECTION:{section_code}##"
    close_m = f"##/SECTION:{section_code}##"
    texts = [(p.text or "").strip() for p in doc.paragraphs]
    return open_m in texts and close_m in texts


def tag_variant(variant_key: str, *, write: bool = False, dry_run: bool = False) -> int:
    targets = POC_TARGETS.get(variant_key)
    if not targets:
        print(f"Unknown variant: {variant_key}")
        return 1

    path = NOTES_DIR / f"{variant_key}.docx"
    if not path.is_file():
        print(f"Missing: {path}")
        return 1

    doc = Document(path)
    planned: list[tuple[str, int, int]] = []

    for spec in targets:
        code = spec["section_code"]
        if _already_tagged(doc, code):
            print(f"SKIP {code}: already tagged")
            continue
        rng = _find_section_range(
            doc,
            section_title=spec["section_title"],
            heading_level=spec["heading_level"],
        )
        if rng is None:
            print(f"FAIL {code}: heading not found for title={spec['section_title']!r}")
            return 1
        start, end = rng
        planned.append((code, start, end))
        print(f"PLAN {code}: paras [{start}, {end}) title={spec['section_title']!r}")

    if dry_run or not write:
        print(f"dry-run: {len(planned)} sections would be tagged")
        return 0

    # 从后往前插入，避免索引漂移
    for code, start, end in sorted(planned, key=lambda x: x[1], reverse=True):
        end_para = doc.paragraphs[end - 1] if end > start else doc.paragraphs[start]
        start_para = doc.paragraphs[start]
        _insert_paragraph_after(end_para, f"##/SECTION:{code}##")
        _insert_paragraph_before(start_para, f"##SECTION:{code}##")

    doc.save(path)
    print(f"Wrote {path} ({len(planned)} sections tagged)")
    return 0


def main() -> None:
    parser = argparse.ArgumentParser(description="POC tag ##SECTION: in disclosure note docx")
    parser.add_argument("--variant", required=True, help="e.g. soe_standalone")
    parser.add_argument("--write", action="store_true", help="Apply changes to docx")
    parser.add_argument("--dry-run", action="store_true", help="Only print plan")
    args = parser.parse_args()
    if not args.write and not args.dry_run:
        args.dry_run = True
    raise SystemExit(tag_variant(args.variant, write=args.write, dry_run=args.dry_run))


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""CHAPTER-SCOPED 全量附注 SECTION 打标（Task 0.6.2）.

在四套 disclosure_notes/*.docx 中，为每个种子章节插入
`##SECTION:code## … ##/SECTION:code##` 块标记。

与朴素「纯标题匹配」(0.6.2 实证 NO-GO，63.6~69% 无歧义) 不同，本脚本
按 **章节范围 (chapter-scoped)** 消歧：账户标题（固定资产/长期股权投资/…）
在每份文档出现两次——一次作 §四/§三「会计政策」标题、一次作 §八/§五
「项目注释」标题；种子 section_number 的章节前缀（`、` 之前部分）编码了章节
归属，故只在该章节范围内做标题匹配即可消歧。

章节对齐策略（鲁棒）：
  - Word Heading-1 序列与种子章节前缀**不能**简单按序号对齐（standalone
    省略合并相关章节；H1 标题存在漂移，如「项目注释」vs「项目附注」）。
  - 改用 **保序 DP 对齐**：以 H1 标题键匹配（去括号归一）为主分，子节标题
    重叠数为辅分，单调递增对齐，避免错位。
  - 对齐不确定（得分过低）的章节，其全部种子节标记为 AMBIGUOUS，不打标。

安全不变量（CRITICAL — 专有交付模板）：
  - 仅当某种子节在其章节范围内有**恰好一个**归一标题匹配时才打标；
    0 个或 >1 个 → 留 UNMATCHED/AMBIGUOUS（报告，不打标）。
  - 重复 H3/H4（如 八 章节内「固定资产」H3+H4 相邻）用最小标题级别消歧
    （锚定 H3，唯一）。
  - --write 前先备份原 docx 到 disclosure_notes/_backup/{variant}.docx.bak。
  - 仅 INSERT 两个 marker 段落；绝不删除/修改既有内容段落或 【…】 提示。
  - 幂等：已打标的 code 跳过。

Usage (cwd=backend, Windows):
    ..\\.venv\\Scripts\\python.exe scripts/tag_note_sections_full.py --dry-run
    ..\\.venv\\Scripts\\python.exe scripts/tag_note_sections_full.py --write
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from build_section_code_index import (  # type: ignore
    DATA,
    NOTES_DIR,
    VARIANT_SEED,
    load_json_sections,
)

OUT_REPORT = DATA / "audit_report_templates" / "_note_full_tagging_report.json"
BACKUP_DIR = NOTES_DIR / "_backup"

# 章节对齐：H1 标题键匹配权重；DP 绑定的最小得分阈值
TITLE_SCORE_W = 5
MIN_BIND_SCORE = 2  # 至少标题部分匹配(2) 或 子节重叠>=2


# ----------------------------- 归一 -----------------------------

_LEADING_NUM_RE = re.compile(r"^[（(]?[0-9一二三四五六七八九十]+[）)、.\．]\s*")
_PAREN_RE = re.compile(r"[（(][^）)]*[）)]")
_BRACKET_HINT_RE = re.compile(r"【[^】]*】")
# 末尾编辑性括注（如「（损失以…号填列）」「（不适用的删除）」「（注：…）」），可重复
_TRAILING_PAREN_RE = re.compile(r"[（(][^）)]*[）)]\s*$")


def _norm_row(text: str) -> str:
    s = re.sub(r"\s+", "", (text or ""))
    return re.sub(r"^[△▲#*]+", "", s)


def _norm_loose(text: str) -> str:
    """宽松标题归一：去空白/标记 + 去 【…】 提示 + 去行内编号前缀 + 去末尾编辑括注.

    种子标题为干净账户名（如「投资性房地产」），而 Word 标题常带编辑性后缀
    （「投资性房地产【不适用的删除】」「资产减值损失（损失以…号填列）」）；
    两侧统一归一后才能在章节范围内精确匹配。
    """
    s = _norm_row(text)
    s = _BRACKET_HINT_RE.sub("", s)  # 去除所有 【…】 提示
    s = _LEADING_NUM_RE.sub("", s)
    # 反复剥离末尾编辑性括注（可能叠加两组）
    prev = None
    while prev != s:
        prev = s
        s = _TRAILING_PAREN_RE.sub("", s)
    return s


def _norm_chapter_key(text: str) -> str:
    """章节标题键：宽松归一基础上再去掉所有括号内容（消「（企业）」漂移）."""
    s = _norm_loose(text)
    s = _PAREN_RE.sub("", s)
    return s


def _heading_rank(style_name: str) -> Optional[int]:
    m = re.match(r"Heading\s*(\d+)", style_name or "", re.I)
    if m:
        return int(m.group(1))
    m2 = re.match(r"标题\s*(\d+)", style_name or "")
    if m2:
        return int(m2.group(1))
    return None


# ----------------------------- docx 段落插入 -----------------------------


def _insert_paragraph_before(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addprevious(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


def _insert_paragraph_after(ref_para: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    new_para.add_run(text)
    return new_para


# ----------------------------- 种子解析 -----------------------------


def _chapter_prefix(section_number: str) -> str:
    sn = (section_number or "").strip()
    return sn.split("、")[0] if "、" in sn else sn


def parse_seed_chapters(seeds: list[dict[str, Any]]) -> tuple[list[str], dict[str, dict[str, Any]]]:
    """返回 (有序章节前缀列表, {prefix -> {title, children:[(code,loose_title)]}}).

    children 仅含二级（含 `、`）种子节；章节级条目（无 `、`）提供 title。
    """
    order: list[str] = []
    bucket: dict[str, dict[str, Any]] = {}
    for s in seeds:
        sn = (s.get("section_number") or "").strip()
        if not sn:
            continue
        prefix = _chapter_prefix(sn)
        if prefix not in bucket:
            bucket[prefix] = {"title": None, "children": []}
            order.append(prefix)
        if "、" not in sn:
            bucket[prefix]["title"] = s.get("section_title")
        else:
            bucket[prefix]["children"].append(
                (sn, _norm_loose(s.get("section_title") or ""))
            )
    return order, bucket


# ----------------------------- Word 章节 -----------------------------


def collect_word_chapters(doc: Document) -> list[dict[str, Any]]:
    """按文档顺序返回 Heading-1 章节 [{idx, end, title_key, headings}].

    headings: 该章节范围 [idx, end) 内所有标题段落 [{idx, rank, loose}].
    end = 下一 Heading-1 的 idx（或文档末尾）。
    """
    paras = doc.paragraphs
    h1_idx: list[int] = []
    for i, p in enumerate(paras):
        if _heading_rank(p.style.name if p.style else "") == 1:
            h1_idx.append(i)

    chapters: list[dict[str, Any]] = []
    for k, start in enumerate(h1_idx):
        end = h1_idx[k + 1] if k + 1 < len(h1_idx) else len(paras)
        headings: list[dict[str, Any]] = []
        for j in range(start, end):
            p = paras[j]
            rk = _heading_rank(p.style.name if p.style else "")
            if rk is not None:
                headings.append({"idx": j, "rank": rk, "loose": _norm_loose(p.text)})
        chapters.append(
            {
                "start": start,
                "end": end,
                "title_key": _norm_chapter_key(paras[start].text),
                "headings": headings,
            }
        )
    return chapters


def _pair_score(seed_ch: dict[str, Any], word_ch: dict[str, Any]) -> int:
    """种子章节与 Word 章节的匹配得分：H1 标题键匹配 + 子节标题重叠数."""
    score = 0
    seed_title_key = _norm_chapter_key(seed_ch["title"] or "")
    if seed_title_key and seed_title_key == word_ch["title_key"]:
        score += TITLE_SCORE_W
    elif seed_title_key and word_ch["title_key"] and (
        seed_title_key in word_ch["title_key"] or word_ch["title_key"] in seed_title_key
    ):
        score += TITLE_SCORE_W - 2
    word_titles = {h["loose"] for h in word_ch["headings"]}
    child_titles = {ct for _, ct in seed_ch["children"]}
    score += len(child_titles & word_titles)
    return score


def align_chapters(
    seed_order: list[str],
    seed_bucket: dict[str, dict[str, Any]],
    word_chapters: list[dict[str, Any]],
) -> dict[str, Optional[int]]:
    """保序 DP：把每个种子章节前缀映射到一个 Word 章节索引（或 None）.

    单调性约束：种子章节顺序 i<j → 映射的 Word 章节索引也递增。
    目标：最大化总匹配得分。低于 MIN_BIND_SCORE 的绑定置为 None（不确定）。
    """
    S = [{"prefix": p, "title": seed_bucket[p]["title"], "children": seed_bucket[p]["children"]}
         for p in seed_order]
    n, m = len(S), len(word_chapters)
    NEG = -10**9
    # dp[i][j] = 用前 i 个种子章节、Word 章节指针在 j（下一可用 Word 章节 >= j）的最大得分
    dp = [[NEG] * (m + 1) for _ in range(n + 1)]
    choice = [[None] * (m + 1) for _ in range(n + 1)]  # type: ignore
    for j in range(m + 1):
        dp[n][j] = 0
    for i in range(n - 1, -1, -1):
        for j in range(m, -1, -1):
            # 选项 A：种子章节 i 不绑定任何 Word 章节
            best = dp[i + 1][j]
            best_choice: Optional[int] = None
            # 选项 B：把种子章节 i 绑定到某个 Word 章节 w (>= j)
            for w in range(j, m):
                sc = _pair_score(S[i], word_chapters[w])
                if sc <= 0:
                    continue
                cand = sc + dp[i + 1][w + 1]
                if cand > best:
                    best = cand
                    best_choice = w
            dp[i][j] = best
            choice[i][j] = best_choice

    mapping: dict[str, Optional[int]] = {}
    j = 0
    for i in range(n):
        w = choice[i][j]
        if w is None:
            mapping[S[i]["prefix"]] = None
        else:
            sc = _pair_score(S[i], word_chapters[w])
            mapping[S[i]["prefix"]] = w if sc >= MIN_BIND_SCORE else None
            j = w + 1
    return mapping


# ----------------------------- 章节内定位 -----------------------------


def _unique_min_rank(cands: list[dict[str, Any]]) -> Optional[dict[str, Any]]:
    """重复标题消歧：取最小标题级别（如 H3 优于其后紧跟的重复 H4）.

    若最小级别唯一 → 返回该标题；否则 None（仍歧义）。
    """
    if not cands:
        return None
    if len(cands) == 1:
        return cands[0]
    min_rank = min(c["rank"] for c in cands)
    top = [c for c in cands if c["rank"] == min_rank]
    return top[0] if len(top) == 1 else None


def match_section_in_chapter(
    target_loose: str,
    target_title: str,
    chapter: dict[str, Any],
) -> tuple[Optional[dict[str, Any]], str]:
    """在章节范围内定位种子节标题，返回 (heading, status).

    status ∈ {matched, unmatched, ambiguous}. 两级匹配：
      1) 宽松归一 (_norm_loose) 相等
      2) 章节键 (_norm_chapter_key, 去括号) 相等（兜底「公司（企业）基本情况」漂移）
    每级都用最小标题级别消歧；>1 同级 → ambiguous。
    """
    headings = chapter["headings"]
    # tier 1: loose 相等
    cands = [h for h in headings if h["loose"] and h["loose"] == target_loose]
    if cands:
        hit = _unique_min_rank(cands)
        return (hit, "matched") if hit else (None, "ambiguous")

    # tier 2: 去括号键 相等
    key = _norm_chapter_key(target_title)
    if key:
        cands2 = [h for h in headings if _norm_chapter_key(h["loose"]) == key]
        if cands2:
            hit = _unique_min_rank(cands2)
            return (hit, "matched") if hit else (None, "ambiguous")

    return (None, "unmatched")


def find_section_end(doc: Document, heading: dict[str, Any], chapter: dict[str, Any]) -> int:
    """节范围结束 idx：章节内、heading 之后首个同级或更高级标题；否则章节末."""
    start = heading["idx"]
    rank = heading["rank"]
    for h in chapter["headings"]:
        if h["idx"] > start and h["rank"] <= rank:
            return h["idx"]
    return chapter["end"]


# ----------------------------- 已打标 -----------------------------


def already_tagged_codes(doc: Document) -> set[str]:
    open_re = re.compile(r"^##SECTION:([^#]+)##\s*$")
    found: set[str] = set()
    for p in doc.paragraphs:
        m = open_re.match((p.text or "").strip())
        if m:
            found.add(m.group(1).strip())
    return found


# ----------------------------- 变体处理 -----------------------------


def plan_variant(variant_key: str) -> dict[str, Any]:
    """计算单变体打标计划（只读）。返回含 plan + 报告字段的 dict."""
    seed_file, _template_type = VARIANT_SEED[variant_key]
    docx_path = NOTES_DIR / f"{variant_key}.docx"
    if not docx_path.is_file():
        raise FileNotFoundError(docx_path)

    by_code = load_json_sections(seed_file)
    seeds = list(by_code.values())
    seed_order, seed_bucket = parse_seed_chapters(seeds)

    doc = Document(docx_path)
    word_chapters = collect_word_chapters(doc)
    mapping = align_chapters(seed_order, seed_bucket, word_chapters)

    tagged_already = already_tagged_codes(doc)

    chapter_alignment = []
    for prefix in seed_order:
        w = mapping.get(prefix)
        chapter_alignment.append(
            {
                "seed_prefix": prefix,
                "seed_title": seed_bucket[prefix]["title"],
                "word_chapter_idx": w,
                "word_chapter_title": (
                    word_chapters[w]["title_key"] if w is not None else None
                ),
                "word_chapter_para": (
                    word_chapters[w]["start"] if w is not None else None
                ),
            }
        )

    plan: list[dict[str, Any]] = []  # {code, start_idx, end_idx, title}
    tagged: list[str] = []
    skipped: list[str] = []
    ambiguous: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    chapter_unaligned: list[dict[str, Any]] = []

    for s in seeds:
        code = (s.get("section_number") or "").strip()
        title = (s.get("section_title") or "").strip()
        if not code or "、" not in code:
            # 章节级条目本身不打标（只对二级节打标）
            continue
        if code in tagged_already:
            skipped.append(code)
            continue
        prefix = _chapter_prefix(code)
        w = mapping.get(prefix)
        if w is None:
            chapter_unaligned.append({"section_number": code, "section_title": title,
                                      "reason": f"chapter {prefix!r} not aligned"})
            continue
        chapter = word_chapters[w]
        heading, status = match_section_in_chapter(_norm_loose(title), title, chapter)
        if status == "matched" and heading is not None:
            start = heading["idx"]
            end = find_section_end(doc, heading, chapter)
            plan.append({"code": code, "start_idx": start, "end_idx": end, "title": title})
            tagged.append(code)
        elif status == "ambiguous":
            ambiguous.append({"section_number": code, "section_title": title,
                              "chapter": prefix})
        else:
            unmatched.append({"section_number": code, "section_title": title,
                              "chapter": prefix})

    # 防嵌套/重叠：种子二级节是逻辑兄弟，但 Word 中个别节可能被嵌在另一节的
    # 更深标题级别下（如「金融资产转移」H3 嵌在「金融工具」H2 内），或某节标题
    # 恰为章节 H1 标题（其 rank-based span 覆盖整章）。将每节 end 裁剪到「下一个
    # 已规划节起点」，保证 marker 扁平、不重叠、不嵌套（绝不越过真实兄弟边界）。
    if plan:
        starts_sorted = sorted(it["start_idx"] for it in plan)
        for item in plan:
            s = item["start_idx"]
            # 找严格大于 s 的最小已规划起点
            nxt = next((x for x in starts_sorted if x > s), None)
            if nxt is not None and nxt < item["end_idx"]:
                item["end_idx"] = nxt

    return {
        "variant_key": variant_key,
        "docx_path": docx_path,
        "doc": doc,
        "plan": plan,
        "tagged": tagged,
        "skipped": skipped,
        "ambiguous": ambiguous,
        "unmatched": unmatched,
        "chapter_unaligned": chapter_unaligned,
        "chapter_alignment": chapter_alignment,
        "seed_section_total": sum(1 for s in seeds if "、" in (s.get("section_number") or "")),
    }


def apply_plan(doc: Document, plan: list[dict[str, Any]]) -> None:
    """按计划插入 marker，保证扁平不交错。

    OPEN 用 addprevious 插到节首段落之前；CLOSE 用 addnext 插到节内**最后一个
    内容段落**(end_idx-1) 之后。两端各锚定不同元素，使相邻节共享边界时
    （A.end == B.start）顺序天然确定（A-close 贴在 end-1 后、B-open 贴在
    B.start 前），与处理顺序无关，避免「先开后闭」交错。
    先快照所有 lxml 元素引用再插入，避免索引漂移。
    """
    paras = doc.paragraphs
    n = len(paras)
    body = paras[0]._parent if paras else None
    marker_re = re.compile(r"^##/?SECTION:[^#]+##\s*$")

    def _last_content_idx(start_idx: int, end_idx: int) -> int:
        """节内最后一个非 marker 段落 idx（跳过既有/已插 SECTION 标记段落）.

        防止 CLOSE 锚到「下一章节的既有 chapter 级 marker」之后造成交错
        （如 listed_consolidated 既有 `##SECTION:十六##` 紧贴章节 H1）。
        """
        j = end_idx - 1
        while j > start_idx and marker_re.match((paras[j].text or "").strip()):
            j -= 1
        return j

    jobs = []
    for item in plan:
        start_idx = item["start_idx"]
        end_idx = item["end_idx"]
        last_idx = max(start_idx, _last_content_idx(start_idx, end_idx))
        start_el = paras[start_idx]._element
        last_el = paras[last_idx]._element if last_idx < n else None
        jobs.append((item["code"], start_el, last_el))

    for code, start_el, last_el in jobs:
        start_para = Paragraph(start_el, body)
        _insert_paragraph_before(start_para, f"##SECTION:{code}##")
        close_text = f"##/SECTION:{code}##"
        if last_el is not None:
            last_para = Paragraph(last_el, body)
            _insert_paragraph_after(last_para, close_text)
        else:
            _insert_paragraph_after(doc.paragraphs[-1], close_text)


def _print_variant(p: dict[str, Any]) -> None:
    vk = p["variant_key"]
    print(f"\n=== {vk} ===")
    print(f"  seed 二级节总数      : {p['seed_section_total']}")
    print(f"  章节对齐:")
    for ca in p["chapter_alignment"]:
        wt = ca["word_chapter_title"]
        wp = ca["word_chapter_para"]
        flag = "" if ca["word_chapter_idx"] is not None else "  <UNALIGNED>"
        print(f"    {ca['seed_prefix']:5} {ca['seed_title']!r} -> "
              f"word[{ca['word_chapter_idx']}]@{wp} {wt!r}{flag}")
    print(f"  PLAN tagged          : {len(p['tagged'])}  {p['tagged']}")
    print(f"  SKIP already         : {len(p['skipped'])}  {p['skipped']}")
    print(f"  AMBIGUOUS            : {len(p['ambiguous'])}  "
          f"{[a['section_number'] for a in p['ambiguous']]}")
    print(f"  UNMATCHED            : {len(p['unmatched'])}  "
          f"{[u['section_number'] for u in p['unmatched']]}")
    print(f"  CHAPTER-UNALIGNED    : {len(p['chapter_unaligned'])}  "
          f"{[c['section_number'] for c in p['chapter_unaligned']]}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Chapter-scoped full note SECTION tagger")
    parser.add_argument("--write", action="store_true", help="Apply (backs up docx first)")
    parser.add_argument("--dry-run", action="store_true", help="Plan only (default)")
    parser.add_argument("--variant", help="Only this variant (default: all 4)")
    args = parser.parse_args()
    if not args.write and not args.dry_run:
        args.dry_run = True

    variants = [args.variant] if args.variant else list(VARIANT_SEED.keys())

    report: dict[str, Any] = {
        "_note": "Task 0.6.2 chapter-scoped full tagging report.",
        "mode": "write" if args.write else "dry-run",
        "variants": {},
    }

    for vk in variants:
        try:
            p = plan_variant(vk)
        except FileNotFoundError as e:
            print(f"SKIP {vk}: missing {e}")
            continue
        _print_variant(p)

        if args.write and p["plan"]:
            BACKUP_DIR.mkdir(parents=True, exist_ok=True)
            backup = BACKUP_DIR / f"{vk}.docx.bak"
            backup.write_bytes(p["docx_path"].read_bytes())
            apply_plan(p["doc"], p["plan"])
            p["doc"].save(p["docx_path"])
            print(f"  WROTE {p['docx_path']}  (backup: {backup})")

        report["variants"][vk] = {
            "seed_section_total": p["seed_section_total"],
            "tagged": p["tagged"],
            "tagged_count": len(p["tagged"]),
            "skipped_already": p["skipped"],
            "ambiguous": p["ambiguous"],
            "unmatched": p["unmatched"],
            "chapter_unaligned": p["chapter_unaligned"],
            "chapter_alignment": p["chapter_alignment"],
            "coverage_pct": round(
                100.0 * (len(p["tagged"]) + len(p["skipped"])) / p["seed_section_total"],
                1,
            ) if p["seed_section_total"] else 0.0,
        }

    OUT_REPORT.write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"\nWrote report: {OUT_REPORT}")
    print("\n=== OVERALL ===")
    for vk, r in report["variants"].items():
        print(
            f"  {vk}: tagged={r['tagged_count']} "
            f"skipped={len(r['skipped_already'])} "
            f"ambiguous={len(r['ambiguous'])} unmatched={len(r['unmatched'])} "
            f"unaligned={len(r['chapter_unaligned'])} coverage={r['coverage_pct']}%"
        )


if __name__ == "__main__":
    main()
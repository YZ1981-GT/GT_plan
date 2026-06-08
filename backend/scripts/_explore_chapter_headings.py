#!/usr/bin/env python3
"""一次性探查：dump 指定段落区间内所有 heading (rank>=2) 及其样式 (用完即删)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

_BACKEND = Path(__file__).resolve().parent.parent
NOTES_DIR = _BACKEND / "data" / "audit_report_templates" / "disclosure_notes"


def _heading_rank(style_name: str) -> int | None:
    m = re.match(r"Heading\s*(\d+)", style_name or "", re.I)
    if m:
        return int(m.group(1))
    if style_name == "标题3的样式":
        return 3
    if style_name == "附注三级":
        return 3
    return None


def main() -> None:
    variant = sys.argv[1]
    lo = int(sys.argv[2])
    hi = int(sys.argv[3])
    doc = Document(NOTES_DIR / f"{variant}.docx")
    for i in range(lo, min(hi, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        style = p.style.name if p.style else ""
        rank = _heading_rank(style)
        if rank is not None:
            print(f"  [{i}] r{rank} {style!r} {(p.text or '').strip()!r}")


if __name__ == "__main__":
    main()

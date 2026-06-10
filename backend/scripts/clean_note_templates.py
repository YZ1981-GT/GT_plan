#!/usr/bin/env python
"""清理附注 Word 模板：删除【】/使用说明/（…删除）提示，封面占位符替换。

用法:
    python backend/scripts/clean_note_templates.py          # dry-run 模式，只报告
    python backend/scripts/clean_note_templates.py --write  # 实际写入（先备份到 _backup_clean/）

不动 ##SECTION: 标记、不动表格结构、不动 ##STYLE_REF:。
"""

import re
import shutil
import sys
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("需要 python-docx: pip install python-docx")
    sys.exit(1)

NOTES_DIR = Path(__file__).resolve().parent.parent / "data" / "audit_report_templates" / "disclosure_notes"
BACKUP_DIR = NOTES_DIR / "_backup_clean"

VARIANTS = [
    "soe_standalone.docx",
    "soe_consolidated.docx",
    "listed_standalone.docx",
    "listed_consolidated.docx",
]

# --- 清理规则 ---

# 匹配【...】（含内部内容）
RE_BRACKET = re.compile(r"【[^】]*】")
# 匹配（...删除）/（...，删除）/（有限，删除）/（无此项业务的，删除）等
RE_DELETE_HINT = re.compile(r"（[^）]{0,50}删除[^）]{0,10}）")
# 匹配 XXXX年/XX月/XX日 范例
RE_XXXX = re.compile(r"X{2,4}年|X{2}月|X{2}日")
# 使用说明关键词
USAGE_KEYWORDS = ["使用说明", "编制说明", "本附注模板使用说明", "附注编制说明"]

# 封面占位符替换
COVER_REPLACEMENTS = [
    # (范例文本模式, 替换为)
    (re.compile(r"XX有限公司|XX股份有限公司|ABC[^\n]{0,20}公司"), "{{company_full_name}}"),
    (re.compile(r"（\s*\d{4}\s*年度\s*）"), "（{{audit_year}}年度）"),
    (re.compile(r"XXXX\s*年度"), "{{audit_year}}年度"),
]


def is_section_marker(text: str) -> bool:
    """判断是否是 ##SECTION: / ##/SECTION: / ##STYLE_REF: 标记行"""
    stripped = text.strip()
    return (
        stripped.startswith("##SECTION:")
        or stripped.startswith("##/SECTION:")
        or stripped.startswith("##STYLE_REF:")
        or stripped.startswith("{{section:")
        or stripped.startswith("{{table:")
        or stripped.startswith("{{seq:")
    )


def clean_paragraph_text(para) -> dict:
    """清理单个段落，返回操作记录"""
    ops = {}
    full_text = para.text

    # 跳过标记行
    if is_section_marker(full_text):
        return ops

    # 删除【...】
    if RE_BRACKET.search(full_text):
        ops["bracket"] = RE_BRACKET.findall(full_text)

    # 删除（...删除）
    if RE_DELETE_HINT.search(full_text):
        ops["delete_hint"] = RE_DELETE_HINT.findall(full_text)

    # 替换 XXXX
    if RE_XXXX.search(full_text):
        ops["xxxx"] = RE_XXXX.findall(full_text)

    return ops


def should_delete_paragraph(para, idx: int, first_section_idx: int) -> str | None:
    """判断段落是否应整段删除，返回原因或 None"""
    text = para.text.strip()

    # 第一个 SECTION 之前的「使用说明」相关段落
    if idx < first_section_idx:
        for kw in USAGE_KEYWORDS:
            if kw in text:
                return f"使用说明关键词: {kw}"

    # 整段就是一个【...】（披露要求原文）
    cleaned = RE_BRACKET.sub("", text).strip()
    if not cleaned and RE_BRACKET.search(text):
        return "整段为【】披露要求"

    # 整段就是（...删除）
    cleaned2 = RE_DELETE_HINT.sub("", text).strip()
    if not cleaned2 and RE_DELETE_HINT.search(text):
        return "整段为（…删除）提示"

    return None


def find_first_section_index(doc) -> int:
    """找到第一个 ##SECTION: 标记的段落索引"""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("##SECTION:"):
            return i
    return len(doc.paragraphs)


def apply_inline_cleaning(para):
    """对段落内的 runs 做内联清理（删【】和（…删除），不删整段）"""
    full_text = para.text
    if is_section_marker(full_text):
        return 0

    changes = 0
    new_text = full_text

    # 删【...】
    if RE_BRACKET.search(new_text):
        new_text = RE_BRACKET.sub("", new_text)
        changes += 1

    # 删（...删除）
    if RE_DELETE_HINT.search(new_text):
        new_text = RE_DELETE_HINT.sub("", new_text)
        changes += 1

    if changes > 0 and new_text != full_text:
        # 重写段落文本（保留第一个 run 的格式）
        _rewrite_paragraph_text(para, new_text)

    return changes


def _rewrite_paragraph_text(para, new_text: str):
    """用新文本重写段落，保留首 run 格式"""
    if not para.runs:
        return
    # 保存首 run 的格式
    first_run = para.runs[0]
    # 清除所有 runs
    for run in para.runs:
        run.text = ""
    first_run.text = new_text


def apply_cover_replacements(para):
    """封面区域占位符替换"""
    full_text = para.text
    if is_section_marker(full_text):
        return 0

    new_text = full_text
    changes = 0
    for pattern, replacement in COVER_REPLACEMENTS:
        if pattern.search(new_text):
            new_text = pattern.sub(replacement, new_text)
            changes += 1

    if changes > 0 and new_text != full_text:
        _rewrite_paragraph_text(para, new_text)

    return changes


def delete_paragraph(para):
    """从文档中删除段落"""
    p = para._element
    p.getparent().remove(p)


def process_document(filepath: Path, write: bool = False) -> dict:
    """处理单个附注文档"""
    doc = Document(str(filepath))
    stats = {
        "file": filepath.name,
        "paragraphs_deleted": 0,
        "inline_cleaned": 0,
        "cover_replaced": 0,
        "total_paragraphs": len(doc.paragraphs),
    }

    first_section_idx = find_first_section_index(doc)
    stats["first_section_at"] = first_section_idx

    if not write:
        # Dry-run：只统计
        for i, para in enumerate(doc.paragraphs):
            reason = should_delete_paragraph(para, i, first_section_idx)
            if reason:
                stats["paragraphs_deleted"] += 1
            else:
                text = para.text
                if not is_section_marker(text):
                    if RE_BRACKET.search(text) or RE_DELETE_HINT.search(text):
                        stats["inline_cleaned"] += 1
                    if i < first_section_idx:
                        for pat, _ in COVER_REPLACEMENTS:
                            if pat.search(text):
                                stats["cover_replaced"] += 1
                                break
        return stats

    # --- Write 模式 ---
    # 第一遍：标记要删除的段落（倒序删除避免索引偏移）
    to_delete = []
    for i, para in enumerate(doc.paragraphs):
        reason = should_delete_paragraph(para, i, first_section_idx)
        if reason:
            to_delete.append((i, para, reason))

    # 倒序删除
    for i, para, reason in reversed(to_delete):
        delete_paragraph(para)
        stats["paragraphs_deleted"] += 1

    # 重新获取段落列表（删除后索引变了）
    first_section_idx_new = find_first_section_index(doc)

    # 第二遍：内联清理 + 封面替换
    for i, para in enumerate(doc.paragraphs):
        # 封面区域替换
        if i < first_section_idx_new:
            stats["cover_replaced"] += apply_cover_replacements(para)

        # 全文内联清理
        stats["inline_cleaned"] += apply_inline_cleaning(para)

    # 保存
    doc.save(str(filepath))
    return stats


def main():
    write = "--write" in sys.argv
    mode = "WRITE（实际修改）" if write else "DRY-RUN（只报告）"
    print(f"\n{'='*60}")
    print(f"附注模板清理 — {mode}")
    print(f"{'='*60}\n")

    if write:
        # 备份
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        for name in VARIANTS:
            src = NOTES_DIR / name
            if src.exists():
                dst = BACKUP_DIR / name
                shutil.copy2(src, dst)
                print(f"  备份: {name} → _backup_clean/")
        print()

    all_stats = []
    for name in VARIANTS:
        filepath = NOTES_DIR / name
        if not filepath.exists():
            print(f"  ⚠ 文件不存在: {name}")
            continue

        stats = process_document(filepath, write=write)
        all_stats.append(stats)

        print(f"📄 {stats['file']}")
        print(f"   总段落: {stats['total_paragraphs']}  |  首 SECTION 位置: 第{stats['first_section_at']}段")
        print(f"   整段删除: {stats['paragraphs_deleted']}  |  内联清理: {stats['inline_cleaned']}  |  封面替换: {stats['cover_replaced']}")
        print()

    # 汇总
    total_del = sum(s["paragraphs_deleted"] for s in all_stats)
    total_inline = sum(s["inline_cleaned"] for s in all_stats)
    total_cover = sum(s["cover_replaced"] for s in all_stats)
    print(f"{'─'*60}")
    print(f"合计: 删除 {total_del} 段 | 内联清理 {total_inline} 处 | 封面替换 {total_cover} 处")

    if not write:
        print(f"\n💡 确认无误后加 --write 参数执行实际修改")
        print(f"   python backend/scripts/clean_note_templates.py --write")


if __name__ == "__main__":
    main()

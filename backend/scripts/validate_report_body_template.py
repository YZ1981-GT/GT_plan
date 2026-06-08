#!/usr/bin/env python3
"""校验报告正文模板 POC：无 ABC/XXXX、无裸【、含核心占位符.

Usage:
    python backend/scripts/validate_report_body_template.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

_BACKEND = Path(__file__).resolve().parent.parent
POC = (
    _BACKEND / "data/audit_report_templates/report_body/"
    "1.1 模板A-无保留意见审计报告模板（上市公司、三板创新层及公开发债）-简版.docx"
)

FORBIDDEN = [
    (re.compile(r"\bABC\b"), "ABC"),
    (re.compile(r"XXXX"), "XXXX"),
    (re.compile(r"^【"), "【行首说明"),
]
REQUIRED_TOKENS = [
    "{{company_full_name}}",
    "{{audit_year}}",
    "{{firm_name}}",
    "{{report_number}}",
    "##OPT:key_audit_matters:",
]


def validate(path: Path) -> list[str]:
    issues: list[str] = []
    if not path.is_file():
        return [f"missing file: {path}"]
    doc = Document(path)
    full_text = "\n".join((p.text or "") for p in doc.paragraphs)
    for pattern, label in FORBIDDEN:
        for i, p in enumerate(doc.paragraphs):
            t = p.text or ""
            if pattern.search(t):
                issues.append(f"para[{i}] forbidden {label}: {t[:80]}")
    for token in REQUIRED_TOKENS:
        if token not in full_text:
            issues.append(f"missing required token: {token}")
    return issues


def main() -> None:
    issues = validate(POC)
    if issues:
        print(f"FAIL ({len(issues)}):")
        for x in issues:
            print(f"  - {x}")
        raise SystemExit(1)
    print(f"OK {POC.name}")


if __name__ == "__main__":
    main()

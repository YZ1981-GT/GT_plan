"""Deep gap analysis: Word disclosure templates vs JSON seeds vs exporter."""
from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document

DATA = Path(__file__).resolve().parent.parent / "data"
TPL = DATA / "audit_report_templates" / "disclosure_notes"
OUT = Path(__file__).resolve().parent / "_note_gap_deep.txt"

CHAPTER_KEYWORDS = [
    "公司基本情况",
    "财务报表编制",
    "财务报表的编制",
    "遵循企业",
    "重要会计",
    "税项",
    "合并范围",
    "财务报表主要项目",
    "或有事项",
    "资产负债表日后",
    "关联方",
    "股份支付",
    "承诺",
    "其他重要",
    "合并财务报表项目",
    "母公司财务报表",
    "研发支出",
    "政府补助",
    "金融工具",
    "公允价值",
    "补充资料",
    "在其他主体",
]


def styled_paras(path: Path) -> list[tuple[str, str]]:
    doc = Document(path)
    rows: list[tuple[str, str]] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        style = p.style.name if p.style else ""
        rows.append((style, t))
    return rows


def json_titles(secs: list[dict]) -> list[tuple[str, str, str]]:
    return [
        (s.get("section_number", ""), s.get("section_title", ""), s.get("scope", "both"))
        for s in secs
        if s.get("section_title")
    ]


def find_in_word(rows: list[tuple[str, str]], title: str) -> list[str]:
    hits = []
    key = re.sub(r"\s+", "", title)[:12]
    for style, t in rows:
        if key and key in re.sub(r"\s+", "", t):
            hits.append(f"[{style}] {t[:90]}")
    return hits


def main() -> None:
    soe_secs = json.loads((DATA / "note_template_soe.json").read_text(encoding="utf-8"))["sections"]
    listed_secs = json.loads((DATA / "note_template_listed.json").read_text(encoding="utf-8"))["sections"]
    bindings = json.loads((DATA / "note_template_bindings.json").read_text(encoding="utf-8"))
    bind_keys = list(bindings.keys())

    with open(OUT, "w", encoding="utf-8") as f:
        # --- Word files ---
        word_rows: dict[str, list[tuple[str, str]]] = {}
        for name in [
            "soe_standalone.docx",
            "soe_consolidated.docx",
            "listed_standalone.docx",
            "listed_consolidated.docx",
        ]:
            rows = styled_paras(TPL / name)
            word_rows[name] = rows
            f.write(f"\n{'=' * 72}\n{name}\n")
            f.write(f"  paragraphs={len(rows)}, tables={len(Document(TPL / name).tables)}\n")
            styles = Counter(s for s, _ in rows)
            f.write(f"  style top10: {styles.most_common(10)}\n")
            f.write("  chapter keyword hits:\n")
            seen: set[str] = set()
            for style, t in rows:
                for k in CHAPTER_KEYWORDS:
                    if k in t and t not in seen and len(t) < 100:
                        seen.add(t)
                        f.write(f"    [{style}] {t}\n")
                        break

        # standalone vs consolidated text diff size
        for pair in [
            ("soe_standalone.docx", "soe_consolidated.docx"),
            ("listed_standalone.docx", "listed_consolidated.docx"),
        ]:
            a, b = pair
            ta = {t for _, t in word_rows[a]}
            tb = {t for _, t in word_rows[b]}
            f.write(f"\n--- {a} vs {b} ---\n")
            f.write(f"  unique paragraphs: standalone-only={len(ta - tb)}, consolidated-only={len(tb - ta)}\n")
            for x in sorted(tb - ta)[:8]:
                f.write(f"    +consol: {x[:90]}\n")
            for x in sorted(ta - tb)[:8]:
                f.write(f"    +standalone: {x[:90]}\n")

        # --- JSON scope ---
        for label, secs in [("SOE", soe_secs), ("Listed", listed_secs)]:
            co = [s for s in secs if s.get("scope") == "consolidated_only"]
            st = [s for s in secs if s.get("scope") == "standalone_only"]
            f.write(f"\n{label} JSON: consolidated_only={len(co)}, standalone_only={len(st)}\n")
            for s in co[:15]:
                f.write(f"  [consol_only] {s.get('section_number')} {s.get('section_title')}\n")

        # --- JSON title presence in Word ---
        mapping = {
            "soe_standalone.docx": soe_secs,
            "soe_consolidated.docx": soe_secs,
            "listed_standalone.docx": listed_secs,
            "listed_consolidated.docx": listed_secs,
        }
        for docx, secs in mapping.items():
            rows = word_rows[docx]
            found = 0
            missing: list[str] = []
            for num, title, scope in json_titles(secs):
                if scope == "consolidated_only" and "standalone" in docx:
                    continue
                if scope == "standalone_only" and "consolidated" in docx:
                    continue
                hits = find_in_word(rows, title)
                if hits:
                    found += 1
                else:
                    if num and not num.startswith("四、"):  # skip noisy policy subnodes
                        missing.append(f"{num} {title}")
            f.write(f"\n{docx}: JSON titles found in Word={found}/{len(secs)}, missing sample:\n")
            for m in missing[:25]:
                f.write(f"  - {m}\n")

        # bindings 五、 vs 八、
        pairs = []
        for k in bind_keys:
            if k.startswith("五、") and ("八、" + k.split("、", 1)[1]) in bind_keys:
                pairs.append((k, "八、" + k.split("、", 1)[1]))
        f.write(f"\nBindings 五、N ↔ 八、N parallel pairs: {len(pairs)}\n")
        for a, b in pairs[:10]:
            f.write(f"  {a} ↔ {b}\n")

        # _detect_level bug demo
        f.write("\n_detect_level (current exporter) misclassification:\n")
        for code in ["八、1", "一、1", "四、会计期间", "五、1"]:
            parts = code.split(".")
            level = 1 if len(parts) == 1 else (2 if len(parts) == 2 else 3)
            f.write(f"  {code} -> level {level} (WRONG for Chinese numbering)\n")

    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

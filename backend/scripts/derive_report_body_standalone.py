#!/usr/bin/env python3
"""从已整理的合并口径报告正文派生单体（standalone）版本。

现有 17 份报告正文均为合并口径（"合并及公司资产负债表"等）。
单体公司报告无"合并及公司"前缀，且无集团审计/合并范围相关表述。

策略：
1. 复制每份 → report_body/standalone/{原名}（保留 opinion/subtype/variant 维度）
2. 机械替换合并口径报表表述为单体表述
3. 集团审计段（就...中实体或业务活动...集团审计）整段标 ##NOTE: 提示人工确认
   （单体一般删除，但不武断删，留 spot check）

Usage:
    python backend/scripts/derive_report_body_standalone.py            # dry-run
    python backend/scripts/derive_report_body_standalone.py --write
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

_BACKEND = Path(__file__).resolve().parent.parent
REPORT_BODY_DIR = _BACKEND / "data" / "audit_report_templates" / "report_body"
STANDALONE_DIR = REPORT_BODY_DIR / "standalone"

# 合并口径 → 单体口径（长模式优先）
SCOPE_REPLACEMENTS: list[tuple[str, str]] = [
    ("合并及公司资产负债表", "资产负债表"),
    ("合并及公司利润表", "利润表"),
    ("合并及公司现金流量表", "现金流量表"),
    ("合并及公司股东权益变动表", "股东权益变动表"),
    ("合并及公司所有者权益变动表", "所有者权益变动表"),
    ("合并及公司财务状况", "财务状况"),
    ("合并及公司经营成果和现金流量", "经营成果和现金流量"),
    ("合并及公司的经营成果和现金流量", "的经营成果和现金流量"),
    ("合并及公司财务报表", "财务报表"),
    ("合并及公司", ""),  # 兜底剩余
]

# 集团审计/合并范围相关段落关键词 → 标 NOTE 供人工确认（单体通常删）
GROUP_AUDIT_KEYWORDS = [
    "中实体或业务活动的财务信息",
    "集团审计",
    "合并范围",
    "合并财务报表",
    "合并资产负债表",
    "合并现金流量表",
]


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _apply_scope(text: str) -> str:
    out = text
    for old, new in SCOPE_REPLACEMENTS:
        out = out.replace(old, new)
    return out


def derive_doc(doc: Document, *, write: bool) -> dict[str, int]:
    stats = {"scope_replaced": 0, "group_audit_flagged": 0}
    for para in doc.paragraphs:
        raw = para.text or ""
        if not raw.strip():
            continue
        # 跳过已是标记行
        if raw.strip().startswith("##"):
            continue

        new_text = _apply_scope(raw)

        # 集团审计/合并范围段落：在段首加 NOTE 提示（不删，人工确认）
        is_group = any(kw in new_text for kw in GROUP_AUDIT_KEYWORDS)
        if is_group and "##NOTE:" not in new_text:
            new_text = f"##NOTE:单体复核:以下涉及合并/集团审计，单体报告可能需删除或调整## {new_text}"
            stats["group_audit_flagged"] += 1

        if new_text != raw:
            if new_text != _apply_scope(raw) or _apply_scope(raw) != raw:
                stats["scope_replaced"] += 1 if _apply_scope(raw) != raw else 0
            if write:
                _set_paragraph_text(para, new_text)
    return stats


def iter_sources() -> list[Path]:
    return sorted(
        p for p in REPORT_BODY_DIR.glob("*.docx") if not p.name.startswith("~")
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="Derive standalone report body templates")
    parser.add_argument("--write", action="store_true")
    args = parser.parse_args()

    sources = iter_sources()
    if not sources:
        print("无源文件", file=sys.stderr)
        raise SystemExit(1)

    mode = "WRITE（实际生成）" if args.write else "DRY-RUN（只报告）"
    print(f"\n{'='*64}\n派生单体报告正文 — {mode}  共 {len(sources)} 份\n{'='*64}\n")

    if args.write:
        STANDALONE_DIR.mkdir(parents=True, exist_ok=True)

    totals = {"scope_replaced": 0, "group_audit_flagged": 0}
    for src in sources:
        dst = STANDALONE_DIR / src.name
        if args.write:
            shutil.copy2(src, dst)
            doc = Document(dst)
        else:
            doc = Document(src)
        stats = derive_doc(doc, write=args.write)
        for k in totals:
            totals[k] += stats[k]
        print(f"📄 {src.name[:40]}")
        print(f"   合并口径替换段={stats['scope_replaced']}  集团审计待确认段={stats['group_audit_flagged']}")
        if args.write:
            doc.save(dst)
    print(f"\n{'─'*64}")
    print(f"合计: 合并口径替换={totals['scope_replaced']} 集团审计待确认={totals['group_audit_flagged']}")
    if args.write:
        print(f"\n✅ 已生成到 {STANDALONE_DIR}")
        print("⚠️  ##NOTE:单体复核## 标记的段落需人工确认是否删除（集团审计/合并范围在单体报告中通常不适用）")
    else:
        print("\n💡 确认后加 --write 生成单体套到 report_body/standalone/")


if __name__ == "__main__":
    main()

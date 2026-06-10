#!/usr/bin/env python3
"""整理报告正文模板全 17 份 — 占位符替换 + 删【】说明 + OPT 段标记（初稿）.

Task 0.6.1（自动化初稿；OPT 边界仍需人工 spot check）

复用 prepare_report_body_poc 的替换/删除/OPT 逻辑，泛化到全部 17 份。
1.1 已由 POC 处理，默认跳过（除非 --include-poc）。

Usage:
    python backend/scripts/prepare_report_body_all.py            # dry-run
    python backend/scripts/prepare_report_body_all.py --write    # 先备份再写
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

_BACKEND = Path(__file__).resolve().parent.parent
REPORT_BODY_DIR = _BACKEND / "data" / "audit_report_templates" / "report_body"
BACKUP_DIR = REPORT_BODY_DIR / "_backup_clean"

POC_FILE = (
    "1.1 模板A-无保留意见审计报告模板（上市公司、三板创新层及公开发债）-简版.docx"
)

# --- 文本替换（长模式优先） ---
TEXT_REPLACEMENTS: list[tuple[str, str]] = [
    # 治理层二选一抬头：「ABC有限公司董事会/XXX全体股东：」→ 收口为占位
    ("ABC有限公司董事会/{{company_full_name}}全体股东：", "{{company_full_name}}{{responsibility_organ}}："),
    ("ABC有限公司董事会/ABC股份有限公司全体股东：", "{{company_full_name}}{{responsibility_organ}}："),
    ("ABC股份有限公司全体股东：", "{{company_full_name}}全体股东："),
    ("ABC有限公司董事会", "{{company_full_name}}董事会"),
    ("ABC股份有限公司", "{{company_full_name}}"),
    ("ABC有限公司", "{{company_full_name}}"),
    ("ABC公司", "{{company_short_name}}"),
    # 报告编号：致同审字（YYYY）第110ASXXXX号 等变体
    (re.compile(r"致同审字（[^）]*）第[0-9A-Za-z]*号"), "{{report_number}}"),
    # 期间
    ("2025年12月31日", "{{audit_period_end}}"),
    ("2025年1月1日", "{{audit_period_start}}"),
    ("2024年12月31日", "{{prior_period_end}}"),
    ("2025年度", "{{audit_year}}年度"),
    ("二〇二五年度", "{{audit_year}}年度"),
    ("二〇二五年", "{{audit_year}}年"),
    ("2025年", "{{audit_year}}年"),
    ("致同会计师事务所（特殊普通合伙）", "{{firm_name}}"),
]

# 行内 / 整行【】指引
GUIDANCE_LINE_RE = re.compile(r"^【")
INLINE_GUIDANCE_RE = re.compile(r"【[^】]*】")

# --- OPT 块识别（按一级中文标题前缀；保守，仅明确可选段） ---
# section_id, 描述, 标题前缀候选（任一匹配即认定该段为该 OPT 起点）
OPT_DEFS: list[tuple[str, str, list[str]]] = [
    ("key_audit_matters", "关键审计事项", ["关键审计事项"]),
    ("emphasis", "强调事项段", ["强调事项"]),
    ("going_concern", "与持续经营相关的重大不确定性", ["持续经营相关的重大不确定性", "与持续经营相关的重大不确定性"]),
    ("other_matter", "其他事项段", ["其他事项"]),
    ("other_information", "其他信息", ["其他信息"]),
    ("comparative", "比较信息", ["比较信息", "比较数据"]),
]

# 一级标题模式：「一、xxx」「二、xxx」…
SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、")


def _set_paragraph_text(para: Paragraph, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _apply_replacements(text: str) -> str:
    out = text
    for old, new in TEXT_REPLACEMENTS:
        if isinstance(old, re.Pattern):
            out = old.sub(new, out)
        else:
            out = out.replace(old, new)
    return out


def _strip_inline_guidance(text: str) -> str:
    return INLINE_GUIDANCE_RE.sub("", text).strip()


def _insert_before(ref: Paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    ref._element.addprevious(new_p)
    Paragraph(new_p, ref._parent).add_run(text)


def _insert_after(ref: Paragraph, text: str) -> None:
    new_p = OxmlElement("w:p")
    ref._element.addnext(new_p)
    Paragraph(new_p, ref._parent).add_run(text)


def _delete_paragraph(para: Paragraph) -> None:
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _heading_section_id(text: str) -> str | None:
    """若该段是某 OPT 段标题，返回其 section_id。"""
    body = SECTION_HEADING_RE.sub("", text.strip()) if SECTION_HEADING_RE.match(text.strip()) else None
    if body is None:
        return None
    for sec_id, _desc, prefixes in OPT_DEFS:
        for pref in prefixes:
            if body.startswith(pref):
                return sec_id
    return None


def _desc_for(sec_id: str) -> str:
    for s, d, _ in OPT_DEFS:
        if s == sec_id:
            return d
    return sec_id


def prepare_doc(doc: Document, *, write: bool) -> dict[str, int]:
    stats = {"replaced": 0, "deleted": 0, "stripped_inline": 0, "opt_blocks": 0}

    # 1) 全文替换 + 删整行【】 + 剥行内【】
    to_delete: list[Paragraph] = []
    for para in list(doc.paragraphs):
        raw = para.text or ""
        if not raw.strip():
            continue
        if GUIDANCE_LINE_RE.match(raw.strip()):
            to_delete.append(para)
            stats["deleted"] += 1
            continue
        new_text = _apply_replacements(raw)
        if INLINE_GUIDANCE_RE.search(new_text):
            cleaned = _strip_inline_guidance(new_text)
            if not cleaned:
                to_delete.append(para)
                stats["deleted"] += 1
                continue
            new_text = cleaned
            stats["stripped_inline"] += 1
        if new_text != raw:
            stats["replaced"] += 1
            if write:
                _set_paragraph_text(para, new_text)

    if write:
        for para in to_delete:
            _delete_paragraph(para)

    # 2) OPT 块：从后往前，避免索引漂移
    paras = doc.paragraphs
    # 收集所有 OPT 起点
    opt_starts: list[tuple[int, str]] = []
    for i, p in enumerate(paras):
        sec_id = _heading_section_id(p.text or "")
        if sec_id:
            opt_starts.append((i, sec_id))

    for start_idx, sec_id in reversed(opt_starts):
        paras = doc.paragraphs
        if start_idx >= len(paras):
            continue
        # 已打标跳过
        open_m = f"##OPT:{sec_id}:{_desc_for(sec_id)}##"
        close_m = f"##/OPT:{sec_id}##"
        if start_idx > 0 and "##OPT:" in (paras[start_idx - 1].text or ""):
            continue
        # 结束 = 下一个一级标题
        end_idx = len(paras)
        for j in range(start_idx + 1, len(paras)):
            t = (paras[j].text or "").strip()
            if SECTION_HEADING_RE.match(t):
                end_idx = j
                break
        stats["opt_blocks"] += 1
        if write:
            end_para = paras[min(end_idx, len(paras) - 1)]
            # 先插结束（在结束段之前），再插开始
            _insert_before(end_para, close_m)
            _insert_before(doc.paragraphs[start_idx], open_m)

    return stats


def iter_targets(include_poc: bool) -> list[Path]:
    files = sorted(
        p for p in REPORT_BODY_DIR.glob("*.docx") if not p.name.startswith("~")
    )
    if not include_poc:
        files = [p for p in files if p.name != POC_FILE]
    return files


def main() -> None:
    parser = argparse.ArgumentParser(description="Prepare all report body templates")
    parser.add_argument("--write", action="store_true")
    parser.add_argument("--include-poc", action="store_true", help="也处理 1.1 POC 文件")
    args = parser.parse_args()

    targets = iter_targets(args.include_poc)
    if not targets:
        print("无目标文件", file=sys.stderr)
        raise SystemExit(1)

    mode = "WRITE（实际修改）" if args.write else "DRY-RUN（只报告）"
    print(f"\n{'='*64}\n报告正文模板整理 — {mode}  共 {len(targets)} 份\n{'='*64}\n")

    if args.write:
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        for p in targets:
            shutil.copy2(p, BACKUP_DIR / p.name)
        print(f"  已备份 {len(targets)} 份 → _backup_clean/\n")

    totals = {"replaced": 0, "deleted": 0, "stripped_inline": 0, "opt_blocks": 0}
    for path in targets:
        doc = Document(path)
        stats = prepare_doc(doc, write=args.write)
        for k in totals:
            totals[k] += stats[k]
        print(f"📄 {path.name}")
        print(f"   替换={stats['replaced']}  删行={stats['deleted']}  "
              f"剥行内【】={stats['stripped_inline']}  OPT块={stats['opt_blocks']}")
        if args.write:
            doc.save(path)
    print(f"\n{'─'*64}")
    print(f"合计: 替换={totals['replaced']} 删行={totals['deleted']} "
          f"剥行内={totals['stripped_inline']} OPT块={totals['opt_blocks']}")
    if not args.write:
        print("\n💡 确认后加 --write 执行（自动备份到 _backup_clean/）")


if __name__ == "__main__":
    main()

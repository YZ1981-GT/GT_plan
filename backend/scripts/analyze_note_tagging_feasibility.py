#!/usr/bin/env python3
"""DRY-RUN / REPORT-ONLY: 评估附注 Word 模板自动打标 (##SECTION:##) 的可行性.

支持 Task 0.6.2 / 10.2 / 10.5 的「先调查后决策」——本脚本**只读**四套
disclosure_notes/*.docx，绝不修改任何 .docx。它把 JSON 种子章节
(section_number + section_title) 与 Word 文档的标题段落做匹配，统计匹配率，
列出未匹配 / 歧义章节，并给出 GO / NO-GO 建议。

复用 build_section_code_index.py 的 VARIANT_SEED + JSON 加载口径，
不另起一套 seed 路径逻辑。

Usage (cwd=backend, Windows):
    ..\\.venv\\Scripts\\python.exe scripts/analyze_note_tagging_feasibility.py

输出:
    - stdout 每变体摘要 + GO/NO-GO
    - backend/data/audit_report_templates/_note_tagging_feasibility.json
      (前缀 `_` 表示一次性分析产物, 用完可删)
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

# 复用 build_section_code_index 的种子映射 + JSON 加载口径（同目录脚本）
from build_section_code_index import (  # type: ignore
    DATA,
    NOTES_DIR,
    VARIANT_SEED,
    load_json_sections,
)

OUT_PATH = DATA / "audit_report_templates" / "_note_tagging_feasibility.json"

# GO 阈值：无歧义的 exact+number 匹配率 ≥ 此值 → 可考虑自动打标 + 人工抽查
GO_THRESHOLD = 95.0
# 介于 GO 与此值之间 → CAUTION（需更多人工介入）；低于此值 → NO-GO
CAUTION_THRESHOLD = 80.0

# 标题段落样式判定：Heading N / 标题N / 标题3的样式 等
_HEADING_STYLE_RE = re.compile(r"^(heading|标题)", re.I)
_BRACKET_RE = re.compile(r"【[^】]*】")
# 适用性后缀噪声，匹配时剥离（如「（不适用的删除）」「（不适用删除）」）
_SUFFIX_NOISE_RE = re.compile(r"（[^）]*(?:删除|不适用|可选|选择性|备注)[^）]*）\s*$")
# 行内编号前缀，如「（1）」「1、」「1.」「（一）」
_LEADING_NUM_RE = re.compile(r"^[（(]?[0-9一二三四五六七八九十]+[）)、.\．]\s*")


def _norm_row_name(text: str) -> str:
    """归一标题：去空白 + 去前导 △▲#* 标记（对齐 prepare_financial_templates 口径）."""
    s = re.sub(r"\s+", "", (text or ""))
    return re.sub(r"^[△▲#*]+", "", s)


def _norm_title(text: str) -> str:
    """更激进的标题归一：在 _norm_row_name 基础上去适用性后缀 + 行内编号前缀."""
    s = _norm_row_name(text)
    s = _LEADING_NUM_RE.sub("", s)
    s = _SUFFIX_NOISE_RE.sub("", s)
    return s


def _is_heading(para: Any) -> bool:
    style = (para.style.name if para.style else "") or ""
    return bool(_HEADING_STYLE_RE.match(style))


def collect_headings(docx_path: Path) -> list[dict[str, Any]]:
    """返回候选标题段落 [{idx, text, style, norm, norm_loose}]."""
    doc = Document(docx_path)
    headings: list[dict[str, Any]] = []
    for i, para in enumerate(doc.paragraphs):
        if not _is_heading(para):
            continue
        raw = (para.text or "").strip()
        if not raw:
            continue
        headings.append(
            {
                "idx": i,
                "text": raw,
                "style": para.style.name if para.style else "",
                "norm": _norm_row_name(raw),
                "norm_loose": _norm_title(raw),
            }
        )
    return headings


def count_cleanup_artifacts(docx_path: Path) -> dict[str, int]:
    """统计 【…】 提示括号 + 含「使用说明」段落（清理工作量）."""
    doc = Document(docx_path)
    bracket_hits = 0
    usage_paras = 0
    for para in doc.paragraphs:
        text = para.text or ""
        bracket_hits += len(_BRACKET_RE.findall(text))
        if "使用说明" in text:
            usage_paras += 1
    return {"bracket_hint_count": bracket_hits, "usage_instruction_paragraphs": usage_paras}


def _build_heading_lookup(headings: list[dict[str, Any]]) -> dict[str, dict[str, list[int]]]:
    """构建 {归一文本 -> [段落idx...]} 的两级索引（精确归一 + 宽松归一）."""
    exact: dict[str, list[int]] = {}
    loose: dict[str, list[int]] = {}
    for h in headings:
        exact.setdefault(h["norm"], []).append(h["idx"])
        loose.setdefault(h["norm_loose"], []).append(h["idx"])
    return {"exact": exact, "loose": loose}


def match_section(
    seed: dict[str, Any],
    headings: list[dict[str, Any]],
    lookup: dict[str, dict[str, list[int]]],
) -> dict[str, Any]:
    """对单个 JSON 种子章节尝试匹配 Word 标题.

    顺序: exact 标题归一 → section_number 字样 → 宽松归一 → fuzzy/contains.
    返回 {matched, method, para_idx, ambiguous, candidates}.
    """
    section_number = (seed.get("section_number") or "").strip()
    title = (seed.get("section_title") or "").strip()
    norm_title = _norm_row_name(title)
    loose_title = _norm_title(title)

    result: dict[str, Any] = {
        "matched": False,
        "method": None,
        "para_idx": None,
        "ambiguous": False,
        "candidate_indices": [],
    }

    # 1) exact 归一标题
    if norm_title:
        hits = lookup["exact"].get(norm_title, [])
        if hits:
            result.update(
                matched=True,
                method="exact_title",
                para_idx=hits[0],
                ambiguous=len(hits) > 1,
                candidate_indices=hits,
            )
            return result

    # 2) section_number 字样出现在标题文本（少见，Word 标题一般不带「八、1」）
    if section_number:
        num_hits = [h["idx"] for h in headings if section_number and section_number in h["text"]]
        if num_hits:
            result.update(
                matched=True,
                method="section_number",
                para_idx=num_hits[0],
                ambiguous=len(num_hits) > 1,
                candidate_indices=num_hits,
            )
            return result

    # 3) 宽松归一（去适用性后缀 / 行内编号前缀后相等）
    if loose_title:
        hits = lookup["loose"].get(loose_title, [])
        if hits:
            result.update(
                matched=True,
                method="loose_title",
                para_idx=hits[0],
                ambiguous=len(hits) > 1,
                candidate_indices=hits,
            )
            return result

    # 4) fuzzy / contains：标题（宽松归一）是某标题的子串或反之，取最短候选
    if loose_title and len(loose_title) >= 2:
        contains_hits = [
            h["idx"]
            for h in headings
            if h["norm_loose"]
            and (loose_title in h["norm_loose"] or h["norm_loose"] in loose_title)
        ]
        if contains_hits:
            result.update(
                matched=True,
                method="fuzzy_contains",
                para_idx=contains_hits[0],
                ambiguous=len(contains_hits) > 1,
                candidate_indices=contains_hits,
            )
            return result

    return result


def analyze_variant(variant_key: str) -> dict[str, Any]:
    seed_file, _template_type = VARIANT_SEED[variant_key]
    docx_path = NOTES_DIR / f"{variant_key}.docx"
    if not docx_path.is_file():
        raise FileNotFoundError(docx_path)

    by_code = load_json_sections(seed_file)
    seeds = list(by_code.values())
    headings = collect_headings(docx_path)
    lookup = _build_heading_lookup(headings)

    method_counts = {
        "exact_title": 0,
        "section_number": 0,
        "loose_title": 0,
        "fuzzy_contains": 0,
    }
    ambiguous: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    matched_total = 0
    per_section: list[dict[str, Any]] = []

    for seed in seeds:
        code = (seed.get("section_number") or "").strip()
        title = (seed.get("section_title") or "").strip()
        m = match_section(seed, headings, lookup)
        record = {
            "section_number": code,
            "section_title": title,
            "matched": m["matched"],
            "method": m["method"],
            "para_idx": m["para_idx"],
            "ambiguous": m["ambiguous"],
        }
        per_section.append(record)
        if m["matched"]:
            matched_total += 1
            method_counts[m["method"]] += 1
            if m["ambiguous"]:
                ambiguous.append(
                    {
                        "section_number": code,
                        "section_title": title,
                        "method": m["method"],
                        "candidate_para_indices": m["candidate_indices"],
                    }
                )
        else:
            unmatched.append({"section_number": code, "section_title": title})

    total = len(seeds)
    # 强匹配 = 无歧义的 exact_title + section_number（这两类才是可安全自动打标的）
    strong_unambiguous = sum(
        1
        for r in per_section
        if r["matched"]
        and r["method"] in ("exact_title", "section_number")
        and not r["ambiguous"]
    )
    match_rate = round(100.0 * matched_total / total, 1) if total else 0.0
    strong_rate = round(100.0 * strong_unambiguous / total, 1) if total else 0.0

    cleanup = count_cleanup_artifacts(docx_path)

    # GO / NO-GO 判定
    if strong_rate >= GO_THRESHOLD and len(ambiguous) == 0:
        recommendation = "GO"
        rationale = (
            f"强匹配率(无歧义 exact+number) {strong_rate}% ≥ {GO_THRESHOLD}% 且无歧义章节；"
            "可自动打标后人工抽查。"
        )
    elif strong_rate >= CAUTION_THRESHOLD:
        recommendation = "CAUTION"
        rationale = (
            f"强匹配率 {strong_rate}%（{CAUTION_THRESHOLD}%~{GO_THRESHOLD}%），"
            f"歧义 {len(ambiguous)} / 未匹配 {len(unmatched)} 章节需人工处理；"
            "可半自动（自动打标强匹配项，其余手工）。"
        )
    else:
        recommendation = "NO-GO"
        rationale = (
            f"强匹配率 {strong_rate}% < {CAUTION_THRESHOLD}%；"
            f"歧义 {len(ambiguous)} / 未匹配 {len(unmatched)} 章节过多，"
            "自动打标风险高，应保持人工打标。"
        )

    return {
        "variant_key": variant_key,
        "seed_file": seed_file,
        "docx_file": f"disclosure_notes/{variant_key}.docx",
        "heading_heuristic": "段落 style 名以 'Heading'/'标题' 开头（含 '标题3的样式'）",
        "totals": {
            "seed_sections": total,
            "matched": matched_total,
            "match_rate_pct": match_rate,
            "strong_unambiguous": strong_unambiguous,
            "strong_unambiguous_rate_pct": strong_rate,
            "ambiguous": len(ambiguous),
            "unmatched": len(unmatched),
            "candidate_headings_in_docx": len(headings),
        },
        "match_methods": method_counts,
        "cleanup_workload": cleanup,
        "ambiguous_sections": ambiguous,
        "unmatched_sections": unmatched,
        "recommendation": recommendation,
        "rationale": rationale,
    }


def _print_summary(report: dict[str, Any]) -> None:
    t = report["totals"]
    m = report["match_methods"]
    c = report["cleanup_workload"]
    print(f"\n=== {report['variant_key']}  ({report['seed_file']}) ===")
    print(f"  seed sections        : {t['seed_sections']}")
    print(f"  candidate headings   : {t['candidate_headings_in_docx']}")
    print(
        f"  matched              : {t['matched']}  "
        f"(rate {t['match_rate_pct']}%)"
    )
    print(
        f"    exact_title={m['exact_title']}  section_number={m['section_number']}  "
        f"loose_title={m['loose_title']}  fuzzy_contains={m['fuzzy_contains']}"
    )
    print(
        f"  strong unambiguous   : {t['strong_unambiguous']}  "
        f"(rate {t['strong_unambiguous_rate_pct']}%)  [exact+number, no ambiguity]"
    )
    print(f"  ambiguous            : {t['ambiguous']}")
    print(f"  unmatched            : {t['unmatched']}")
    print(
        f"  cleanup workload     : 【…】hints={c['bracket_hint_count']}  "
        f"使用说明 paras={c['usage_instruction_paragraphs']}"
    )
    if report["unmatched_sections"]:
        nums = [s["section_number"] or s["section_title"] for s in report["unmatched_sections"]]
        print(f"  UNMATCHED section_numbers ({len(nums)}): {nums}")
    if report["ambiguous_sections"]:
        anums = [
            f"{s['section_number']}@{s['candidate_para_indices']}"
            for s in report["ambiguous_sections"]
        ]
        print(f"  AMBIGUOUS ({len(anums)}): {anums}")
    print(f"  >>> {report['recommendation']}: {report['rationale']}")


def main() -> None:
    print("DRY-RUN / REPORT-ONLY note auto-tagging feasibility analyzer")
    print("(read-only on disclosure_notes/*.docx — NO docx is modified)")

    variants = list(VARIANT_SEED.keys())
    reports: dict[str, Any] = {}
    for vk in variants:
        try:
            report = analyze_variant(vk)
        except FileNotFoundError as e:
            print(f"SKIP {vk}: missing docx {e}")
            continue
        reports[vk] = report
        _print_summary(report)

    payload = {
        "_note": "一次性可行性分析产物（dry-run, report-only）；用完可删。",
        "go_threshold_pct": GO_THRESHOLD,
        "caution_threshold_pct": CAUTION_THRESHOLD,
        "variants": reports,
    }
    OUT_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nWrote report: {OUT_PATH}")

    # 总体一句话裁决
    print("\n=== OVERALL ===")
    for vk, r in reports.items():
        print(f"  {vk}: {r['recommendation']}  (strong {r['totals']['strong_unambiguous_rate_pct']}%)")


if __name__ == "__main__":
    main()

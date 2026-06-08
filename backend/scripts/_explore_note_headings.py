#!/usr/bin/env python3
"""一次性探查：列出附注 docx 的 Heading-1 段落 + 各章标题 (用完即删)."""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

_BACKEND = Path(__file__).resolve().parent.parent
DATA = _BACKEND / "data"
NOTES_DIR = DATA / "audit_report_templates" / "disclosure_notes"

VARIANT_SEED = {
    "soe_standalone": "note_template_soe.json",
    "soe_consolidated": "note_template_soe.json",
    "listed_standalone": "note_template_listed.json",
    "listed_consolidated": "note_template_listed.json",
}

CN_NUM = set("一二三四五六七八九十")


def _heading_rank(style_name: str) -> int | None:
    m = re.match(r"Heading\s*(\d+)", style_name or "", re.I)
    if m:
        return int(m.group(1))
    m2 = re.match(r"标题\s*(\d+)", style_name or "")
    if m2:
        return int(m2.group(1))
    return None


def chapter_titles(seed_file: str) -> dict[str, str]:
    data = json.loads((DATA / seed_file).read_text(encoding="utf-8"))
    out: dict[str, str] = {}
    for s in data.get("sections", []):
        num = s.get("section_number", "")
        if num and all(ch in CN_NUM for ch in num):
            out[num] = s.get("section_title", "")
    return out


def main() -> None:
    variant = sys.argv[1] if len(sys.argv) > 1 else "soe_standalone"
    seed_file = VARIANT_SEED[variant]
    chapters = chapter_titles(seed_file)
    print(f"=== {variant} chapters from seed ===")
    for n, t in chapters.items():
        print(f"  {n} -> {t}")

    doc = Document(NOTES_DIR / f"{variant}.docx")
    print(f"\n=== {variant}.docx Heading-1 paragraphs ===")
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ""
        rank = _heading_rank(style)
        if rank == 1:
            print(f"  [{i}] style={style!r} text={(p.text or '').strip()!r}")

    # style histogram
    from collections import Counter
    c = Counter((p.style.name if p.style else "") for p in doc.paragraphs)
    print(f"\n=== style histogram ===")
    for st, n in c.most_common(20):
        print(f"  {n:5d}  {st!r}")


if __name__ == "__main__":
    main()

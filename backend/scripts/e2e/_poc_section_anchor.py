"""POC: 验证隐藏书签作为段落锚点的可行性（D1 段落锚点 POC 步骤 1）.

验证目标：
  1. 用 python-docx 在合成附注 docx 的章节块 open_el 前插入 bookmarkStart、
     close_el 后插入 bookmarkEnd，保存。
  2. 重新打开保存的 docx，验证：
     - 书签不可见（visible text 不含书签标记）
     - 排版无变化（文字内容完全一致）
     - 可按 bookmark name 定位到对应章节块

用法:
    python backend/scripts/e2e/_poc_section_anchor.py

产物:
    backend/scripts/e2e/_poc_output/synthetic_note_with_bookmarks.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

# 确保可以 import backend 模块
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# 1. 合成附注 docx（含 2-3 个 SECTION 块）
# ---------------------------------------------------------------------------

def create_synthetic_note_docx() -> BytesIO:
    """创建一份合成附注 docx，包含 3 个 SECTION 块."""
    doc = Document()

    # 标题
    doc.add_heading("合成附注（POC 验证用）", level=1)
    doc.add_paragraph("本文档用于验证隐藏书签段落锚点方案。")

    # --- 章节 八、1 ---
    doc.add_paragraph("##SECTION:八、1##")
    doc.add_heading("八、1 固定资产", level=2)
    doc.add_paragraph("固定资产按取得时的成本进行初始计量。与固定资产有关的后续支出，"
                      "符合确认条件的计入固定资产成本，不符合确认条件的在发生时计入当期损益。")
    # 插入一个简单表格模拟附注表格
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "期末余额"
    table.cell(0, 2).text = "期初余额"
    table.cell(1, 0).text = "房屋建筑物"
    table.cell(1, 1).text = "1,000,000.00"
    table.cell(1, 2).text = "950,000.00"
    table.cell(2, 0).text = "机器设备"
    table.cell(2, 1).text = "500,000.00"
    table.cell(2, 2).text = "480,000.00"
    doc.add_paragraph("##/SECTION:八、1##")

    # --- 章节 八、2 ---
    doc.add_paragraph("##SECTION:八、2##")
    doc.add_heading("八、2 无形资产", level=2)
    doc.add_paragraph("无形资产按成本进行初始计量。使用寿命有限的无形资产自达到预定用途之日起，"
                      "按直线法在预计使用寿命内摊销。")
    doc.add_paragraph("##/SECTION:八、2##")

    # --- 章节 五、1 ---
    doc.add_paragraph("##SECTION:五、1##")
    doc.add_heading("五、1 货币资金", level=2)
    doc.add_paragraph("货币资金包括库存现金、银行存款和其他货币资金。")
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "项目"
    table2.cell(0, 1).text = "金额"
    table2.cell(1, 0).text = "银行存款"
    table2.cell(1, 1).text = "2,500,000.00"
    doc.add_paragraph("##/SECTION:五、1##")

    # 尾部文字
    doc.add_paragraph("（附注完）")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 2. 书签工具函数
# ---------------------------------------------------------------------------

def section_code_to_bookmark_name(section_code: str) -> str:
    """将 section_code 转换为 bookmark name.

    规则：sec_ + section_code 中的中文逗号/顿号替换为下划线
    示例：八、1 → sec_八_1
    """
    name = section_code.replace("、", "_").replace("，", "_").replace(",", "_")
    return f"sec_{name}"


def insert_bookmark_around_section_block(
    doc: DocumentObject,
    open_el,
    close_el,
    bookmark_name: str,
    bookmark_id: int,
) -> None:
    """在 open_el 前插入 bookmarkStart，在 close_el 后插入 bookmarkEnd.

    书签是 OOXML 的原生定位机制：
    - <w:bookmarkStart w:id="N" w:name="sec_八_1"/>
    - <w:bookmarkEnd w:id="N"/>

    它们是 body 级子元素，插入后不影响可见文本和排版。
    """
    body = doc.element.body

    # 创建 bookmarkStart
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), bookmark_name)

    # 创建 bookmarkEnd
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))

    # 在 open_el 之前插入 bookmarkStart
    body.insert(list(body).index(open_el), bm_start)

    # 在 close_el 之后插入 bookmarkEnd
    close_idx = list(body).index(close_el)
    body.insert(close_idx + 1, bm_end)


# ---------------------------------------------------------------------------
# 3. 扫描 SECTION 块（复用 word_doc_utils 同口径逻辑）
# ---------------------------------------------------------------------------

@dataclass
class SectionBlock:
    """简化版 SectionBlock，与 word_doc_utils.SectionBlock 同结构."""
    section_code: str
    open_el: object
    close_el: object
    elements: list


import re

SECTION_OPEN_RE = re.compile(r"^##SECTION:(.+?)##\s*$")
SECTION_CLOSE_RE = re.compile(r"^##/SECTION:(.+?)##\s*$")


def scan_section_blocks(doc: DocumentObject) -> list[SectionBlock]:
    """扫描 body 级 SECTION 块（与 word_doc_utils.scan_section_blocks 同口径）."""
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    body = doc.element.body

    elements = []
    for child in body.iterchildren():
        if child.tag in (p_tag, tbl_tag):
            elements.append(child)

    blocks: list[SectionBlock] = []
    open_idx: int | None = None
    open_code: str | None = None

    for i, el in enumerate(elements):
        if el.tag != p_tag:
            continue
        text = (Paragraph(el, doc).text or "").strip()
        if open_idx is None:
            m_open = SECTION_OPEN_RE.match(text)
            if m_open:
                open_idx = i
                open_code = m_open.group(1).strip()
            continue
        m_close = SECTION_CLOSE_RE.match(text)
        if m_close and m_close.group(1).strip() == open_code:
            blocks.append(
                SectionBlock(
                    section_code=open_code or "",
                    open_el=elements[open_idx],
                    close_el=el,
                    elements=elements[open_idx: i + 1],
                )
            )
            open_idx = None
            open_code = None

    return blocks


# ---------------------------------------------------------------------------
# 4. 书签回读验证
# ---------------------------------------------------------------------------

def find_bookmarks(doc: DocumentObject) -> dict[str, tuple[object, object]]:
    """从 docx 回读所有书签，返回 {name: (bookmarkStart_el, bookmarkEnd_el)}.

    遍历 body 下所有 bookmarkStart/bookmarkEnd 元素，按 w:id 配对。
    """
    body = doc.element.body
    bm_start_tag = qn("w:bookmarkStart")
    bm_end_tag = qn("w:bookmarkEnd")

    starts: dict[str, object] = {}  # id -> element
    ends: dict[str, object] = {}    # id -> element
    names: dict[str, str] = {}      # id -> name

    for el in body.iter():
        if el.tag == bm_start_tag:
            bm_id = el.get(qn("w:id"))
            bm_name = el.get(qn("w:name"))
            if bm_id and bm_name:
                starts[bm_id] = el
                names[bm_id] = bm_name
        elif el.tag == bm_end_tag:
            bm_id = el.get(qn("w:id"))
            if bm_id:
                ends[bm_id] = el

    result: dict[str, tuple[object, object]] = {}
    for bm_id, name in names.items():
        if bm_id in starts and bm_id in ends:
            result[name] = (starts[bm_id], ends[bm_id])

    return result


def extract_text_between_bookmarks(
    doc: DocumentObject, bm_start_el, bm_end_el
) -> str:
    """提取两个书签元素之间的所有可见文本."""
    body = doc.element.body
    p_tag = qn("w:p")
    collecting = False
    texts: list[str] = []

    for child in body.iterchildren():
        if child is bm_start_el:
            collecting = True
            continue
        if child is bm_end_el:
            break
        if collecting and child.tag == p_tag:
            para_text = Paragraph(child, doc).text
            if para_text:
                texts.append(para_text)

    return "\n".join(texts)


def get_all_visible_text(doc: DocumentObject) -> str:
    """获取文档所有可见文本（段落文字）."""
    texts = []
    for para in doc.paragraphs:
        if para.text:
            texts.append(para.text)
    return "\n".join(texts)


# ---------------------------------------------------------------------------
# 5. 主流程
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("D1 段落锚点 POC 步骤 1：隐藏书签写入与验证")
    print("=" * 60)

    # Step 1: 创建合成附注 docx
    print("\n[1/5] 创建合成附注 docx（含 3 个 SECTION 块）...")
    buf = create_synthetic_note_docx()
    doc = Document(buf)
    blocks = scan_section_blocks(doc)
    print(f"  扫描到 {len(blocks)} 个 SECTION 块:")
    for b in blocks:
        print(f"    - {b.section_code}")
    assert len(blocks) == 3, f"期望 3 个块，实际 {len(blocks)}"

    # 记录写入前的可见文本
    text_before = get_all_visible_text(doc)

    # Step 2: 为每个 SECTION 块写入隐藏书签
    print("\n[2/5] 为每个 SECTION 块写入隐藏书签...")
    for i, block in enumerate(blocks):
        bm_name = section_code_to_bookmark_name(block.section_code)
        insert_bookmark_around_section_block(
            doc=doc,
            open_el=block.open_el,
            close_el=block.close_el,
            bookmark_name=bm_name,
            bookmark_id=i + 100,  # 避免与 Word 内部书签 ID 冲突
        )
        print(f"  ✓ {block.section_code} → bookmark name={bm_name}, id={i + 100}")

    # Step 3: 保存
    output_dir = Path(__file__).parent / "_poc_output"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "synthetic_note_with_bookmarks.docx"
    doc.save(str(output_path))
    print(f"\n[3/5] 保存至: {output_path}")

    # Step 4: 重新打开并验证
    print("\n[4/5] 重新打开并验证...")
    doc2 = Document(str(output_path))

    # 验证 A: 书签不可见（可见文本不变）
    text_after = get_all_visible_text(doc2)
    assert text_before == text_after, (
        "❌ 验证失败：写入书签后可见文本发生变化！\n"
        f"  变化前长度={len(text_before)}, 变化后长度={len(text_after)}"
    )
    print("  ✓ 验证 A 通过：书签不可见，可见文本未变化")

    # 验证 B: 可按 name 定位书签
    bookmarks = find_bookmarks(doc2)
    expected_names = ["sec_八_1", "sec_八_2", "sec_五_1"]
    for name in expected_names:
        assert name in bookmarks, f"❌ 验证失败：未找到书签 {name}"
    print(f"  ✓ 验证 B 通过：成功找到 {len(expected_names)} 个书签: {expected_names}")

    # 验证 C: 书签区间内文字可正确提取
    print("\n[5/5] 验证书签区间文字提取...")
    bm_start_el, bm_end_el = bookmarks["sec_八_1"]
    section_text = extract_text_between_bookmarks(doc2, bm_start_el, bm_end_el)
    # 区间内应包含章节标题和正文
    assert "固定资产" in section_text, f"❌ sec_八_1 区间未包含'固定资产': {section_text[:100]}"
    assert "初始计量" in section_text, f"❌ sec_八_1 区间未包含'初始计量': {section_text[:100]}"
    print(f"  ✓ sec_八_1 区间文字提取正确（含'固定资产'+'初始计量'）")

    bm_start_el, bm_end_el = bookmarks["sec_八_2"]
    section_text = extract_text_between_bookmarks(doc2, bm_start_el, bm_end_el)
    assert "无形资产" in section_text, f"❌ sec_八_2 区间未包含'无形资产': {section_text[:100]}"
    print(f"  ✓ sec_八_2 区间文字提取正确（含'无形资产'）")

    bm_start_el, bm_end_el = bookmarks["sec_五_1"]
    section_text = extract_text_between_bookmarks(doc2, bm_start_el, bm_end_el)
    assert "货币资金" in section_text, f"❌ sec_五_1 区间未包含'货币资金': {section_text[:100]}"
    print(f"  ✓ sec_五_1 区间文字提取正确（含'货币资金'）")

    # --- 最终结论 ---
    print("\n" + "=" * 60)
    print("🎉 POC 步骤 1 全部验证通过！")
    print("=" * 60)
    print("\n结论：")
    print("  ① 隐藏书签写入后不影响可见文本和排版 ✓")
    print("  ② python-docx 可按 bookmark name 回读定位 ✓")
    print("  ③ 书签区间可正确切出对应章节块文字 ✓")
    print(f"\n产物: {output_path}")
    print("\n下一步: 将 docx 上传到 OnlyOffice Docker 进行往返保真验证（步骤 2）")


if __name__ == "__main__":
    main()

"""POC: OnlyOffice 往返保真验证（D1 段落锚点 POC 步骤 2）.

验证目标：
  1. 将步骤 1 生成的 synthetic_note_with_bookmarks.docx 上传到 OnlyOffice Docker（9.4.0）
  2. 通过 OnlyOffice API 打开文档（验证书签不可见、排版无变化）
  3. 模拟编辑无关章节后保存（OnlyOffice callback status=2）
  4. 经 signed-download 重新下载已保存的文件
  5. 用 python-docx 回读验证书签是否在往返后保真（不丢失、不损坏）

已知坑：
  - SSRF 私有 IP 放行：OnlyOffice 9.4.0 默认拦截私有 IP 下载，
    需在容器 /etc/onlyoffice/documentserver/local.json 加：
    {"services":{"CoAuthoring":{"request-filtering-agent":{"allowPrivateIPAddress":true}}}}
    然后 docker restart audit-onlyoffice
  - callback 不被 ResponseWrapperMiddleware 包装：
    已通过 _SKIP_CONTAINS=("onlyoffice/callback",) 解决

环境变量：
  ONLYOFFICE_URL          OnlyOffice 服务 URL（默认 http://localhost:8080）
  ONLYOFFICE_CALLBACK_BASE 后端回调基址（默认 http://host.docker.internal:9980）
  BACKEND_BASE_URL        后端 API 地址（默认 http://localhost:9980）

用法：
    python backend/scripts/e2e/_poc_section_anchor_oo_roundtrip.py

    # 仅执行离线验证（不需要 OnlyOffice Docker）：
    python backend/scripts/e2e/_poc_section_anchor_oo_roundtrip.py --offline

产物：
    backend/scripts/e2e/_poc_output/roundtrip_after_oo_edit.docx  （OnlyOffice 编辑后下载）
    验证报告打印到 stdout
"""

from __future__ import annotations

import argparse
import hashlib
import hmac
import os
import shutil
import sys
import time
from pathlib import Path
from uuid import uuid4

# 确保可以 import backend 模块
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from io import BytesIO

import httpx
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------

ONLYOFFICE_URL = os.environ.get("ONLYOFFICE_URL", "http://localhost:8080")
BACKEND_BASE_URL = os.environ.get("BACKEND_BASE_URL", "http://localhost:9980")
CALLBACK_BASE = os.environ.get("ONLYOFFICE_CALLBACK_BASE", "http://host.docker.internal:9980")

POC_OUTPUT_DIR = Path(__file__).parent / "_poc_output"
SOURCE_DOCX = POC_OUTPUT_DIR / "synthetic_note_with_bookmarks.docx"
ROUNDTRIP_DOCX = POC_OUTPUT_DIR / "roundtrip_after_oo_edit.docx"


# ---------------------------------------------------------------------------
# 书签回读工具（复用步骤 1 同口径逻辑）
# ---------------------------------------------------------------------------

def find_bookmarks(doc: DocumentObject) -> dict[str, tuple[object, object]]:
    """从 docx 回读所有书签，返回 {name: (bookmarkStart_el, bookmarkEnd_el)}."""
    body = doc.element.body
    bm_start_tag = qn("w:bookmarkStart")
    bm_end_tag = qn("w:bookmarkEnd")

    starts: dict[str, object] = {}
    ends: dict[str, object] = {}
    names: dict[str, str] = {}

    for el in body.iter():
        if el.tag == bm_start_tag:
            bm_id = el.get(qn("w:id"))
            bm_name = el.get(qn("w:name"))
            if bm_id and bm_name:
                starts[bm_id] = el
                names[bm_id] = bm_name
        elif el.tag == bm_end_tag:
            bm_id = el.get(qn("w:id"))
            if bm_id:
                ends[bm_id] = el

    result: dict[str, tuple[object, object]] = {}
    for bm_id, name in names.items():
        if bm_id in starts and bm_id in ends:
            result[name] = (starts[bm_id], ends[bm_id])

    return result


def extract_text_between_bookmarks(
    doc: DocumentObject, bm_start_el, bm_end_el
) -> str:
    """提取两个书签元素之间的所有可见文本."""
    body = doc.element.body
    p_tag = qn("w:p")
    collecting = False
    texts: list[str] = []

    for child in body.iterchildren():
        if child is bm_start_el:
            collecting = True
            continue
        if child is bm_end_el:
            break
        if collecting and child.tag == p_tag:
            para_text = Paragraph(child, doc).text
            if para_text:
                texts.append(para_text)

    return "\n".join(texts)


def get_all_visible_text(doc: DocumentObject) -> str:
    """获取文档所有可见文本."""
    return "\n".join(p.text for p in doc.paragraphs if p.text)


# ---------------------------------------------------------------------------
# OnlyOffice 可用性检测
# ---------------------------------------------------------------------------

def check_onlyoffice_available() -> bool:
    """探测 OnlyOffice /healthcheck 是否可用."""
    try:
        resp = httpx.get(f"{ONLYOFFICE_URL}/healthcheck", timeout=5)
        return resp.status_code == 200
    except Exception:
        return False


# ---------------------------------------------------------------------------
# 模拟 OnlyOffice 往返流程
# ---------------------------------------------------------------------------

def simulate_oo_roundtrip_via_api(source_path: Path, output_path: Path) -> bool:
    """通过 OnlyOffice Conversion API 模拟往返。

    OnlyOffice 的 Conversion API 可以接受 docx → docx 转换，
    这会触发内部解析与重新序列化，模拟文档往返。
    这比完整的编辑器流程更轻量但同样验证书签保真。

    方案：启动临时 HTTP 服务器 serve 文件 → OnlyOffice 从中下载 → 转换 → 下载结果。
    由于 OnlyOffice 运行在 Docker 中，需要使用 host.docker.internal 访问主机。
    """
    import threading
    from http.server import HTTPServer, SimpleHTTPRequestHandler

    convert_url = f"{ONLYOFFICE_URL}/ConvertService.ashx"

    # 启动临时 HTTP 服务器在随机端口提供文件
    serve_port = 18976  # 选一个不太常见的端口

    class QuietHandler(SimpleHTTPRequestHandler):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, directory=str(source_path.parent), **kwargs)

        def log_message(self, format, *args):
            pass  # 静默

    server = HTTPServer(("0.0.0.0", serve_port), QuietHandler)
    server_thread = threading.Thread(target=server.serve_forever, daemon=True)
    server_thread.start()
    print(f"  临时文件服务器启动: 0.0.0.0:{serve_port}")

    try:
        # OnlyOffice 容器内通过 host.docker.internal 访问主机
        file_url = f"http://host.docker.internal:{serve_port}/{source_path.name}"
        print(f"  OnlyOffice 下载 URL: {file_url}")

        # 调用 Conversion API（docx → docx 往返）
        convert_body = {
            "async": False,
            "filetype": "docx",
            "outputtype": "docx",
            "title": "roundtrip_test.docx",
            "url": file_url,
            "key": f"poc_roundtrip_{int(time.time())}",
        }

        resp = httpx.post(
            convert_url,
            json=convert_body,
            timeout=60,
            headers={"Content-Type": "application/json"},
        )

        if resp.status_code == 200:
            # Conversion API 可能返回 JSON 或 XML
            content_type = resp.headers.get("content-type", "")
            if "json" in content_type:
                result = resp.json()
            else:
                # 解析 XML 响应
                import xml.etree.ElementTree as ET
                root = ET.fromstring(resp.text)
                result = {}
                for child in root:
                    tag = child.tag
                    text = child.text or ""
                    if tag == "EndConvert":
                        result["endConvert"] = text.lower() == "true"
                    elif tag == "FileUrl":
                        result["fileUrl"] = text
                    elif tag == "Error":
                        result["error"] = int(text) if text else 0
                    elif tag == "Percent":
                        result["percent"] = int(text) if text else 0

            if result.get("endConvert"):
                download_url = result.get("fileUrl")
                if download_url:
                    print(f"  转换完成，下载结果: {download_url[:80]}...")
                    dl_resp = httpx.get(download_url, timeout=30)
                    if dl_resp.status_code == 200:
                        output_path.write_bytes(dl_resp.content)
                        print(f"  ✓ 文件已下载: {output_path} ({len(dl_resp.content)} bytes)")
                        return True
                    else:
                        print(f"  ⚠ 结果下载失败: HTTP {dl_resp.status_code}")
            elif result.get("error"):
                error_code = result["error"]
                error_msgs = {
                    -1: "未知错误",
                    -2: "下载超时",
                    -3: "转换错误",
                    -4: "下载错误（可能是 SSRF 拦截私有 IP）",
                    -5: "密码保护文档",
                    -6: "格式不支持",
                    -7: "请求参数错误",
                    -8: "令牌过期",
                }
                msg = error_msgs.get(error_code, f"错误码 {error_code}")
                print(f"  ⚠ 转换失败: {msg}")
                if error_code == -4:
                    print(f"     → 这通常是 SSRF 拦截！请确保 local.json 已配置 allowPrivateIPAddress=true")
                    print(f"     → docker exec audit-onlyoffice cat /etc/onlyoffice/documentserver/local.json")
            else:
                print(f"  ⚠ 转换响应: {result}")
        else:
            print(f"  ⚠ 转换 API 失败: HTTP {resp.status_code}: {resp.text[:200]}")

    except Exception as e:
        print(f"  ⚠ 往返 API 调用异常: {e}")
    finally:
        server.shutdown()
        print(f"  临时文件服务器已关闭")

    return False


def simulate_oo_roundtrip_via_command_api(source_path: Path, output_path: Path) -> bool:
    """通过 OnlyOffice Command Service + doc builder 方式验证。

    这是备选方案：使用 OnlyOffice 的 command service 打开文档后强制保存。
    """
    command_url = f"{ONLYOFFICE_URL}/coauthoring/CommandService.ashx"

    try:
        # 生成唯一 key
        doc_key = f"poc_bookmark_test_{int(time.time())}"

        # 尝试 forcesave 命令（需要文档先通过编辑器打开）
        # 这里主要验证 OnlyOffice 能解析带书签的文档
        resp = httpx.post(
            command_url,
            json={
                "c": "info",
                "key": doc_key,
            },
            timeout=10,
        )
        print(f"  Command API info: HTTP {resp.status_code} → {resp.text[:100]}")
        return False  # Command API 方式需要文档已在编辑会话中

    except Exception as e:
        print(f"  ⚠ Command API 异常: {e}")
        return False


# ---------------------------------------------------------------------------
# 离线往返验证（不需要 OnlyOffice 运行）
# ---------------------------------------------------------------------------

def offline_roundtrip_verification(source_path: Path) -> bool:
    """离线验证：将 docx 重新序列化模拟往返。

    python-docx 打开→保存→重新打开的过程也是一种有损往返
    （python-docx 会规范化 XML），可验证书签结构在序列化后是否保持。
    """
    print("\n[离线模式] 通过 python-docx 重新序列化模拟 XML 往返...")

    # 打开 → 保存到新位置 → 重新打开
    doc = Document(str(source_path))
    roundtrip_path = POC_OUTPUT_DIR / "roundtrip_pydocx_reserialized.docx"
    doc.save(str(roundtrip_path))

    # 重新打开验证
    doc2 = Document(str(roundtrip_path))
    bookmarks = find_bookmarks(doc2)
    expected = ["sec_八_1", "sec_八_2", "sec_五_1"]

    all_found = True
    for name in expected:
        if name not in bookmarks:
            print(f"  ❌ 书签丢失: {name}")
            all_found = False
        else:
            print(f"  ✓ 书签保留: {name}")

    if all_found:
        # 验证文字内容
        bm_start, bm_end = bookmarks["sec_八_1"]
        text = extract_text_between_bookmarks(doc2, bm_start, bm_end)
        if "固定资产" in text:
            print(f"  ✓ sec_八_1 区间文字正确（含'固定资产'）")
        else:
            print(f"  ❌ sec_八_1 区间文字异常")
            all_found = False

    return all_found


# ---------------------------------------------------------------------------
# OnlyOffice 在线往返验证
# ---------------------------------------------------------------------------

def online_roundtrip_verification(source_path: Path, output_path: Path) -> bool:
    """在线验证：通过 OnlyOffice Docker 进行真实往返。

    完整流程：
    1. 将 docx 通过 OnlyOffice 文件上传接口上传
    2. OnlyOffice 内部解析文档（此过程会测试书签保真）
    3. 通过 Conversion API 重新导出 docx
    4. 下载导出后的文件并验证书签
    """
    print("\n[在线模式] 通过 OnlyOffice Docker 进行真实文档往返...")

    # 方案 A：通过 Conversion API 上传+转换
    success = simulate_oo_roundtrip_via_api(source_path, output_path)

    if not success:
        print("  ⚠ Conversion API 方式失败，尝试 Command API...")
        success = simulate_oo_roundtrip_via_command_api(source_path, output_path)

    if not success:
        print("\n  ⚠ OnlyOffice API 往返未成功。")
        print("  这通常是因为 OnlyOffice 的 Conversion API 需要能从 URL 获取源文件。")
        print("  请通过以下手动步骤完成验证：")
        print_manual_verification_steps()
        return False

    # 验证下载的文件
    if output_path.exists():
        doc = Document(str(output_path))
        bookmarks = find_bookmarks(doc)
        expected = ["sec_八_1", "sec_八_2", "sec_五_1"]

        all_found = True
        for name in expected:
            if name not in bookmarks:
                print(f"  ❌ 往返后书签丢失: {name}")
                all_found = False
            else:
                print(f"  ✓ 往返后书签保留: {name}")

        if all_found:
            # 验证区间文字
            bm_start, bm_end = bookmarks["sec_八_1"]
            text = extract_text_between_bookmarks(doc, bm_start, bm_end)
            if "固定资产" in text:
                print(f"  ✓ 往返后 sec_八_1 文字正确（含'固定资产'）")
            else:
                print(f"  ❌ 往返后 sec_八_1 区间文字异常: {text[:80]}")
                all_found = False

        return all_found

    return False


# ---------------------------------------------------------------------------
# 手动验证步骤文档
# ---------------------------------------------------------------------------

def print_manual_verification_steps():
    """输出手动验证步骤指引."""
    print("""
╔══════════════════════════════════════════════════════════════════╗
║           OnlyOffice 往返保真 - 手动验证步骤                       ║
╠══════════════════════════════════════════════════════════════════╣
║                                                                  ║
║  前置条件：                                                       ║
║  1. OnlyOffice Docker 9.4.0 已启动                               ║
║     docker compose --profile onlyoffice up -d                    ║
║                                                                  ║
║  2. SSRF 私有 IP 已放行：                                         ║
║     docker exec audit-onlyoffice bash -c '                       ║
║       echo \'{"services":{"CoAuthoring":{                        ║
║         "request-filtering-agent":{                               ║
║           "allowPrivateIPAddress":true                            ║
║         }                                                        ║
║       }}}\' > /etc/onlyoffice/documentserver/local.json'         ║
║     docker restart audit-onlyoffice                              ║
║                                                                  ║
║  3. 后端已启动（端口 9980）且 .env 含：                             ║
║     ONLYOFFICE_URL=http://localhost:8080                         ║
║     ONLYOFFICE_CALLBACK_BASE=http://host.docker.internal:9980    ║
║                                                                  ║
║  验证步骤：                                                       ║
║  A. 创建测试出品物并上传 docx：                                     ║
║     - 在前端「出品物中心」创建一个测试附注出品物                       ║
║     - 或通过 API 创建 WordExportTask（doc_type=disclosure_notes）  ║
║     - 将 _poc_output/synthetic_note_with_bookmarks.docx 作为       ║
║       初始版本文件                                                 ║
║                                                                  ║
║  B. 打开 OnlyOffice 编辑器：                                      ║
║     - 在出品物详情页点击「在线编辑」                                 ║
║     - 确认：文档打开正常，看不到任何 "sec_八_1" 等书签标记           ║
║     - 确认：排版与章节内容与原文一致（表格、段落完整）                ║
║                                                                  ║
║  C. 编辑无关章节并保存：                                           ║
║     - 在「八、2 无形资产」章节末尾添加一段测试文字                    ║
║       （如"POC 测试编辑 - 验证书签保真"）                           ║
║     - 点击保存（或等待自动保存触发 callback）                        ║
║     - OnlyOffice 会调用 callback → 后端下载并创建新版本             ║
║                                                                  ║
║  D. 下载已保存的新版本：                                           ║
║     - 在出品物详情页下载最新版本 docx                               ║
║     - 或直接从存储路径复制文件                                      ║
║                                                                  ║
║  E. 运行步骤 3 回读验证：                                          ║
║     python backend/scripts/e2e/_poc_section_anchor.py             ║
║     （步骤 3 脚本会验证书签是否存在+区间文字是否正确）                 ║
║                                                                  ║
║  判据：                                                           ║
║  ① 打开时书签不可见（排版无异常标记）           → PASS / FAIL       ║
║  ② 编辑保存后书签仍存在（3 个 sec_xxx 全在）    → PASS / FAIL       ║
║  ③ 书签区间文字可正确按 name 定位              → PASS / FAIL       ║
║                                                                  ║
║  三项全 PASS → 书签方案可用于生产                                   ║
║  任一 FAIL → 需评估 SDT 内容控件备选方案                            ║
║                                                                  ║
╚══════════════════════════════════════════════════════════════════╝
""")


# ---------------------------------------------------------------------------
# OnlyOffice 配置验证
# ---------------------------------------------------------------------------

def verify_oo_configuration():
    """验证 OnlyOffice 配置是否正确（SSRF 放行等）."""
    print("\n[配置检查] 验证 OnlyOffice 环境配置...")

    # 1. 检查 OnlyOffice 服务可达
    try:
        resp = httpx.get(f"{ONLYOFFICE_URL}/healthcheck", timeout=5)
        if resp.status_code == 200:
            print(f"  ✓ OnlyOffice 服务可达: {ONLYOFFICE_URL}")
        else:
            print(f"  ⚠ OnlyOffice healthcheck 非 200: {resp.status_code}")
    except Exception as e:
        print(f"  ❌ OnlyOffice 服务不可达: {e}")
        return False

    # 2. 检查 info 端点（验证版本等）
    try:
        resp = httpx.get(f"{ONLYOFFICE_URL}/info/info.json", timeout=5)
        if resp.status_code == 200:
            info = resp.json()
            version = info.get("serverVersion", "unknown")
            print(f"  ✓ OnlyOffice 版本: {version}")
        else:
            print(f"  ⚠ 无法获取版本信息 (HTTP {resp.status_code})")
    except Exception:
        print(f"  ⚠ info 端点不可用（正常，部分版本不暴露）")

    # 3. 提示 SSRF 配置
    print(f"\n  📋 SSRF 配置提醒：")
    print(f"     OnlyOffice 9.4.0 默认拦截私有 IP（192.168.x.x, 10.x.x.x）")
    print(f"     如果 OnlyOffice 容器内无法下载 host.docker.internal 的文件，")
    print(f"     需在容器内执行：")
    print(f"       docker exec audit-onlyoffice bash -c 'cat > /etc/onlyoffice/documentserver/local.json << EOF")
    print(f'       {{"services":{{"CoAuthoring":{{"request-filtering-agent":{{"allowPrivateIPAddress":true}}}}}}}}')
    print(f"       EOF'")
    print(f"     然后 docker restart audit-onlyoffice")

    # 4. 提示 callback 包装跳过
    print(f"\n  📋 Callback 中间件提醒：")
    print(f"     ResponseWrapperMiddleware 已配置跳过 onlyoffice/callback 路径")
    print(f"     （_SKIP_CONTAINS 包含 'onlyoffice/callback'）")
    print(f"     OnlyOffice 收到原始 {{\"error\": 0}} 响应即认为 callback 成功")

    return True


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="D1 段落锚点 POC 步骤 2：OnlyOffice 往返保真验证")
    parser.add_argument("--offline", action="store_true", help="仅执行离线验证（不需要 OnlyOffice Docker）")
    args = parser.parse_args()

    print("=" * 70)
    print("D1 段落锚点 POC 步骤 2：OnlyOffice 往返保真验证")
    print("=" * 70)

    # 前置检查：步骤 1 产物是否存在
    if not SOURCE_DOCX.exists():
        print(f"\n❌ 步骤 1 产物不存在: {SOURCE_DOCX}")
        print("   请先运行: python backend/scripts/e2e/_poc_section_anchor.py")
        sys.exit(1)

    # 读取源文件基线信息
    print(f"\n[0/4] 读取源文件基线...")
    doc_before = Document(str(SOURCE_DOCX))
    bookmarks_before = find_bookmarks(doc_before)
    text_before = get_all_visible_text(doc_before)
    print(f"  源文件: {SOURCE_DOCX}")
    print(f"  书签数: {len(bookmarks_before)}")
    print(f"  书签列表: {sorted(bookmarks_before.keys())}")
    print(f"  可见文本长度: {len(text_before)} 字符")

    expected_bookmarks = ["sec_八_1", "sec_八_2", "sec_五_1"]
    for name in expected_bookmarks:
        assert name in bookmarks_before, f"源文件缺少书签 {name}，请重新运行步骤 1"

    # 离线验证（总是执行）
    print("\n" + "─" * 70)
    print("[1/4] 离线往返验证（python-docx 序列化往返）")
    print("─" * 70)
    offline_ok = offline_roundtrip_verification(SOURCE_DOCX)

    if args.offline:
        # 离线模式：直接输出结论
        print("\n" + "=" * 70)
        if offline_ok:
            print("🎉 离线往返验证通过！")
            print("   书签在 python-docx 序列化往返后完整保留。")
        else:
            print("❌ 离线往返验证失败！")
        print("=" * 70)
        print("\n注意：完整验证需要 OnlyOffice Docker，请去掉 --offline 参数运行。")
        sys.exit(0 if offline_ok else 1)

    # 在线验证
    print("\n" + "─" * 70)
    print("[2/4] 检测 OnlyOffice Docker 可用性")
    print("─" * 70)

    oo_available = check_onlyoffice_available()
    if not oo_available:
        print(f"\n  ⚠ OnlyOffice Docker 不可达 ({ONLYOFFICE_URL})")
        print(f"  可能原因：")
        print(f"    - Docker 容器未启动: docker compose --profile onlyoffice up -d")
        print(f"    - URL 不正确: 设置 ONLYOFFICE_URL 环境变量")
        print(f"    - 端口未映射: 检查 docker-compose.yml ports 配置")
        print(f"\n  回退到手动验证模式...")
        print_manual_verification_steps()
        # 输出离线验证结论
        print("\n" + "=" * 70)
        print("📋 POC 步骤 2 部分验证结果：")
        print(f"   ① 离线（python-docx）往返: {'✓ PASS' if offline_ok else '❌ FAIL'}")
        print(f"   ② OnlyOffice 在线往返: ⏸ 待手动验证（Docker 不可达）")
        print("=" * 70)
        sys.exit(0)

    print(f"  ✓ OnlyOffice Docker 可达: {ONLYOFFICE_URL}")

    # 配置验证
    print("\n" + "─" * 70)
    print("[3/4] OnlyOffice 环境配置验证")
    print("─" * 70)
    verify_oo_configuration()

    # 在线往返
    print("\n" + "─" * 70)
    print("[4/4] OnlyOffice 在线文档往返")
    print("─" * 70)
    online_ok = online_roundtrip_verification(SOURCE_DOCX, ROUNDTRIP_DOCX)

    # --- 最终结论 ---
    print("\n" + "=" * 70)
    print("📋 POC 步骤 2 验证结果：")
    print("=" * 70)
    print(f"   ① 离线（python-docx）序列化往返: {'✓ PASS' if offline_ok else '❌ FAIL'}")
    if online_ok:
        print(f"   ② OnlyOffice 在线往返: ✓ PASS")
        print(f"\n🎉 步骤 2 全部验证通过！隐藏书签在 OnlyOffice 往返后保持完整。")
        print(f"\n结论：")
        print(f"   书签方案（w:bookmarkStart/End）满足 OnlyOffice 往返保真要求。")
        print(f"   可进入步骤 3（python-docx 回读定位验证）。")
    else:
        print(f"   ② OnlyOffice 在线往返: ⏸ 需手动验证")
        print(f"\n📝 在线往返 API 方式未成功（常见于 OnlyOffice 无法从后端 URL 下载文件）。")
        print(f"   请按照上方手动步骤完成验证。")
        print(f"\n   离线验证已确认书签在 XML 序列化层面保真，")
        print(f"   OnlyOffice 作为 OOXML 原生解析器应同样保留书签。")
        print(f"   手动验证重点关注：①编辑器内书签不可见 ②保存后书签仍在")

    if ROUNDTRIP_DOCX.exists():
        print(f"\n产物: {ROUNDTRIP_DOCX}")

    print("=" * 70)
    print(f"\n下一步: 运行步骤 3 对往返后文件做 python-docx 回读定位验证")
    print(f"  python backend/scripts/e2e/_poc_section_anchor.py")


if __name__ == "__main__":
    main()

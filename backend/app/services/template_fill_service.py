"""报告正文模板填充服务（两阶段 preview/confirm 流水线）.

audit-report-template-integration Phase 2 / task 7。

设计参考 design.md §4 / §4.2 / §5 / §6.8：
- preview：resolve company_subtype → manifest → copy 模板 → 替换占位符 →
  扫描 OPT → 写 ``fill_preview_sessions`` 行（不落库交付件）。
- confirm：载入 session（校验未过期 + user_id）→ apply_optional_sections →
  保存 guidance 副本（with_notes_v{n}.docx）→ strip_guidance_notes →
  ``DeliverableService.export_or_new_deliverable`` + ``render_and_store``
  （version_no 递增）→ 更新 ``audit_report.report_body_json`` → 删除 session。

铁律：本服务仅 ``flush`` 不 ``commit``（router 统一 commit）；异步 AsyncSession。
旧 ``word_template_filler.py`` 暂保留（task 17 下线），本服务不修改它。
"""

from __future__ import annotations

import logging
import shutil
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from uuid import UUID

import sqlalchemy as sa
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.phase13_models import WordExportDocType
from app.models.report_models import (
    AuditReport,
    CompanyType,
    FillPreviewSession,
    OpinionType,
)
from app.services.deliverable_service import STORAGE_ROOT, DeliverableService
from app.services.placeholder_registry import (
    PlaceholderRegistry,
    get_placeholder_registry,
)
from app.services.report_body_service import ReportBodyService
from app.services.template_manifest_loader import (
    TemplateManifestLoader,
    get_template_manifest_loader,
)
from app.services.word_doc_utils import (
    apply_optional_sections,
    replace_placeholders_in_doc,
    scan_optional_sections,
    strip_guidance_notes,
)

logger = logging.getLogger(__name__)

PREVIEW_SUBDIR = "preview"

# disclaimer（无法表示意见）映射到单一通用模板键（manifest report_body.disclaimer._all）
DISCLAIMER_SUBTYPE_KEY = "_all"


@dataclass
class OptionalSectionView:
    """preview 返回的单个可选段落视图（对应 design §11 Response）。"""

    section_id: str
    description: str
    preview: str
    default_keep: bool
    group: str


@dataclass
class ReportBodyPreviewResult:
    """preview_report_body 结果（不落库交付件）。"""

    preview_session_id: UUID
    optional_sections: list[OptionalSectionView]
    missing_fields: list[str]
    template_version: str
    company_subtype_resolved: str


@dataclass
class ReportBodyConfirmResult:
    """confirm_report_body 结果（与现有 ReportBodyRenderResponse 兼容）。"""

    task_id: UUID
    version_no: int
    download_url: str
    report_body_json: dict
    validation_warning: str | None = None
    guidance_version_path: str | None = None


def validate_kam_word_mode(
    *,
    optional_sections: dict[str, bool],
    company_type: str,
    is_pie: bool,
    opinion_type: str,
    report_body_service: ReportBodyService | None = None,
) -> str | None:
    """KAM 合规校验（Word 模式，design §4.2）。

    步骤：
    1. 复用 ``ReportBodyService.kam_required(...)``；若 False → 返回 None。
    2. 若 ``optional_sections.get("key_audit_matters") is False`` → 返回 KAM 缺失警告。
    3. 警告文案复用现 ``ReportBodyService.validate_kam`` 的措辞。

    **不**扫描 docx 正文判空（选用决策已在 OPT 弹窗完成；段落内容空由 EQCR 复核发现）。
    """
    svc = report_body_service or ReportBodyService.__new__(ReportBodyService)
    required = ReportBodyService.kam_required(
        svc,
        company_type=company_type,
        is_pie=is_pie,
        opinion_type=opinion_type,
    )
    if not required:
        return None
    if optional_sections.get("key_audit_matters") is False:
        # 文案与 ReportBodyService.validate_kam 一致
        return (
            "上市公司或公共利益实体审计报告必须包含至少一个关键审计事项(KAM)，"
            "请编辑「关键审计事项段」后再定稿"
        )
    return None


class TemplateFillService:
    """报告正文 Word 填充流水线（preview / confirm）。"""

    def __init__(
        self,
        db: AsyncSession,
        loader: TemplateManifestLoader | None = None,
        registry: PlaceholderRegistry | None = None,
    ):
        self.db = db
        self.loader = loader or get_template_manifest_loader()
        self.registry = registry or get_placeholder_registry()

    # ------------------------------------------------------------------
    # company_subtype 解析
    # ------------------------------------------------------------------
    async def _resolve_company_subtype(
        self,
        project_id: UUID,
        explicit: str | None,
    ) -> str:
        """解析 company_subtype。

        优先级（design §4 + 需求 7.7/7.8）：
        ① explicit（API 显式传入，用户当前选择）
        ② project.company_subtype（用户已确认/手动修改，需求 7.8 优先于自动推断）
        ③ matching_rules 规则推荐（需求 7.7：优先于 1.4 fallback）
        ④ fallback（需求 1.4：listed→type_a / 其余→type_d）。
        """
        if explicit:
            return explicit.strip().lower()

        from app.models.core import Project

        project = await self.db.get(Project, project_id)
        if project is not None and project.company_subtype:
            return project.company_subtype.strip().lower()

        # ③ matching_rules 规则推荐（需求 7.7 优先于 1.4 fallback）。
        # 仅在有明确推荐（rule 命中，confidence != none）时采用，避免默认值掩盖 fallback。
        report = await self._get_any_report(project_id)
        company_type_val: str | None = None
        if report is not None:
            company_type_val = (
                report.company_type.value
                if isinstance(report.company_type, CompanyType)
                else str(report.company_type)
            )

        if project is not None:
            from app.services.matching_rules_service import recommend_company_subtype

            std = project.applicable_standard_v2 or {}
            rec = recommend_company_subtype(
                {
                    "entity_type": std.get("entity_type"),
                    "scope": std.get("scope"),
                    "scenario": project.scenario,
                    "template_type": project.template_type,
                    "report_scope": project.report_scope,
                    "company_name": project.name,
                    "client_name": project.client_name,
                    "company_type": company_type_val,
                    "applicable_standard_v2": project.applicable_standard_v2,
                }
            )
            # 规则命中（rule/fallback 来源且有候选）→ 采用推荐；仅 default 来源时不采用，
            # 留给下方 company_type 兜底（保持原 1.4 语义一致）。
            if rec.subtype and rec.source in ("rule", "fallback"):
                return rec.subtype

        # ④ fallback（需求 1.4）：按 company_type 兜底 listed→type_a / 其余→type_d。
        # company_type 权威来源是 audit_report（project 表无该枚举）。
        return "type_a" if company_type_val == CompanyType.listed.value else "type_d"

    async def _get_any_report(self, project_id: UUID) -> AuditReport | None:
        """取项目最近一条 audit_report（仅用于 company_type 兜底推断）。"""
        result = await self.db.execute(
            sa.select(AuditReport)
            .where(
                AuditReport.project_id == project_id,
                AuditReport.is_deleted == sa.false(),
            )
            .order_by(AuditReport.year.desc())
            .limit(1)
        )
        return result.scalar_one_or_none()

    async def _resolve_report_scope(self, project_id: UUID) -> str:
        """解析项目报表口径（consolidated/standalone），默认 consolidated。

        权威源 project.report_scope；wizard_state.basic_info 回退；缺失默认合并。
        """
        from app.models.core import Project

        project = await self.db.get(Project, project_id)
        if project is None:
            return "consolidated"
        raw = (getattr(project, "report_scope", None) or "").strip().lower()
        if raw not in ("consolidated", "standalone"):
            basic = (project.wizard_state or {}).get("basic_info", {}).get("data", {})
            raw = str(basic.get("report_scope") or "").strip().lower()
        return raw if raw in ("consolidated", "standalone") else "consolidated"

    def _manifest_subtype_key(self, opinion_type: str, company_subtype: str) -> str:
        """disclaimer 意见映射到单一通用模板键 ``_all``（manifest）。"""
        if opinion_type == OpinionType.disclaimer.value:
            return DISCLAIMER_SUBTYPE_KEY
        return company_subtype

    # ------------------------------------------------------------------
    # preview
    # ------------------------------------------------------------------
    async def preview_report_body(
        self,
        project_id: UUID,
        year: int,
        *,
        opinion_type: str,
        company_subtype: str | None,
        template_variant: str = "simple",
        user_id: UUID,
    ) -> ReportBodyPreviewResult:
        """preview 阶段：copy 模板 → 替换占位符 → 扫描 OPT → 写 session（不落库）。"""
        resolved_subtype = await self._resolve_company_subtype(
            project_id, company_subtype
        )
        manifest_key = self._manifest_subtype_key(opinion_type, resolved_subtype)
        report_scope = await self._resolve_report_scope(project_id)

        entry = self.loader.resolve_report_body(
            opinion_type, manifest_key, template_variant, report_scope
        )
        if not entry.exists:
            raise ValueError(
                f"报告正文模板文件缺失: {entry.abs_path}"
                f"（opinion_type={opinion_type}, company_subtype={manifest_key}, "
                f"variant={template_variant}）"
            )

        # preview session id 先生成，作为 preview 工作目录名
        session = FillPreviewSession(
            project_id=project_id,
            user_id=user_id,
            year=year,
            opinion_type=opinion_type,
            company_subtype=resolved_subtype,
            template_variant=template_variant,
            template_version=self.loader.version(),
            expires_at=self._compute_expiry(),
        )
        self.db.add(session)
        await self.db.flush()  # 取得 session.id

        # copy 模板到 preview 工作目录
        work_dir = self._preview_dir(session.id)
        work_dir.mkdir(parents=True, exist_ok=True)
        working_path = work_dir / "working.docx"
        shutil.copy2(entry.abs_path, working_path)

        # 占位符替换
        doc = Document(str(working_path))
        mapping = await self.registry.build_placeholder_map(project_id, self.db)
        replace_placeholders_in_doc(doc, mapping)
        doc.save(str(working_path))

        # 扫描 OPT + 待补充字段
        opt_views = self._build_optional_views(doc, resolved_subtype)
        missing_fields = self.registry.detect_missing_fields(mapping)

        # 写回 session 缓存
        session.working_path = str(working_path)
        session.optional_sections_json = {
            v.section_id: {
                "description": v.description,
                "preview": v.preview,
                "default_keep": v.default_keep,
                "group": v.group,
            }
            for v in opt_views
        }
        session.missing_fields = {"fields": missing_fields}
        await self.db.flush()

        return ReportBodyPreviewResult(
            preview_session_id=session.id,
            optional_sections=opt_views,
            missing_fields=missing_fields,
            template_version=self.loader.version(),
            company_subtype_resolved=resolved_subtype,
        )

    def _build_optional_views(
        self, doc: Document, company_subtype: str
    ) -> list[OptionalSectionView]:
        """扫描 OPT 块 + 结合 opt_defaults / section_group_map 组装视图。"""
        sections = scan_optional_sections(doc)
        opt_defaults = self.registry.get_opt_defaults(company_subtype)
        group_map = self.registry.get_section_group_map()
        group_labels = self._opt_group_labels()

        paras = doc.paragraphs
        views: list[OptionalSectionView] = []
        for sec in sections:
            # 预览取块内首个非标记段落前 50 字
            preview_text = ""
            for idx in range(sec.start_index + 1, min(sec.end_index - 1, len(paras))):
                text = (paras[idx].text or "").strip()
                if text:
                    preview_text = text[:50]
                    break
            group_key = group_map.get(sec.section_id, "body")
            views.append(
                OptionalSectionView(
                    section_id=sec.section_id,
                    description=sec.description,
                    preview=preview_text,
                    default_keep=bool(opt_defaults.get(sec.section_id, False)),
                    group=group_labels.get(group_key, group_key),
                )
            )
        return views

    def _opt_group_labels(self) -> dict[str, str]:
        """opt_groups（group_key → 中文标题）。registry 未暴露则用默认。"""
        data = getattr(self.registry, "_data", {}) or {}
        groups = data.get("opt_groups")
        if isinstance(groups, dict) and groups:
            return {str(k): str(v) for k, v in groups.items()}
        return {"body": "报告正文段落", "supplement": "补充信息段落"}

    # ------------------------------------------------------------------
    # confirm
    # ------------------------------------------------------------------
    async def confirm_report_body(
        self,
        project_id: UUID,
        year: int,
        *,
        preview_session_id: UUID,
        optional_sections: dict[str, bool],
        user_id: UUID,
    ) -> ReportBodyConfirmResult:
        """confirm 阶段：应用 OPT → guidance 副本 → 清理 NOTE → 入库 → 更新 schema。"""
        session = await self.db.get(FillPreviewSession, preview_session_id)
        if session is None:
            raise ValueError(f"preview 会话不存在或已过期: {preview_session_id}")
        if session.project_id != project_id:
            raise ValueError("preview 会话与项目不匹配")
        if session.user_id != user_id:
            raise ValueError("preview 会话不属于当前用户")
        if self._is_expired(session):
            await self._purge_session(session)
            raise ValueError("preview 会话已过期，请重新生成")

        working_path = Path(session.working_path) if session.working_path else None
        if working_path is None or not working_path.exists():
            raise ValueError("preview 工作副本缺失，请重新生成")

        # 1. 确定交付件 task（终态再导出 → 新建独立交付物）
        deliverable_svc = DeliverableService(self.db)
        task, _is_new = await deliverable_svc.export_or_new_deliverable(
            project_id,
            WordExportDocType.audit_report.value,
            session.template_variant,
            user_id,
        )

        # 2. 应用可选段落（删除未勾选块，保留块去标记）
        doc = Document(str(working_path))
        apply_optional_sections(doc, optional_sections)

        # 3. 保存 guidance 副本（含 NOTE，剥除前）到交付件目录
        version_no = await self._next_version_no(deliverable_svc, task.id)
        deliverable_dir = (
            STORAGE_ROOT / "deliverables" / str(project_id) / str(task.id)
        )
        deliverable_dir.mkdir(parents=True, exist_ok=True)
        guidance_path = deliverable_dir / f"with_notes_v{version_no}.docx"
        doc.save(str(guidance_path))

        # 4. 剥除 NOTE 指引注释（工作副本）
        strip_guidance_notes(doc)
        clean_path = self._preview_dir(session.id) / "clean.docx"
        doc.save(str(clean_path))

        # 5. 入库：render_and_store（version_no 递增）
        store_result = await deliverable_svc.render_and_store(
            task.id,
            docx_path=clean_path,
            user_id=user_id,
            selected_sections=[
                sid for sid, keep in optional_sections.items() if keep
            ],
            file_name=f"audit_report_v{version_no}.docx",
            created_via="generate",
        )

        # 6. KAM 校验（Word 模式）
        validation_warning = await self._validate_kam(
            project_id, year, optional_sections
        )

        # 7. 更新 audit_report.report_body_json schema（需求 6.8）
        report_body_json = await self._update_report_body_json(
            project_id,
            year,
            optional_sections=optional_sections,
            guidance_version_path=str(guidance_path),
            template_version=session.template_version or self.loader.version(),
            company_subtype=session.company_subtype,
            template_variant=session.template_variant,
            missing_fields=(session.missing_fields or {}).get("fields", []),
        )

        # 8. 删除 preview session（TTL 清理：confirm 即删）
        await self._purge_session(session)

        return ReportBodyConfirmResult(
            task_id=task.id,
            version_no=store_result.version.version_no,
            download_url=store_result.download_url,
            report_body_json=report_body_json,
            validation_warning=validation_warning,
            guidance_version_path=str(guidance_path),
        )

    async def _next_version_no(
        self, deliverable_svc: DeliverableService, task_id: UUID
    ) -> int:
        """计算本次 confirm 将产生的 version_no（与 render_and_store 一致）。"""
        latest = await deliverable_svc._latest_version(task_id)
        return (latest.version_no + 1) if latest else 1

    async def _validate_kam(
        self,
        project_id: UUID,
        year: int,
        optional_sections: dict[str, bool],
    ) -> str | None:
        """读取 audit_report 的 company_type/is_pie/opinion_type 后执行 KAM 校验。"""
        report = await self._get_report(project_id, year)
        if report is None:
            # 无 audit_report 行时无法判定，跳过校验（不阻断）
            return None
        return validate_kam_word_mode(
            optional_sections=optional_sections,
            company_type=report.company_type.value
            if isinstance(report.company_type, CompanyType)
            else str(report.company_type),
            is_pie=bool(report.is_pie),
            opinion_type=report.opinion_type.value
            if isinstance(report.opinion_type, OpinionType)
            else str(report.opinion_type),
            report_body_service=ReportBodyService(self.db),
        )

    # ------------------------------------------------------------------
    # audit_report.report_body_json 更新（需求 6.8 schema）
    # ------------------------------------------------------------------
    async def _update_report_body_json(
        self,
        project_id: UUID,
        year: int,
        *,
        optional_sections: dict[str, bool],
        guidance_version_path: str,
        template_version: str,
        company_subtype: str | None,
        template_variant: str | None,
        missing_fields: list[str],
    ) -> dict:
        """写入需求 6.8 schema 到 audit_report.report_body_json。"""
        report = await self._get_report(project_id, year)
        body_json = {
            "optional_sections": optional_sections,
            "guidance_version_path": guidance_version_path,
            "template_version": template_version,
            "company_subtype": company_subtype,
            "template_variant": template_variant,
            "missing_fields": missing_fields,
        }
        if report is not None:
            report.report_body_json = body_json
            # 同步 Word 模式相关元数据列
            if company_subtype:
                report.company_subtype = company_subtype
            if template_variant:
                report.template_variant = template_variant
            report.template_version = template_version
            await self.db.flush()
        return body_json

    async def _get_report(self, project_id: UUID, year: int) -> AuditReport | None:
        result = await self.db.execute(
            sa.select(AuditReport).where(
                AuditReport.project_id == project_id,
                AuditReport.year == year,
                AuditReport.is_deleted == sa.false(),
            )
        )
        return result.scalar_one_or_none()

    # ------------------------------------------------------------------
    # preview session TTL 清理
    # ------------------------------------------------------------------
    def _compute_expiry(self) -> datetime:
        hours = int(getattr(settings, "FILL_PREVIEW_TTL_HOURS", 24) or 24)
        return datetime.now(timezone.utc) + timedelta(hours=hours)

    @staticmethod
    def _is_expired(session: FillPreviewSession) -> bool:
        if session.expires_at is None:
            return False
        expires = session.expires_at
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        return datetime.now(timezone.utc) > expires

    async def _purge_session(self, session: FillPreviewSession) -> None:
        """删除单个 preview session（DB 行 + 工作目录）。"""
        work_dir = self._preview_dir(session.id)
        await self.db.delete(session)
        await self.db.flush()
        self._safe_rmtree(work_dir)

    async def purge_expired_sessions(self) -> int:
        """清理所有已过期 preview session（DB 行 + 工作目录）。

        可由定时任务调用；本服务不自行注册 scheduler（design §5）。
        返回清理的会话数。
        """
        now = datetime.now(timezone.utc)
        result = await self.db.execute(
            sa.select(FillPreviewSession).where(
                FillPreviewSession.expires_at.is_not(None),
                FillPreviewSession.expires_at < now,
            )
        )
        sessions = list(result.scalars().all())
        for session in sessions:
            work_dir = self._preview_dir(session.id)
            await self.db.delete(session)
            self._safe_rmtree(work_dir)
        if sessions:
            await self.db.flush()
        return len(sessions)

    # ------------------------------------------------------------------
    # 路径辅助
    # ------------------------------------------------------------------
    @staticmethod
    def _preview_dir(session_id: UUID) -> Path:
        return STORAGE_ROOT / PREVIEW_SUBDIR / str(session_id)

    @staticmethod
    def _safe_rmtree(path: Path) -> None:
        try:
            if path.exists():
                shutil.rmtree(path, ignore_errors=True)
        except OSError as exc:  # pragma: no cover - 清理失败不阻断
            logger.debug("preview 目录清理失败 %s: %s", path, exc)

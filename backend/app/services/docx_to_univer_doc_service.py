"""docx → Univer Doc IDocumentData JSON snapshot 转换器。

R10 复盘补丁 — wp_templates 109 个 docx 类底稿前端编辑支持。

设计决策（用户确认 B 选项 = 80-90% 还原 + mammoth+TipTap 兜底）：
- 主路径：python-docx 解析段落/标题/表格 → 生成最小 IDocumentData snapshot
- 损失：页眉页脚 / 字段域 / 嵌套表格 / 编号样式（回退到纯文本）
- 兜底：转换失败时由 caller fallback 到 mammoth → HTML → TipTap

参考：
- IDocumentData schema：https://reference.univer.ai/en-US/interfaces/IDocumentData
- dataStream 用 \\r 标记段落结束、\\n 标记段落组结束、\\f 标记节结束
"""
from __future__ import annotations

import logging
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any, BinaryIO

logger = logging.getLogger(__name__)


# Univer dataStream 控制字符
_PARAGRAPH_BREAK = "\r"     # 段落结束
_SECTION_BREAK = "\n"       # 节结束
_DOC_END = "\r\n"           # 文档结束（必须）


def _safe_text(value: Any) -> str:
    """把任意值转成 unicode 文本，过滤控制字符（保留换行/制表符）。"""
    if value is None:
        return ""
    text = str(value)
    # 过滤掉用作 Univer 控制字符的 \r \n \f 避免冲突
    text = text.replace("\r", "").replace("\n", " ").replace("\f", "")
    return text


def _heading_level_from_style(style_name: str | None) -> int | None:
    """从 docx 段落 style 名称推断标题级别。

    docx 标准样式名：'Heading 1' / '标题 1' / '一级标题' 等。
    返回 1-9（标题级别）或 None（普通段落）。
    """
    if not style_name:
        return None
    sn = style_name.lower().strip()
    # 英文
    if sn.startswith("heading "):
        try:
            return int(sn.split()[1])
        except (ValueError, IndexError):
            return None
    # 中文样式
    for i in range(1, 10):
        if f"标题 {i}" in style_name or f"标题{i}" in style_name:
            return i
        cn_num = "一二三四五六七八九"[i - 1]
        if f"{cn_num}级标题" in style_name:
            return i
    return None


def docx_bytes_to_univer_doc(
    docx_bytes: bytes | BinaryIO,
    *,
    doc_id: str | None = None,
) -> dict:
    """把 docx 字节流转换为 Univer Doc IDocumentData JSON snapshot。

    Parameters
    ----------
    docx_bytes : bytes | BinaryIO
        docx 文件字节流或可读流
    doc_id : str | None
        文档 ID，缺省自动生成

    Returns
    -------
    dict
        IDocumentData JSON dict，可直接给 Univer createUnit 使用

    Raises
    ------
    ImportError
        python-docx 未安装
    ValueError
        docx 解析失败（损坏 / 格式不识别）
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise ImportError("python-docx not installed; pip install python-docx") from e

    if isinstance(docx_bytes, (bytes, bytearray)):
        stream = BytesIO(docx_bytes)
    else:
        stream = docx_bytes  # 假设是可读流

    try:
        docx_doc = Document(stream)
    except Exception as e:
        raise ValueError(f"docx parse failed: {type(e).__name__}: {e}") from e

    # 构建 dataStream + paragraphs (元数据：每个段落起止位置 + 样式)
    data_stream_parts: list[str] = []
    paragraphs_meta: list[dict] = []   # paragraphs 字段：{ startIndex, paragraphStyle, ... }
    text_runs: list[dict] = []         # textRuns 字段（可选，简化版仅记录粗体/标题）

    cursor = 0  # 当前 dataStream 位置

    # 1. 处理段落（Heading/Paragraph）
    for p in docx_doc.paragraphs:
        text = _safe_text(p.text)
        # 段落起始位置
        para_start = cursor

        # 段落文本
        data_stream_parts.append(text)
        cursor += len(text)

        # 段落标记（必须）
        data_stream_parts.append(_PARAGRAPH_BREAK)
        cursor += 1

        # 标题级别
        heading_level = _heading_level_from_style(p.style.name if p.style else None)
        para_meta: dict[str, Any] = {"startIndex": cursor - 1}  # endIndex of paragraph mark
        if heading_level is not None:
            # Univer 用 namedStyleType: 'Heading1'..'Heading9'
            para_meta["paragraphStyle"] = {
                "namedStyleType": f"HEADING_{heading_level}",
            }
        paragraphs_meta.append(para_meta)

        # textRuns: 检测加粗/斜体（最小化处理，不处理颜色/字号）
        run_cursor = para_start
        for run in p.runs:
            run_text = _safe_text(run.text)
            run_len = len(run_text)
            if run_len == 0:
                run_cursor += run_len
                continue
            text_style: dict[str, Any] = {}
            if run.bold:
                text_style["bl"] = 1   # Univer 内部 bold 缩写
            if run.italic:
                text_style["it"] = 1
            if run.underline:
                text_style["ul"] = {"s": 1}
            if text_style:
                text_runs.append({
                    "st": run_cursor,
                    "ed": run_cursor + run_len,
                    "ts": text_style,
                })
            run_cursor += run_len

    # 2. 处理表格（简化为：每行 cell 文本拼接，行间用 PARAGRAPH_BREAK）
    #    Univer 0.21 表格 schema 复杂，先降级为段落形式（80% 还原核心，复杂表格走 mammoth 兜底）
    for table in docx_doc.tables:
        for row in table.rows:
            row_texts = [_safe_text(cell.text) for cell in row.cells]
            row_line = " | ".join(row_texts)  # 简化：用 " | " 分隔列
            data_stream_parts.append(row_line)
            cursor += len(row_line)
            data_stream_parts.append(_PARAGRAPH_BREAK)
            cursor += 1
            paragraphs_meta.append({"startIndex": cursor - 1})

    # 3. 节结束 + 文档结束（IDocumentData 必须以 \r\n 结尾）
    data_stream_parts.append(_SECTION_BREAK)
    cursor += 1
    # Univer 要求 dataStream 末尾不再加 \r（最后一个段落的 \r 已经收尾）

    final_data_stream = "".join(data_stream_parts)

    # 构建 IDocumentData snapshot
    snapshot: dict[str, Any] = {
        "id": doc_id or f"doc-{uuid.uuid4().hex[:12]}",
        "body": {
            "dataStream": final_data_stream,
            "paragraphs": paragraphs_meta,
        },
        "documentStyle": {
            "pageSize": {"width": 595, "height": 842},  # A4 默认
            "marginTop": 72,
            "marginBottom": 72,
            "marginLeft": 90,
            "marginRight": 90,
        },
    }
    if text_runs:
        snapshot["body"]["textRuns"] = text_runs

    return snapshot


def docx_path_to_univer_doc(
    docx_path: Path | str,
    *,
    doc_id: str | None = None,
) -> dict:
    """从 docx 文件路径加载并转换为 Univer Doc snapshot。"""
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"docx file not found: {path}")
    return docx_bytes_to_univer_doc(path.read_bytes(), doc_id=doc_id)


__all__ = [
    "docx_bytes_to_univer_doc",
    "docx_path_to_univer_doc",
]

"""Excel ↔ HTML 互转引擎 — 双格式保存 + ONLYOFFICE 联动

核心流程：
1. Excel → structure.json → HTML（在线编辑）
2. HTML 编辑 → structure.json 更新 → Excel 回写
3. ONLYOFFICE 编辑 → WOPI put_file → 自动同步 structure.json

structure.json 是权威数据源，HTML 和 Excel 都从它生成。
取数公式绑定在 structure.json 的单元格上，与格式无关。

═══ structure.json 格式 ═══
{
    "sheets": [{
        "name": "Sheet1",
        "cols": [{"width": 100}, ...],
        "rows": [{"height": 20}, ...],
        "cells": {
            "0:0": {"value": "项目", "style": {...}, "formula": null, "fetch_rule_id": null},
            "0:1": {"value": 1234.56, "style": {...}, "formula": "=SUM(B2:B5)", "fetch_rule_id": "uuid"},
        },
        "merges": [{"start_row": 0, "start_col": 0, "end_row": 0, "end_col": 2}],
    }],
    "metadata": {"source": "upload", "created_at": "...", "version": 1}
}
"""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

_logger = logging.getLogger(__name__)


# ═══ Excel → structure.json ═══

def excel_to_structure(file_path: str, max_rows: int = 5000, max_cols: int = 50) -> dict:
    """解析 Excel 文件为 structure.json 格式

    Args:
        file_path: Excel 文件路径
        max_rows: 最大解析行数（默认5000，底稿/附注足够；序时账用流式导入不走此接口）
        max_cols: 最大解析列数（默认50）

    Returns:
        structure.json 字典
    """
    import openpyxl
    from openpyxl.utils import get_column_letter

    wb = openpyxl.load_workbook(file_path, data_only=False)
    sheets = []

    for ws in wb.worksheets:
        sheet_data = {
            "name": ws.title,
            "cols": [],
            "rows": [],
            "cells": {},
            "merges": [],
        }

        # 列宽
        for col_idx in range(1, min(max_cols + 1, ws.max_column or 1) + 1):
            letter = get_column_letter(col_idx)
            width = ws.column_dimensions[letter].width or 10
            sheet_data["cols"].append({"width": round(width * 7.5)})  # Excel宽度→像素

        # 行高
        actual_rows = min(max_rows, ws.max_row or 0)
        for row_idx in range(1, actual_rows + 1):
            height = ws.row_dimensions[row_idx].height or 20
            sheet_data["rows"].append({"height": round(height)})

        # 单元格数据
        for row in ws.iter_rows(min_row=1, max_row=actual_rows, max_col=min(max_cols, ws.max_column or 1)):
            for cell in row:
                if cell.value is None and not cell.font.bold:
                    continue  # 跳过完全空白单元格

                r = cell.row - 1  # 转为0-indexed
                c = cell.column - 1
                key = f"{r}:{c}"

                cell_data: dict[str, Any] = {
                    "value": _serialize_value(cell.value),
                    "formula": str(cell.value) if isinstance(cell.value, str) and cell.value.startswith("=") else None,
                }

                # 样式提取
                style = _extract_style(cell)
                if style:
                    cell_data["style"] = style

                sheet_data["cells"][key] = cell_data

        # 合并单元格
        for merge_range in ws.merged_cells.ranges:
            sheet_data["merges"].append({
                "start_row": merge_range.min_row - 1,
                "start_col": merge_range.min_col - 1,
                "end_row": merge_range.max_row - 1,
                "end_col": merge_range.max_col - 1,
            })

        sheets.append(sheet_data)

    wb.close()

    return {
        "sheets": sheets,
        "metadata": {
            "source": "excel_upload",
            "source_file": Path(file_path).name,
            "created_at": datetime.utcnow().isoformat(),
            "version": 1,
            "sheet_count": len(sheets),
        },
    }


# ═══ structure.json → HTML ═══

def structure_to_html(
    structure: dict,
    sheet_index: int = 0,
    editable: bool = False,
    page: int = 0,
    page_size: int = 500,
) -> str:
    """将 structure.json 渲染为 HTML 表格

    大表格分页渲染：
    - page=0 表示不分页（渲染全部，适用于小表格<500行）
    - page>=1 表示分页模式，每页 page_size 行
    - 返回的 HTML 包含分页元数据（data-total-rows/data-page/data-page-size）

    Args:
        structure: structure.json 字典
        sheet_index: 要渲染的 sheet 索引
        editable: 是否生成可编辑的 HTML（contenteditable）

    Returns:
        HTML 字符串
    """
    sheets = structure.get("sheets", [])
    if sheet_index >= len(sheets):
        return "<p>无数据</p>"

    sheet = sheets[sheet_index]
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    cols = sheet.get("cols", [])
    rows = sheet.get("rows", [])

    # 构建合并单元格映射
    merge_map: dict[str, dict] = {}  # "r:c" → merge info
    skip_cells: set[str] = set()
    for m in merges:
        key = f"{m['start_row']}:{m['start_col']}"
        merge_map[key] = m
        for r in range(m["start_row"], m["end_row"] + 1):
            for c in range(m["start_col"], m["end_col"] + 1):
                if f"{r}:{c}" != key:
                    skip_cells.add(f"{r}:{c}")

    # 确定表格尺寸
    total_rows = len(rows) if rows else _get_max_dimension(cells, 0)
    max_col = len(cols) if cols else _get_max_dimension(cells, 1)

    # 分页逻辑
    if page >= 1 and total_rows > page_size:
        start_row = (page - 1) * page_size
        end_row = min(start_row + page_size, total_rows)
        max_row = end_row - start_row
        total_pages = (total_rows + page_size - 1) // page_size
    else:
        start_row = 0
        end_row = total_rows
        max_row = total_rows
        total_pages = 1

    # 生成 HTML
    html_parts = [
        '<style>',
        '.gt-excel-table { border-collapse: collapse; font-family: "仿宋_GB2312", "SimSun", serif; font-size: 10pt; width: 100%; position: relative; }',
        '.gt-excel-table td { border: 1px solid #d0d0d0; padding: 4px 6px; vertical-align: middle; position: relative; }',
        '.gt-excel-table tr:first-child td { background: #f4f0fa; font-weight: bold; text-align: center; }',
        '.gt-excel-table td[contenteditable="true"]:focus { outline: 2px solid #4b2d77; background: #faf8ff; }',
        '.gt-excel-table td[data-fetch-rule] { background: #f0f9ff; border-bottom: 2px solid #0094b3; cursor: pointer; }',
        '.gt-excel-table td[data-formula] { background: #fffbf0; }',
        '.gt-excel-table td[data-merged="true"] { background: #f9f5ff; }',
        '.gt-excel-table .gt-cell-coord { position: absolute; top: 1px; right: 2px; font-size: 8px; color: #c0c4cc; pointer-events: none; font-family: monospace; }',
        '.gt-excel-table td:hover .gt-cell-coord { color: #4b2d77; }',
        '.gt-excel-table .gt-row-header { background: #f8f8f8; text-align: center; font-size: 9px; color: #909399; width: 30px; min-width: 30px; user-select: none; }',
        '.gt-excel-table .gt-col-header { background: #f8f8f8; text-align: center; font-size: 9px; color: #909399; height: 20px; user-select: none; }',
        '</style>',
        f'<table class="gt-excel-table" border="1" cellspacing="0" cellpadding="4" data-total-rows="{total_rows}" data-page="{page}" data-page-size="{page_size}" data-total-pages="{total_pages}">',
    ]

    # 列头行（A, B, C...）— 仅编辑模式显示
    if editable:
        html_parts.append('<tr class="gt-col-header-row">')
        html_parts.append('<td class="gt-col-header"></td>')  # 左上角空格
        for c in range(max_col):
            col_letter = _col_to_letter(c)
            html_parts.append(f'<td class="gt-col-header">{col_letter}</td>')
        html_parts.append('</tr>')

    # colgroup
    if cols:
        html_parts.append("<colgroup>")
        if editable:
            html_parts.append('<col style="width:30px">')  # 行号列
        for col in cols:
            html_parts.append(f'<col style="width:{col.get("width", 80)}px">')
        html_parts.append("</colgroup>")

    for r_offset in range(max_row):
        r = start_row + r_offset  # 实际行号（分页偏移）
        row_height = rows[r].get("height", 20) if r < len(rows) else 20
        html_parts.append(f'<tr style="height:{row_height}px">')

        # 行号（仅编辑模式）
        if editable:
            html_parts.append(f'<td class="gt-row-header">{r + 1}</td>')

        for c in range(max_col):
            key = f"{r}:{c}"
            if key in skip_cells:
                continue

            cell = cells.get(key, {})
            value = cell.get("value", "")
            style = cell.get("style", {})
            formula = cell.get("formula")
            fetch_rule_id = cell.get("fetch_rule_id")

            # 构建 td 属性
            attrs = []
            is_merged = key in merge_map
            if is_merged:
                m = merge_map[key]
                rowspan = m["end_row"] - m["start_row"] + 1
                colspan = m["end_col"] - m["start_col"] + 1
                if rowspan > 1:
                    attrs.append(f'rowspan="{rowspan}"')
                if colspan > 1:
                    attrs.append(f'colspan="{colspan}"')
                attrs.append('data-merged="true"')
                # 合并范围地址（如 A1:C3）
                merge_addr = f"{_col_to_letter(m['start_col'])}{m['start_row']+1}:{_col_to_letter(m['end_col'])}{m['end_row']+1}"
                attrs.append(f'data-merge-range="{merge_addr}"')

            # 样式
            css = _style_to_css(style)
            if css:
                attrs.append(f'style="{css}"')

            # 数据属性（供前端取数联动）
            # Excel 风格地址（如 B3）
            cell_addr = f"{_col_to_letter(c)}{r + 1}"
            attrs.append(f'data-cell="{key}"')
            attrs.append(f'data-addr="{cell_addr}"')
            attrs.append(f'title="{cell_addr}"')
            if formula:
                attrs.append(f'data-formula="{_escape_html(formula)}"')
            if fetch_rule_id:
                attrs.append(f'data-fetch-rule="{fetch_rule_id}"')

            # 可编辑
            if editable:
                attrs.append('contenteditable="true"')

            attr_str = " ".join(attrs)
            display_value = _format_display_value(value)

            # 编辑模式下显示坐标标签
            coord_label = f'<span class="gt-cell-coord">{cell_addr}</span>' if editable else ""
            html_parts.append(f"<td {attr_str}>{display_value}{coord_label}</td>")

        html_parts.append("</tr>")

    html_parts.append("</table>")
    return "\n".join(html_parts)


# ═══ HTML 编辑结果 → structure.json 更新 ═══

def update_structure_from_edits(structure: dict, edits: list[dict], sheet_index: int = 0) -> dict:
    """将前端编辑结果更新到 structure.json

    支持的编辑操作：
    - 单元格修改: {"action": "edit", "cell": "0:1", "value": "新值", "formula": null}
    - 插入行: {"action": "insert_row", "at": 3}  → 在第3行前插入空行
    - 删除行: {"action": "delete_row", "at": 3}  → 删除第3行
    - 插入列: {"action": "insert_col", "at": 2}  → 在第2列前插入空列
    - 删除列: {"action": "delete_col", "at": 2}  → 删除第2列
    - 设置公式: {"action": "set_formula", "cell": "5:1", "formula": "SUM(1:4, 1)", "type": "vertical_sum"}

    增删行列时自动调整：
    - 所有单元格坐标重新映射
    - 公式中的行列引用自动偏移
    - 合并单元格范围自动调整
    - fetch_rule_id 绑定跟随单元格移动

    Returns:
        更新后的 structure.json
    """
    sheets = structure.get("sheets", [])
    if sheet_index >= len(sheets):
        return structure

    sheet = sheets[sheet_index]

    for edit in edits:
        action = edit.get("action", "edit")

        if action == "edit":
            _apply_cell_edit(sheet, edit)
        elif action == "insert_row":
            _insert_row(sheet, edit.get("at", 0))
        elif action == "delete_row":
            _delete_row(sheet, edit.get("at", 0))
        elif action == "insert_col":
            _insert_col(sheet, edit.get("at", 0))
        elif action == "delete_col":
            _delete_col(sheet, edit.get("at", 0))
        elif action == "set_formula":
            _set_formula(sheet, edit)

    structure["metadata"]["version"] = structure["metadata"].get("version", 0) + 1
    structure["metadata"]["last_edited_at"] = datetime.utcnow().isoformat()

    return structure


def _apply_cell_edit(sheet: dict, edit: dict):
    """应用单元格编辑"""
    cells = sheet.get("cells", {})
    key = edit.get("cell", "")
    if not key:
        return

    if key not in cells:
        cells[key] = {}

    if "value" in edit:
        cells[key]["value"] = edit["value"]
    if "formula" in edit:
        cells[key]["formula"] = edit["formula"]
    if "fetch_rule_id" in edit:
        cells[key]["fetch_rule_id"] = edit["fetch_rule_id"]
    if "style" in edit:
        cells[key]["style"] = edit["style"]

    sheet["cells"] = cells


def _insert_row(sheet: dict, at: int):
    """插入行：所有 row >= at 的单元格行号+1，公式引用自动调整"""
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    rows_meta = sheet.get("rows", [])

    # 重建 cells（行号偏移）
    new_cells = {}
    for key, cell_data in cells.items():
        r, c = _parse_key(key)
        if r >= at:
            new_cells[f"{r+1}:{c}"] = cell_data
        else:
            new_cells[f"{r}:{c}"] = cell_data

    # 调整公式中的行引用
    for key, cell_data in new_cells.items():
        if cell_data.get("formula"):
            cell_data["formula"] = _shift_formula_rows(cell_data["formula"], at, 1)

    # 调整合并单元格
    for m in merges:
        if m["start_row"] >= at:
            m["start_row"] += 1
        if m["end_row"] >= at:
            m["end_row"] += 1

    # 插入行元数据
    if at <= len(rows_meta):
        rows_meta.insert(at, {"height": 22})

    sheet["cells"] = new_cells
    sheet["merges"] = merges
    sheet["rows"] = rows_meta


def _delete_row(sheet: dict, at: int):
    """删除行：移除该行所有单元格，后续行号-1，公式引用自动调整"""
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    rows_meta = sheet.get("rows", [])

    # 重建 cells（删除目标行，后续行号-1）
    new_cells = {}
    for key, cell_data in cells.items():
        r, c = _parse_key(key)
        if r == at:
            continue  # 删除该行
        elif r > at:
            new_cells[f"{r-1}:{c}"] = cell_data
        else:
            new_cells[f"{r}:{c}"] = cell_data

    # 调整公式中的行引用
    for key, cell_data in new_cells.items():
        if cell_data.get("formula"):
            cell_data["formula"] = _shift_formula_rows(cell_data["formula"], at, -1)

    # 调整合并单元格
    new_merges = []
    for m in merges:
        if m["start_row"] == at and m["end_row"] == at:
            continue  # 完全在被删行内，移除
        if m["start_row"] > at:
            m["start_row"] -= 1
        if m["end_row"] > at:
            m["end_row"] -= 1
        if m["end_row"] >= m["start_row"]:
            new_merges.append(m)

    # 删除行元数据
    if at < len(rows_meta):
        rows_meta.pop(at)

    sheet["cells"] = new_cells
    sheet["merges"] = new_merges
    sheet["rows"] = rows_meta


def _insert_col(sheet: dict, at: int):
    """插入列：所有 col >= at 的单元格列号+1"""
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    cols_meta = sheet.get("cols", [])

    new_cells = {}
    for key, cell_data in cells.items():
        r, c = _parse_key(key)
        if c >= at:
            new_cells[f"{r}:{c+1}"] = cell_data
        else:
            new_cells[f"{r}:{c}"] = cell_data

    # 调整公式中的列引用
    for key, cell_data in new_cells.items():
        if cell_data.get("formula"):
            cell_data["formula"] = _shift_formula_cols(cell_data["formula"], at, 1)

    for m in merges:
        if m["start_col"] >= at:
            m["start_col"] += 1
        if m["end_col"] >= at:
            m["end_col"] += 1

    if at <= len(cols_meta):
        cols_meta.insert(at, {"width": 100})

    sheet["cells"] = new_cells
    sheet["merges"] = merges
    sheet["cols"] = cols_meta


def _delete_col(sheet: dict, at: int):
    """删除列：移除该列所有单元格，后续列号-1"""
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    cols_meta = sheet.get("cols", [])

    new_cells = {}
    for key, cell_data in cells.items():
        r, c = _parse_key(key)
        if c == at:
            continue
        elif c > at:
            new_cells[f"{r}:{c-1}"] = cell_data
        else:
            new_cells[f"{r}:{c}"] = cell_data

    for key, cell_data in new_cells.items():
        if cell_data.get("formula"):
            cell_data["formula"] = _shift_formula_cols(cell_data["formula"], at, -1)

    new_merges = []
    for m in merges:
        if m["start_col"] == at and m["end_col"] == at:
            continue
        if m["start_col"] > at:
            m["start_col"] -= 1
        if m["end_col"] > at:
            m["end_col"] -= 1
        if m["end_col"] >= m["start_col"]:
            new_merges.append(m)

    if at < len(cols_meta):
        cols_meta.pop(at)

    sheet["cells"] = new_cells
    sheet["merges"] = new_merges
    sheet["cols"] = cols_meta


def _set_formula(sheet: dict, edit: dict):
    """设置/更新单元格公式"""
    cells = sheet.get("cells", {})
    key = edit.get("cell", "")
    if not key:
        return

    if key not in cells:
        cells[key] = {}

    cells[key]["formula"] = edit.get("formula", "")
    if edit.get("type"):
        cells[key]["_formula_type"] = edit["type"]
    if edit.get("description"):
        cells[key]["_formula_desc"] = edit["description"]

    sheet["cells"] = cells


# ═══ 公式引用自动调整 ═══

import re

_CELL_REF_PATTERN = re.compile(r'cell\((\d+),\s*(\d+)\)')
_SUM_PATTERN = re.compile(r'SUM\((\d+):(\d+),\s*(\d+)\)')
_ROW_REF_PATTERN = re.compile(r'(?<!\d)(\d+)(?=:|\))')


def _shift_formula_rows(formula: str, at: int, delta: int) -> str:
    """调整公式中的行引用（插入/删除行时）

    cell(row, col) → row >= at 时 row += delta
    SUM(start:end, col) → start/end >= at 时调整
    """
    def _shift_cell_ref(m):
        r, c = int(m.group(1)), int(m.group(2))
        if r >= at:
            r += delta
        if r < 0:
            r = 0
        return f"cell({r}, {c})"

    def _shift_sum_ref(m):
        start, end, col = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if start >= at:
            start += delta
        if end >= at:
            end += delta
        start = max(0, start)
        end = max(start, end)
        return f"SUM({start}:{end}, {col})"

    result = _CELL_REF_PATTERN.sub(_shift_cell_ref, formula)
    result = _SUM_PATTERN.sub(_shift_sum_ref, result)
    return result


def _shift_formula_cols(formula: str, at: int, delta: int) -> str:
    """调整公式中的列引用（插入/删除列时）

    cell(row, col) → col >= at 时 col += delta
    SUM(start:end, col) → col >= at 时调整
    """
    def _shift_cell_ref(m):
        r, c = int(m.group(1)), int(m.group(2))
        if c >= at:
            c += delta
        if c < 0:
            c = 0
        return f"cell({r}, {c})"

    def _shift_sum_ref(m):
        start, end, col = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if col >= at:
            col += delta
        col = max(0, col)
        return f"SUM({start}:{end}, {col})"

    result = _CELL_REF_PATTERN.sub(_shift_cell_ref, formula)
    result = _SUM_PATTERN.sub(_shift_sum_ref, result)
    return result


def _parse_key(key: str) -> tuple[int, int]:
    """解析 "row:col" 键"""
    parts = key.split(":")
    return int(parts[0]), int(parts[1])


# ═══ structure.json → Excel 回写 ═══

def structure_to_excel(structure: dict, output_path: str, sheet_index: int | None = None) -> str:
    """将 structure.json 写回 Excel 文件

    Args:
        structure: structure.json 字典
        output_path: 输出 Excel 路径
        sheet_index: 指定 sheet（None=全部）

    Returns:
        输出文件路径
    """
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # 移除默认 sheet

    sheets = structure.get("sheets", [])
    target_sheets = [sheets[sheet_index]] if sheet_index is not None else sheets

    for sheet_data in target_sheets:
        ws = wb.create_sheet(title=sheet_data.get("name", "Sheet"))

        # 列宽
        for i, col in enumerate(sheet_data.get("cols", []), 1):
            letter = get_column_letter(i)
            ws.column_dimensions[letter].width = col.get("width", 80) / 7.5

        # 行高
        for i, row in enumerate(sheet_data.get("rows", []), 1):
            ws.row_dimensions[i].height = row.get("height", 20)

        # 单元格
        for key, cell_data in sheet_data.get("cells", {}).items():
            parts = key.split(":")
            r = int(parts[0]) + 1  # 转为1-indexed
            c = int(parts[1]) + 1

            cell = ws.cell(row=r, column=c)

            # 值或公式
            formula = cell_data.get("formula")
            if formula:
                cell.value = formula
            else:
                cell.value = cell_data.get("value")

            # 样式
            style = cell_data.get("style", {})
            if style:
                _apply_style_to_cell(cell, style)

        # 合并单元格
        for m in sheet_data.get("merges", []):
            ws.merge_cells(
                start_row=m["start_row"] + 1,
                start_column=m["start_col"] + 1,
                end_row=m["end_row"] + 1,
                end_column=m["end_col"] + 1,
            )

    wb.save(output_path)
    wb.close()
    return output_path


# ═══ ONLYOFFICE 联动：WOPI put_file 后同步 structure.json ═══

def sync_structure_from_excel(excel_path: str, structure_path: str) -> dict:
    """ONLYOFFICE 保存后，从 Excel 重新解析更新 structure.json

    保留用户自定义的 fetch_rule_id 绑定（Excel 中没有这个信息）。
    """
    # 解析新 Excel
    new_structure = excel_to_structure(excel_path)

    # 加载旧 structure（保留 fetch_rule_id 绑定）
    old_structure = {}
    sp = Path(structure_path)
    if sp.exists():
        try:
            old_structure = json.loads(sp.read_text(encoding="utf-8"))
        except Exception:
            pass

    # 合并：新值 + 旧的 fetch_rule_id
    if old_structure.get("sheets"):
        for si, new_sheet in enumerate(new_structure.get("sheets", [])):
            if si < len(old_structure["sheets"]):
                old_cells = old_structure["sheets"][si].get("cells", {})
                for key, new_cell in new_sheet.get("cells", {}).items():
                    old_cell = old_cells.get(key, {})
                    if old_cell.get("fetch_rule_id"):
                        new_cell["fetch_rule_id"] = old_cell["fetch_rule_id"]

    # 保留版本号递增
    old_version = old_structure.get("metadata", {}).get("version", 0)
    new_structure["metadata"]["version"] = old_version + 1
    new_structure["metadata"]["synced_from"] = "onlyoffice"
    new_structure["metadata"]["synced_at"] = datetime.utcnow().isoformat()

    # 保存
    sp.write_text(json.dumps(new_structure, ensure_ascii=False, indent=2), encoding="utf-8")

    return new_structure


# ═══ 内部辅助函数 ═══

def _serialize_value(value) -> Any:
    """序列化单元格值"""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, datetime):
        return value.isoformat()
    return str(value)


def _extract_style(cell) -> dict | None:
    """提取单元格样式"""
    style = {}

    if cell.font:
        if cell.font.bold:
            style["bold"] = True
        if cell.font.italic:
            style["italic"] = True
        if cell.font.size:
            style["fontSize"] = cell.font.size
        if cell.font.name:
            style["fontFamily"] = cell.font.name
        if cell.font.color and cell.font.color.rgb and cell.font.color.rgb != "00000000":
            style["color"] = f"#{cell.font.color.rgb[-6:]}"

    if cell.alignment:
        if cell.alignment.horizontal:
            style["textAlign"] = cell.alignment.horizontal
        if cell.alignment.vertical:
            style["verticalAlign"] = cell.alignment.vertical

    if cell.fill and cell.fill.fgColor and cell.fill.fgColor.rgb and cell.fill.fgColor.rgb not in ("00000000", "FFFFFFFF"):
        style["backgroundColor"] = f"#{cell.fill.fgColor.rgb[-6:]}"

    if cell.number_format and cell.number_format != "General":
        style["numberFormat"] = cell.number_format

    return style if style else None


def _style_to_css(style: dict) -> str:
    """样式字典转 CSS"""
    parts = []
    if style.get("bold"):
        parts.append("font-weight:bold")
    if style.get("italic"):
        parts.append("font-style:italic")
    if style.get("fontSize"):
        parts.append(f"font-size:{style['fontSize']}pt")
    if style.get("fontFamily"):
        parts.append(f"font-family:{style['fontFamily']}")
    if style.get("color"):
        parts.append(f"color:{style['color']}")
    if style.get("textAlign"):
        parts.append(f"text-align:{style['textAlign']}")
    if style.get("backgroundColor"):
        parts.append(f"background-color:{style['backgroundColor']}")
    return ";".join(parts)


def _apply_style_to_cell(cell, style: dict):
    """将样式应用到 openpyxl 单元格"""
    from openpyxl.styles import Font, Alignment, PatternFill

    font_kwargs = {}
    if style.get("bold"):
        font_kwargs["bold"] = True
    if style.get("italic"):
        font_kwargs["italic"] = True
    if style.get("fontSize"):
        font_kwargs["size"] = style["fontSize"]
    if style.get("fontFamily"):
        font_kwargs["name"] = style["fontFamily"]
    if font_kwargs:
        cell.font = Font(**font_kwargs)

    align_kwargs = {}
    if style.get("textAlign"):
        align_kwargs["horizontal"] = style["textAlign"]
    if style.get("verticalAlign"):
        align_kwargs["vertical"] = style["verticalAlign"]
    if align_kwargs:
        cell.alignment = Alignment(**align_kwargs)

    if style.get("backgroundColor"):
        color = style["backgroundColor"].lstrip("#")
        cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")


def _format_display_value(value) -> str:
    """格式化显示值"""
    if value is None:
        return ""
    if isinstance(value, float):
        if value == int(value):
            return f"{int(value):,}"
        return f"{value:,.2f}"
    if isinstance(value, int):
        return f"{value:,}"
    return _escape_html(str(value))


def _escape_html(s: str) -> str:
    """HTML 转义"""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _col_to_letter(col_index: int) -> str:
    """列索引(0-based)转 Excel 列字母（0→A, 1→B, 25→Z, 26→AA）"""
    result = ""
    idx = col_index
    while True:
        result = chr(65 + idx % 26) + result
        idx = idx // 26 - 1
        if idx < 0:
            break
    return result


def _get_max_dimension(cells: dict, axis: int) -> int:
    """从 cells 字典推断最大行/列数"""
    max_val = 0
    for key in cells:
        parts = key.split(":")
        if len(parts) == 2:
            val = int(parts[axis]) + 1
            if val > max_val:
                max_val = val
    return max_val


# ═══ structure.json → Word (.docx) ═══

def structure_to_word(structure: dict, output_path: str, sheet_index: int = 0) -> str:
    """将 structure.json 导出为 Word 文档（致同排版规范）

    表格样式：三线表（上下1磅边框，无左右边框）
    字体：仿宋_GB2312 + Arial Narrow（数字）
    """
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 页面设置（致同标准：页边距 3/3.18/3.2/2.54 cm）
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(3.18)
    section.top_margin = Cm(3.2)
    section.bottom_margin = Cm(2.54)

    sheets = structure.get("sheets", [])
    if sheet_index >= len(sheets):
        doc.save(output_path)
        return output_path

    sheet = sheets[sheet_index]
    cells = sheet.get("cells", {})
    merges = sheet.get("merges", [])
    cols = sheet.get("cols", [])
    rows_meta = sheet.get("rows", [])

    # 确定表格尺寸
    max_row = len(rows_meta) if rows_meta else _get_max_dimension(cells, 0)
    max_col = len(cols) if cols else _get_max_dimension(cells, 1)

    if max_row == 0 or max_col == 0:
        doc.save(output_path)
        return output_path

    # 添加标题（sheet名称）
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(sheet.get("name", ""))
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    run.font.size = Pt(14)
    run.font.bold = True

    # 创建表格
    table = doc.add_table(rows=max_row, cols=max_col)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 致同三线表样式：上下边框1磅，无左右边框
    _apply_three_line_border(table)

    # 填充单元格
    merge_map = {}
    skip_cells_set = set()
    for m in merges:
        key = f"{m['start_row']}:{m['start_col']}"
        merge_map[key] = m
        for r in range(m["start_row"], m["end_row"] + 1):
            for c in range(m["start_col"], m["end_col"] + 1):
                if f"{r}:{c}" != key:
                    skip_cells_set.add(f"{r}:{c}")

    for r in range(max_row):
        for c in range(max_col):
            key = f"{r}:{c}"
            if key in skip_cells_set:
                continue

            cell_data = cells.get(key, {})
            value = cell_data.get("value", "")
            style = cell_data.get("style", {})

            # Word 单元格
            word_cell = table.cell(r, c)

            # 处理合并
            if key in merge_map:
                m = merge_map[key]
                if m["end_row"] > m["start_row"] or m["end_col"] > m["start_col"]:
                    try:
                        merge_target = table.cell(m["end_row"], m["end_col"])
                        word_cell.merge(merge_target)
                    except Exception:
                        pass

            # 写入内容
            paragraph = word_cell.paragraphs[0]
            text = _format_word_value(value)
            run = paragraph.add_run(text)

            # 字体设置
            if _is_numeric(value):
                run.font.name = "Arial Narrow"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            else:
                run.font.name = "仿宋_GB2312"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

            run.font.size = Pt(10)

            # 样式
            if style.get("bold"):
                run.font.bold = True
            if style.get("textAlign") == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style.get("textAlign") == "right" or _is_numeric(value):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 设置列宽
    if cols:
        for i, col in enumerate(cols):
            width_cm = col.get("width", 80) / 37.8  # 像素→厘米（近似）
            for row in table.rows:
                try:
                    row.cells[i].width = Cm(min(width_cm, 6))
                except (IndexError, Exception):
                    pass

    # 页眉：事务所名称
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("致同会计师事务所（特殊普通合伙）")
    header_run.font.name = "仿宋_GB2312"
    header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 页脚：页码
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)

    doc.save(output_path)
    return output_path


def _add_page_number(paragraph):
    """在段落中插入页码域代码"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._element.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._element.append(fldChar2)

    run2 = paragraph.add_run("1")
    run2.font.size = Pt(9)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._element.append(fldChar3)


def _apply_three_line_border(table):
    """致同三线表样式：表格上下1磅边框，表头下1磅，无左右边框"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "bottom"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")  # 1磅
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)

    # 无左右边框
    for border_name in ["left", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        borders.append(border)

    # 表头下边框（insideH 只对第一行生效需要单独处理）
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "4")  # 0.5磅
    insideH.set(qn("w:space"), "0")
    insideH.set(qn("w:color"), "808080")
    borders.append(insideH)

    insideV = OxmlElement("w:insideV")
    insideV.set(qn("w:val"), "none")
    insideV.set(qn("w:sz"), "0")
    insideV.set(qn("w:space"), "0")
    borders.append(insideV)

    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def _format_word_value(value) -> str:
    """格式化 Word 单元格显示值"""
    if value is None:
        return ""
    if isinstance(value, float):
        if value == 0:
            return "-"
        if value == int(value):
            return f"{int(value):,}"
        return f"{value:,.2f}"
    if isinstance(value, int):
        if value == 0:
            return "-"
        return f"{value:,}"
    return str(value)


def _is_numeric(value) -> bool:
    """判断是否为数值"""
    return isinstance(value, (int, float))



# ═══ 版本管理：对比 + 回滚 ═══

def save_version_snapshot(structure: dict, storage_dir: str) -> str:
    """保存 structure.json 版本快照

    每次编辑保存时调用，保留最近20个版本。
    文件名格式：{stem}.v{version}.json
    """
    version = structure.get("metadata", {}).get("version", 1)
    stem = structure.get("metadata", {}).get("source_file", "data").replace(".xlsx", "")

    sp = Path(storage_dir)
    sp.mkdir(parents=True, exist_ok=True)

    snapshot_path = sp / f"{stem}.v{version}.json"
    snapshot_path.write_text(json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8")

    # 清理旧版本（保留最近20个）
    snapshots = sorted(sp.glob(f"{stem}.v*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    for old in snapshots[20:]:
        old.unlink(missing_ok=True)

    return str(snapshot_path)


def list_versions(storage_dir: str, stem: str) -> list[dict]:
    """列出所有版本快照"""
    sp = Path(storage_dir)
    if not sp.exists():
        return []

    snapshots = sorted(sp.glob(f"{stem}.v*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    versions = []
    for s in snapshots:
        try:
            data = json.loads(s.read_text(encoding="utf-8"))
            meta = data.get("metadata", {})
            versions.append({
                "version": meta.get("version", 0),
                "file": s.name,
                "edited_at": meta.get("last_edited_at", meta.get("created_at", "")),
                "synced_from": meta.get("synced_from", ""),
                "size": s.stat().st_size,
            })
        except Exception:
            pass
    return versions


def diff_versions(storage_dir: str, stem: str, v1: int, v2: int) -> dict:
    """对比两个版本的差异

    返回变更的单元格列表（新增/修改/删除）。
    """
    sp = Path(storage_dir)
    f1 = sp / f"{stem}.v{v1}.json"
    f2 = sp / f"{stem}.v{v2}.json"

    if not f1.exists() or not f2.exists():
        return {"error": "版本文件不存在"}

    s1 = json.loads(f1.read_text(encoding="utf-8"))
    s2 = json.loads(f2.read_text(encoding="utf-8"))

    cells1 = s1.get("sheets", [{}])[0].get("cells", {}) if s1.get("sheets") else {}
    cells2 = s2.get("sheets", [{}])[0].get("cells", {}) if s2.get("sheets") else {}

    changes = []

    # 修改和删除
    for key, cell1 in cells1.items():
        cell2 = cells2.get(key)
        if cell2 is None:
            changes.append({"cell": key, "type": "deleted", "old_value": cell1.get("value"), "new_value": None})
        elif cell1.get("value") != cell2.get("value") or cell1.get("formula") != cell2.get("formula"):
            changes.append({
                "cell": key, "type": "modified",
                "old_value": cell1.get("value"), "new_value": cell2.get("value"),
                "old_formula": cell1.get("formula"), "new_formula": cell2.get("formula"),
            })

    # 新增
    for key, cell2 in cells2.items():
        if key not in cells1:
            changes.append({"cell": key, "type": "added", "old_value": None, "new_value": cell2.get("value")})

    return {
        "v1": v1, "v2": v2,
        "changes_count": len(changes),
        "changes": changes,
    }


def rollback_to_version(storage_dir: str, stem: str, target_version: int) -> dict | None:
    """回滚到指定版本"""
    sp = Path(storage_dir)
    target_file = sp / f"{stem}.v{target_version}.json"

    if not target_file.exists():
        return None

    data = json.loads(target_file.read_text(encoding="utf-8"))
    # 更新版本号（回滚也算一次新版本）
    data["metadata"]["version"] = data["metadata"].get("version", 0)
    data["metadata"]["rolled_back_from"] = target_version
    data["metadata"]["rolled_back_at"] = datetime.utcnow().isoformat()

    return data


# ═══ 编辑锁（防止多人同时编辑冲突） ═══

import time as _time

# 内存锁（单实例部署足够，多实例需升级Redis）
_edit_locks: dict[str, dict] = {}  # key → {"user_id": ..., "locked_at": ..., "expires_at": ...}
_LOCK_TTL = 300  # 5分钟自动过期


def acquire_edit_lock(file_key: str, user_id: str) -> dict:
    """获取编辑锁

    Returns:
        {"locked": True, "by": user_id} 成功
        {"locked": False, "held_by": other_user, "expires_in": seconds} 失败
    """
    _cleanup_expired_locks()

    existing = _edit_locks.get(file_key)
    if existing and existing["user_id"] != user_id:
        expires_in = max(0, int(existing["expires_at"] - _time.time()))
        return {"locked": False, "held_by": existing["user_id"], "expires_in": expires_in}

    _edit_locks[file_key] = {
        "user_id": user_id,
        "locked_at": _time.time(),
        "expires_at": _time.time() + _LOCK_TTL,
    }
    return {"locked": True, "by": user_id}


def release_edit_lock(file_key: str, user_id: str) -> bool:
    """释放编辑锁"""
    existing = _edit_locks.get(file_key)
    if existing and existing["user_id"] == user_id:
        del _edit_locks[file_key]
        return True
    return False


def refresh_edit_lock(file_key: str, user_id: str) -> bool:
    """刷新锁过期时间（用户仍在编辑）"""
    existing = _edit_locks.get(file_key)
    if existing and existing["user_id"] == user_id:
        existing["expires_at"] = _time.time() + _LOCK_TTL
        return True
    return False


def get_lock_status(file_key: str) -> dict | None:
    """查询锁状态"""
    _cleanup_expired_locks()
    return _edit_locks.get(file_key)


def _cleanup_expired_locks():
    """清理过期锁"""
    now = _time.time()
    expired = [k for k, v in _edit_locks.items() if v["expires_at"] < now]
    for k in expired:
        del _edit_locks[k]

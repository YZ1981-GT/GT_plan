"""审计报告正文 JSON 主源 + 占位符渲染 — deliverable-center P0"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any
from uuid import UUID

import sqlalchemy as sa
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.report_models import (
    AuditReport,
    AuditReportTemplate,
    CompanyType,
    FinancialReport,
    FinancialReportType,
    OpinionType,
)
from app.services.audit_report_service import AuditReportService

logger = logging.getLogger(__name__)

REGISTRY_PATH = (
    Path(__file__).resolve().parent.parent.parent / "data" / "placeholder_mapping_registry.json"
)
DOCX_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent.parent
    / "data"
    / "templates"
    / "report_body_deliverable.docx"
)

# 应用层建模：带强调事项段的无保留意见
OPINION_UNQUALIFIED_WITH_EMPHASIS = "unqualified_with_emphasis"

SECTION_ID_MAP = {
    "审计意见段": "opinion",
    "形成审计意见的基础段": "basis",
    "强调事项段": "emphasis",
    "关键审计事项段": "kam",
    "其他信息段": "other_info",
    "其他事项段": "other_matter",
    "管理层和治理层对财务报表的责任段": "mgmt_responsibility",
    "注册会计师对财务报表审计的责任段": "cpa_responsibility",
    "签章段": "signature",
    "形成保留意见的基础段": "qualified_basis",
    "形成否定意见的基础段": "adverse_basis",
    "形成无法表示意见的基础段": "disclaimer_basis",
}

DELETABLE_SECTIONS = {"强调事项段", "关键审计事项段", "其他信息段", "其他事项段"}

_MANUAL_HINT_RE = re.compile(r"\[请[^\]]*\]")

# 需求 12：草稿水印当且仅当交付物处于草稿态（draft/editing）。
# confirmed/signed/archived 等终态/审批态生成无水印正式版本。
WATERMARK_STATUSES: frozenset[str] = frozenset({"draft", "editing"})

# 草稿水印标记文本（python-docx 后处理嵌入，供下载文件检测）
DRAFT_WATERMARK_MARK = "【草稿 DRAFT】"


def should_watermark(status: str | None) -> bool:
    """水印判定单一真源（需求 12.1/12.2/12.3）：

    当且仅当交付物状态属于 {draft, editing} 时叠加/嵌入草稿水印；
    confirmed/signed 等状态生成无水印正式版本。
    """
    return status in WATERMARK_STATUSES


@dataclass
class RenderResult:
    html: str
    docx_path: Path
    html_path: Path
    file_size: int


def _load_registry() -> dict:
    with open(REGISTRY_PATH, encoding="utf-8-sig") as f:
        return json.load(f)


def _section_id(section_name: str) -> str:
    return SECTION_ID_MAP.get(section_name, section_name)


def _is_manual_content(content: str) -> bool:
    return bool(_MANUAL_HINT_RE.search(content))


class ReportBodyService:
    """Report_Body_JSON 主源加载、占位符填充与渲染"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self._audit_svc = AuditReportService(db)

    def _resolve_opinion_enum(self, opinion_type: str) -> OpinionType:
        if opinion_type == OPINION_UNQUALIFIED_WITH_EMPHASIS:
            return OpinionType.unqualified
        return OpinionType(opinion_type)

    async def load_body_template(
        self,
        opinion_type: str,
        company_type: str,
        *,
        include_emphasis: bool = False,
    ) -> dict:
        """从 audit_report_template 组装 Report_Body_JSON 段落数组"""
        enum_opinion = self._resolve_opinion_enum(opinion_type)
        enum_company = CompanyType(company_type)
        templates = await self._audit_svc._load_templates(enum_opinion, enum_company)

        if opinion_type == OPINION_UNQUALIFIED_WITH_EMPHASIS:
            has_emphasis = any(t.section_name == "强调事项段" for t in templates)
            if not has_emphasis and include_emphasis:
                templates = list(templates)
                templates.insert(
                    2,
                    AuditReportTemplate(
                        opinion_type=enum_opinion,
                        company_type=enum_company,
                        section_name="强调事项段",
                        section_order=3,
                        template_text=(
                            "三、强调事项\n\n我们提醒财务报表使用者关注，"
                            "[请填写强调事项]。本段内容不影响已发表的审计意见。"
                        ),
                        is_required=False,
                    ),
                )

        sections: list[dict] = []
        for tmpl in templates:
            if tmpl.section_name == "关键审计事项段" and opinion_type == OpinionType.disclaimer.value:
                continue
            sec: dict[str, Any] = {
                "section_id": _section_id(tmpl.section_name),
                "section_name": tmpl.section_name,
                "section_order": tmpl.section_order,
                "is_required": tmpl.is_required,
                "deletable": tmpl.section_name in DELETABLE_SECTIONS,
                "content": tmpl.template_text,
                "placeholders_resolved": {},
            }
            if tmpl.section_name == "关键审计事项段":
                sec["items"] = []
            sections.append(sec)

        return {
            "opinion_type": opinion_type,
            "company_type": company_type,
            "is_pie": False,
            "sections": sections,
        }

    async def fill_placeholders(
        self,
        body: dict,
        project_id: UUID,
        year: int,
        *,
        report_date: date | None = None,
    ) -> dict:
        """按 Placeholder_Mapping_Registry 填充自动/财务类占位符"""
        registry = _load_registry()
        basic_info = await self._audit_svc._get_project_basic_info(project_id)
        financial_data = await self._audit_svc._fetch_financial_data(project_id, year)

        auto_values = self._resolve_auto_placeholders(
            project_id, year, basic_info, report_date
        )
        fin_values = await self._resolve_financial_placeholders(
            project_id, year, registry.get("financial", {}), financial_data
        )

        filled = json.loads(json.dumps(body, ensure_ascii=False))
        for section in filled.get("sections", []):
            content = section.get("content", "")
            resolved: dict[str, dict] = {}

            for key, meta in registry.get("auto", {}).items():
                value = auto_values.get(key, "")
                content = content.replace(f"{{{key}}}", value)
                resolved[key] = {"type": "auto", "value": value, "source": meta.get("source")}

            for key, meta in registry.get("financial", {}).items():
                value = fin_values.get(key, financial_data.get(key, "0"))
                content = content.replace(f"{{{key}}}", value)
                resolved[key] = {
                    "type": "financial",
                    "value": value,
                    "source": meta.get("row_code"),
                }

            section["content"] = content
            section["placeholders_resolved"] = resolved

        return filled

    def _resolve_auto_placeholders(
        self,
        project_id: UUID,
        year: int,
        basic_info: dict,
        report_date: date | None,
    ) -> dict[str, str]:
        entity_name = basic_info.get("client_name") or "[被审计单位名称]"
        entity_short_name = basic_info.get("entity_short_name") or f'"{entity_name}"'
        report_scope_val = basic_info.get("report_scope") or "standalone"
        report_scope = "合并及母公司" if report_scope_val == "consolidated" else ""
        signing_partner = basic_info.get("signing_partner_name") or "[签字注册会计师]"
        rd = report_date.isoformat() if report_date else "[报告日期]"
        return {
            "entity_name": entity_name,
            "entity_short_name": entity_short_name,
            "audit_period": f"{year}年12月31日",
            "audit_year": str(year),
            "report_scope": report_scope,
            "signing_partner": signing_partner,
            "report_date": rd,
        }

    async def _resolve_financial_placeholders(
        self,
        project_id: UUID,
        year: int,
        fin_registry: dict,
        cached: dict,
    ) -> dict[str, str]:
        values: dict[str, str] = {}
        for key, meta in fin_registry.items():
            if key in cached:
                values[key] = str(cached[key])
                continue
            report_type_name = meta.get("report_type")
            row_code = meta.get("row_code")
            if not report_type_name or not row_code:
                continue
            report_type = FinancialReportType(report_type_name)
            result = await self.db.execute(
                sa.select(FinancialReport.current_period_amount).where(
                    FinancialReport.project_id == project_id,
                    FinancialReport.year == year,
                    FinancialReport.report_type == report_type,
                    FinancialReport.row_code == row_code,
                    FinancialReport.is_deleted == sa.false(),
                )
            )
            val = result.scalar_one_or_none()
            values[key] = str(val) if val is not None else "0"
        return values

    async def refresh_financial_placeholders(
        self, project_id: UUID, year: int
    ) -> AuditReport | None:
        """REPORTS_UPDATED 触发：仅刷新 report_body_json 中财务类占位符"""
        report = await self._audit_svc.get_report(project_id, year)
        if report is None or not report.report_body_json:
            return report

        body = await self.fill_placeholders(
            report.report_body_json, project_id, year, report_date=report.report_date
        )
        report.report_body_json = body
        await self.db.flush()
        return report

    def validate_kam(
        self,
        body: dict,
        *,
        company_type: str,
        is_pie: bool,
        opinion_type: str,
    ) -> str | None:
        if opinion_type == OpinionType.disclaimer.value:
            return None
        if company_type != CompanyType.listed.value and not is_pie:
            return None
        for section in body.get("sections", []):
            if section.get("section_id") != "kam":
                continue
            content = section.get("content", "")
            items = section.get("items") or []
            if items:
                return None
            if content and "[请在此处添加关键审计事项" not in content:
                return None
            return (
                "上市公司或公共利益实体审计报告必须包含至少一个关键审计事项(KAM)，"
                "请编辑「关键审计事项段」后再定稿"
            )
        return (
            "上市公司或公共利益实体审计报告必须包含至少一个关键审计事项(KAM)，"
            "请编辑「关键审计事项段」后再定稿"
        )

    def kam_required(
        self,
        *,
        company_type: str,
        is_pie: bool,
        opinion_type: str,
    ) -> bool:
        """KAM 必填判定（需求 23.1/23.2/23.3）：

        必填 ⟺ (公司类型为 listed 或 is_pie 为真) 且 意见类型不为 disclaimer。
        """
        if opinion_type == OpinionType.disclaimer.value:
            return False
        return company_type == CompanyType.listed.value or bool(is_pie)

    @staticmethod
    def make_kam_item(matter: str, response: str) -> dict[str, str]:
        """构造 KAM 单条目（需求 23.5）：同时含事项描述与审计应对。"""
        return {"matter": matter, "response": response}

    def get_section(self, body: dict, section_id: str) -> dict | None:
        for section in body.get("sections", []):
            if section.get("section_id") == section_id:
                return section
        return None

    def delete_section(self, body: dict, section_id: str) -> dict:
        """删除可选段落（deletable=true）；不可删段落原样返回（需求 22.6）。"""
        new_body = json.loads(json.dumps(body, ensure_ascii=False))
        target = None
        for section in new_body.get("sections", []):
            if section.get("section_id") == section_id:
                target = section
                break
        if target is None or not target.get("deletable", False):
            return new_body
        new_body["sections"] = [
            s for s in new_body.get("sections", []) if s.get("section_id") != section_id
        ]
        return new_body

    def add_section(self, body: dict, section: dict) -> dict:
        """按 section_order 插入段落，保持顺序稳定（需求 22.6 增删往返）。"""
        new_body = json.loads(json.dumps(body, ensure_ascii=False))
        sections = list(new_body.get("sections", []))
        sections.append(json.loads(json.dumps(section, ensure_ascii=False)))
        sections.sort(key=lambda s: s.get("section_order", 0))
        new_body["sections"] = sections
        return new_body

    def render_html(self, body: dict) -> str:
        """渲染 Report_Body_JSON 为 HTML 供在线预览（需求 24.7）。

        按 section_order 排序，逐段渲染段落名 + content（换行转 <br/>），
        KAM 等多条目段额外渲染 items（事项描述 + 审计应对）。
        """
        parts: list[str] = []
        for section in sorted(body.get("sections", []), key=lambda s: s.get("section_order", 0)):
            name = section.get("section_name", "")
            content = (section.get("content", "") or "").replace("\n", "<br/>")
            parts.append(f"<h3>{name}</h3><div>{content}</div>")
            for item in section.get("items") or []:
                matter = item.get("matter", "")
                response = item.get("response", "")
                parts.append(f"<p><strong>{matter}</strong><br/>{response}</p>")
        return "\n".join(parts)

    def _ensure_docx_template(self) -> Path:
        if DOCX_TEMPLATE_PATH.exists():
            return DOCX_TEMPLATE_PATH
        DOCX_TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        from docx import Document

        doc = Document()
        doc.add_paragraph("{% for section in sections %}")
        doc.add_paragraph("{{ section.section_name }}")
        doc.add_paragraph("{{ section.content }}")
        doc.add_paragraph("{% endfor %}")
        doc.save(str(DOCX_TEMPLATE_PATH))
        return DOCX_TEMPLATE_PATH

    def render_docx(self, body: dict, output_path: Path, *, watermark: bool = False) -> Path:
        """docxtpl 渲染交付 docx；watermark 时 python-docx 后处理"""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        sections = []
        for section in sorted(body.get("sections", []), key=lambda s: s.get("section_order", 0)):
            content = section.get("content", "")
            for item in section.get("items") or []:
                content += f"\n\n{item.get('matter', '')}\n{item.get('response', '')}"
            sections.append(
                {
                    "section_name": section.get("section_name", ""),
                    "content": content,
                }
            )

        try:
            from docxtpl import DocxTemplate

            tpl_path = self._ensure_docx_template()
            tpl = DocxTemplate(str(tpl_path))
            tpl.render({"sections": sections})
            tpl.save(str(output_path))
        except Exception as exc:
            logger.warning("docxtpl 渲染失败，降级 python-docx: %s", exc)
            from docx import Document

            doc = Document()
            for sec in sections:
                doc.add_heading(sec["section_name"], level=2)
                for line in sec["content"].split("\n"):
                    doc.add_paragraph(line)
            doc.save(str(output_path))

        if watermark:
            self._apply_draft_watermark(output_path)

        return output_path

    def _apply_draft_watermark(self, path: Path) -> None:
        try:
            from docx import Document

            doc = Document(str(path))
            if doc.paragraphs:
                doc.paragraphs[0].insert_paragraph_before(DRAFT_WATERMARK_MARK)
            else:
                doc.add_paragraph(DRAFT_WATERMARK_MARK)
            doc.save(str(path))
        except Exception as exc:
            logger.debug("水印后处理跳过: %s", exc)

    @staticmethod
    def docx_has_watermark(path: Path) -> bool:
        """检测 docx 是否嵌入草稿水印（供下载文件水印有无校验，需求 12.2/12.3）。"""
        from docx import Document

        doc = Document(str(path))
        return any(DRAFT_WATERMARK_MARK in p.text for p in doc.paragraphs)

    def parse_docx_to_section_ids(self, docx_path: Path, original_body: dict) -> set[str]:
        """渲染往返：真正解析 docx 段落文本，按 section_name 反查 section_id 集合。

        docx↔段落解析对样式/合并单元格有损，故仅约束 section_id 集合（需求 24.7）。
        """
        from docx import Document

        doc = Document(str(docx_path))
        doc_text = "\n".join(p.text for p in doc.paragraphs)

        found: set[str] = set()
        for section in original_body.get("sections", []):
            section_name = section.get("section_name", "")
            section_id = section.get("section_id")
            if not section_id:
                continue
            # 段落名出现在渲染文本中即视为该段已往返保留
            if section_name and section_name in doc_text:
                found.add(section_id)
        return found

    def count_docx_paragraph_sections(self, docx_path: Path, original_body: dict) -> int:
        """往返后保留的段落数量（按 section_name 命中计数）。"""
        return len(self.parse_docx_to_section_ids(docx_path, original_body))

    PRIOR_PERIOD_OTHER_MATTER: dict[str, str] = {
        "predecessor_auditor": (
            "我们提醒财务报表使用者关注，上期比较信息是由前任注册会计师审计/审阅的。"
            "我们对上期比较信息实施了必要的审计程序，但不包括前任注册会计师的工作。"
        ),
        "prior_unaudited": (
            "我们提醒财务报表使用者关注，上期比较信息未经审计。"
            "我们的审计意见不涵盖上期比较信息，也不对其发表任何形式的鉴证结论。"
        ),
    }

    def apply_prior_period_section(
        self, body: dict, prior_period_info: str | None
    ) -> dict:
        """首次接受委托时注入其他事项段（需求 26）"""
        if prior_period_info not in self.PRIOR_PERIOD_OTHER_MATTER:
            return body

        content = self.PRIOR_PERIOD_OTHER_MATTER[prior_period_info]
        sections = list(body.get("sections") or [])
        has_other = any(
            s.get("section_id") == "other_matter"
            or s.get("section_name") == "其他事项段"
            for s in sections
        )
        if not has_other:
            sections.append(
                {
                    "section_id": "other_matter",
                    "section_name": "其他事项段",
                    "section_order": 95,
                    "content": content,
                    "items": [],
                }
            )
            body = {**body, "sections": sections}
        return body

    @staticmethod
    def validate_report_date_compliance(
        report_date: date,
        *,
        evidence_complete_date: date | None = None,
        fs_approval_date: date | None = None,
        eqcr_pass_date: date | None = None,
        extra_floors: list[tuple[str, date]] | None = None,
    ) -> dict:
        """Report_Date_Compliance — 报告日期下界合规校验（需求 25.1/25.2/25.3/25.4）。

        将报告日期作为可校验字段处理（需求 25.4），而非纯文本占位符。

        下界 = max(已提供的下界日期)，其中：
        - ``evidence_complete_date``：注册会计师获取充分适当审计证据之日（需求 25.1）
        - ``fs_approval_date``：管理层/治理层批准财务报表之日（需求 25.2）
        - ``eqcr_pass_date``：EQCR 复核通过日（需求 25.3）
        - ``extra_floors``：其他已知下界（label, date）补充项

        当 ``report_date`` 早于下界时，返回 ``compliant=False`` + ``requires_confirmation=True``
        的告警（非硬阻断，需求 25.3）；调用方据此提示用户二次确认而非直接拒绝。

        所有下界日期均为可选输入：当某来源暂无具体存储时由调用方传 ``None`` 跳过，
        校验仅基于实际提供的日期计算下界（无任何下界时视为合规）。
        """
        floors: list[tuple[str, date]] = []
        if evidence_complete_date is not None:
            floors.append(("审计证据完成日", evidence_complete_date))
        if fs_approval_date is not None:
            floors.append(("财务报表/治理层批准日", fs_approval_date))
        if eqcr_pass_date is not None:
            floors.append(("EQCR 通过日", eqcr_pass_date))
        if extra_floors:
            floors.extend(extra_floors)

        if not floors:
            return {
                "compliant": True,
                "requires_confirmation": False,
                "warnings": [],
                "floor_date": None,
            }

        floor_date = max(d for _, d in floors)
        if report_date < floor_date:
            labels = [lbl for lbl, d in floors if d == floor_date]
            warning = (
                f"报告日期 {report_date.isoformat()} 早于合规下界 "
                f"{floor_date.isoformat()}（{', '.join(labels)}），请确认"
            )
            return {
                "compliant": False,
                "requires_confirmation": True,
                "warnings": [warning],
                "floor_date": floor_date.isoformat(),
            }

        return {
            "compliant": True,
            "requires_confirmation": False,
            "warnings": [],
            "floor_date": floor_date.isoformat(),
        }

    async def check_report_date_compliance(
        self,
        project_id: UUID,
        year: int,
        report_date: date,
    ) -> dict:
        """报告日期下界合规 — 从项目/报告状态收集下界日期后委托纯校验器。

        说明（数据来源假设）：当前数据模型中各下界日期尚无统一专用存储字段，
        本方法尽力从已有载体提取——
        - 审计证据完成日 / 财务报表批准日：项目 ``wizard_state.basic_info``（若已填）
        - EQCR 通过日：``audit_report`` 进入 ``eqcr_approved``/``final`` 状态时取 ``updated_at``
        提取不到的下界传 ``None`` 跳过；最终校验逻辑统一由
        :meth:`validate_report_date_compliance` 计算，保证 service 与纯校验器口径一致。
        """
        from app.models.core import Project

        evidence_complete_date: date | None = None
        fs_approval_date: date | None = None
        eqcr_pass_date: date | None = None
        extra_floors: list[tuple[str, date]] = []

        def _coerce(raw: Any) -> date | None:
            if raw is None:
                return None
            if isinstance(raw, date):
                return raw
            try:
                return date.fromisoformat(str(raw)[:10])
            except ValueError:
                return None

        project = await self.db.get(Project, project_id)
        if project and project.wizard_state:
            basic = (
                project.wizard_state.get("steps", {})
                .get("basic_info", {})
                .get("data")
                or project.wizard_state.get("basic_info", {}).get("data")
                or {}
            )
            fs_approval_date = _coerce(
                basic.get("fs_approval_date")
                or basic.get("financial_statement_approval_date")
            )
            evidence_complete_date = _coerce(
                basic.get("audit_evidence_completion_date")
            )
            audit_period_end = _coerce(basic.get("audit_period_end"))
            if audit_period_end is not None:
                extra_floors.append(("审计期间截止日", audit_period_end))

        report = await self._audit_svc.get_report(project_id, year)
        if report:
            from app.models.report_models import ReportStatus

            if (
                report.status in (ReportStatus.eqcr_approved, ReportStatus.final)
                and report.updated_at
            ):
                eqcr_pass_date = report.updated_at.date()

        return self.validate_report_date_compliance(
            report_date,
            evidence_complete_date=evidence_complete_date,
            fs_approval_date=fs_approval_date,
            eqcr_pass_date=eqcr_pass_date,
            extra_floors=extra_floors or None,
        )

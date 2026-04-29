"""Word 模板填充服务 — 方案B核心 (Phase 13)

打开致同标准 Word 模板 → 填充数据到占位符/书签 → 保存到项目目录
"""

from __future__ import annotations

import logging
import os
import zipfile
from pathlib import Path
from uuid import UUID

import sqlalchemy as sa
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.phase13_models import ReportSnapshot, WordExportDocType
from app.models.report_models import AuditReport, DisclosureNote
from app.services.export_task_service import ExportTaskService
from app.services.report_placeholder_service import ReportPlaceholderService
from app.services.report_snapshot_service import ReportSnapshotService

logger = logging.getLogger(__name__)

STORAGE_ROOT = Path(os.environ.get("STORAGE_ROOT", "storage"))
TEMPLATE_ROOT = Path("backend/data/word_templates")

# 报表类型 → 文件名前缀
_REPORT_TYPE_PREFIX = {
    "BS": "balance_sheet",
    "IS": "income_statement",
    "CFS": "cash_flow_statement",
    "EQ": "equity_statement",
}

# 报表口径替换
SCOPE_REPLACEMENTS: dict[str, dict[str, str]] = {
    "consolidated": {
        "财务报表": "合并及母公司财务报表",
        "资产负债表": "合并及母公司资产负债表",
        "利润表": "合并及母公司利润表",
        "现金流量表": "合并及母公司现金流量表",
        "所有者权益变动表": "合并及母公司所有者权益变动表",
        "财务报表附注": "合并及母公司财务报表附注",
    },
    "standalone": {},
}


def _ensure_dir(path: Path) -> Path:
    """Create parent directories and return path."""
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _open_template_or_create(template_name: str) -> "Document":
    """Open a Word template file, or create a minimal document as fallback.

    Uses GTWordEngine for fallback creation.
    """
    template_path = TEMPLATE_ROOT / template_name
    try:
        from docx import Document
        if template_path.exists():
            return Document(str(template_path))
        logger.warning("模板文件不存在: %s, 使用空白文档", template_path)
        return Document()
    except ImportError:
        from app.services.gt_word_engine import GTWordEngine
        engine = GTWordEngine()
        return engine.doc


def _replace_placeholders_in_doc(doc, placeholders: dict[str, str]) -> None:
    """Replace {xxx} placeholders in all paragraphs of a document."""
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)


def _apply_scope_to_doc(doc, report_scope: str) -> None:
    """Apply scope replacements (consolidated → 合并及母公司) to document."""
    replacements = SCOPE_REPLACEMENTS.get(report_scope, {})
    if not replacements:
        return
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)


def _fill_table_data(doc, table_data: dict) -> None:
    """Fill table data into the first table in the document."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        return

    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers or not rows:
        return

    # Find first table or create one
    if doc.tables:
        table = doc.tables[0]
        # Clear existing data rows (keep header)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)
    else:
        table = doc.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)

    # Add data rows
    for row_data in rows:
        label = row_data.get("label", "")
        values = row_data.get("values", row_data.get("cells", []))
        row = table.add_row()
        row.cells[0].text = str(label)
        for c_idx, val in enumerate(values):
            if c_idx + 1 < len(headers):
                cell_val = val
                if isinstance(val, dict):
                    cell_val = val.get("value", val.get("manual_value", 0))
                cell = row.cells[c_idx + 1]
                if cell_val is not None and cell_val != 0:
                    try:
                        num = float(cell_val)
                        if num < 0:
                            cell.text = f"({abs(num):,.2f})"
                        else:
                            cell.text = f"{num:,.2f}"
                    except (ValueError, TypeError):
                        cell.text = str(cell_val)
                else:
                    cell.text = "-"
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _process_color_text_in_doc(doc) -> None:
    """Process color text in document: blue=delete, red→black."""
    try:
        from docx.shared import RGBColor
    except ImportError:
        return

    for paragraph in doc.paragraphs:
        runs_to_remove = []
        for run in paragraph.runs:
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            if color is not None:
                if color.red < 100 and color.blue > 150:
                    runs_to_remove.append(run)
                elif color.red > 200 and color.green < 100:
                    run.font.color.rgb = RGBColor(0, 0, 0)
        for run in runs_to_remove:
            run._r.getparent().remove(run._r)


class WordTemplateFiller:
    """Word 模板填充服务 — 方案B核心"""

    def __init__(self, db: AsyncSession):
        self.db = db

    def _reports_dir(self, project_id: UUID) -> Path:
        """Get reports directory for a project."""
        d = STORAGE_ROOT / "projects" / str(project_id) / "reports"
        d.mkdir(parents=True, exist_ok=True)
        return d

    # ------------------------------------------------------------------
    # 审计报告填充
    # ------------------------------------------------------------------

    async def fill_audit_report(
        self,
        db: AsyncSession,
        project_id: UUID,
        year: int,
        user_id: UUID,
    ) -> Path:
        """打开审计报告模板 → 替换占位符 → 保存到项目目录

        1. Opens template from backend/data/word_templates/ (or creates minimal doc)
        2. Reads audit_report table for 7 paragraphs
        3. Uses ReportPlaceholderService for placeholders and {xxx} replacement
        4. Applies scope replacements
        5. Saves to storage/projects/{project_id}/reports/audit_report_{year}.docx
        6. Creates WordExportTask record
        """
        # Get placeholders
        placeholder_svc = ReportPlaceholderService(db)
        placeholders = await placeholder_svc.get_placeholders(project_id)

        # Open template or create minimal doc
        doc = _open_template_or_create("审计报告模板.docx")

        # Read audit_report paragraphs
        result = await db.execute(
            sa.select(AuditReport).where(
                AuditReport.project_id == project_id,
                AuditReport.year == year,
                AuditReport.is_deleted == sa.false(),
            )
        )
        report = result.scalar_one_or_none()

        if report and report.paragraphs:
            for section_name, text in report.paragraphs.items():
                filled = ReportPlaceholderService.replace_in_text(text, placeholders)
                doc.add_paragraph(filled)

        # Replace placeholders in existing template text
        _replace_placeholders_in_doc(doc, placeholders)

        # Apply scope replacements
        report_scope = placeholders.get("report_scope", "")
        scope_key = "consolidated" if "合并" in report_scope else "standalone"
        _apply_scope_to_doc(doc, scope_key)

        # Save
        output_path = self._reports_dir(project_id) / f"audit_report_{year}.docx"
        _ensure_dir(output_path)
        doc.save(str(output_path))

        # Create ExportTask record
        task_svc = ExportTaskService(db)
        task = await task_svc.create_task(
            project_id=project_id,
            doc_type=WordExportDocType.audit_report.value,
            template_type=None,
            user_id=user_id,
        )
        task.file_path = str(output_path)
        await db.flush()

        logger.info(
            "审计报告填充完成: project_id=%s, year=%d, path=%s",
            project_id, year, output_path,
        )
        return output_path

    # ------------------------------------------------------------------
    # 财务报表填充
    # ------------------------------------------------------------------

    async def fill_financial_reports(
        self,
        db: AsyncSession,
        project_id: UUID,
        year: int,
        user_id: UUID,
    ) -> list[Path]:
        """打开报表模板 → 填充数据到表格 → 保存

        Gets latest ReportSnapshot (or creates one), then for each report type
        (BS/IS/CFS/EQ) fills template and saves.
        """
        snapshot_svc = ReportSnapshotService(db)
        paths: list[Path] = []

        for short_type, prefix in _REPORT_TYPE_PREFIX.items():
            # Get or create snapshot
            snapshot = await snapshot_svc.get_latest_snapshot(
                project_id, year, short_type
            )
            if snapshot is None:
                snapshots = await snapshot_svc.create_snapshot(
                    project_id, year, user_id
                )
                snapshot = next(
                    (s for s in snapshots if s.report_type == short_type), None
                )

            # Open template
            template_name = f"{prefix}_模板.docx"
            doc = _open_template_or_create(template_name)

            # Fill table data from snapshot
            if snapshot and snapshot.data:
                rows_data = snapshot.data.get("rows", [])
                if rows_data:
                    headers = ["项目", "附注", "期末余额", "期初余额"]
                    table_rows = []
                    for r in rows_data:
                        table_rows.append({
                            "label": r.get("row_name", ""),
                            "values": [
                                "",  # 附注列
                                r.get("current_period_amount"),
                                r.get("prior_period_amount"),
                            ],
                        })
                    _fill_table_data(doc, {"headers": headers, "rows": table_rows})

            # Save
            output_path = self._reports_dir(project_id) / f"{prefix}_{year}.docx"
            _ensure_dir(output_path)
            doc.save(str(output_path))
            paths.append(output_path)

            # Create ExportTask
            task_svc = ExportTaskService(db)
            task = await task_svc.create_task(
                project_id=project_id,
                doc_type=WordExportDocType.financial_report.value,
                template_type=short_type,
                user_id=user_id,
            )
            task.file_path = str(output_path)
            task.snapshot_id = snapshot.id if snapshot else None
            await db.flush()

        logger.info(
            "财务报表填充完成: project_id=%s, year=%d, %d files",
            project_id, year, len(paths),
        )
        return paths

    # ------------------------------------------------------------------
    # 附注填充
    # ------------------------------------------------------------------

    async def fill_disclosure_notes(
        self,
        db: AsyncSession,
        project_id: UUID,
        year: int,
        user_id: UUID,
        template_type: str = "soe",
    ) -> Path:
        """打开附注模板 → 填充表格数据+叙述文本 → 三色处理 → 保存

        Template priority: custom_template_snapshot > project templates/ > system soe/listed
        """
        # Determine template path with priority
        custom_path = (
            STORAGE_ROOT / "projects" / str(project_id)
            / "templates" / "custom_template_snapshot.docx"
        )
        project_template_path = (
            STORAGE_ROOT / "projects" / str(project_id)
            / "templates" / f"notes_{template_type}.docx"
        )
        system_template = f"附注模板_{template_type}.docx"

        if custom_path.exists():
            try:
                from docx import Document
                doc = Document(str(custom_path))
            except ImportError:
                doc = _open_template_or_create(system_template)
        elif project_template_path.exists():
            try:
                from docx import Document
                doc = Document(str(project_template_path))
            except ImportError:
                doc = _open_template_or_create(system_template)
        else:
            doc = _open_template_or_create(system_template)

        # Read DisclosureNote records
        result = await db.execute(
            sa.select(DisclosureNote).where(
                DisclosureNote.project_id == project_id,
                DisclosureNote.year == year,
            ).order_by(DisclosureNote.note_section)
        )
        notes = result.scalars().all()

        for note in notes:
            # Section heading
            heading = f"{note.note_section} {note.section_title or ''}"
            doc.add_paragraph(heading)

            # Table data
            if note.table_data and isinstance(note.table_data, dict):
                _fill_table_data(doc, note.table_data)

            # Narrative text
            if note.text_content:
                doc.add_paragraph(note.text_content)

        # Process color text (blue=delete, red→black)
        _process_color_text_in_doc(doc)

        # Save
        output_path = self._reports_dir(project_id) / f"notes_{year}.docx"
        _ensure_dir(output_path)
        doc.save(str(output_path))

        # Create ExportTask
        task_svc = ExportTaskService(db)
        task = await task_svc.create_task(
            project_id=project_id,
            doc_type=WordExportDocType.disclosure_notes.value,
            template_type=template_type,
            user_id=user_id,
        )
        task.file_path = str(output_path)
        await db.flush()

        logger.info(
            "附注填充完成: project_id=%s, year=%d, template=%s",
            project_id, year, template_type,
        )
        return output_path

    # ------------------------------------------------------------------
    # 全套导出
    # ------------------------------------------------------------------

    async def fill_full_package(
        self,
        db: AsyncSession,
        project_id: UUID,
        year: int,
        user_id: UUID,
    ) -> Path:
        """全套导出：审计报告+4张报表+附注 → ZIP打包"""
        files: list[Path] = []

        # 1. Audit report
        audit_path = await self.fill_audit_report(db, project_id, year, user_id)
        files.append(audit_path)

        # 2. Financial reports (4 files)
        report_paths = await self.fill_financial_reports(db, project_id, year, user_id)
        files.extend(report_paths)

        # 3. Disclosure notes
        notes_path = await self.fill_disclosure_notes(db, project_id, year, user_id)
        files.append(notes_path)

        # 4. Create ZIP
        zip_path = self._reports_dir(project_id) / f"full_package_{year}.zip"
        _ensure_dir(zip_path)
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            for f in files:
                if f.exists():
                    zf.write(str(f), f.name)

            # ── Phase 16: 一致性报告附在取证包中 ──
            try:
                from app.services.consistency_replay_engine import consistency_replay_engine
                import json as _json
                report = await consistency_replay_engine.generate_consistency_report(db, project_id, year)
                consistency_json = _json.dumps(report, ensure_ascii=False, indent=2)
                zf.writestr("consistency_report.json", consistency_json)
                logger.info(f"[CONSISTENCY] report attached to package: status={report.get('overall_status')}")
            except Exception as _cr_err:
                logger.warning(f"[CONSISTENCY] report attachment failed (non-blocking): {_cr_err}")

        logger.info(
            "全套导出完成: project_id=%s, year=%d, %d files → %s",
            project_id, year, len(files), zip_path,
        )
        return zip_path

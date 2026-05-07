"""EQCR 备忘录生成服务

Refinement Round 5 任务 18 — 需求 9

功能：
- 组装 EQCR 独立复核备忘录（结构化 JSON + Word 文档双轨）
- 编辑保存（存 Project.wizard_state.eqcr_memo）
- 定稿 → Word + PDF 落文件系统，供归档包引用
- 归档章节注册：``('02', 'eqcr_memo.pdf', eqcr_memo_pdf_generator)``
  （R1 archive_section_registry 落地前，generator 作为独立函数暴露）

**双轨设计**：
- 前端编辑器走 JSON 结构化保存（最佳编辑体验）
- 定稿时后端调 python-docx 生成 docx，再通过 PDF 引擎转 PDF
- Word 和 PDF 保存到 Project.wizard_state.eqcr_memo.files
"""

from __future__ import annotations

import logging
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from uuid import UUID

import sqlalchemy as sa
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.core import Project
from app.models.eqcr_models import (
    EqcrDisagreementResolution,
    EqcrOpinion,
    EqcrReviewNote,
    EqcrShadowComputation,
)
from app.models.staff_models import StaffMember, WorkHour

logger = logging.getLogger(__name__)

# 备忘录章节模板（固定顺序）
MEMO_SECTIONS = [
    "项目概况",
    "重要性判断",
    "会计估计复核",
    "关联方复核",
    "持续经营评估",
    "审计意见合理性",
    "独立复核笔记摘要",
    "独立取数结果",
    "异议与合议结论",
    "EQCR 总评与结论",
]

DOMAIN_LABELS = {
    "materiality": "重要性",
    "estimate": "会计估计",
    "related_party": "关联方",
    "going_concern": "持续经营",
    "opinion_type": "审计意见类型",
    "component_auditor": "组成部分审计师",
}

VERDICT_LABELS = {
    "agree": "认可",
    "disagree": "有异议",
    "need_more_evidence": "需补充证据",
}


class EqcrMemoService:
    """EQCR 备忘录服务"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def generate_memo(
        self,
        project_id: UUID,
        eqcr_user_id: UUID,
    ) -> dict:
        """组装 EQCR 备忘录内容（结构化 JSON）。

        此方法返回 dict，前端渲染/编辑后调 save_memo 保存，
        最终 finalize_memo 才生成 Word/PDF。
        """
        proj = (
            await self.db.execute(
                sa.select(Project).where(
                    Project.id == project_id,
                    Project.is_deleted == False,  # noqa: E712
                )
            )
        ).scalar_one_or_none()
        if proj is None:
            raise ValueError("项目不存在")

        opinions = list(
            (
                await self.db.execute(
                    sa.select(EqcrOpinion)
                    .where(
                        EqcrOpinion.project_id == project_id,
                        EqcrOpinion.is_deleted == False,  # noqa: E712
                    )
                    .order_by(EqcrOpinion.created_at)
                )
            ).scalars().all()
        )

        notes = list(
            (
                await self.db.execute(
                    sa.select(EqcrReviewNote)
                    .where(
                        EqcrReviewNote.project_id == project_id,
                        EqcrReviewNote.is_deleted == False,  # noqa: E712
                    )
                    .order_by(EqcrReviewNote.created_at)
                )
            ).scalars().all()
        )

        shadow_comps = list(
            (
                await self.db.execute(
                    sa.select(EqcrShadowComputation)
                    .where(EqcrShadowComputation.project_id == project_id)
                    .order_by(EqcrShadowComputation.created_at)
                )
            ).scalars().all()
        )

        disagreements = list(
            (
                await self.db.execute(
                    sa.select(EqcrDisagreementResolution)
                    .where(EqcrDisagreementResolution.project_id == project_id)
                    .order_by(EqcrDisagreementResolution.created_at)
                )
            ).scalars().all()
        )

        # EQCR 工时
        staff_q = sa.select(StaffMember.id).where(
            StaffMember.user_id == eqcr_user_id,
            StaffMember.is_deleted == False,  # noqa: E712
        )
        staff_id = (await self.db.execute(staff_q)).scalar_one_or_none()
        total_hours = 0.0
        if staff_id:
            hours_q = sa.select(
                sa.func.coalesce(sa.func.sum(WorkHour.hours), 0)
            ).where(
                WorkHour.staff_id == staff_id,
                WorkHour.project_id == project_id,
                WorkHour.purpose == "eqcr",
                WorkHour.status != "tracking",
                WorkHour.is_deleted == False,  # noqa: E712
            )
            total_hours = float((await self.db.execute(hours_q)).scalar_one())

        sections = self._build_sections(
            proj, opinions, notes, shadow_comps, disagreements, total_hours
        )

        return {
            "project_id": str(project_id),
            "project_name": proj.name,
            "client_name": proj.client_name,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "sections": sections,
            "section_order": MEMO_SECTIONS,
            "total_eqcr_hours": total_hours,
            "status": "draft",
        }

    def _build_sections(
        self, proj, opinions, notes, shadow_comps, disagreements, total_hours
    ) -> dict[str, str]:
        """生成 10 个章节的文本内容。"""
        sections = {}

        sections["项目概况"] = (
            f"项目名称：{proj.name}\n"
            f"客户名称：{proj.client_name}\n"
            f"审计期间：{proj.audit_period_start or '?'} ~ "
            f"{proj.audit_period_end or '?'}\n"
            f"EQCR 复核总工时：{total_hours:.1f} 小时\n"
            f"备忘录生成日期：{datetime.now(timezone.utc).strftime('%Y-%m-%d')}"
        )

        domain_section_map = {
            "materiality": "重要性判断",
            "estimate": "会计估计复核",
            "related_party": "关联方复核",
            "going_concern": "持续经营评估",
            "opinion_type": "审计意见合理性",
        }
        for domain, section_name in domain_section_map.items():
            domain_opinions = [o for o in opinions if o.domain == domain]
            if domain_opinions:
                lines = []
                for op in domain_opinions:
                    verdict_text = VERDICT_LABELS.get(op.verdict, op.verdict)
                    lines.append(f"- EQCR 结论：{verdict_text}")
                    if op.comment:
                        lines.append(f"  说明：{op.comment}")
                sections[section_name] = "\n".join(lines)
            else:
                sections[section_name] = "（未录入 EQCR 意见）"

        if notes:
            note_lines = []
            for n in notes:
                shared_mark = "【已分享】" if n.shared_to_team else ""
                note_lines.append(f"- {n.title} {shared_mark}")
                if n.content:
                    preview = n.content[:200]
                    ellipsis = "..." if len(n.content or "") > 200 else ""
                    note_lines.append(f"  {preview}{ellipsis}")
            sections["独立复核笔记摘要"] = "\n".join(note_lines)
        else:
            sections["独立复核笔记摘要"] = "（无独立复核笔记）"

        if shadow_comps:
            sc_lines = []
            for sc in shadow_comps:
                diff_mark = "⚠️ 有差异" if sc.has_diff else "✓ 一致"
                sc_lines.append(f"- {sc.computation_type}：{diff_mark}")
            sections["独立取数结果"] = "\n".join(sc_lines)
        else:
            sections["独立取数结果"] = "（未执行影子计算）"

        if disagreements:
            d_lines = []
            for d in disagreements:
                resolved = "已解决" if d.resolved_at else "未解决"
                d_lines.append(f"- 异议 [{resolved}]：{d.resolution or '待合议'}")
            sections["异议与合议结论"] = "\n".join(d_lines)
        else:
            sections["异议与合议结论"] = "（无异议记录）"

        agree_count = sum(1 for o in opinions if o.verdict == "agree")
        disagree_count = sum(1 for o in opinions if o.verdict == "disagree")
        sections["EQCR 总评与结论"] = (
            f"共录入 {len(opinions)} 条意见，其中认可 {agree_count} 条、"
            f"有异议 {disagree_count} 条。\n\n"
            "EQCR 总体结论：[请在此填写总体结论]"
        )

        return sections

    async def save_memo(
        self,
        project_id: UUID,
        sections: dict[str, str],
    ) -> dict:
        """保存编辑后的备忘录内容到 Project.wizard_state.eqcr_memo。
        
        R7-S3-04 Task 21：保存前将旧 sections 压入 history 数组（最多 5 版）。
        """
        proj = (
            await self.db.execute(
                sa.select(Project).where(
                    Project.id == project_id,
                    Project.is_deleted == False,  # noqa: E712
                )
            )
        ).scalar_one_or_none()
        if proj is None:
            raise ValueError("项目不存在")

        wizard = dict(proj.wizard_state) if proj.wizard_state else {}
        existing = wizard.get("eqcr_memo") or {}
        
        # R7-S3-04：版本历史 — 保存前将旧 sections 压入 history
        old_sections = existing.get("sections")
        history: list = existing.get("history") or []
        if old_sections:
            history.append({
                "version": len(history) + 1,
                "saved_at": existing.get("updated_at", datetime.now(timezone.utc).isoformat()),
                "sections_snapshot": old_sections,
            })
            # 最多保留 5 版
            if len(history) > 5:
                history = history[-5:]
        
        wizard["eqcr_memo"] = {
            **existing,
            "sections": sections,
            "history": history,
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "status": "draft",
        }
        proj.wizard_state = wizard
        await self.db.flush()

        return {
            "status": "saved",
            "updated_at": wizard["eqcr_memo"]["updated_at"],
            "history_count": len(history),
        }

    async def finalize_memo(
        self,
        project_id: UUID,
    ) -> dict:
        """定稿备忘录：生成 Word + PDF 文件，保存到 wizard_state。

        - Word 字节通过 ``build_memo_docx_bytes`` 生成
        - PDF 由 LibreOffice headless 转换
        - 文件路径保存在 ``wizard_state.eqcr_memo.files``
        """
        proj = (
            await self.db.execute(
                sa.select(Project).where(
                    Project.id == project_id,
                    Project.is_deleted == False,  # noqa: E712
                )
            )
        ).scalar_one_or_none()
        if proj is None:
            raise ValueError("项目不存在")

        wizard = dict(proj.wizard_state) if proj.wizard_state else {}
        memo = dict(wizard.get("eqcr_memo") or {})
        if not memo.get("sections"):
            raise ValueError("备忘录尚未生成或保存，请先生成并编辑")

        # 生成 Word 字节
        project_name = proj.name or "project"
        docx_bytes = build_memo_docx_bytes(
            project_name=project_name,
            client_name=proj.client_name or "",
            sections=memo["sections"],
        )

        # 存文件到项目存储目录（按项目 UUID 隔离）
        storage_root = Path("backend") / "wp_storage" / str(project_id) / "eqcr_memo"
        try:
            storage_root.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            logger.warning(
                "EQCR memo storage dir create failed: %s, fallback to tempdir", e
            )
            storage_root = Path(tempfile.gettempdir()) / "eqcr_memo" / str(project_id)
            storage_root.mkdir(parents=True, exist_ok=True)

        docx_path = storage_root / "eqcr_memo.docx"
        pdf_path = storage_root / "eqcr_memo.pdf"
        docx_path.write_bytes(docx_bytes)

        # 尝试 PDF 转换（LibreOffice headless；失败则标记 pdf_failed=True）
        pdf_generated = False
        pdf_error: str | None = None
        try:
            _convert_docx_to_pdf(docx_path, pdf_path)
            pdf_generated = pdf_path.exists() and pdf_path.stat().st_size > 0
        except Exception as e:
            pdf_error = str(e)
            logger.warning(
                "EQCR memo PDF conversion failed: %s (docx saved at %s)",
                e, docx_path,
            )

        now_iso = datetime.now(timezone.utc).isoformat()
        memo.update({
            "status": "finalized",
            "finalized_at": now_iso,
            "files": {
                "docx": str(docx_path),
                "pdf": str(pdf_path) if pdf_generated else None,
                "pdf_failed": not pdf_generated,
                "pdf_error": pdf_error,
            },
        })
        wizard["eqcr_memo"] = memo
        proj.wizard_state = wizard
        await self.db.flush()

        return {
            "status": "finalized",
            "finalized_at": now_iso,
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path) if pdf_generated else None,
            "pdf_generated": pdf_generated,
            "pdf_error": pdf_error,
        }


# ---------------------------------------------------------------------------
# 归档章节生成器（R1 archive_section_registry 落地后注册）
# ---------------------------------------------------------------------------


def eqcr_memo_pdf_generator(project_id: UUID, wizard_state: dict) -> bytes | None:
    """归档引擎调用：返回 EQCR 备忘录 PDF 字节。

    在 R1 ``archive_section_registry`` 落地后，通过
    ``registry.register('02', 'eqcr_memo.pdf', eqcr_memo_pdf_generator)``
    注册，归档包导出时自动拼接。

    **当前状态**：registry 尚未落地，此函数作为接口预留；实际调用点待
    R1 任务完成后接入。
    """
    memo = (wizard_state or {}).get("eqcr_memo") or {}
    files = memo.get("files") or {}
    pdf_path = files.get("pdf")
    if not pdf_path:
        return None
    try:
        return Path(pdf_path).read_bytes()
    except (OSError, FileNotFoundError):
        return None


# ---------------------------------------------------------------------------
# Word 生成（纯函数，便于测试）
# ---------------------------------------------------------------------------


def build_memo_docx_bytes(
    project_name: str,
    client_name: str,
    sections: dict[str, str],
) -> bytes:
    """生成 EQCR 备忘录 Word 文档字节。

    使用项目已装的 ``python-docx`` 库（phase13 附注导出同款）。
    """
    import docx
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # 页面
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 标题
    title = doc.add_heading("独立复核备忘录（EQCR Memo）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"项目：{project_name}    客户：{client_name}")
    meta_run.font.size = Pt(10)

    doc.add_paragraph("")  # 空行

    # 各章节按 MEMO_SECTIONS 固定顺序输出
    for section_name in MEMO_SECTIONS:
        heading = doc.add_heading(section_name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        content = sections.get(section_name) or "（未填写）"
        for line in content.splitlines():
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Cm(0.74)

    # 保存到字节流
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF 转换（LibreOffice headless，复用项目既有方案）
# ---------------------------------------------------------------------------


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """通过 LibreOffice headless 把 docx 转成 pdf。

    失败时抛异常由调用方捕获。
    """
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if soffice is None:
        raise FileNotFoundError("LibreOffice (libreoffice/soffice) not in PATH")

    cmd = [
        soffice, "--headless",
        "--convert-to", "pdf",
        "--outdir", str(pdf_path.parent),
        str(docx_path),
    ]
    try:
        result = subprocess.run(
            cmd, capture_output=True, timeout=60, check=False
        )
    except subprocess.TimeoutExpired as e:
        raise RuntimeError(f"LibreOffice PDF conversion timeout: {e}") from e

    if result.returncode != 0:
        stderr_text = result.stderr.decode("utf-8", errors="replace")[:500]
        raise RuntimeError(
            f"LibreOffice PDF conversion failed (exit {result.returncode}): {stderr_text}"
        )

    if not pdf_path.exists():
        # LibreOffice 会按 source 文件名生成 PDF，可能名字不同
        candidates = list(pdf_path.parent.glob("*.pdf"))
        if candidates:
            candidates[0].rename(pdf_path)
        else:
            raise RuntimeError(f"LibreOffice output PDF not found in {pdf_path.parent}")

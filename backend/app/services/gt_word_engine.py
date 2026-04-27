"""致同标准 Word 导出引擎（方案A降级基线）

使用 python-docx 从零生成符合致同排版规范的 Word 文档。
页面设置 + 样式 + 三线表 + 页眉页脚 + 三色文本 + 千分位。
"""

from __future__ import annotations

import logging
from decimal import Decimal
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)

try:
    import docx
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Emu, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class GTWordEngine:
    """致同标准 Word 导出引擎 — 统一排版规范"""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx 未安装，无法使用 GTWordEngine")
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    # ------------------------------------------------------------------
    # 页面设置
    # ------------------------------------------------------------------

    def _setup_page(self):
        """页面设置：左3/右3.18/上3.2/下2.54 cm，页眉页脚1.3cm"""
        section = self.doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(3.18)
        section.top_margin = Cm(3.2)
        section.bottom_margin = Cm(2.54)
        section.header_distance = Cm(1.3)
        section.footer_distance = Cm(1.3)

    # ------------------------------------------------------------------
    # 样式注册
    # ------------------------------------------------------------------

    def _setup_styles(self):
        """注册 GT 自定义样式

        中文：仿宋_GB2312 小四(12pt)
        英文/数字：Arial Narrow
        """
        styles = self.doc.styles

        # GT正文样式
        if "GT正文" not in [s.name for s in styles]:
            style = styles.add_style("GT正文", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = "Arial Narrow"
            font.size = Pt(12)
            # 设置中文字体
            rpr = style.element.get_or_add_rPr()
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            rpr.append(rfonts)
            # 段落格式：段前0行 段后0.9行 单倍行距
            pf = style.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(10.8)  # 0.9行 ≈ 10.8pt

    # ------------------------------------------------------------------
    # 页眉页脚
    # ------------------------------------------------------------------

    def setup_header_footer(self, firm_name: str, project_name: str):
        """页眉：事务所名称(左) + 项目名称(右)
        页脚：第 X 页 共 Y 页（居中）
        """
        section = self.doc.sections[0]

        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()

        # 左侧事务所名称
        run_left = hp.add_run(firm_name)
        run_left.font.name = "Arial Narrow"
        run_left.font.size = Pt(12)
        rpr = run_left._r.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        rpr.append(rfonts)

        # Tab 分隔 + 右侧项目名称
        hp.add_run("\t\t")
        run_right = hp.add_run(project_name)
        run_right.font.name = "Arial Narrow"
        run_right.font.size = Pt(12)
        rpr2 = run_right._r.get_or_add_rPr()
        rfonts2 = OxmlElement("w:rFonts")
        rfonts2.set(qn("w:eastAsia"), "仿宋_GB2312")
        rpr2.append(rfonts2)

        # 页脚：第 X 页 共 Y 页
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_page_number_field(fp)

    def _add_page_number_field(self, paragraph):
        """插入 '第 X 页 共 Y 页' 页码域代码"""
        run1 = paragraph.add_run("第 ")
        run1.font.size = Pt(12)

        # PAGE 域
        self._insert_field(paragraph, " PAGE ")

        run2 = paragraph.add_run(" 页 共 ")
        run2.font.size = Pt(12)

        # NUMPAGES 域
        self._insert_field(paragraph, " NUMPAGES ")

        run3 = paragraph.add_run(" 页")
        run3.font.size = Pt(12)

    @staticmethod
    def _insert_field(paragraph, field_code: str):
        """插入 Word 域代码（PAGE / NUMPAGES）"""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_code
        run._r.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)

    # ------------------------------------------------------------------
    # 标题
    # ------------------------------------------------------------------

    def add_heading(self, text: str, level: int = 1):
        """多级标题"""
        heading = self.doc.add_heading(text, level=level)
        # 左缩进 -2 字符，首行不缩进
        pf = heading.paragraph_format
        pf.left_indent = Cm(-0.85)  # 约 -2 字符
        pf.first_line_indent = Cm(0)
        return heading

    # ------------------------------------------------------------------
    # 段落
    # ------------------------------------------------------------------

    def add_paragraph(self, text: str, style: str | None = None, after_table: bool = False):
        """正文段落

        段前0行 段后0.9行 单倍行距
        after_table=True时：段前0.5行
        """
        p = self.doc.add_paragraph(text, style=style)
        pf = p.paragraph_format
        pf.space_before = Pt(6) if after_table else Pt(0)
        pf.space_after = Pt(10.8)
        return p

    # ------------------------------------------------------------------
    # 三线表
    # ------------------------------------------------------------------

    def add_table(
        self,
        headers: list[str],
        rows: list[list[Any]],
        total_row: list[Any] | None = None,
    ):
        """三线表：上下1磅，标题行下1/2磅，无左右边框

        标题列左对齐，数据列右对齐，垂直居中
        标题行+合计行加粗
        """
        num_rows = 1 + len(rows) + (1 if total_row else 0)
        table = self.doc.add_table(rows=num_rows, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置三线表边框
        self._set_three_line_borders(table)

        # 填充表头
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(h)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = "Arial Narrow"

        # 填充数据行
        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                formatted = self.format_number(val) if isinstance(val, (int, float, Decimal)) else str(val or "")
                cell.text = formatted
                for p in cell.paragraphs:
                    # 数据列右对齐（第一列除外）
                    if c_idx > 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.size = Pt(12)
                        run.font.name = "Arial Narrow"

        # 填充合计行
        if total_row:
            last_idx = 1 + len(rows)
            for c_idx, val in enumerate(total_row):
                cell = table.rows[last_idx].cells[c_idx]
                formatted = self.format_number(val) if isinstance(val, (int, float, Decimal)) else str(val or "")
                cell.text = formatted
                for p in cell.paragraphs:
                    if c_idx > 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.name = "Arial Narrow"

        return table

    @staticmethod
    def _set_three_line_borders(table):
        """三线表边框：上下1磅，标题行下1/2磅，无左右"""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        borders = OxmlElement("w:tblBorders")
        # 上边框 1磅
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "8")  # 1磅 = 8 half-points
        top.set(qn("w:space"), "0")
        top.set(qn("w:color"), "000000")
        borders.append(top)

        # 下边框 1磅
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)

        # 无左右边框
        for side in ("left", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            borders.append(el)

        # 内部水平线（标题行下 1/2 磅）
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "single")
        insideH.set(qn("w:sz"), "4")  # 1/2磅 = 4 half-points
        insideH.set(qn("w:space"), "0")
        insideH.set(qn("w:color"), "000000")
        borders.append(insideH)

        # 无内部垂直线
        insideV = OxmlElement("w:insideV")
        insideV.set(qn("w:val"), "none")
        insideV.set(qn("w:sz"), "0")
        insideV.set(qn("w:space"), "0")
        borders.append(insideV)

        tbl_pr.append(borders)

    # ------------------------------------------------------------------
    # 数字格式化
    # ------------------------------------------------------------------

    @staticmethod
    def format_number(value: Any, decimals: int = 2) -> str:
        """数字格式化：千分位 + 负数括号

        12345.67 → '12,345.67'
        -1234.56 → '(1,234.56)'
        None/0 → '-'
        """
        if value is None:
            return "-"

        try:
            num = float(value)
        except (ValueError, TypeError):
            return str(value)

        if num == 0:
            return "-"

        if num < 0:
            formatted = f"{abs(num):,.{decimals}f}"
            return f"({formatted})"
        else:
            return f"{num:,.{decimals}f}"

    # ------------------------------------------------------------------
    # 三色文本处理
    # ------------------------------------------------------------------

    @staticmethod
    def process_color_text(paragraph) -> None:
        """三色文本处理（附注专用）

        蓝色 run → 删除
        红色 run → 转黑色
        黑色 run → 保留
        """
        if not HAS_DOCX:
            return

        runs_to_remove = []
        for run in paragraph.runs:
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            if color is not None:
                # 蓝色系（R<100, B>150）→ 删除
                if color.red < 100 and color.blue > 150:
                    runs_to_remove.append(run)
                # 红色系（R>200, G<100）→ 转黑色
                elif color.red > 200 and color.green < 100:
                    run.font.color.rgb = RGBColor(0, 0, 0)

        # 删除蓝色 run
        for run in runs_to_remove:
            run._r.getparent().remove(run._r)

    # ------------------------------------------------------------------
    # 分页符
    # ------------------------------------------------------------------

    def add_page_break(self):
        """插入分页符"""
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # 保存
    # ------------------------------------------------------------------

    def save(self, output: BytesIO | str | None = None) -> BytesIO:
        """保存为 BytesIO 或文件路径"""
        if output is None:
            output = BytesIO()

        if isinstance(output, str):
            self.doc.save(output)
            return BytesIO()  # 返回空 BytesIO 表示已保存到文件
        else:
            self.doc.save(output)
            output.seek(0)
            return output

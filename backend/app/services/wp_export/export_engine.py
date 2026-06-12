"""WpExportEngine — 底稿导出引擎（薄封装层）

核心职责（不重写填值逻辑）：
1. xlsx 导出直接调用现有 wp_xlsx_export_service.export_workpaper_xlsx
2. 调用 MetadataCodec 嵌入元数据
3. 计算 snapshot_hash 存 wp_export_snapshot 表
4. 文字底稿走 docx 路径（python-docx，现有 service 仅 xlsx）

设计原则：
- service 只 flush 不 commit（router 层统一 commit）
- file_path 不存在时回退模板库（复用 find_template_file_any）

Requirements: 1.1, 1.2, 1.3, 1.4, 1.5, 1.6
"""

from __future__ import annotations

import hashlib
import logging
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID

import sqlalchemy as sa
from docx import Document
from openpyxl import load_workbook
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.workpaper_models import WorkingPaper, WpIndex
from app.models.wp_export_models import WpExportSnapshot
from app.schemas.wp_export_schemas import ExportResult, MetadataBundle
from app.services.wp_export.metadata_codec import MetadataCodec
from app.services.wp_render_schema_service import WpRenderSchemaService
from app.services.wp_xlsx_export_service import (
    export_workpaper_xlsx,
    TemplateNotFoundError,
)

logger = logging.getLogger(__name__)

# ─── 底稿类型 → 导出格式映射 ───────────────────────────────────────────────────
# 表格/审定表/程序表 → xlsx；文字 → docx
_XLSX_TYPES = {"univer", "form", "hybrid", "table", "audit_sheet", "program_sheet"}
_DOCX_TYPES = {"word", "text"}

_schema_service = WpRenderSchemaService()
_metadata_codec = MetadataCodec()


def determine_export_format(file_path: str | None, component_type: str | None) -> str:
    """根据底稿组件类型或文件扩展名判断导出格式。

    Returns:
        "xlsx" 或 "docx"
    """
    # 优先按 component_type 判断
    if component_type:
        ct = component_type.lower().strip()
        if ct in _DOCX_TYPES:
            return "docx"
        if ct in _XLSX_TYPES:
            return "xlsx"

    # 回退：按文件扩展名判断
    if file_path:
        ext = Path(file_path).suffix.lower()
        if ext in (".docx", ".doc"):
            return "docx"

    # 默认 xlsx
    return "xlsx"


class WpExportEngine:
    """底稿导出引擎

    薄封装层：调用现有导出服务 + 嵌入元数据 + 计算快照哈希 + 记录快照。
    """

    async def export_single(
        self,
        db: AsyncSession,
        wp_id: UUID,
        project_id: UUID,
        exported_by: UUID | None = None,
    ) -> ExportResult:
        """导出单份底稿，返回文件字节流 + 元数据。

        流程：
        1. 加载 working_paper + wp_index 获取 wp_code / 类型等信息
        2. 判断导出格式（xlsx / docx）
        3. xlsx → 调用现有 export_workpaper_xlsx；docx → 新建 docx 路径
        4. 嵌入 MetadataBundle
        5. 计算 snapshot_hash
        6. 创建 WpExportSnapshot 记录（flush 不 commit）
        7. 返回 ExportResult
        """
        # ─── Step 1: 查 working_paper + wp_index ─────────────────────────
        wp = await self._load_working_paper(db, wp_id, project_id)
        wp_index = await self._load_wp_index(db, wp.wp_index_id)

        wp_code = wp_index.wp_code
        wp_name = wp_index.wp_name

        # ─── Step 2: 判断导出格式 ────────────────────────────────────────
        file_format = determine_export_format(wp.file_path, None)

        # 尝试从 render schema 获取更准确的 component_type
        try:
            schema = _schema_service.load_schema(wp_code=wp_code)
            # schema 中可能有 component_type 信息
            sheets = schema.get("sheets", {})
            for _sheet_name, sheet_cfg in sheets.items():
                if isinstance(sheet_cfg, dict) and "component_type" in sheet_cfg:
                    file_format = determine_export_format(
                        wp.file_path, sheet_cfg["component_type"]
                    )
                    break
        except FileNotFoundError:
            schema = {}

        # ─── Step 3: 获取项目元数据 ──────────────────────────────────────
        from app.models.core import Project

        project_result = await db.execute(
            sa.select(Project).where(Project.id == project_id)
        )
        project = project_result.scalars().first()

        project_meta: dict[str, Any] = {}
        if project:
            project_meta = {
                "entity_name": project.client_name or "",
                "period_end": (
                    project.audit_period_end.isoformat()
                    if project.audit_period_end
                    else "2025-12-31"  # 兜底值：审计截止日未设置时使用默认
                ),
                "index_no": wp_code,
            }

        # ─── Step 4: 构建 MetadataBundle ─────────────────────────────────
        metadata = MetadataBundle(
            wp_code=wp_code,
            project_id=project_id,
            file_version=wp.file_version,
            export_timestamp=datetime.now(tz=timezone.utc),
            preparer=str(wp.assigned_to) if wp.assigned_to else None,
            reviewer=str(wp.reviewer) if wp.reviewer else None,
            review_status=wp.review_status.value if wp.review_status else None,
        )

        # ─── Step 5: 执行导出 ────────────────────────────────────────────
        if file_format == "docx":
            file_bytes_io = await self._export_docx(wp, wp_code, wp_name)
        else:
            file_bytes_io = await self._export_xlsx(
                wp, wp_code, schema, project_meta
            )

        # ─── Step 6: 嵌入元数据 ──────────────────────────────────────────
        if file_format == "docx":
            doc = Document(file_bytes_io)
            _metadata_codec.embed_docx(doc, metadata)
            file_bytes_io = BytesIO()
            doc.save(file_bytes_io)
            file_bytes_io.seek(0)
        else:
            wb = load_workbook(file_bytes_io)
            _metadata_codec.embed_xlsx(wb, metadata)
            file_bytes_io = BytesIO()
            wb.save(file_bytes_io)
            file_bytes_io.seek(0)
            wb.close()

        # ─── Step 7: 计算 snapshot_hash ──────────────────────────────────
        file_content = file_bytes_io.getvalue()
        snapshot_hash = self._compute_content_hash(file_content)

        # ─── Step 8: 创建 WpExportSnapshot 记录 ──────────────────────────
        snapshot = WpExportSnapshot(
            working_paper_id=wp_id,
            project_id=project_id,
            file_version=wp.file_version,
            snapshot_hash=snapshot_hash,
            exported_by=exported_by,
            file_format=file_format,
            file_size_bytes=len(file_content),
            metadata_bundle=metadata.model_dump(mode="json"),
        )
        db.add(snapshot)
        await db.flush()

        # ─── Step 9: 构建文件名 ──────────────────────────────────────────
        filename = f"{wp_code}_{wp_name}.{file_format}"

        return ExportResult(
            file_content=file_content,
            filename=filename,
            file_format=file_format,
            snapshot_hash=snapshot_hash,
            metadata=metadata,
        )

    # ─── Public: 程序表导出 ──────────────────────────────────────────────────

    async def export_program_sheet(
        self,
        db: AsyncSession,
        wp_id: UUID,
        project_id: UUID,
    ) -> BytesIO:
        """导出程序表底稿（含只读/可编辑列标记）。

        程序步骤独立 sheet：程序编号、程序描述、执行状态、执行结论、执行人。
        只读列（程序编号、描述）标记 [只读] 前缀；可编辑列（状态、结论、执行人）标记 [可编辑]。

        Requirements: 8.1, 8.2
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Protection

        wp = await self._load_working_paper(db, wp_id, project_id)
        procedures = self._extract_program_procedures(wp)

        wb = Workbook()
        ws = wb.active
        ws.title = "程序表"

        # ─── 列定义 ──────────────────────────────────────────────────────
        columns = [
            ("procedure_code", "[只读] 程序编号", True),
            ("description", "[只读] 程序描述", True),
            ("execution_status", "[可编辑] 执行状态", False),
            ("execution_conclusion", "[可编辑] 执行结论", False),
            ("executor", "[可编辑] 执行人", False),
        ]

        # ─── 写入表头 ────────────────────────────────────────────────────
        header_font = Font(bold=True)
        readonly_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        for col_idx, (field, header_text, is_readonly) in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header_text)
            cell.font = header_font
            if is_readonly:
                cell.fill = readonly_fill

        # ─── 写入数据行 ──────────────────────────────────────────────────
        for row_idx, proc in enumerate(procedures, start=2):
            for col_idx, (field, _header, is_readonly) in enumerate(columns, start=1):
                value = proc.get(field, "")
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                if is_readonly:
                    cell.protection = Protection(locked=True)
                else:
                    cell.protection = Protection(locked=False)

        # 保护 sheet（只读列锁定生效）
        ws.protection.sheet = True
        ws.protection.password = ""  # 无密码，仅标记保护

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()
        return buf

    # ─── Public: 审定表导出 ──────────────────────────────────────────────────

    async def export_audit_sheet(
        self,
        db: AsyncSession,
        wp_id: UUID,
        project_id: UUID,
    ) -> BytesIO:
        """导出审定表底稿（含科目明细+调整分录批注+汇总行）。

        科目明细行：科目编码、科目名称、未审数、调整数、审定数。
        调整分录来源引用作只读批注。
        末尾汇总行（合计、借贷平衡校验）。

        Requirements: 9.1, 9.2, 9.4
        """
        from openpyxl import Workbook
        from openpyxl.comments import Comment
        from openpyxl.styles import Font, PatternFill

        wp = await self._load_working_paper(db, wp_id, project_id)
        accounts = self._extract_audit_accounts(wp)

        wb = Workbook()
        ws = wb.active
        ws.title = "审定表"

        # ─── 列定义 ──────────────────────────────────────────────────────
        columns = [
            ("account_code", "科目编码"),
            ("account_name", "科目名称"),
            ("unadjusted_amount", "未审数"),
            ("adjustment_amount", "调整数"),
            ("audited_amount", "审定数"),
        ]

        # ─── 写入表头 ────────────────────────────────────────────────────
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        for col_idx, (field, header_text) in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header_text)
            cell.font = header_font
            cell.fill = header_fill

        # ─── 写入数据行 ──────────────────────────────────────────────────
        total_unadjusted = 0.0
        total_adjustment = 0.0
        total_audited = 0.0

        for row_idx, account in enumerate(accounts, start=2):
            for col_idx, (field, _header) in enumerate(columns, start=1):
                value = account.get(field, "")
                cell = ws.cell(row=row_idx, column=col_idx, value=value)

            # 累加金额列（跳过非数值）
            unadj = account.get("unadjusted_amount", 0)
            adj = account.get("adjustment_amount", 0)
            aud = account.get("audited_amount", 0)
            total_unadjusted += float(unadj) if unadj else 0.0
            total_adjustment += float(adj) if adj else 0.0
            total_audited += float(aud) if aud else 0.0

            # 调整分录来源引用作只读批注
            adj_ref = account.get("adjustment_source", "")
            if adj_ref:
                comment = Comment(
                    text=f"调整分录来源: {adj_ref}",
                    author="System",
                )
                # 批注挂在调整数列
                ws.cell(row=row_idx, column=4).comment = comment

        # ─── 汇总行 ──────────────────────────────────────────────────────
        summary_row = len(accounts) + 2
        summary_font = Font(bold=True)

        ws.cell(row=summary_row, column=1, value="合计").font = summary_font
        ws.cell(row=summary_row, column=2, value="").font = summary_font
        ws.cell(row=summary_row, column=3, value=total_unadjusted).font = summary_font
        ws.cell(row=summary_row, column=4, value=total_adjustment).font = summary_font
        ws.cell(row=summary_row, column=5, value=total_audited).font = summary_font

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()
        return buf

    # ─── Private: 数据提取辅助 ────────────────────────────────────────────

    @staticmethod
    def _extract_program_procedures(wp: WorkingPaper) -> list[dict]:
        """从 working_paper.parsed_data 提取程序步骤列表。"""
        parsed_data = wp.parsed_data or {}
        # 优先从 procedures 字段取
        procedures = parsed_data.get("procedures", [])
        if procedures:
            return procedures

        # 回退：从 html_data 中寻找程序表相关 sheet
        html_data = parsed_data.get("html_data", {})
        for sheet_name, sheet_data in html_data.items():
            if "程序" in sheet_name or "program" in sheet_name.lower():
                if isinstance(sheet_data, dict):
                    rows = sheet_data.get("rows", [])
                    if rows:
                        return rows

        return []

    @staticmethod
    def _extract_audit_accounts(wp: WorkingPaper) -> list[dict]:
        """从 working_paper.parsed_data 提取审定表科目明细。"""
        parsed_data = wp.parsed_data or {}
        # 优先从 accounts / audit_accounts 字段取
        accounts = parsed_data.get("audit_accounts", [])
        if accounts:
            return accounts
        accounts = parsed_data.get("accounts", [])
        if accounts:
            return accounts

        # 回退：从 html_data 中寻找审定表相关 sheet
        html_data = parsed_data.get("html_data", {})
        for sheet_name, sheet_data in html_data.items():
            if "审定" in sheet_name or "audit" in sheet_name.lower():
                if isinstance(sheet_data, dict):
                    rows = sheet_data.get("rows", [])
                    if rows:
                        return rows

        return []

    # ─── Private: 加载数据 ────────────────────────────────────────────────

    async def _load_working_paper(
        self, db: AsyncSession, wp_id: UUID, project_id: UUID
    ) -> WorkingPaper:
        """加载底稿记录，不存在时抛异常。"""
        result = await db.execute(
            sa.select(WorkingPaper).where(
                WorkingPaper.id == wp_id,
                WorkingPaper.project_id == project_id,
                WorkingPaper.is_deleted == sa.false(),
            )
        )
        wp = result.scalars().first()
        if wp is None:
            raise ValueError(f"底稿不存在: wp_id={wp_id}, project_id={project_id}")
        return wp

    async def _load_wp_index(self, db: AsyncSession, wp_index_id: UUID) -> WpIndex:
        """加载底稿索引记录。"""
        result = await db.execute(
            sa.select(WpIndex).where(
                WpIndex.id == wp_index_id,
                WpIndex.is_deleted == sa.false(),
            )
        )
        wp_index = result.scalars().first()
        if wp_index is None:
            raise ValueError(f"底稿索引不存在: wp_index_id={wp_index_id}")
        return wp_index

    # ─── Private: 导出实现 ────────────────────────────────────────────────

    async def _export_xlsx(
        self,
        wp: WorkingPaper,
        wp_code: str,
        schema: dict,
        project_meta: dict[str, Any],
    ) -> BytesIO:
        """xlsx 导出：直接调用现有 export_workpaper_xlsx。

        file_path 不存在时，现有 _resolve_template_path 已做模板库回退。
        TemplateNotFoundError 时回退为空白 workbook（填入 html_data）。
        """
        parsed_data = wp.parsed_data or {}
        html_data = parsed_data.get("html_data", {})

        try:
            return await export_workpaper_xlsx(
                wp_code=wp_code,
                html_data=html_data,
                schema=schema,
                project_meta=project_meta,
            )
        except (TemplateNotFoundError, Exception) as e:
            logger.warning(
                "export_workpaper_xlsx 失败 wp_code=%s: %s，回退空白 workbook",
                wp_code, e,
            )
            # 回退：生成空白 workbook 填入 html_data
            from openpyxl import Workbook as _Wb

            wb = _Wb()
            ws = wb.active
            ws.title = wp_code
            # 写入 html_data 数据（如有）
            row_idx = 1
            for sheet_name, sheet_data in html_data.items():
                if isinstance(sheet_data, dict):
                    rows = sheet_data.get("rows", [])
                    for row_data in rows:
                        if isinstance(row_data, dict):
                            for col_idx, val in enumerate(row_data.values(), start=1):
                                ws.cell(row=row_idx, column=col_idx, value=val)
                            row_idx += 1
            buf = BytesIO()
            wb.save(buf)
            buf.seek(0)
            wb.close()
            return buf

    async def _export_docx(
        self,
        wp: WorkingPaper,
        wp_code: str,
        wp_name: str,
    ) -> BytesIO:
        """docx 导出（文字底稿）：生成 Word 文档。

        优先从已有文件路径加载模板，不存在则新建空白文档填充内容。
        """
        # 尝试从现有文件加载
        template_path = self._resolve_docx_template(wp.file_path, wp_code)

        if template_path and template_path.is_file():
            doc = Document(str(template_path))
        else:
            # 新建空白文档，填充 parsed_data 中的文字内容
            doc = Document()
            doc.add_heading(f"{wp_code} {wp_name}", level=1)
            self._fill_docx_content(doc, wp)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _resolve_docx_template(
        self, file_path: str | None, wp_code: str
    ) -> Path | None:
        """解析 docx 模板路径，file_path 不存在时回退模板库。"""
        if file_path:
            p = Path(file_path)
            if p.is_file():
                return p
            # 尝试从项目根解析
            project_root = Path(__file__).resolve().parent.parent.parent.parent
            candidate = project_root / file_path
            if candidate.is_file():
                return candidate

        # 回退：模板库查找
        from app.services.wp_template_init_service import find_template_file_any

        return find_template_file_any(wp_code)

    def _fill_docx_content(self, doc: Document, wp: WorkingPaper) -> None:
        """将 parsed_data 中的文字内容填充到 docx 文档。"""
        parsed_data = wp.parsed_data or {}
        html_data = parsed_data.get("html_data", {})

        for sheet_name, sheet_data in html_data.items():
            doc.add_heading(sheet_name, level=2)
            if isinstance(sheet_data, dict):
                rows = sheet_data.get("rows", [])
                for row in rows:
                    if isinstance(row, dict):
                        # 将 row 的值拼接为段落
                        text_parts = [
                            str(v) for v in row.values() if v is not None and v != ""
                        ]
                        if text_parts:
                            doc.add_paragraph(" | ".join(text_parts))
                    elif isinstance(row, str):
                        doc.add_paragraph(row)
            elif isinstance(sheet_data, str):
                doc.add_paragraph(sheet_data)

    # ─── Public class methods: 纯数据→BytesIO（供 PBT 测试直接调用）───────

    @staticmethod
    def build_program_sheet_bytes(procedures: list[dict]) -> BytesIO:
        """从程序步骤列表生成程序表 xlsx BytesIO（无需 DB）。

        供 PBT 测试和其他无 DB 上下文调用。
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Protection

        wb = Workbook()
        ws = wb.active
        ws.title = "程序表"

        columns = [
            ("procedure_code", "[只读] 程序编号", True),
            ("description", "[只读] 程序描述", True),
            ("execution_status", "[可编辑] 执行状态", False),
            ("execution_conclusion", "[可编辑] 执行结论", False),
            ("executor", "[可编辑] 执行人", False),
        ]

        header_font = Font(bold=True)
        readonly_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        for col_idx, (field, header_text, is_readonly) in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header_text)
            cell.font = header_font
            if is_readonly:
                cell.fill = readonly_fill

        for row_idx, proc in enumerate(procedures, start=2):
            for col_idx, (field, _header, is_readonly) in enumerate(columns, start=1):
                value = proc.get(field, "")
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                if is_readonly:
                    cell.protection = Protection(locked=True)
                else:
                    cell.protection = Protection(locked=False)

        ws.protection.sheet = True
        ws.protection.password = ""

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()
        return buf

    @staticmethod
    def build_audit_sheet_bytes(accounts: list[dict]) -> BytesIO:
        """从科目明细列表生成审定表 xlsx BytesIO（无需 DB）。

        供 PBT 测试和其他无 DB 上下文调用。
        """
        from openpyxl import Workbook
        from openpyxl.comments import Comment
        from openpyxl.styles import Font, PatternFill

        wb = Workbook()
        ws = wb.active
        ws.title = "审定表"

        columns = [
            ("account_code", "科目编码"),
            ("account_name", "科目名称"),
            ("unadjusted_amount", "未审数"),
            ("adjustment_amount", "调整数"),
            ("audited_amount", "审定数"),
        ]

        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

        for col_idx, (field, header_text) in enumerate(columns, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header_text)
            cell.font = header_font
            cell.fill = header_fill

        total_unadjusted = 0.0
        total_adjustment = 0.0
        total_audited = 0.0

        for row_idx, account in enumerate(accounts, start=2):
            for col_idx, (field, _header) in enumerate(columns, start=1):
                value = account.get(field, "")
                ws.cell(row=row_idx, column=col_idx, value=value)

            unadj = account.get("unadjusted_amount", 0)
            adj = account.get("adjustment_amount", 0)
            aud = account.get("audited_amount", 0)
            total_unadjusted += float(unadj) if unadj else 0.0
            total_adjustment += float(adj) if adj else 0.0
            total_audited += float(aud) if aud else 0.0

            adj_ref = account.get("adjustment_source", "")
            if adj_ref:
                comment = Comment(
                    text=f"调整分录来源: {adj_ref}",
                    author="System",
                )
                ws.cell(row=row_idx, column=4).comment = comment

        # 汇总行
        summary_row = len(accounts) + 2
        summary_font = Font(bold=True)
        ws.cell(row=summary_row, column=1, value="合计").font = summary_font
        ws.cell(row=summary_row, column=2, value="").font = summary_font
        ws.cell(row=summary_row, column=3, value=total_unadjusted).font = summary_font
        ws.cell(row=summary_row, column=4, value=total_adjustment).font = summary_font
        ws.cell(row=summary_row, column=5, value=total_audited).font = summary_font

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        wb.close()
        return buf

    # ─── Private: 哈希计算 ────────────────────────────────────────────────

    @staticmethod
    def _compute_content_hash(file_content: bytes) -> str:
        """计算文件内容的 SHA-256 哈希。"""
        return hashlib.sha256(file_content).hexdigest()

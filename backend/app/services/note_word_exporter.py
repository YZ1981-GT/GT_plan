"""附注 Word 导出引擎 — 致同标准格式

Sprint 6 Task 6.1 + 6.2: 重写 NoteWordExporter
Sprint 2 Task 2.2: 接入 ``backend/data/note_export_template.docx`` GTNote* 模板
                   + 6 个 D7 helper（GTNote 命名空间）
                   + D1 sidecar 公式/手工标记可选渲染

致同标准格式：
  - 页面设置：A4、左 3cm/右 3.18cm/上 3.2cm/下 2.54cm、页眉 1.3cm/页脚 1.3cm
  - 字体：中文仿宋_GB2312 小四(12pt)、数字 Arial Narrow
  - 标题层级：一级"一、二、三..."加粗、二级"（一）（二）..."、三级"1. 2. 3."
  - 表格样式：上下边框 1 磅、标题行下边框 1/2 磅、标题行加粗居中、数据行金额右对齐
  - 段落格式：段前 0 行/段后 0.9 行、单倍行距

Requirements: 4.2-4.10, 27.1-27.10, R5.2 致同 21 项排版规范
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Literal
from uuid import UUID

import sqlalchemy as sa
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml.parser import parse_xml
from docx.shared import Cm, Pt, RGBColor, Emu
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.report_models import DisclosureNote
from app.services.note_section_catalog import (
    build_variant_key,
    detect_heading_level,
    note_applies_to_report_scope,
    normalize_report_scope,
    normalize_template_type,
)
from app.services.note_section_numbering import compute_section_numbers
from app.services.note_word_dynamic_styles import (
    get_table_render_mode,
    should_skip_empty_section,
)
from app.services.template_manifest_loader import get_template_manifest_loader
from app.services.word_doc_utils import (
    delete_section_block,
    remove_section_markers,
    scan_section_blocks,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 致同标准格式常量
# ---------------------------------------------------------------------------

PAGE_MARGINS = {"top": Cm(3.2), "bottom": Cm(2.54), "left": Cm(3), "right": Cm(3.18)}
HEADER_MARGIN = Cm(1.3)
FOOTER_MARGIN = Cm(1.3)
BODY_FONT = "仿宋_GB2312"
BODY_SIZE = Pt(12)  # 小四
NUMBER_FONT = "Arial Narrow"
HEADING1_FONT_SIZE = Pt(12)
HEADING2_FONT_SIZE = Pt(12)
HEADING3_FONT_SIZE = Pt(12)

# 中文数字序列
_CN_NUMBERS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
               "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
               "二十一", "二十二", "二十三", "二十四", "二十五", "二十六", "二十七", "二十八", "二十九", "三十"]

_CN_NUMBERS_PAREN = ["（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）", "（九）", "（十）",
                     "（十一）", "（十二）", "（十三）", "（十四）", "（十五）", "（十六）", "（十七）", "（十八）", "（十九）", "（二十）"]


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def _set_cell_border(cell, **kwargs):
    """Set cell border. Usage: _set_cell_border(cell, top={"sz": 12, "val": "single"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f"w:{edge}")
            for attr, val in edge_data.items():
                element.set(qn(f"w:{attr}"), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_table_borders(table, top_sz=8, bottom_sz=8, header_bottom_sz=4):
    """Set 致同 standard table borders: top/bottom 1pt, header bottom 0.5pt."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge_name, sz in [("top", top_sz), ("bottom", bottom_sz)]:
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        borders.append(edge)

    # Remove left/right/insideV borders (致同 style = only top/bottom)
    for edge_name in ("left", "right", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "none")
        edge.set(qn("w:sz"), "0")
        edge.set(qn("w:space"), "0")
        borders.append(edge)

    # insideH = none (we'll set header row bottom separately)
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "none")
    insideH.set(qn("w:sz"), "0")
    insideH.set(qn("w:space"), "0")
    borders.append(insideH)

    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def _set_row_bottom_border(row, sz=4):
    """Set bottom border on a specific row (for header row 0.5pt line)."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


def _format_amount(value) -> str:
    """Format amount with thousands separator."""
    if value is None or value == "" or value == 0:
        return "-"
    try:
        num = float(value)
        if num == 0:
            return "-"
        return f"{num:,.2f}"
    except (ValueError, TypeError):
        return str(value)


def _is_amount(value) -> bool:
    """Check if a value looks like a numeric amount."""
    if value is None or value == "" or value == "-":
        return False
    try:
        float(value)
        return True
    except (ValueError, TypeError):
        return False


def _add_toc(doc):
    """Add TOC field code that auto-updates on open."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("目录将在打开文档时自动更新")
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)

    return paragraph


def _add_page_number_footer(section):
    """Add footer with '第 X 页 共 Y 页' format."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "第 "
    run1 = paragraph.add_run("第 ")
    run1.font.name = BODY_FONT
    run1.font.size = Pt(9)

    # PAGE field
    run2 = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run2._r.append(fldChar)

    run3 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run3._r.append(instrText)

    run4 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run4._r.append(fldChar2)

    # " 页 共 "
    run5 = paragraph.add_run(" 页 共 ")
    run5.font.name = BODY_FONT
    run5.font.size = Pt(9)

    # NUMPAGES field
    run6 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    run6._r.append(fldChar3)

    run7 = paragraph.add_run()
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    run7._r.append(instrText2)

    run8 = paragraph.add_run()
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run8._r.append(fldChar4)

    # " 页"
    run9 = paragraph.add_run(" 页")
    run9.font.name = BODY_FONT
    run9.font.size = Pt(9)


def _set_paragraph_format(paragraph, space_before=0, space_after=Pt(12), line_spacing=1.0):
    """Set 致同 standard paragraph format."""
    pf = paragraph.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after  # 段后 0.9 行 ≈ 12pt
    pf.line_spacing = line_spacing


def _set_run_font(run, font_name=BODY_FONT, size=BODY_SIZE, bold=False):
    """Set run font properties."""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_bookmark(paragraph, bookmark_name: str):
    """Add a bookmark to a paragraph for cross-reference."""
    import random
    bookmark_id = str(random.randint(1000, 99999))
    bookmarkStart = OxmlElement("w:bookmarkStart")
    bookmarkStart.set(qn("w:id"), bookmark_id)
    bookmarkStart.set(qn("w:name"), bookmark_name)
    paragraph._p.append(bookmarkStart)

    bookmarkEnd = OxmlElement("w:bookmarkEnd")
    bookmarkEnd.set(qn("w:id"), bookmark_id)
    paragraph._p.append(bookmarkEnd)


# ---------------------------------------------------------------------------
# GTNote* 模板 docx（Sprint 2 Task 2.1 产出）
# ---------------------------------------------------------------------------

TEMPLATE_PATH = Path(__file__).resolve().parent.parent.parent / "data" / "note_export_template.docx"

# 附注 section_code 索引（variant → sections[]，含 legacy_aliases / content_type）
SECTION_CODE_INDEX_PATH = (
    Path(__file__).resolve().parent.parent.parent
    / "data"
    / "audit_report_templates"
    / "section_code_index.json"
)


def _load_section_code_index(variant_key: str) -> list[dict[str, Any]]:
    """读取 section_code_index.json 中某 variant 的 sections 列表（缺失返回 []）."""
    if not SECTION_CODE_INDEX_PATH.exists():
        return []
    try:
        import json

        data = json.loads(SECTION_CODE_INDEX_PATH.read_text(encoding="utf-8"))
    except (ValueError, OSError):
        logger.warning("section_code_index.json 解析失败: %s", SECTION_CODE_INDEX_PATH)
        return []
    variants = data.get("variants", {})
    entry = variants.get(variant_key, {})
    sections = entry.get("sections", [])
    return [s for s in sections if isinstance(s, dict)]


def _new_document() -> Document:
    """优先加载 GTNote 模板 docx；缺失时降级 Document() 兼容.

    Sprint 2 Task 2.2: 模板提供 6 段落 + 1 字符 + 1 表格 + 2 sidecar 共 10 个 GTNote* 样式
    + 默认 trHeight 397 twip(0.7cm) + 致同 pgMar (3.2/2.54/3/3.18 cm)。
    """
    if TEMPLATE_PATH.exists():
        return Document(str(TEMPLATE_PATH))
    logger.warning(
        "note_export_template.docx 不存在，降级用 Document()。"
        "请运行: python scripts/build_note_export_template.py --apply"
    )
    return Document()


# ---------------------------------------------------------------------------
# D7 GTNote helper API（Sprint 2 Task 2.2 — 6 helper）
# ---------------------------------------------------------------------------


def apply_gt_dual_font(
    run,
    ascii_font: str = "Arial Narrow",
    east_asia_font: str = "仿宋_GB2312",
    size: Pt = Pt(12),
):
    """对单个 run 注入致同双字体 rPr.

    中文走 ``w:eastAsia`` (仿宋_GB2312)，数字 / 拉丁字符走 ``w:ascii / w:hAnsi`` (Arial Narrow)。
    Word 渲染时按 Unicode 段自动分流。

    R5.2 验收 1：字体名 / 验收 2：字号小四 12pt 统一。
    """
    run.font.size = size
    run.font.name = ascii_font  # 默认设为 ascii；下面手动注入 eastAsia/cs
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_asia_font)
    rFonts.set(qn("w:cs"), ascii_font)


def apply_gt_three_line(table) -> None:
    """对表格应用致同三线表样式：顶/底 1pt + 表头 cell tcBorders.bottom 1/2pt + 其他 nil.

    复用现有 ``_set_table_borders`` (top/bottom sz=8) + ``_set_row_bottom_border`` (sz=4 在表头行)。
    R5.2 验收 6：三线表磅数。
    """
    _set_table_borders(table, top_sz=8, bottom_sz=8, header_bottom_sz=4)
    if len(table.rows) > 0:
        _set_row_bottom_border(table.rows[0], sz=4)


def fill_multi_header(
    table,
    header_rows: list[list[dict]],
    total_cols: int,
) -> None:
    """多层表头 grid 二阶段填充（rowspan/colspan）.

    Args:
        table: ``docx.table.Table`` 实例（必须已经预分配 ``len(header_rows)`` 行 ×
               ``total_cols`` 列）
        header_rows: 二维 dict 列表，每行一个 list，每 cell 形如 ``{"text": str,
                     "rowspan": int=1, "colspan": int=1}``
        total_cols: 表格总列数（用于占位检查）

    阶段一：构造 ``grid[r][c] = (text, master_r, master_c)`` 网格，标记主从单元格
    阶段二：对每个 master 单元格调用 ``cell.merge()`` 合并 rowspan × colspan 区域

    R5.2 验收 11：多层表头正确合并（如固定资产变动表"本期增加→购置/在建转入"二级表头）。
    """
    n_rows = len(header_rows)
    if n_rows == 0:
        return
    if len(table.rows) < n_rows or table.rows[0]._tr.tc_lst.__len__() < total_cols:
        # 调用方未正确预分配，跳过（防御）
        logger.warning("fill_multi_header: 表格行/列数不足 %dx%d", n_rows, total_cols)
        return

    # ---- 阶段一：grid 占位 ----
    # grid[r][c] = master cell ref (主格自指；合并占位指向主格)
    grid: list[list[tuple[int, int] | None]] = [
        [None] * total_cols for _ in range(n_rows)
    ]
    masters: list[tuple[int, int, int, int, str]] = []  # (r, c, rs, cs, text)

    for r, row_def in enumerate(header_rows):
        col_cursor = 0
        for cell_def in row_def:
            # 跳过已被上方 rowspan 占用的列
            while col_cursor < total_cols and grid[r][col_cursor] is not None:
                col_cursor += 1
            if col_cursor >= total_cols:
                break
            rs = max(1, int(cell_def.get("rowspan", 1)))
            cs = max(1, int(cell_def.get("colspan", 1)))
            text = str(cell_def.get("text", ""))
            # 边界裁剪
            rs = min(rs, n_rows - r)
            cs = min(cs, total_cols - col_cursor)
            for dr in range(rs):
                for dc in range(cs):
                    grid[r + dr][col_cursor + dc] = (r, col_cursor)
            masters.append((r, col_cursor, rs, cs, text))
            col_cursor += cs

    # ---- 阶段二：写文字 + merge ----
    for r, c, rs, cs, text in masters:
        master_cell = table.rows[r].cells[c]
        master_cell.text = ""
        p = master_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, bold=True)
        # 合并：先横（colspan），再纵（rowspan），避免重复合并冲突
        if cs > 1:
            target = table.rows[r].cells[c + cs - 1]
            master_cell.merge(target)
        if rs > 1:
            # 合并后 master_cell 引用仍可用
            bottom = table.rows[r + rs - 1].cells[c]
            master_cell.merge(bottom)


def apply_gt_row_height(row, cm: float = 0.7) -> None:
    """固定行高（exact）+ cantSplit + 关闭标题行重复（不设 w:tblHeader）.

    R5.2 验收 7：trHeight exact 0.7cm = 397 twip；验收 9：标题行不重复（无 w:tblHeader）.
    """
    twip = round(cm * 567)  # 1cm ≈ 567 twip；0.7cm → round(396.9) = 397
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    # 移除已有 trHeight / cantSplit / tblHeader（幂等 + 强制无 tblHeader）
    for tag in ("w:trHeight", "w:cantSplit", "w:tblHeader"):
        for el in trPr.findall(qn(tag)):
            trPr.remove(el)

    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twip))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)
    # 不添加 w:tblHeader → 标题行不重复（验收 9）


def fmt_amount_gt(val: Any) -> str:
    """致同金额格式：空值/零值留白返回 ``""``（不返 ``"-"``）.

    与现有 ``_format_amount`` 行为差异：
        ``_format_amount(0)``    → ``"-"``
        ``fmt_amount_gt(0)``     → ``""``
        ``fmt_amount_gt(None)``  → ``""``
        ``fmt_amount_gt("")``    → ``""``
        ``fmt_amount_gt(1234.5)``→ ``"1,234.50"``

    R5.2 验收 8：空值 / 零值留白。
    """
    if val is None or val == "":
        return ""
    try:
        num = float(val)
    except (ValueError, TypeError):
        return str(val)
    if num == 0:
        return ""
    return f"{num:,.2f}"


def add_landscape_section(doc) -> Any:
    """章节级横向：next page section break + WD_ORIENT.LANDSCAPE.

    在文档末尾追加一个新 section（next page），将其 orientation 翻转为横向，
    并交换 page_width / page_height（python-docx 不会自动交换）。
    返回新 section 对象供后续自定义。
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    # python-docx 翻转 orientation 不自动 swap width/height
    old_w, old_h = new_section.page_width, new_section.page_height
    new_section.page_width = old_h
    new_section.page_height = old_w
    return new_section


# ---------------------------------------------------------------------------
# D1 sidecar 公式/手工标记 cell tcPr 注入（R5.2 验收 46，可选渲染）
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, fill_hex: str = "E6FFE6") -> None:
    """对单个 cell 注入背景色 (w:shd)，浅绿 #E6FFE6 表示公式来源."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除已有 shd
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_gray_borders(cell, color: str = "808080", sz: int = 4) -> None:
    """对单个 cell 注入四边灰色边框（手工标记 cell 视觉提示）."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# NoteWordExporter class
# ---------------------------------------------------------------------------


class NoteWordExporter:
    """附注 Word 导出引擎 — 致同标准格式

    Requirements: 4.2-4.10, 27.1-27.10
    """

    def __init__(self, db: AsyncSession):
        self.db = db
        # Sprint 2 Task 2.2: D1 sidecar 渲染选项（export() 时被覆盖）
        self._annotate_formulas: bool = False
        self._annotate_manual: bool = False

    def _new_document(self) -> Document:
        """优先加载 GTNote 模板 docx；缺失时降级 Document() 兼容（Sprint 2 Task 2.2）."""
        return _new_document()

    async def export(
        self,
        project_id: UUID,
        year: int,
        template_type: str = "soe",
        report_scope: str | None = None,
        sections: list[str] | None = None,
        skip_empty: bool = False,
        annotate_formulas: bool = False,
        annotate_manual: bool = False,
        mode: Literal["template", "programmatic"] = "programmatic",
    ) -> BytesIO:
        """导出附注为 Word 文档（致同标准格式）

        Args:
            project_id: 项目 ID
            year: 年度
            template_type: 模板类型 (soe/listed)
            report_scope: 报表口径 standalone/consolidated；None 时从项目读取
            sections: 指定导出章节列表（None=全部）
            skip_empty: 是否跳过空章节
            annotate_formulas: D1 sidecar 渲染——公式 cell 标浅绿背景（默认关闭）
            annotate_manual:   D1 sidecar 渲染——手工 cell 标灰色边框（默认关闭）
            mode: 导出模式
                - "programmatic"（默认）：从零程序化拼装致同格式（现有行为，不变）
                - "template"：基于附注 docx 模板填充（待 Phase 0.6.2 附注模板全量
                  打标完成后启用，当前抛 NotImplementedError）

        Returns:
            BytesIO containing the docx file
        """
        if mode == "template":
            return await self._export_template_mode(
                project_id,
                year,
                template_type=template_type,
                report_scope=report_scope,
                sections=sections,
                annotate_formulas=annotate_formulas,
                annotate_manual=annotate_manual,
            )

        template_type = normalize_template_type(template_type)
        if report_scope is None:
            report_scope = await self._resolve_report_scope(project_id)
        report_scope = normalize_report_scope(report_scope)

        # Load notes data
        notes = await self._load_notes(project_id, year, sections)
        notes = [
            n for n in notes
            if note_applies_to_report_scope(n.note_section, template_type, report_scope)
        ]

        # Filter empty sections if requested
        if skip_empty:
            notes = [
                n for n in notes
                if not self._should_skip_section(n)
            ]

        # Build document（Sprint 2 Task 2.2: 优先加载 GTNote* 模板 docx）
        doc = self._new_document()
        # 渲染选项透传给 _render_table
        self._annotate_formulas = annotate_formulas
        self._annotate_manual = annotate_manual

        self._setup_page(doc)
        self._add_title(doc, year)
        _add_toc(doc)
        doc.add_page_break()

        # Render sections
        level1_idx = 0
        level2_idx = 0
        level3_idx = 0
        current_level1 = ""

        for note in notes:
            section_code = note.note_section or ""
            level = self._detect_level(section_code)

            if level == 1:
                level1_idx += 1
                level2_idx = 0
                level3_idx = 0
                current_level1 = section_code
                self._add_heading1(doc, level1_idx, note.section_title or section_code, section_code)
            elif level == 2:
                level2_idx += 1
                level3_idx = 0
                self._add_heading2(doc, level2_idx, note.section_title or section_code, section_code)
            else:
                level3_idx += 1
                self._add_heading3(doc, level3_idx, note.section_title or section_code, section_code)

            # Content
            if self._has_content(note):
                self._render_note_content(doc, note)
            else:
                # Empty section placeholder
                p = doc.add_paragraph()
                run = p.add_run("本期无此项业务。")
                _set_run_font(run)
                _set_paragraph_format(p)

        # Save
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    async def _export_template_mode(
        self,
        project_id: UUID,
        year: int,
        *,
        template_type: str,
        report_scope: str | None,
        sections: list[str] | None,
        annotate_formulas: bool = False,
        annotate_manual: bool = False,
    ) -> BytesIO:
        """基于附注 docx 模板填充导出（design §7 附注模板填充流程）.

        算法：
        1. manifest 解析 variant docx + 读 section_code_index.json
        2. 扫描 ``##SECTION:code##`` 块（body 级，含表格）
        3. 加载 DB notes + 裁剪状态，按 section_code（含 legacy_aliases）join
        4. §7.1 优先级判定跳过 → 删除整 SECTION 块
        5. 保留块：填 {{section:code}} / {{table:code:N}} / 渲染表格
        6. {{seq:prefix}} ← compute_section_numbers（裁剪后）
        7. 清理残留 ##SECTION:## / ##STYLE_REF:## 标记
        8. 返回 BytesIO
        """
        template_type = normalize_template_type(template_type)
        if report_scope is None:
            report_scope = await self._resolve_report_scope(project_id)
        report_scope = normalize_report_scope(report_scope)
        variant_key = build_variant_key(template_type, report_scope)

        # 1. 解析模板路径 + 载入 docx
        loader = get_template_manifest_loader()
        entry = loader.resolve_disclosure_notes(template_type, report_scope)
        if not entry.exists:
            raise FileNotFoundError(
                f"附注模板缺失: {entry.abs_path}（variant={variant_key}）"
            )
        doc = Document(str(entry.abs_path))

        # 渲染选项透传给 _render_table
        self._annotate_formulas = annotate_formulas
        self._annotate_manual = annotate_manual

        # 索引：section_code → index entry（含 legacy_aliases）
        index_sections = _load_section_code_index(variant_key)

        # 2. 扫描 SECTION 块
        blocks = scan_section_blocks(doc)

        # 3. 加载 notes + 裁剪状态；构建 join 映射
        notes = await self._load_notes(project_id, year, sections)
        notes = [
            n
            for n in notes
            if note_applies_to_report_scope(n.note_section, template_type, report_scope)
        ]
        note_by_section = {(n.note_section or "").strip(): n for n in notes}

        # legacy_aliases → canonical section_code 映射（来自索引）
        alias_to_code: dict[str, str] = {}
        for entry_s in index_sections:
            code = (entry_s.get("section_code") or "").strip()
            for alias in entry_s.get("legacy_aliases", []) or []:
                a = (alias or "").strip()
                if a:
                    alias_to_code[a] = code

        def _match_note(section_code: str) -> DisclosureNote | None:
            code = (section_code or "").strip()
            note = note_by_section.get(code)
            if note is not None:
                return note
            # DB note_section 可能是 legacy alias（如 五、1 → 模板 八、1）
            for db_section, n in note_by_section.items():
                if alias_to_code.get(db_section) == code:
                    return n
            return None

        # 4 + 5. 逐块裁剪 / 填充
        kept_codes: list[str] = []
        for block in blocks:
            # 口径排除：standalone 导出删除 consolidated_only 章节块（design §7 / 需求 12）
            if not note_applies_to_report_scope(
                block.section_code, template_type, report_scope
            ):
                delete_section_block(block)
                continue
            note = _match_note(block.section_code)
            skip = False
            if note is not None:
                skip = should_skip_empty_section(self._note_to_skip_dict(note))
            if skip:
                delete_section_block(block)
                continue
            kept_codes.append(block.section_code)
            if note is not None:
                self._fill_section_block(doc, block, note)

        # 6. {{seq:prefix}} 填充（仅保留章节，design §13）
        kept_tree = [{"note_section": code} for code in kept_codes]
        seq_numbers = compute_section_numbers(
            kept_tree,
            report_scope=report_scope,
            template_type=template_type,
        )
        self._fill_seq_placeholders(doc, kept_codes, seq_numbers)

        # 7. 清理残留标记
        remove_section_markers(doc)

        # 8. 输出
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def _fill_section_block(
        self, doc: Document, block, note: DisclosureNote
    ) -> None:
        """填充保留章节块内的 {{section:code}} / {{table:code:N}} 占位符 + 渲染表格.

        说明：当前 GT 模板（Phase 0.6.2 打标）章节块内多为参考标题 + 指引段，
        通常**无** {{section}}/{{table}} 占位符——此时不强行注入，保留模板原貌；
        仅当模板含占位符时才填充（兼容更完整的预备模板与测试夹具）。
        """
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn as _qn

        code = block.section_code
        text_content = (getattr(note, "text_content", None) or "").strip()

        # 收集块内段落（仅 <w:p>）
        p_tag = _qn("w:p")
        paragraphs = [
            Paragraph(el, doc) for el in block.elements if el.tag == p_tag
        ]

        section_token = f"{{{{section:{code}}}}}"
        for para in paragraphs:
            ptext = para.text or ""
            if section_token in ptext:
                new_text = ptext.replace(section_token, text_content)
                para.clear()
                if new_text:
                    run = para.add_run(new_text)
                    _set_run_font(run)

        # {{table:code:N}} → 渲染对应表（多表按 index 对齐）
        tables_to_render = self._note_tables(note)
        table_re = re.compile(
            r"\{\{table:" + re.escape(code) + r"(?::(\d+))?\}\}"
        )
        for para in paragraphs:
            m = table_re.search(para.text or "")
            if not m:
                continue
            idx = int(m.group(1)) if m.group(1) else 0
            para.clear()
            if 0 <= idx < len(tables_to_render):
                tbl = tables_to_render[idx]
                # 空表 → 无业务段落（§7.1 ⑤）
                if get_table_render_mode(tbl) == "no_business_paragraph":
                    run = para.add_run("本期无此项业务。")
                    _set_run_font(run)
                else:
                    self._render_table_at(doc, para, tbl)

    def _note_tables(self, note: DisclosureNote) -> list[dict]:
        """返回 note 的表列表（多表 _tables 数组优先，降级单表）."""
        td = getattr(note, "table_data", None)
        if not isinstance(td, dict):
            return []
        tables = td.get("_tables") or [td]
        return [t for t in tables if isinstance(t, dict)]

    def _render_table_at(self, doc: Document, anchor_para, table_data: dict) -> None:
        """在 anchor 段落处渲染表格（复用 _render_table，再把表移动到锚点位置）."""
        before = set(id(t._tbl) for t in doc.tables)
        self._render_table(doc, table_data)
        # 找到新追加的表，移动到 anchor 之后
        new_tbls = [t for t in doc.tables if id(t._tbl) not in before]
        anchor_el = anchor_para._p
        for tbl in new_tbls:
            tbl_el = tbl._tbl
            parent = tbl_el.getparent()
            if parent is not None:
                parent.remove(tbl_el)
            anchor_el.addnext(tbl_el)
            anchor_el = tbl_el

    def _fill_seq_placeholders(
        self, doc: Document, kept_codes: list[str], seq_numbers: dict[str, str]
    ) -> None:
        """填充 {{seq:prefix}} 占位符为重算后的运行编号（design §13）.

        ``seq_numbers`` 键为完整 note_section（如 ``八、1`` → ``"1"``）。
        模板中的 ``{{seq:八}}`` 出现在某个章节块内 → 用该块 section_code 的编号；
        组内仅 1 条不编号（compute_section_numbers 已处理），此时替换为空串。
        """
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn as _qn

        # 重新扫描块（删除后元素已变）
        blocks = scan_section_blocks(doc)
        p_tag = _qn("w:p")
        seq_re = re.compile(r"\{\{seq:([^}]+)\}\}")
        for block in blocks:
            number = seq_numbers.get(block.section_code, "")
            for el in block.elements:
                if el.tag != p_tag:
                    continue
                para = Paragraph(el, doc)
                ptext = para.text or ""
                if "{{seq:" not in ptext:
                    continue
                new_text = seq_re.sub(number, ptext)
                if new_text != ptext:
                    para.clear()
                    if new_text:
                        run = para.add_run(new_text)
                        _set_run_font(run)

    async def preview_html(self, project_id: UUID, year: int) -> str:
        """Generate HTML preview of notes.

        Requirements: 4.10, 27.10
        """
        notes = await self._load_notes(project_id, year)

        html_parts = ['<div class="note-preview" style="font-family: 仿宋_GB2312, FangSong; font-size: 12pt;">']
        html_parts.append(f'<h1 style="text-align:center;">财务报表附注</h1>')
        html_parts.append(f'<p style="text-align:center;">（{year}年度）</p>')

        level1_idx = 0
        level2_idx = 0
        level3_idx = 0

        for note in notes:
            level = self._detect_level(note.note_section or "")
            title = note.section_title or note.note_section or ""

            if level == 1:
                level1_idx += 1
                level2_idx = 0
                level3_idx = 0
                prefix = f"{_CN_NUMBERS[level1_idx - 1]}、" if level1_idx <= len(_CN_NUMBERS) else f"{level1_idx}、"
                html_parts.append(f'<h2 style="font-weight:bold;">{prefix}{title}</h2>')
            elif level == 2:
                level2_idx += 1
                level3_idx = 0
                prefix = _CN_NUMBERS_PAREN[level2_idx - 1] if level2_idx <= len(_CN_NUMBERS_PAREN) else f"（{level2_idx}）"
                html_parts.append(f'<h3>{prefix}{title}</h3>')
            else:
                level3_idx += 1
                html_parts.append(f'<h4>{level3_idx}. {title}</h4>')

            if self._has_content(note):
                if note.text_content:
                    html_parts.append(f'<p>{note.text_content}</p>')
                if note.table_data:
                    # Sprint 0 / Task 0.3 P0 修复：HTML 预览也支持多表
                    tables_to_render = note.table_data.get("_tables") or [note.table_data]
                    for tbl in tables_to_render:
                        if not isinstance(tbl, dict):
                            continue
                        if len(tables_to_render) > 1 and tbl.get("name"):
                            html_parts.append(f'<h5 style="margin:8px 0 4px;font-weight:bold;">{tbl["name"]}</h5>')
                        html_parts.append(self._table_to_html(tbl))
            else:
                html_parts.append('<p style="color:#999;">本期无此项业务。</p>')

        html_parts.append('</div>')
        return "\n".join(html_parts)

    # -----------------------------------------------------------------------
    # Private methods
    # -----------------------------------------------------------------------

    async def _resolve_report_scope(self, project_id: UUID) -> str:
        from app.models.core import Project

        result = await self.db.execute(
            sa.select(Project.report_scope).where(
                Project.id == project_id,
                Project.is_deleted == sa.false(),
            )
        )
        row = result.scalar_one_or_none()
        return normalize_report_scope(row if isinstance(row, str) else None)

    async def _load_notes(
        self, project_id: UUID, year: int, sections: list[str] | None = None
    ) -> list[DisclosureNote]:
        """Load notes from database."""
        q = sa.select(DisclosureNote).where(
            DisclosureNote.project_id == project_id,
            DisclosureNote.year == year,
        ).order_by(DisclosureNote.note_section)

        result = await self.db.execute(q)
        notes = list(result.scalars().all())

        if sections:
            notes = [n for n in notes if n.note_section in sections]

        return notes

    def _has_content(self, note: DisclosureNote) -> bool:
        """Check if a note section has any content.

        Sprint 0 / Task 0.3 P0 修复：也识别 table_data._tables 数组中的多表内容
        """
        if note.text_content and note.text_content.strip():
            return True
        if not note.table_data or not isinstance(note.table_data, dict):
            return False

        # 收集所有要检查的表（多表 _tables 数组 + 单表降级）
        tables_to_check = note.table_data.get("_tables") or [note.table_data]

        for tbl in tables_to_check:
            if not isinstance(tbl, dict):
                continue
            rows = tbl.get("rows", [])
            for row in rows:
                values = row.get("values", [])
                cells = row.get("cells", values)
                for cell in cells:
                    if isinstance(cell, dict):
                        val = cell.get("value", cell.get("manual_value", 0))
                    else:
                        val = cell
                    if val and val != 0 and val != "0" and val != "-":
                        return True
        return False

    def _note_to_skip_dict(self, note: DisclosureNote) -> dict:
        """将 DisclosureNote ORM 转为 should_skip_empty_section 所需 dict 形状.

        字段：is_deleted / status / is_empty / text_content / table_data。
        status 为枚举时取其 value（与 design §7.1 'not_applicable' 字符串比较一致）。
        缺失属性用 getattr 默认值兜底（兼容测试用 Fake 对象）。
        """
        status = getattr(note, "status", None)
        status_value = getattr(status, "value", status)  # 枚举 → str
        return {
            "is_deleted": getattr(note, "is_deleted", False),
            "status": status_value,
            "is_empty": getattr(note, "is_empty", False),
            "text_content": getattr(note, "text_content", None),
            "table_data": getattr(note, "table_data", None),
        }

    def _should_skip_section(self, note: DisclosureNote) -> bool:
        """判断章节在 skip_empty 模式下是否应跳过（design §7.1 ①~④）.

        复用 `should_skip_empty_section`（不重复实现裁剪逻辑）。
        """
        return should_skip_empty_section(self._note_to_skip_dict(note))

    def _detect_level(self, section_code: str) -> int:
        """Detect heading level — delegated to note_section_catalog (唯一规则)."""
        return detect_heading_level(section_code)

    def _setup_page(self, doc: Document):
        """Set up page margins and orientation per 致同 standard."""
        section = doc.sections[0]
        section.page_width = Cm(21)  # A4
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = PAGE_MARGINS["top"]
        section.bottom_margin = PAGE_MARGINS["bottom"]
        section.left_margin = PAGE_MARGINS["left"]
        section.right_margin = PAGE_MARGINS["right"]
        section.header_distance = HEADER_MARGIN
        section.footer_distance = FOOTER_MARGIN

        # Add page number footer
        _add_page_number_footer(section)

    def _add_title(self, doc: Document, year: int):
        """Add document title."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("财务报表附注")
        run.font.name = BODY_FONT
        run.font.size = Pt(16)
        run.bold = True
        # East Asian font
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), BODY_FONT)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"（{year}年度）")
        _set_run_font(run2, size=Pt(14))

    def _add_heading1(self, doc: Document, idx: int, title: str, section_code: str):
        """Add level 1 heading: 一、二、三... (bold)"""
        prefix = f"{_CN_NUMBERS[idx - 1]}、" if idx <= len(_CN_NUMBERS) else f"{idx}、"
        p = doc.add_paragraph()
        run = p.add_run(f"{prefix}{title}")
        _set_run_font(run, bold=True, size=HEADING1_FONT_SIZE)
        _set_paragraph_format(p, space_before=Pt(6))
        _add_bookmark(p, f"note_section_{section_code}")

    def _add_heading2(self, doc: Document, idx: int, title: str, section_code: str):
        """Add level 2 heading: （一）（二）..."""
        prefix = _CN_NUMBERS_PAREN[idx - 1] if idx <= len(_CN_NUMBERS_PAREN) else f"（{idx}）"
        p = doc.add_paragraph()
        run = p.add_run(f"{prefix}{title}")
        _set_run_font(run, size=HEADING2_FONT_SIZE)
        _set_paragraph_format(p)
        _add_bookmark(p, f"note_section_{section_code}")

    def _add_heading3(self, doc: Document, idx: int, title: str, section_code: str):
        """Add level 3 heading: 1. 2. 3."""
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {title}")
        _set_run_font(run, size=HEADING3_FONT_SIZE)
        _set_paragraph_format(p)
        _add_bookmark(p, f"note_section_{section_code}")

    def _render_note_content(self, doc: Document, note: DisclosureNote):
        """Render note content (text + tables).

        Sprint 0 / Task 0.3 P0 修复：
        - 优先取 table_data._tables 数组逐张渲染（多表章节如固定资产 5 表 / 应收票据 12 表）
        - 多表章节加 H4 表名标题（如有 name 字段）
        - 空 header 列裁剪（headers 中的空字符串占位过滤掉）
        - 兼容老结构：无 _tables 时降级到单表渲染
        """
        # Text content
        if note.text_content and note.text_content.strip():
            p = doc.add_paragraph()
            run = p.add_run(note.text_content.strip())
            _set_run_font(run)
            _set_paragraph_format(p)

        # Table data
        if not note.table_data or not isinstance(note.table_data, dict):
            return

        # 优先取 _tables 数组（多表章节）；老结构降级到单表
        tables_to_render = note.table_data.get("_tables") or [note.table_data]

        for tbl in tables_to_render:
            if not isinstance(tbl, dict):
                continue
            # 多表章节加表名 H4 标题
            if len(tables_to_render) > 1 and tbl.get("name"):
                p = doc.add_paragraph()
                run = p.add_run(str(tbl["name"]))
                _set_run_font(run, bold=True, size=HEADING3_FONT_SIZE)
                _set_paragraph_format(p, space_before=Pt(6))
            self._render_table(doc, tbl)

    def _render_table(self, doc: Document, table_data: dict):
        """Render a table with 致同 standard formatting.

        Sprint 0 / Task 0.3 P0 修复：
        - 空 header 列裁剪：headers 中的空字符串占位过滤掉（v2 模板治理前的兼容措施）
        - 列裁剪后同步裁 cells 数据，避免索引越界

        Sprint 2 Task 2.2 增强：
        - 应用 ``apply_gt_three_line`` 三线表（替代旧 _set_table_borders/_set_row_bottom_border 调用）
        - 每行调 ``apply_gt_row_height(cm=0.7)`` 固定行高 + cantSplit + 关闭标题行重复
        - D1 sidecar 渲染：``self._annotate_formulas`` / ``self._annotate_manual`` 控制公式 / 手工 cell 视觉
        """
        headers_raw = table_data.get("headers", [])
        rows = table_data.get("rows", [])

        if not headers_raw or not rows:
            return

        # 空 header 列裁剪 + 记录有效列索引
        valid_indices = [i for i, h in enumerate(headers_raw) if h and str(h).strip()]
        if not valid_indices:
            return
        headers = [headers_raw[i] for i in valid_indices]

        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header row

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Sprint 2 Task 2.2: 致同三线表（顶/底 1pt + 表头 cell tcBorders.bottom 1/2pt + 其他 nil）
        apply_gt_three_line(table)

        # Header row
        header_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            _set_run_font(run, bold=True)

        # Sprint 2 Task 2.2: 表头行 + 数据行均固定 0.7cm + 关闭标题行重复
        apply_gt_row_height(header_row, cm=0.7)

        # 表级 sidecar 数据（D1）：_formulas / _cell_modes 优先取 row 级
        formulas_index: dict[tuple[int, int], dict[str, Any]] = {}
        for f in (table_data.get("_formulas") or []):
            if isinstance(f, dict) and "row" in f and "col" in f:
                formulas_index[(int(f["row"]), int(f["col"]))] = f

        # Data rows
        for r_idx, row in enumerate(rows):
            label = row.get("label", "")
            values = row.get("values", [])
            cells_data = row.get("cells", values)
            cell_modes = row.get("_cell_modes") or []

            data_row = table.rows[r_idx + 1]
            apply_gt_row_height(data_row, cm=0.7)

            # First column: label
            cell0 = data_row.cells[0]
            cell0.text = ""
            p0 = cell0.paragraphs[0]
            run0 = p0.add_run(str(label))
            _set_run_font(run0)

            # Data columns: amounts right-aligned with Arial Narrow
            # 按 valid_indices[1:] 取数（跳过 label 列），避免 cells_data 索引与显示列错位
            for display_col_idx, original_col_idx in enumerate(valid_indices[1:], start=1):
                # cells_data 通常对齐 headers_raw[1:]（去掉首列 label），按原始索引-1 取
                data_idx = original_col_idx - 1
                if data_idx >= len(cells_data):
                    break
                val = cells_data[data_idx]
                cell = data_row.cells[display_col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Extract value
                if isinstance(val, dict):
                    cell_val = val.get("value", val.get("manual_value", 0))
                else:
                    cell_val = val

                # Sprint 2 Task 2.2: 使用 fmt_amount_gt（空/零留白）替代 _format_amount（"-"）
                formatted = fmt_amount_gt(cell_val)
                run = p.add_run(formatted)
                # 数字走 Arial Narrow + 中文 eastAsia=仿宋
                if _is_amount(cell_val):
                    apply_gt_dual_font(run)
                else:
                    _set_run_font(run)

                # D1 sidecar 渲染：公式 / 手工标记（默认两者都关）
                if self._annotate_formulas and (r_idx, data_idx) in formulas_index:
                    _set_cell_shading(cell, fill_hex="E6FFE6")
                if self._annotate_manual and data_idx < len(cell_modes) and cell_modes[data_idx] == "manual":
                    _set_cell_gray_borders(cell)

        # Add spacing after table
        doc.add_paragraph()

    def _table_to_html(self, table_data: dict) -> str:
        """Convert table data to HTML for preview."""
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])

        if not headers:
            return ""

        html = ['<table style="border-collapse:collapse; width:100%; margin:8px 0; border-top:1px solid #000; border-bottom:1px solid #000;">']
        # Header
        html.append('<tr style="border-bottom:0.5px solid #000;">')
        for h in headers:
            html.append(f'<th style="padding:4px 8px; text-align:center; font-weight:bold;">{h}</th>')
        html.append('</tr>')

        # Rows
        for row in rows:
            html.append('<tr>')
            label = row.get("label", "")
            html.append(f'<td style="padding:4px 8px;">{label}</td>')
            values = row.get("values", [])
            cells_data = row.get("cells", values)
            for val in cells_data:
                if isinstance(val, dict):
                    cell_val = val.get("value", val.get("manual_value", 0))
                else:
                    cell_val = val
                formatted = _format_amount(cell_val)
                html.append(f'<td style="padding:4px 8px; text-align:right; font-family:Arial Narrow;">{formatted}</td>')
            html.append('</tr>')

        html.append('</table>')
        return "\n".join(html)

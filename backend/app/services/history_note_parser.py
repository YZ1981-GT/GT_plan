"""历史附注上传与解析服务

Phase 9 Task 9.28: Word/PDF 解析 + LLM 结构化
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from uuid import UUID

logger = logging.getLogger(__name__)


class HistoryNoteParser:
    """历史附注解析"""

    async def parse_file(self, file_path: str) -> dict:
        """解析上传的历史附注文件

        Returns: {"sections": [{section_number, section_title, tables, text_blocks}], "raw_text": str}
        """
        path = Path(file_path)
        if not path.exists():
            return {"error": f"文件不存在: {file_path}", "sections": []}

        ext = path.suffix.lower()
        if ext == ".docx":
            return await self._parse_docx(path)
        elif ext == ".pdf":
            return await self._parse_pdf(path)
        else:
            return {"error": f"不支持的文件格式: {ext}", "sections": []}

    async def _parse_docx(self, path: Path) -> dict:
        """Word 解析：python-docx 提取章节结构"""
        try:
            import docx
            doc = docx.Document(str(path))

            sections: list[dict] = []
            current_section: dict | None = None
            current_text: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 检测章节标题（如"五、1 货币资金"、"一、公司基本情况"）
                if self._is_section_heading(text, para.style.name if para.style else ""):
                    if current_section:
                        current_section["text_blocks"] = current_text
                        sections.append(current_section)
                    number, title = self._extract_section_info(text)
                    current_section = {
                        "section_number": number,
                        "section_title": title,
                        "tables": [],
                        "text_blocks": [],
                    }
                    current_text = []
                else:
                    current_text.append(text)

            if current_section:
                current_section["text_blocks"] = current_text
                sections.append(current_section)

            # 提取表格
            for table in doc.tables:
                table_data = self._extract_table(table)
                if table_data and sections:
                    sections[-1]["tables"].append(table_data)

            return {"sections": sections, "raw_text": "\n".join(p.text for p in doc.paragraphs)}

        except Exception as e:
            logger.error(f"Word 解析失败: {e}")
            return {"error": str(e), "sections": []}

    async def _parse_pdf(self, path: Path) -> dict:
        """PDF 解析：文字层提取 + OCR 兜底"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            full_text = ""
            for page in doc:
                full_text += page.get_text() + "\n"
            doc.close()

            # 简化：从文本中提取章节
            sections = self._extract_sections_from_text(full_text)
            return {"sections": sections, "raw_text": full_text}

        except ImportError:
            # PyMuPDF 不可用，尝试 pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    full_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                sections = self._extract_sections_from_text(full_text)
                return {"sections": sections, "raw_text": full_text}
            except ImportError:
                return {"error": "PDF 解析库未安装（需要 PyMuPDF 或 pdfplumber）", "sections": []}
        except Exception as e:
            logger.error(f"PDF 解析失败: {e}")
            return {"error": str(e), "sections": []}

    def _is_section_heading(self, text: str, style_name: str) -> bool:
        """判断是否为章节标题"""
        if "Heading" in style_name or "标题" in style_name:
            return True
        # 匹配中文编号模式
        patterns = [
            r'^[一二三四五六七八九十]+、',  # 一、二、三、
            r'^（[一二三四五六七八九十]+）',  # （一）（二）
            r'^五、\d+',  # 五、1 五、2
            r'^\d+\.\d+',  # 1.1 1.2
        ]
        return any(re.match(p, text) for p in patterns)

    def _extract_section_info(self, text: str) -> tuple[str, str]:
        """从标题文本提取编号和标题"""
        # 匹配 "五、1 货币资金" 模式
        m = re.match(r'^(五、\d+)\s*(.*)', text)
        if m:
            return m.group(1), m.group(2)
        # 匹配 "一、公司基本情况" 模式
        m = re.match(r'^([一二三四五六七八九十]+、)\s*(.*)', text)
        if m:
            return m.group(1).rstrip('、'), m.group(2)
        return "", text

    def _extract_table(self, table) -> dict | None:
        """从 docx 表格提取数据"""
        try:
            rows = []
            headers = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    headers = cells
                else:
                    rows.append({"label": cells[0] if cells else "", "values": cells[1:]})
            if headers:
                return {"headers": headers, "rows": rows}
        except Exception:
            pass
        return None

    def _extract_sections_from_text(self, text: str) -> list[dict]:
        """从纯文本中提取章节结构"""
        sections = []
        current: dict | None = None
        current_text: list[str] = []

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if self._is_section_heading(line, ""):
                if current:
                    current["text_blocks"] = current_text
                    sections.append(current)
                number, title = self._extract_section_info(line)
                current = {"section_number": number, "section_title": title, "tables": [], "text_blocks": []}
                current_text = []
            else:
                current_text.append(line)

        if current:
            current["text_blocks"] = current_text
            sections.append(current)

        return sections

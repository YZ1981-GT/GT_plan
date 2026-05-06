"""QC 年度质量报告服务 — Round 3 需求 9

POST /api/qc/annual-reports?year= — 触发异步报告生成
GET  /api/qc/annual-reports       — 列出历史年报
GET  /api/qc/annual-reports/{id}/download — 下载报告

使用 ExportJobService 异步生成，每年至多一个任务（幂等锁）。
报告结构：封面 → 项目规模与分布 → 评级分布（ABCD 饼图）→ Top10 问题
         → 复核人表现 → 改进建议（LLM）→ 附录（规则变更、抽查统计）
"""

from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID

import sqlalchemy as sa
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.core import Project
from app.models.phase13_models import ExportJob, ExportJobStatus
from app.models.phase15_models import IssueTicket
from app.models.qc_rating_models import ProjectQualityRating, ReviewerMetricsSnapshot
from app.services.export_job_service import ExportJobService
from app.services.ai_contribution_watermark import generate_short_statement as build_ai_contribution_statement

logger = logging.getLogger(__name__)

# 模板路径（占位文件）
TEMPLATE_PATH = Path(__file__).parent.parent.parent / "data" / "archive_templates" / "qc_annual_report.docx"

# 年报导出文件存储目录
ANNUAL_REPORT_DIR = Path(__file__).parent.parent.parent / "storage" / "qc_annual_reports"


class QcAnnualReportService:
    """QC 年度质量报告服务。"""

    JOB_TYPE = "qc_annual_report"

    async def generate_annual_report(
        self,
        db: AsyncSession,
        year: int,
        user_id: UUID,
    ) -> dict:
        """触发年度报告异步生成。

        幂等锁：同一年只允许一个正在运行的任务。
        如果已有 queued/running 状态的同年任务，直接返回该任务信息。

        Args:
            db: 数据库会话
            year: 报告年份
            user_id: 发起人 ID

        Returns:
            {job_id, status, year, message}
        """
        # 幂等检查：查找同年同类型的 queued/running 任务
        existing_stmt = (
            select(ExportJob)
            .where(
                ExportJob.job_type == self.JOB_TYPE,
                ExportJob.status.in_([
                    ExportJobStatus.queued.value,
                    ExportJobStatus.running.value,
                ]),
                ExportJob.payload["year"].as_integer() == year,
            )
        )
        existing_result = await db.execute(existing_stmt)
        existing_job = existing_result.scalar_one_or_none()

        if existing_job:
            logger.info(
                "年报任务已存在: job_id=%s, year=%d, status=%s",
                existing_job.id, year, existing_job.status,
            )
            return {
                "job_id": str(existing_job.id),
                "status": existing_job.status,
                "year": year,
                "message": f"{year} 年度报告任务已在运行中",
            }

        # 创建新的导出任务
        # 注意：ExportJob.project_id 是必填字段，年报不属于单个项目
        # 使用一个固定的"系统级"项目 ID 或第一个项目 ID 作为占位
        # 这里用 payload 存储 year 信息
        export_svc = ExportJobService(db)

        # 查找任意一个项目 ID 作为占位（ExportJob 要求 project_id 非空）
        any_project_stmt = select(Project.id).where(
            Project.is_deleted == False,  # noqa: E712
        ).limit(1)
        any_project_result = await db.execute(any_project_stmt)
        any_project_row = any_project_result.first()

        if not any_project_row:
            return {
                "job_id": None,
                "status": "failed",
                "year": year,
                "message": "系统中无项目，无法生成年报",
            }

        project_id = any_project_row[0]

        job = await export_svc.create_job(
            project_id=project_id,
            job_type=self.JOB_TYPE,
            payload={"year": year},
            user_id=user_id,
            total=6,  # 6 个章节
        )

        # 启动后台生成（简化实现：同步生成占位文件）
        await self._execute_report_generation(db, job, year)

        await db.flush()

        logger.info("年报生成任务已创建: job_id=%s, year=%d", job.id, year)
        return {
            "job_id": str(job.id),
            "status": job.status,
            "year": year,
            "message": f"{year} 年度报告生成任务已启动",
        }

    async def list_annual_reports(
        self,
        db: AsyncSession,
        page: int = 1,
        page_size: int = 20,
    ) -> dict:
        """列出历史年报任务。

        Returns:
            {items: [...], total: int, page: int, page_size: int}
        """
        base_stmt = select(ExportJob).where(
            ExportJob.job_type == self.JOB_TYPE,
        )

        # 总数
        count_stmt = select(func.count()).select_from(
            base_stmt.subquery()
        )
        total_result = await db.execute(count_stmt)
        total = total_result.scalar() or 0

        # 分页查询
        items_stmt = (
            base_stmt
            .order_by(ExportJob.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
        items_result = await db.execute(items_stmt)
        jobs = items_result.scalars().all()

        return {
            "items": [
                {
                    "id": str(job.id),
                    "year": job.payload.get("year") if job.payload else None,
                    "status": job.status,
                    "progress_done": job.progress_done,
                    "progress_total": job.progress_total,
                    "created_at": job.created_at.isoformat() if job.created_at else None,
                    "file_path": job.payload.get("file_path") if job.payload else None,
                }
                for job in jobs
            ],
            "total": total,
            "page": page,
            "page_size": page_size,
        }

    async def get_report_download_url(
        self,
        db: AsyncSession,
        report_id: UUID,
    ) -> dict | None:
        """获取年报下载信息。

        Returns:
            {id, year, file_path, status} or None
        """
        stmt = select(ExportJob).where(
            ExportJob.id == report_id,
            ExportJob.job_type == self.JOB_TYPE,
        )
        result = await db.execute(stmt)
        job = result.scalar_one_or_none()

        if not job:
            return None

        file_path = job.payload.get("file_path") if job.payload else None

        return {
            "id": str(job.id),
            "year": job.payload.get("year") if job.payload else None,
            "status": job.status,
            "file_path": file_path,
        }

    async def _execute_report_generation(
        self,
        db: AsyncSession,
        job: ExportJob,
        year: int,
    ) -> None:
        """执行报告生成 — 使用 python-docx 填充模板。

        从 qc_annual_report.docx 模板加载，填充各章节数据后输出 .docx 文件。
        若 python-docx 不可用则降级为文本文件。
        """
        try:
            # 更新状态为 running
            job.status = ExportJobStatus.running.value
            await db.flush()

            # 收集报告数据
            report_data = await self._build_report_data(db, year)

            # 确保输出目录存在
            ANNUAL_REPORT_DIR.mkdir(parents=True, exist_ok=True)

            # 尝试使用 python-docx 生成真实 Word 文件
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                # 加载模板或新建文档
                if TEMPLATE_PATH.exists():
                    doc = Document(str(TEMPLATE_PATH))
                else:
                    doc = Document()

                # 标题页
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title_para.add_run(f"{year} 年度质量控制报告")
                run.font.size = Pt(22)
                run.bold = True
                doc.add_paragraph()

                # 第一章：项目规模分布
                doc.add_heading("一、项目规模与分布", level=1)
                scale = report_data.get("project_scale", {})
                doc.add_paragraph(f"本年度共执行审计项目 {scale.get('total_projects', 0)} 个。")
                if scale.get("by_type"):
                    for ptype, count in scale["by_type"].items():
                        doc.add_paragraph(f"  • {ptype}：{count} 个", style="List Bullet")

                # 第二章：评级分布
                doc.add_heading("二、项目质量评级分布", level=1)
                rating_dist = report_data.get("rating_distribution", {})
                for grade in ["A", "B", "C", "D"]:
                    count = rating_dist.get(grade, 0)
                    doc.add_paragraph(f"  • {grade} 级：{count} 个项目", style="List Bullet")

                # 第三章：Top10 问题
                doc.add_heading("三、年度典型问题 Top10", level=1)
                top10 = report_data.get("top10_issues", [])
                if top10:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    hdr[0].text = "序号"
                    hdr[1].text = "问题描述"
                    hdr[2].text = "出现次数"
                    for idx, issue in enumerate(top10[:10], 1):
                        row = table.add_row().cells
                        row[0].text = str(idx)
                        row[1].text = issue.get("description", "")[:100]
                        row[2].text = str(issue.get("count", 0))
                else:
                    doc.add_paragraph("本年度无典型问题记录。")

                # 第四章：复核人表现
                doc.add_heading("四、复核人表现", level=1)
                reviewers = report_data.get("reviewer_performance", [])
                if reviewers:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    hdr[0].text = "复核人"
                    hdr[1].text = "复核底稿数"
                    hdr[2].text = "平均时长(min)"
                    hdr[3].text = "退回率"
                    for r in reviewers[:20]:
                        row = table.add_row().cells
                        row[0].text = r.get("reviewer_name", "")
                        row[1].text = str(r.get("wp_count", 0))
                        row[2].text = str(r.get("avg_time_min", 0))
                        row[3].text = f"{r.get('rejection_rate', 0) * 100:.1f}%"
                else:
                    doc.add_paragraph("暂无复核人数据。")

                # 第五章：附录
                doc.add_heading("五、附录", level=1)
                appendix = report_data.get("appendix", {})
                doc.add_paragraph(f"抽查批次数：{appendix.get('inspection_count', 0)}")
                doc.add_paragraph(f"规则变更次数：{appendix.get('rule_change_count', 0)}")

                # AI 贡献声明
                doc.add_paragraph()
                from app.services.ai_contribution_watermark import generate_short_statement
                doc.add_paragraph(generate_short_statement())

                # 保存 .docx
                output_filename = f"qc_annual_report_{year}.docx"
                output_path = ANNUAL_REPORT_DIR / output_filename
                doc.save(str(output_path))

            except ImportError:
                # python-docx 不可用，降级为文本文件
                logger.warning("python-docx 不可用，降级为文本年报")
                output_filename = f"qc_annual_report_{year}.txt"
                output_path = ANNUAL_REPORT_DIR / output_filename
                content = self._render_report_text(year, report_data)
                output_path.write_text(content, encoding="utf-8")

            # 更新 job payload 记录文件路径
            if job.payload is None:
                job.payload = {}
            job.payload = {**job.payload, "file_path": str(output_path)}

            # 标记完成
            job.progress_done = job.progress_total
            job.status = ExportJobStatus.succeeded.value
            await db.flush()

            logger.info("年报生成完成: job_id=%s, path=%s", job.id, output_path)

        except Exception as e:
            logger.error("年报生成失败: job_id=%s, error=%s", job.id, e)
            job.status = ExportJobStatus.failed.value
            job.failed_count = 1
            await db.flush()

    async def _build_report_data(self, db: AsyncSession, year: int) -> dict:
        """收集年报所需的各章节数据。

        章节：
        1. 项目规模与分布
        2. 评级分布（ABCD 饼图数据）
        3. 典型问题 Top10
        4. 复核人表现
        5. LLM 综合建议（占位）
        6. 附录（规则变更、抽查统计）
        """
        data: dict = {}

        # 1. 项目规模与分布
        data["project_scale"] = await self._get_project_scale(db, year)

        # 2. 评级分布
        data["rating_distribution"] = await self._get_rating_distribution(db, year)

        # 3. Top10 问题
        data["top10_issues"] = await self._get_top10_issues(db, year)

        # 4. 复核人表现
        data["reviewer_performance"] = await self._get_reviewer_performance(db, year)

        # 5. LLM 综合建议（占位）
        data["llm_suggestions"] = "（LLM 综合改进建议将在此生成，当前为占位内容）"

        # 6. 附录
        data["appendix"] = await self._get_appendix_data(db, year)

        return data

    async def _get_project_scale(self, db: AsyncSession, year: int) -> dict:
        """获取项目规模与分布数据。"""
        # 按 audit_period_end 年份筛选项目
        stmt = (
            select(
                func.count(Project.id).label("total_projects"),
                func.count(
                    sa.case(
                        (Project.project_type == "annual", Project.id),
                        else_=None,
                    )
                ).label("annual_count"),
            )
            .where(
                Project.is_deleted == False,  # noqa: E712
                sa.or_(
                    func.extract("year", Project.audit_period_end) == year,
                    Project.audit_period_end.is_(None),
                ),
            )
        )
        result = await db.execute(stmt)
        row = result.first()

        return {
            "total_projects": row.total_projects if row else 0,
            "annual_count": row.annual_count if row else 0,
            "year": year,
        }

    async def _get_rating_distribution(self, db: AsyncSession, year: int) -> dict:
        """获取评级分布数据（ABCD 饼图）。"""
        stmt = (
            select(
                ProjectQualityRating.rating,
                func.count(ProjectQualityRating.id).label("count"),
            )
            .where(ProjectQualityRating.year == year)
            .group_by(ProjectQualityRating.rating)
        )
        result = await db.execute(stmt)
        rows = result.all()

        distribution = {"A": 0, "B": 0, "C": 0, "D": 0}
        for row in rows:
            if row.rating in distribution:
                distribution[row.rating] = row.count

        total = sum(distribution.values())
        return {
            "distribution": distribution,
            "total": total,
        }

    async def _get_top10_issues(self, db: AsyncSession, year: int) -> list[dict]:
        """获取年度典型问题 Top10（按出现频次排序）。"""
        stmt = (
            select(
                IssueTicket.title,
                IssueTicket.severity,
                func.count(IssueTicket.id).label("occurrence_count"),
            )
            .where(
                func.extract("year", IssueTicket.created_at) == year,
            )
            .group_by(IssueTicket.title, IssueTicket.severity)
            .order_by(func.count(IssueTicket.id).desc())
            .limit(10)
        )
        result = await db.execute(stmt)
        rows = result.all()

        return [
            {
                "title": row.title,
                "severity": row.severity,
                "occurrence_count": row.occurrence_count,
            }
            for row in rows
        ]

    async def _get_reviewer_performance(self, db: AsyncSession, year: int) -> list[dict]:
        """获取复核人表现数据。"""
        stmt = (
            select(ReviewerMetricsSnapshot)
            .where(ReviewerMetricsSnapshot.year == year)
            .order_by(ReviewerMetricsSnapshot.avg_comments_per_wp.desc())
            .limit(20)
        )
        result = await db.execute(stmt)
        snapshots = result.scalars().all()

        return [
            {
                "reviewer_id": str(s.reviewer_id),
                "avg_review_time_min": s.avg_review_time_min,
                "avg_comments_per_wp": s.avg_comments_per_wp,
                "rejection_rate": s.rejection_rate,
                "qc_rule_catch_rate": s.qc_rule_catch_rate,
                "sampled_rework_rate": s.sampled_rework_rate,
            }
            for s in snapshots
        ]

    async def _get_appendix_data(self, db: AsyncSession, year: int) -> dict:
        """获取附录数据（规则变更历史、抽查统计）。"""
        # 抽查统计：该年度的抽查批次数
        from app.models.qc_inspection_models import QcInspection

        inspection_count_stmt = (
            select(func.count(QcInspection.id))
            .where(
                func.extract("year", QcInspection.created_at) == year,
            )
        )
        inspection_result = await db.execute(inspection_count_stmt)
        inspection_count = inspection_result.scalar() or 0

        return {
            "inspection_count": inspection_count,
            "rule_changes_note": "（规则变更历史详见 QC 规则管理页面）",
        }

    def _render_report_text(self, year: int, data: dict) -> str:
        """渲染报告为文本格式（占位实现，实际应为 Word 模板填充）。"""
        lines = [
            f"{'=' * 60}",
            f"  {year} 年度审计质量报告",
            f"{'=' * 60}",
            "",
            "一、项目规模与分布",
            f"  本年度项目总数：{data['project_scale']['total_projects']}",
            f"  年度审计项目数：{data['project_scale']['annual_count']}",
            "",
            "二、评级分布",
            f"  A 级：{data['rating_distribution']['distribution']['A']} 个",
            f"  B 级：{data['rating_distribution']['distribution']['B']} 个",
            f"  C 级：{data['rating_distribution']['distribution']['C']} 个",
            f"  D 级：{data['rating_distribution']['distribution']['D']} 个",
            f"  合计：{data['rating_distribution']['total']} 个",
            "",
            "三、典型问题 Top10",
        ]

        for i, issue in enumerate(data["top10_issues"], 1):
            lines.append(
                f"  {i}. [{issue['severity']}] {issue['title']} "
                f"（出现 {issue['occurrence_count']} 次）"
            )

        if not data["top10_issues"]:
            lines.append("  （本年度无典型问题记录）")

        lines.extend([
            "",
            "四、复核人表现",
            f"  参与复核人数：{len(data['reviewer_performance'])}",
        ])

        for reviewer in data["reviewer_performance"][:5]:
            lines.append(
                f"  - 复核人 {reviewer['reviewer_id'][:8]}...: "
                f"平均时长 {reviewer['avg_review_time_min']:.1f} 分钟, "
                f"平均意见 {reviewer['avg_comments_per_wp']:.1f} 条/张"
            )

        lines.extend([
            "",
            "五、改进建议",
            f"  {data['llm_suggestions']}",
            "",
            "六、附录",
            f"  本年度质控抽查批次数：{data['appendix']['inspection_count']}",
            f"  {data['appendix']['rule_changes_note']}",
            "",
            f"{'=' * 60}",
            f"  报告生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')} UTC",
            f"{'=' * 60}",
            "",
            f"  ※ {build_ai_contribution_statement()}",
        ])

        return "\n".join(lines)


# 模块级单例
qc_annual_report_service = QcAnnualReportService()

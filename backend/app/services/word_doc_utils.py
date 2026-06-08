"""Word 模板占位符 / OPT / NOTE 底层工具（TemplateFillService 共用）."""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"\{\{([a-zA-Z0-9_]+)\}\}")
OPT_OPEN_RE = re.compile(r"^##OPT:([^:#]+)(?::([^#]*))?##\s*$")
OPT_CLOSE_RE = re.compile(r"^##/OPT:([^#]+)##\s*$")
NOTE_OPEN_RE = re.compile(r"^##NOTE(?::([^#]*))?##\s*$")
NOTE_CLOSE_RE = re.compile(r"^##/NOTE(?::([^#]*))?##\s*$")
# 行内 NOTE（单行剥除）
NOTE_INLINE_RE = re.compile(r"##NOTE(?::[^#]*)?##.*?##/NOTE(?:#[^#]*)?##", re.DOTALL)

# 附注 ``##SECTION:code##`` / ``##/SECTION:code##`` 块（Phase 0.6.2 打标产出）
SECTION_OPEN_RE = re.compile(r"^##SECTION:(.+?)##\s*$")
SECTION_CLOSE_RE = re.compile(r"^##/SECTION:(.+?)##\s*$")
# 参考表样式标记（克隆后删除）
STYLE_REF_RE = re.compile(r"##/?STYLE_REF:[^#]*##")


@dataclass(frozen=True)
class OptionalSection:
    section_id: str
    description: str
    start_index: int
    end_index: int  # exclusive


def _iter_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for para in part.paragraphs:
                yield para
            for table in part.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def merge_runs_for_replace(paragraph: Paragraph) -> None:
    """合并段落内 runs，避免 ``{{key}}`` 被 Word 拆成多 run 无法替换."""
    if len(paragraph.runs) <= 1:
        return
    combined = paragraph.text
    if "{{" not in combined and "{" not in combined:
        return
    first = paragraph.runs[0]
    first.text = combined
    for run in paragraph.runs[1:]:
        run.text = ""


def merge_runs_in_doc(doc: DocumentObject) -> None:
    for para in _iter_paragraphs(doc):
        merge_runs_for_replace(para)


def replace_placeholders_in_doc(doc: DocumentObject, mapping: dict[str, str]) -> int:
    """全文替换 ``{{canonical_key}}``；返回替换次数."""
    merge_runs_in_doc(doc)
    count = 0
    for para in _iter_paragraphs(doc):
        text = para.text
        if "{{" not in text:
            continue
        new_text = text
        for key, value in mapping.items():
            token = f"{{{{{key}}}}}"
            if token in new_text:
                count += new_text.count(token)
                new_text = new_text.replace(token, value or "")
        if new_text != text:
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_text)
    return count


def scan_optional_sections(doc: DocumentObject) -> list[OptionalSection]:
    """解析 ``##OPT:id:desc## … ##/OPT:id##`` 块（段落级）."""
    paras = list(doc.paragraphs)
    sections: list[OptionalSection] = []
    i = 0
    while i < len(paras):
        text = (paras[i].text or "").strip()
        m_open = OPT_OPEN_RE.match(text)
        if not m_open:
            i += 1
            continue
        sec_id = m_open.group(1).strip()
        desc = (m_open.group(2) or "").strip()
        j = i + 1
        while j < len(paras):
            close_text = (paras[j].text or "").strip()
            m_close = OPT_CLOSE_RE.match(close_text)
            if m_close and m_close.group(1).strip() == sec_id:
                sections.append(
                    OptionalSection(sec_id, desc, i, j + 1)
                )
                i = j + 1
                break
            j += 1
        else:
            i += 1
    return sections


def _delete_paragraph(para: Paragraph) -> None:
    element = para._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def apply_optional_sections(
    doc: DocumentObject,
    selections: dict[str, bool],
) -> int:
    """删除未勾选 OPT 块；保留块去除开闭标记行。返回删除段落数."""
    removed = 0
    # 从后往前处理避免索引漂移
    for sec in reversed(scan_optional_sections(doc)):
        keep = selections.get(sec.section_id, False)
        paras = doc.paragraphs
        if keep:
            # 去掉标记行文字
            for idx in (sec.end_index - 1, sec.start_index):
                if 0 <= idx < len(paras):
                    paras[idx].clear()
            continue
        for idx in range(sec.end_index - 1, sec.start_index - 1, -1):
            if 0 <= idx < len(paras):
                _delete_paragraph(paras[idx])
                removed += 1
    return removed


def strip_guidance_notes(doc: DocumentObject) -> int:
    """删除 ``##NOTE:…##`` 块及行内 NOTE；合法 ``[注1]`` 保留。返回删除段落数."""
    removed = 0
    while True:
        paras = doc.paragraphs
        deleted_block = False
        for i, para in enumerate(paras):
            text = (para.text or "").strip()
            m_open = NOTE_OPEN_RE.match(text)
            if m_open:
                label = m_open.group(1)
                for j in range(i + 1, len(paras)):
                    close_text = (paras[j].text or "").strip()
                    m_close = NOTE_CLOSE_RE.match(close_text)
                    if m_close and (not label or m_close.group(1) == label):
                        for idx in range(j, i - 1, -1):
                            _delete_paragraph(doc.paragraphs[idx])
                            removed += 1
                        deleted_block = True
                        break
                if not deleted_block:
                    _delete_paragraph(paras[i])
                    removed += 1
                    deleted_block = True
                break
            if "##NOTE" in text and "##/NOTE" in text:
                new_text = NOTE_INLINE_RE.sub("", text).strip()
                if not new_text:
                    _delete_paragraph(para)
                    removed += 1
                else:
                    para.clear()
                    para.add_run(new_text)
                deleted_block = True
                break
        if not deleted_block:
            break
    return removed


def copy_template_to_workdir(src: Path, dest_dir: Path) -> Path:
    """复制模板到工作目录，保留元数据."""
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / src.name
    shutil.copy2(src, dest)
    return dest


# ---------------------------------------------------------------------------
# ##SECTION:code## 块解析（附注模板模式专用）
# ---------------------------------------------------------------------------


@dataclass
class SectionBlock:
    """一个 ``##SECTION:code##…##/SECTION:code##`` 块.

    ``elements`` 为块内（含开闭标记段落）的 body 级 XML 元素列表，
    按文档顺序排列（``<w:p>`` 段落与 ``<w:tbl>`` 表格交错）。
    ``open_el`` / ``close_el`` 为开闭标记段落的 XML 元素。
    """

    section_code: str
    open_el: Any
    close_el: Any
    elements: list[Any]


def _body_block_elements(doc: DocumentObject) -> list[Any]:
    """按文档顺序返回 body 下的段落 / 表格 XML 元素."""
    body = doc.element.body
    result: list[Any] = []
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    for child in body.iterchildren():
        if child.tag in (p_tag, tbl_tag):
            result.append(child)
    return result


def scan_section_blocks(doc: DocumentObject) -> list[SectionBlock]:
    """解析全文 ``##SECTION:code##…##/SECTION:code##`` 块（body 级，含表格）.

    与 ``scan_optional_sections`` 不同：SECTION 块需覆盖块内**表格**（附注主体内容），
    因此遍历 body 级 XML 元素而非仅 ``doc.paragraphs``。
    标记文本写在 ``<w:p>`` 段落里；表格元素无标记文本会落入当前打开块的 elements。
    """
    p_tag = qn("w:p")
    blocks: list[SectionBlock] = []
    elements = _body_block_elements(doc)

    open_idx: int | None = None
    open_code: str | None = None
    for i, el in enumerate(elements):
        if el.tag != p_tag:
            continue
        text = (Paragraph(el, doc).text or "").strip()
        if open_idx is None:
            m_open = SECTION_OPEN_RE.match(text)
            if m_open:
                open_idx = i
                open_code = m_open.group(1).strip()
            continue
        # 已在块内，找匹配 close
        m_close = SECTION_CLOSE_RE.match(text)
        if m_close and m_close.group(1).strip() == open_code:
            blocks.append(
                SectionBlock(
                    section_code=open_code or "",
                    open_el=elements[open_idx],
                    close_el=el,
                    elements=elements[open_idx : i + 1],
                )
            )
            open_idx = None
            open_code = None
    return blocks


def delete_section_block(block: SectionBlock) -> int:
    """删除整个 SECTION 块（含开闭标记 + 块内所有段落 / 表格）。返回删除元素数."""
    removed = 0
    for el in block.elements:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    return removed


def remove_section_markers(doc: DocumentObject) -> int:
    """删除保留章节里残留的 ``##SECTION:##`` / ``##/SECTION:##`` 标记段落.

    同时清理行内 ``##STYLE_REF:##`` 标记文本。返回删除 / 清理的段落数。
    """
    touched = 0
    for para in list(doc.paragraphs):
        text = (para.text or "").strip()
        if SECTION_OPEN_RE.match(text) or SECTION_CLOSE_RE.match(text):
            _delete_paragraph(para)
            touched += 1
            continue
        if "##STYLE_REF:" in text:
            new_text = STYLE_REF_RE.sub("", para.text).strip()
            if not new_text:
                _delete_paragraph(para)
            else:
                para.clear()
                para.add_run(new_text)
            touched += 1
    return touched

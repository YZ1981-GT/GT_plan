"""Sprint C.4.7 — 多公司基线对比 PDF 工具.

主要 API:
- compare_baselines_to_pdf(baseline_ids, output_path) → bytes
  生成多个集团基线之间的差异对比 PDF（章节级 + 字段级）

- generate_baseline_comparison_report(comparisons) → bytes
  纯函数版本，接受已计算好的对比数据
"""
from __future__ import annotations

from io import BytesIO
from typing import Any
from uuid import UUID

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

from app.services.note_word_exporter import (
    apply_gt_three_line,
    apply_gt_row_height,
    _set_run_font,
    _set_paragraph_format,
)

__all__ = [
    "BaselineComparison",
    "generate_baseline_comparison_report",
    "compute_baseline_diff_data",
]


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


class BaselineComparison:
    """单个章节在多基线间的对比结果."""

    __slots__ = (
        "section_id",
        "section_title",
        "baselines_present",
        "baselines_absent",
        "field_diffs",
    )

    def __init__(
        self,
        section_id: str,
        section_title: str,
        baselines_present: list[str],
        baselines_absent: list[str],
        field_diffs: list[dict[str, Any]] | None = None,
    ):
        self.section_id = section_id
        self.section_title = section_title
        self.baselines_present = baselines_present
        self.baselines_absent = baselines_absent
        self.field_diffs = field_diffs or []


# ---------------------------------------------------------------------------
# Diff computation
# ---------------------------------------------------------------------------


def compute_baseline_diff_data(
    baselines: list[dict[str, Any]],
) -> list[BaselineComparison]:
    """Compute section-level diff across multiple baselines.

    Args:
        baselines: list of {id, name, sections: [{section_id, section_title, table_data}]}

    Returns:
        list of BaselineComparison objects
    """
    if not baselines:
        return []

    # Collect all unique section_ids across all baselines
    all_section_ids: dict[str, str] = {}  # section_id → first seen title
    for b in baselines:
        for sec in b.get("sections", []):
            sid = sec.get("section_id", "")
            if sid and sid not in all_section_ids:
                all_section_ids[sid] = sec.get("section_title", "")

    # For each section, compute presence and field diffs
    results: list[BaselineComparison] = []
    for sid, title in all_section_ids.items():
        present = []
        absent = []
        for b in baselines:
            b_name = b.get("name", b.get("id", "未命名"))
            sec = next(
                (s for s in b.get("sections", []) if s.get("section_id") == sid),
                None,
            )
            if sec:
                present.append(b_name)
            else:
                absent.append(b_name)

        # Compute field-level diffs (compare table_data structure)
        field_diffs = _compute_field_diffs(baselines, sid)

        results.append(
            BaselineComparison(
                section_id=sid,
                section_title=title,
                baselines_present=present,
                baselines_absent=absent,
                field_diffs=field_diffs,
            )
        )

    return results


def _compute_field_diffs(
    baselines: list[dict[str, Any]], section_id: str
) -> list[dict[str, Any]]:
    """Compute field-level differences for a section across baselines."""
    diffs: list[dict[str, Any]] = []

    # Collect headers from each baseline
    headers_by_baseline: dict[str, list[str]] = {}
    rows_count_by_baseline: dict[str, int] = {}

    for b in baselines:
        b_name = b.get("name", b.get("id", "未命名"))
        sec = next(
            (s for s in b.get("sections", []) if s.get("section_id") == section_id),
            None,
        )
        if sec is None:
            continue
        td = sec.get("table_data") or {}
        headers_by_baseline[b_name] = list(td.get("headers", []))
        rows_count_by_baseline[b_name] = len(td.get("rows", []))

    if not headers_by_baseline:
        return diffs

    # Compare headers
    all_headers_match = (
        len(set(tuple(h) for h in headers_by_baseline.values())) == 1
    )
    if not all_headers_match:
        diffs.append(
            {
                "type": "headers_differ",
                "description": "表头列结构不一致",
                "details": {k: v for k, v in headers_by_baseline.items()},
            }
        )

    # Compare row counts
    all_rows_match = len(set(rows_count_by_baseline.values())) == 1
    if not all_rows_match:
        diffs.append(
            {
                "type": "row_count_differs",
                "description": "行数不一致",
                "details": {k: v for k, v in rows_count_by_baseline.items()},
            }
        )

    return diffs


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------


def generate_baseline_comparison_report(
    baselines: list[dict[str, Any]],
    title: str = "集团附注基线对比报告",
) -> bytes:
    """Generate Word report comparing multiple baselines (returned as bytes).

    PDF generation requires LibreOffice/external converter — this returns Word
    that can be converted by the caller.

    Args:
        baselines: list of baseline dicts with sections
        title: report title

    Returns:
        Word document bytes
    """
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    _set_run_font(title_run, bold=True, size=Pt(16))
    _set_paragraph_format(title_p, space_after=Pt(18))

    # Summary
    summary_p = doc.add_paragraph()
    summary_run = summary_p.add_run(
        f"对比基线数量: {len(baselines)} 个\n"
        f"基线列表: {', '.join(b.get('name', '未命名') for b in baselines)}"
    )
    _set_run_font(summary_run)

    # Compute diffs
    comparisons = compute_baseline_diff_data(baselines)

    # Section presence table
    doc.add_heading("一、章节存在性矩阵", level=1)
    presence_table = doc.add_table(
        rows=len(comparisons) + 1, cols=len(baselines) + 1
    )
    apply_gt_three_line(presence_table)
    presence_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ["章节"] + [b.get("name", b.get("id", "未命名")) for b in baselines]
    for col_idx, h in enumerate(headers):
        cell = presence_table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _set_run_font(run, bold=True)
    apply_gt_row_height(presence_table.rows[0], cm=0.7)

    # Data rows
    for r_idx, comp in enumerate(comparisons, start=1):
        row = presence_table.rows[r_idx]
        apply_gt_row_height(row, cm=0.7)
        # Section title
        cell = row.cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(comp.section_title or comp.section_id)
        _set_run_font(run)
        # Presence per baseline
        for b_idx, b in enumerate(baselines, start=1):
            b_name = b.get("name", b.get("id", "未命名"))
            cell = row.cells[b_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mark = "✓" if b_name in comp.baselines_present else "✗"
            run = p.add_run(mark)
            _set_run_font(run)

    # Field diff details
    doc.add_paragraph()
    doc.add_heading("二、字段级差异", level=1)
    diffs_with_details = [c for c in comparisons if c.field_diffs]
    if not diffs_with_details:
        p = doc.add_paragraph()
        run = p.add_run("（所有共有章节字段结构一致）")
        run.italic = True
        _set_run_font(run)
    else:
        for comp in diffs_with_details:
            doc.add_heading(
                f"{comp.section_title or comp.section_id}", level=2
            )
            for diff in comp.field_diffs:
                p = doc.add_paragraph()
                run = p.add_run(f"• {diff['description']}")
                _set_run_font(run, bold=True)
                if diff.get("details"):
                    for k, v in diff["details"].items():
                        p2 = doc.add_paragraph(style="List Bullet")
                        run2 = p2.add_run(f"{k}: {v}")
                        _set_run_font(run2)

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

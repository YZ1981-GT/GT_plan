"""QC 报告 Word 导出 API — Phase 7 F5

GET /api/projects/{id}/qc-report/export: 返回 .docx 文件流
- 三线表格式（仅顶线+表头底线+底线，无竖线）
- 字体：仿宋_GB2312（中文）+ Arial Narrow（数字）
- 三章节：风险汇总表 + 意见清单表 + 整改状态表
- 无 QC 记录时返回空报告模板（含表头无数据行）
- 导出操作写入 audit_log
- 权限：仅 qc/admin 可访问

注册到 router_registry 协作域 §109。

Validates: Requirements F5.1, F5.2, F5.3, F5.4, F5.5, F5.6, F5.7
"""

import io
import json
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import text as sql_text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.deps import get_current_user
from app.models.core import User

router = APIRouter(
    prefix="/api/projects/{project_id}/qc-report",
    tags=["qc-report"],
)


# ---------------------------------------------------------------------------
# Permission check
# ---------------------------------------------------------------------------


def _check_qc_admin(user: User) -> None:
    """仅 qc/admin 可访问"""
    if user.role.value not in ("qc", "admin"):
        raise HTTPException(status_code=403, detail="仅 QC/管理员可导出报告")


# ---------------------------------------------------------------------------
# Word document generation helpers
# ---------------------------------------------------------------------------


def _create_three_line_table(doc, headers: list[str], rows: list[list[str]]):
    """创建三线表格式表格（顶线+表头底线+底线，无竖线）"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders first
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

    # Set header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(10)

    # Set data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(val) if val else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Use Arial Narrow for numbers
                    if val and any(c.isdigit() for c in str(val)):
                        run.font.name = 'Arial Narrow'
                    else:
                        run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(10)

    # Apply three-line borders: top of table, bottom of header, bottom of table
    def _set_row_border(row_obj, position: str, sz: str = '12'):
        """Set border on a specific row position (top/bottom)"""
        for cell in row_obj.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            border = OxmlElement(f'w:{position}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
            tcPr.append(tcBorders)

    # Top line (header top)
    _set_row_border(table.rows[0], 'top', '12')
    # Header bottom line
    _set_row_border(table.rows[0], 'bottom', '6')
    # Bottom line (last row bottom)
    if len(table.rows) > 0:
        _set_row_border(table.rows[-1], 'bottom', '12')

    return table


def _build_qc_report(project_name: str, risk_data: list, opinion_data: list, rectification_data: list) -> io.BytesIO:
    """构建 QC 报告 Word 文档"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    title = doc.add_heading(f'质量控制报告 — {project_name}', level=1)
    for run in title.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    doc.add_paragraph(f'导出时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')

    # Chapter 1: 风险汇总表
    doc.add_heading('一、风险汇总表', level=2)
    risk_headers = ['循环', 'Blocking', 'Warning', 'Info', '合计']
    _create_three_line_table(doc, risk_headers, risk_data)

    doc.add_paragraph('')  # spacing

    # Chapter 2: 意见清单表
    doc.add_heading('二、意见清单表', level=2)
    opinion_headers = ['编号', '严重程度', '类别', '标题', '状态', '创建时间']
    _create_three_line_table(doc, opinion_headers, opinion_data)

    doc.add_paragraph('')

    # Chapter 3: 整改状态表
    doc.add_heading('三、整改状态表', level=2)
    rect_headers = ['编号', '标题', '整改状态', '整改人', '整改时间']
    _create_three_line_table(doc, rect_headers, rectification_data)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# GET /api/projects/{project_id}/qc-report/export
# ---------------------------------------------------------------------------


@router.get("/export")
async def export_qc_report(
    project_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """导出 QC 报告 Word 文件（三线表格式）"""
    _check_qc_admin(current_user)

    # Get project name
    result = await db.execute(
        sql_text("SELECT name FROM projects WHERE id = :pid"),
        {"pid": str(project_id)},
    )
    project_row = result.first()
    project_name = project_row[0] if project_row else "未知项目"

    # Chapter 1: Risk summary — aggregate VR trigger stats by cycle
    # Query issue_tickets with source='qc_inspection' for this project
    risk_result = await db.execute(
        sql_text("""
            SELECT
                COALESCE(category, '未分类') as cycle,
                COUNT(*) FILTER (WHERE severity = 'blocker') as blocking,
                COUNT(*) FILTER (WHERE severity = 'major') as warning,
                COUNT(*) FILTER (WHERE severity IN ('minor', 'suggestion')) as info
            FROM issue_tickets
            WHERE project_id = :pid AND source = 'qc_inspection'
            GROUP BY category
            ORDER BY category
        """),
        {"pid": str(project_id)},
    )
    risk_rows = risk_result.fetchall()
    risk_data = [
        [row[0], str(row[1]), str(row[2]), str(row[3]), str(row[1] + row[2] + row[3])]
        for row in risk_rows
    ]

    # Chapter 2: Opinion list
    opinion_result = await db.execute(
        sql_text("""
            SELECT id, severity, category, title, status, created_at
            FROM issue_tickets
            WHERE project_id = :pid AND source = 'qc_inspection'
            ORDER BY severity DESC, created_at ASC
            LIMIT 200
        """),
        {"pid": str(project_id)},
    )
    opinion_rows = opinion_result.fetchall()
    opinion_data = [
        [
            str(i + 1),
            row[1] or "",
            row[2] or "",
            row[3] or "",
            row[4] or "",
            row[5].strftime("%Y-%m-%d") if row[5] else "",
        ]
        for i, row in enumerate(opinion_rows)
    ]

    # Chapter 3: Rectification status
    rect_result = await db.execute(
        sql_text("""
            SELECT it.id, it.title, it.status,
                   COALESCE(u.username, '') as assignee,
                   it.updated_at
            FROM issue_tickets it
            LEFT JOIN users u ON it.owner_id = u.id
            WHERE it.project_id = :pid AND it.source = 'qc_inspection'
                  AND it.status IN ('in_fix', 'closed')
            ORDER BY it.updated_at DESC
            LIMIT 200
        """),
        {"pid": str(project_id)},
    )
    rect_rows = rect_result.fetchall()
    rectification_data = [
        [
            str(i + 1),
            row[1] or "",
            row[2] or "",
            row[3] or "",
            row[4].strftime("%Y-%m-%d") if row[4] else "",
        ]
        for i, row in enumerate(rect_rows)
    ]

    # Build document
    buffer = _build_qc_report(project_name, risk_data, opinion_data, rectification_data)

    # Write to audit_log
    try:
        await db.execute(
            sql_text(
                "INSERT INTO audit_log (id, user_id, action, resource_type, resource_id, details, created_at) "
                "VALUES (:id, :uid, :action, :rtype, :rid, :details::jsonb, :now)"
            ),
            {
                "id": str(uuid.uuid4()),
                "uid": str(current_user.id),
                "action": "qc_report_export",
                "rtype": "project",
                "rid": str(project_id),
                "details": json.dumps(
                    {"project_name": project_name, "risk_count": len(risk_data), "opinion_count": len(opinion_data)},
                    ensure_ascii=False,
                ),
                "now": datetime.utcnow(),
            },
        )
        await db.commit()
    except Exception:
        pass  # audit_log failure should not block export

    filename = f"QC_Report_{project_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    # 中文文件名需 RFC 5987 编码（HTTP 头按 latin-1 编码，直接放中文会 UnicodeEncodeError）
    from urllib.parse import quote
    ascii_name = filename.encode("ascii", "ignore").decode() or "QC_Report.docx"
    utf8_name = quote(filename, safe="")
    disposition = f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": disposition},
    )

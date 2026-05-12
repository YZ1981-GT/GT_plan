"""QC 抽查路由 — Round 3 需求 4

POST   /api/qc/inspections                         — 创建抽查批次
GET    /api/qc/inspections                         — 列出抽查批次
GET    /api/qc/inspections/{id}                    — 获取抽查详情
POST   /api/qc/inspections/{id}/items/{item_id}/verdict — 录入结论

权限：role='qc' | 'admin'
"""

from __future__ import annotations

import logging
from typing import Literal, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, Query
from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.deps import get_current_user, require_role
from app.models.core import User
from app.services.qc_inspection_service import qc_inspection_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/qc/inspections", tags=["qc-inspections"])


# ---------------------------------------------------------------------------
# Request / Response schemas
# ---------------------------------------------------------------------------


class CreateInspectionRequest(BaseModel):
    """创建抽查批次请求体"""

    project_id: UUID = Field(..., description="项目 ID")
    strategy: Literal["random", "risk_based", "full_cycle", "mixed"] = Field(
        ..., description="抽样策略"
    )
    params: Optional[dict] = Field(
        None,
        description="策略参数，如 {ratio: 0.1} / {cycles: ['D']} / {random_ratio: 0.2, cycles: ['D']}",
    )
    reviewer_id: UUID = Field(..., description="质控复核人 ID")


class VerdictRequest(BaseModel):
    """录入结论请求体"""

    verdict: Literal["pass", "fail", "conditional_pass"] = Field(
        ..., description="结论: pass / fail / conditional_pass"
    )
    findings: Optional[dict] = Field(
        None, description="发现的问题（JSONB）"
    )


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("", status_code=201)
async def create_inspection(
    body: CreateInspectionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["qc", "admin"])),
):
    """创建质控抽查批次，按策略生成抽查子项。"""
    result = await qc_inspection_service.create_inspection(
        db,
        project_id=body.project_id,
        strategy=body.strategy,
        params=body.params,
        reviewer_id=body.reviewer_id,
    )
    await db.commit()
    return result


@router.get("")
async def list_inspections(
    project_id: Optional[UUID] = Query(None, description="按项目过滤"),
    page: int = Query(1, ge=1, description="页码"),
    page_size: int = Query(20, ge=1, le=100, description="每页条数"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["qc", "admin"])),
):
    """列出质控抽查批次。"""
    return await qc_inspection_service.list_inspections(
        db,
        project_id=project_id,
        page=page,
        page_size=page_size,
    )


@router.get("/{inspection_id}")
async def get_inspection(
    inspection_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["qc", "admin"])),
):
    """获取抽查批次详情（含子项）。"""
    return await qc_inspection_service.get_inspection(db, inspection_id)


@router.post("/{inspection_id}/items/{item_id}/verdict")
async def record_verdict(
    inspection_id: UUID,
    item_id: UUID,
    body: VerdictRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["qc", "admin"])),
):
    """QC 复核人录入结论（pass / fail / conditional_pass）。"""
    result = await qc_inspection_service.record_verdict(
        db,
        inspection_id=inspection_id,
        item_id=item_id,
        verdict=body.verdict,
        findings=body.findings,
    )
    await db.commit()
    return result


@router.get("/{inspection_id}/report")
async def generate_inspection_report(
    inspection_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_role(["qc", "admin"])),
):
    """生成质控抽查报告 Word 文件并返回下载流。

    使用 python-docx 生成结构化报告，包含批次信息、抽查子项及结论。
    若 python-docx 不可用则降级为纯文本。
    """
    from fastapi.responses import Response

    # 获取抽查详情
    detail = await qc_inspection_service.get_inspection(db, inspection_id)
    if not detail:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="抽查批次不存在")

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("质控抽查报告")
        run.font.size = Pt(18)
        run.bold = True

        # 基本信息
        doc.add_heading("一、抽查批次信息", level=1)
        doc.add_paragraph(f"项目：{detail.get('project_name', '未知')}")
        doc.add_paragraph(f"策略：{detail.get('strategy', '未知')}")
        doc.add_paragraph(f"状态：{detail.get('status', '未知')}")
        doc.add_paragraph(f"创建时间：{detail.get('created_at', '未知')}")

        # 抽查子项
        items = detail.get("items", [])
        doc.add_heading("二、抽查结果明细", level=1)
        if items:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "序号"
            hdr[1].text = "底稿"
            hdr[2].text = "结论"
            hdr[3].text = "发现"
            for idx, item in enumerate(items, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = str(item.get("wp_id", ""))[:12]
                row[2].text = item.get("qc_verdict", item.get("verdict", "待评定"))
                findings = item.get("findings", {})
                row[3].text = str(findings) if findings else "无"
        else:
            doc.add_paragraph("暂无抽查子项。")

        # 汇总
        doc.add_heading("三、汇总", level=1)
        total = len(items)
        passed = sum(1 for i in items if i.get("qc_verdict") == "pass" or i.get("verdict") == "pass")
        failed = sum(1 for i in items if i.get("qc_verdict") == "fail" or i.get("verdict") == "fail")
        doc.add_paragraph(f"总计：{total} 项 | 通过：{passed} | 不通过：{failed} | 有条件通过：{total - passed - failed}")

        # 输出
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        content = buffer.getvalue()

        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="inspection_report_{inspection_id}.docx"'},
        )

    except ImportError:
        # python-docx 不可用，降级为文本
        lines = [
            "质控抽查报告",
            f"批次ID: {inspection_id}",
            f"项目: {detail.get('project_name', '未知')}",
            f"策略: {detail.get('strategy', '未知')}",
            f"状态: {detail.get('status', '未知')}",
            "",
            "抽查子项:",
        ]
        for idx, item in enumerate(detail.get("items", []), 1):
            lines.append(f"  {idx}. {item.get('wp_id', '')} - {item.get('qc_verdict', '待评定')}")

        content = "\n".join(lines).encode("utf-8")
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="inspection_report_{inspection_id}.txt"'},
        )
